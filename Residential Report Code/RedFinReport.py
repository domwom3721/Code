#Date: 5/2/2022
#Author: Mike Leahy
#Summary: Injests RedFin residential real estate data and produces report documents on the selcted areas

import os
from turtle import fillcolor
from numpy import True_
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from tkinter import *
from tkinter import ttk
import us
from datetime import date, datetime
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from validator_collection import none

#Define file pre-paths
dropbox_root                   =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
project_location               =  os.path.join(dropbox_root,'Research','Projects','Research Report Automation Project') 
data_location                  =  os.path.join(project_location,'Data\Residential Reports Data\RedFin Data\Clean') 

#Data Related functions
def DetermineSubjectAndComp():
    #This function presents the user with options for what geographic level, area, and report type (Condo vs SF) they would like to create a report for as well as the comparison area \
    #Creates 2 pandas dataframes, 1 with data on the subject, and another with data on the comparsion area
    global df_subject, df_comparison
    
    #Import our clean RedFin data
    df = pd.read_csv(os.path.join(data_location, 'Clean RedFin Data.csv'), 
                dtype={         'Type': str,
                                'Region':str,
                                'Month of Period End':object,
                                'Median Sale Price':float	,
                                'Median Sale Price MoM':float ,	
                                'Median Sale Price YoY':float ,	
                                'Homes Sold':float,
                                'Homes Sold MoM':float,
                                'Homes Sold YoY':float,	
                                'New Listings':float,
                                'New Listings MoM':float, 	
                                'New Listings YoY':float,	
                                'Inventory':float,
                                'Inventory MoM':float,	 
                                'Inventory YoY':float,	
                                'Days on Market':float,	
                                'Days on Market MoM':float ,	
                                'Days on Market YoY':float	,
                                'Average Sale To List':float ,
                                'Average Sale To List MoM':float ,	
                                'Average Sale To List YoY':float,
                                },
                                parse_dates=['Month of Period End'],
                            )
    
  
    
    
    #Create a list with all possible subject areas
    possible_subject_areas = df['Unique Subject Name'].unique().tolist()


    #Launches GUI for user to select subject area
    def select_market(event):
        global  selected_subject
        selected_subject = comboExample.get()
        
    app = Tk() 
    app.geometry('600x300')
    app.config(bg='#404858')
    app.title('Research Automation Project - Residential Reports') 

    labelTop = Label(app, text = ("Choose your subject area"))
    labelTop.grid(column=0, row=0)

    comboExample = ttk.Combobox(app, values = possible_subject_areas, width=50)

    comboExample.grid(column=0, row=1)
    comboExample.current(0)
    comboExample.bind("<<ComboboxSelected>>", select_market)
    app.mainloop()



    #Launches GUI for user to select comparsion area
    def select_comp(event):
        global  selected_comparsion
        selected_comparsion = comboExample.get()
        
    app = Tk() 
    app.geometry('600x300')
    app.config(bg='#404858')
    app.title('Research Automation Project - Residential Reports') 

    labelTop = Label(app, text = ("Choose your comparison area"))
    labelTop.grid(column=0, row=0)

    comboExample = ttk.Combobox(app, values = possible_subject_areas, width=50)

    comboExample.grid(column=0, row=1)
    comboExample.current(0)
    comboExample.bind("<<ComboboxSelected>>", select_comp)
    app.mainloop()


    df_subject    = df.loc[df['Unique Subject Name']== selected_subject].copy()
    df_comparison = df.loc[df['Unique Subject Name']== selected_comparsion].copy()

def GetDataForOverviewTable():
    #Each list in the list of lists is a row in the overivew table
    row1  = ['',subject_name,'YoY','MoM',comparison_name,'YoY','MoM']
    row2  = ['Median Sales Price', 2, 3, 4, 5, 6, 7]
    row3  = ['Price Per Sqft', 2,3,4,5,6,7]
    row4  = ['Homes Sold', 2,3,4,5,6,7]
    row5  = ['Inventory', 2,3,4,5,6,7]
    row6  = ['Days on Market',2,3,4,5,6,7]
    row7  = ['Average Sale To List',2,3,4,5,6,7]

    
    return([row1, row2, row3, row4, row5, row6, row7])

def CreateDirectory():
    global report_path, report_folder
    #This function creates a folder for the report within a 2 letter state folder, it also returns a file path for the report document
    if subject_geo_level == 'Place' or  subject_geo_level == "County":
        state_code  = subject_name.split(', ')[1]
        folder_name = subject_name.split(', ')[0]
    
    elif subject_geo_level == 'Metro':
        state_code  = subject_name.split(', ')[1][0:2]
        folder_name = subject_name.split(', ')[0] + " Metro Area"


    elif subject_geo_level == 'State':
        state_code  =us.states.lookup(subject_name).abbr  
    else:
        assert False
    
    #Make State Folder
    assert len(state_code) == 2
    state_folder_path = os.path.join(output_location,state_code) 
    if os.path.exists(state_folder_path) == False:
        os.mkdir(state_folder_path)

    #Make Report folder
    if subject_geo_level != 'State':
        report_folder = os.path.join(state_folder_path,folder_name) 
        if os.path.exists(report_folder) == False:
             os.mkdir(report_folder)
        

    elif subject_geo_level == 'State':
        report_folder = state_folder_path 


    document_name =  ('2022 Q1' + ' - ' + state_code + ' - ' +  subject_name.split(', ')[0].replace('/','')  + ' - ' + subject_property_type + '_draft.docx')
    report_path = os.path.join(report_folder,document_name)

def CleanUpPNGs():
    #Report writing done, delete figures
    files = os.listdir(report_folder)
    for image in files:
        if image.endswith(".png"):
            while os.path.exists(os.path.join(report_folder, image)):
                try:
                    os.remove(os.path.join(report_folder, image))
                except Exception as e: 
                    print(e)

#Language Related functions
def OverviewLanguage():
    try:
        return(['Overview language'])
    except Exception as e:
        print(e, 'Unable to create overview language')

def SupplyDemandLanguage():
    try:
        return(['Supply and Demand language'])
    except Exception as e:
        print(e, 'Unable to create supply and demand language')

def ValuesLanguage():
    try:
        return(['Values language'])
    except Exception as e:
        print(e, 'Unable to create values language')

def ConclusionLanguage():
    try:
        return(['Outlook language'])
    except Exception as e:
        print(e, 'Unable to create conclusion language')

def CreateLanguage():
    print('Creating Language')
    global overview_language, supply_and_demand_language, values_language, conclusion_language
    overview_language          = OverviewLanguage()
    supply_and_demand_language = SupplyDemandLanguage()
    values_language            = ValuesLanguage()
    conclusion_language        = ConclusionLanguage()

#Graph related functions
def CreateGraphs():
    print('Creating Graphs')
    CreateHomesSoldGraph()
    CreateDaysOnMarketGraph()
    CreateMedianSalePriceGraph()

def CreateHomesSoldGraph():


    #Create figure with secondary y-axis
    fig = make_subplots(specs=[[{"secondary_y": True}]])
    #Add Inventory
    fig.add_trace(
    go.Scatter(x        = df_subject['Month of Period End'],
           y            = df_subject['Inventory'],
           name         = 'Inventory (L)',
           mode         = 'none',
           fill         = 'tozeroy',
           fillcolor    = bowery_dark_blue,
           ),
           secondary_y=False
                )

    #Add Sales Volume
    fig.add_trace(
    go.Scatter(x        = df_subject['Month of Period End'],
           y            = df_subject['Homes Sold'],
           name         = '# of Homes Sold (L)',
           marker_color = bowery_grey,
           mode         = 'none',
            fill        = 'tozeroy',
            fillcolor   = bowery_grey
           ),
           secondary_y=False
                )
    
    #Add Avg Sale to List Ratio
    fig.add_trace(
    go.Scatter(x        = df_subject['Month of Period End'],
           y            = df_subject['Average Sale To List'],
           name         = 'Average Sale To List (R)',
           mode         = 'lines',
           line = dict(color = bowery_black, dash = 'dash') 
          
           ),
           secondary_y=True_
                )

    #Set formatting 
    fig.update_layout(
        title_text    = "",    
        font_family   = font_family,
        font_color    = font_color,
        font_size     = font_size,
        height        = graph_height,
        width         = graph_width,
        margin        = dict(l = left_margin, r = right_margin, t = top_margin, b = bottom_margin),
        paper_bgcolor = backgroundcolor,
        plot_bgcolor  = backgroundcolor,
        
        title = {
            'y':       title_position,
            'x':       0.5,
            'xanchor': 'center',
            'yanchor': 'top'
                },

        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

                  )

    fig.update_xaxes(
        tickmode = 'array',
        tickfont = dict(size=tickfont_size)
                    )
    
    #Set y axis format
    fig.update_yaxes(tickfont = dict(size=tickfont_size), secondary_y = False,)  #left axis
    fig.update_yaxes(tickfont = dict(size=tickfont_size), secondary_y = True, ticksuffix = '%')  #right
    
    #Export figure as PNG file
    fig.write_image(os.path.join(report_folder,'sales_volume.png'), engine = 'kaleido', scale = scale)

def CreateDaysOnMarketGraph():
    #Create figure with secondary y-axis
    fig = make_subplots(specs=[[{"secondary_y": True}]])
    
    #Add Days on market for subject area
    fig.add_trace(
    go.Scatter(x        = df_subject['Month of Period End'],
           y            = df_subject['Days on Market'],
           name         = subject_name,
           mode         = 'lines',
          line = dict(color = bowery_dark_blue, ) 
           ),
           secondary_y=False
                )

    #Add Days on market for comparsion area
    fig.add_trace(
    go.Scatter(x        = df_comparison['Month of Period End'],
           y            = df_comparison['Days on Market'],
           name         = comparison_name,
           mode         = 'lines',
           line = dict(color = bowery_light_blue,) 
          
           ),
           secondary_y=False
                )

    #Set formatting 
    fig.update_layout(
        title_text    = "Average Days on Market",    
        font_family   = font_family,
        font_color    = font_color,
        font_size     = font_size,
        height        = graph_height,
        width         = graph_width,
        margin        = dict(l = left_margin, r = right_margin, t = top_margin, b = bottom_margin),
        paper_bgcolor = backgroundcolor,
        plot_bgcolor  = backgroundcolor,
        
        title = {
            'y':       title_position,
            'x':       0.5,
            'xanchor': 'center',
            'yanchor': 'top'
                },

        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

                  )

    fig.update_xaxes(
        tickmode = 'array',
        tickfont = dict(size=tickfont_size)
                    )
    
    #Set y axis format
    fig.update_yaxes(tickfont = dict(size=tickfont_size), secondary_y = False,)  #left axis
    fig.update_yaxes(tickfont = dict(size=tickfont_size), secondary_y = True, )  #right
    
    #Export figure as PNG file
    fig.write_image(os.path.join(report_folder,'days_on_market.png'), engine = 'kaleido', scale = scale)

def CreateMedianSalePriceGraph():
    #Create figure with secondary y-axis
    fig = make_subplots(specs=[[{"secondary_y": True}]])
    
    #Add Price Per Sqft for subject area
    fig.add_trace(
    go.Scatter(x        = df_subject['Month of Period End'],
           y            = df_subject['Median Price Per Sqft'],
           name         = subject_name,
           mode         = 'lines',
            line = dict(color = bowery_dark_blue, ) 
          
           ),
           secondary_y=False,
                )
    
        
    #Add Price Per Sqft for comparison area
    fig.add_trace(
    go.Scatter(x        = df_comparison['Month of Period End'],
           y            = df_comparison['Median Price Per Sqft'],
           name         = comparison_name,
           mode         = 'lines',
           line = dict(color = bowery_light_blue,) 
          
           ),
           secondary_y=False,
                )
    
    #Add bars with YoY Median Price/SF % Growth
    fig.add_trace(
        go.Bar(
            x            = df_subject['Month of Period End'],
            y            = df_subject['YoY Median Sale Price/SF Growth'],
            name         = 'YoY Median Price Per Sqft Growth',
            marker_color = bowery_grey,
            base         = dict(layer = 'Below')        
               ),
        secondary_y = True
            )

    #Set formatting 
    fig.update_layout(
        title_text    = "Median Sale Price/SF",    
        font_family   = font_family,
        font_color    = font_color,
        font_size     = font_size,
        height        = graph_height,
        width         = graph_width,
        margin        = dict(l = left_margin, r = right_margin, t = top_margin, b = bottom_margin),
        paper_bgcolor = backgroundcolor,
        plot_bgcolor  = backgroundcolor,
        
        title = {
            'y':       title_position,
            'x':       0.5,
            'xanchor': 'center',
            'yanchor': 'top'
                },

        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

                  )

    fig.update_xaxes(
        tickmode = 'array',
        tickfont = dict(size=tickfont_size)
                    )
    
    #Set y axis format
    fig.update_yaxes(tickfont = dict(size=tickfont_size), secondary_y = False,tickprefix = '$')  #left axis
    fig.update_yaxes(tickfont = dict(size=tickfont_size), secondary_y = True, ticksuffix = '%')  #right
    
    #Export figure as PNG file
    fig.write_image(os.path.join(report_folder,'sales_price.png'), engine = 'kaleido', scale = scale)

#Document related functions
def SetPageMargins(document, margin_size):
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin   = Inches(margin_size)
        section.right_margin  = Inches(margin_size)

def SetDocumentStyle(document):
    style     = document.styles['Normal']
    font      = style.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(9)

def AddTitle(document):
    title_text = subject_name + ' ' + subject_property_type + ' Market Analysis'
    title                               = document.add_heading(title_text,level=1)
    title.style                         = document.styles['Heading 2']
    title.paragraph_format.space_after  = Pt(6)
    title.paragraph_format.space_before = Pt(12)
    title_style                         = title.style
    title_style.font.name               = "Avenir Next LT Pro Light"
    title_style.font.size               = Pt(14)
    title_style.font.bold               = False
    title_style.font.color.rgb          = RGBColor.from_string('3F65AB')
    title_style.element.xml
    rFonts                              = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

    above_map_paragraph = document.add_paragraph("""This report was created using data from Redfin, a national real estate brokerage. Data represents """ + "{property_type}".format(property_type = """Condos""" if subject_property_type == 'Condo' else 'Single Family Homes') + """ in """ + subject_name + """ with monthly data through """ +  subject_latest_period.strftime('%m/%d/%Y') + """.""")
    above_map_style                                   = above_map_paragraph.style
    above_map_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.JUSTIFY
    above_map_style.font.size                         = Pt(9)
    above_map_paragraph.paragraph_format.space_after  = Pt(primary_space_after_paragraph)

def AddMap(document):
    #Add image of map if we already have one
    map_path = os.path.join(output_location, 'map.png')
    if os.path.exists(map_path):
        print('Adding map png to document')
        map = document.add_picture(map_path, width=Inches(6.5))
    Citation(document=document,text= 'Google Maps')

def AddHeading(document, title, heading_level): 
    #Function we use to insert the headers other than the title header
    heading                               = document.add_heading(title, level = heading_level)
    heading.style                         = document.styles['Heading 3']
    heading_style                         = heading.style
    heading_style.font.name               = "Avenir Next LT Pro"
    heading_style.font.size               = Pt(11)
    heading_style.font.bold               = False
    heading.paragraph_format.space_after  = Pt(6)
    heading.paragraph_format.space_before = Pt(12)

    #Color
    heading_style.font.color.rgb          = RGBColor.from_string('3F65AB')            
    heading_style.element.xml
    rFonts                                = heading_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def AddOverviewTable(document, number_cols, row_data): #Function we use to insert our overview table into the report document

    #Make sure each row has the same number of items 
    for row in row_data:
        for row2 in row_data:
            assert len(row) == len(row2)


    #create table object
    tab = document.add_table(rows=len(row_data), cols=len(row_data[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True
    #decide if we use standard style or custom

    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,row_data)): 
        assert (len(row_data_list) == number_cols) and (isinstance(row_data_list, list)) #make sure there is an item for each column in this row
        
        #Set height for row
        row.height = Inches(0.17)
        
        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):

            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)

            elif current_column == 1:
                cell.width = Inches(1.19)

            elif current_column == 2:
                cell.width = Inches(0.8)

            elif current_column == 3:
                cell.width = Inches(0.8)

            elif current_column == 4:
                cell.width = Inches(1.18)

            elif current_column == 5:
                cell.width = Inches(0.8)

            elif current_column == 6:
                cell.width = Inches(0.8)

            #add border to top row
            if current_row == 1:
                    tcPr      = cell._element.tcPr
                    tcBorders = OxmlElement("w:tcBorders")
                    top       = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')

                    tcBorders.append(top)
                    tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                #Paragaph spacing before and after
                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                
                for run in paragraph.runs:
                    font          = run.font
                    font.size     = Pt(7)
                    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    #make first row bold
                    if current_row == 0: 
                        font      = run.font
                        font.size = Pt(8)
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'

def AddDocumentParagraph(document, language_variable):
    assert type(language_variable) == list

    for paragraph in language_variable:
        
        #Skip blank paragraphs
        if paragraph == '':
            continue
        
        par                                               = document.add_paragraph(str(paragraph))
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
        style                                             = document.styles['Normal']
        font                                              = style.font
        font.name                                         = primary_font
        par.style                                         = document.styles['Normal']

def AddDocumentPicture(document, image_path, citation):
    if os.path.exists(image_path):
        fig                                         = document.add_picture(os.path.join(image_path),width=Inches(6.5))
        last_paragraph                              = document.paragraphs[-1] 
        last_paragraph.paragraph_format.space_after = Pt(0)
        last_paragraph.alignment                    = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document, citation)

def Citation(document, text):
    citation_paragraph                               = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(6)
    citation_paragraph.paragraph_format.space_before = Pt(0)
    run                                              = citation_paragraph.add_run('Source: ' + text)
    font                                             = run.font
    font.name                                        = primary_font
    font.size                                        = Pt(8)
    font.italic                                      = True
    font.color.rgb                                   = RGBColor.from_string('929292')
    citation_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.RIGHT

def AddTableTitle(document, title):
    table_title_paragraph                               = document.add_paragraph(title)
    table_title_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.CENTER
    table_title_paragraph.paragraph_format.space_after  = Pt(6)
    table_title_paragraph.paragraph_format.space_before = Pt(12)

    for run in table_title_paragraph.runs:
                    font      = run.font
                    font.name = 'Avenir Next LT Pro Medium'

def OverviewSection(document):
    print('Writing Overview Section')
    AddHeading(document = document, title = 'At a Glance', heading_level = 2)

    #Add Overview langauge
    AddDocumentParagraph(document = document, language_variable = overview_language)

    AddTableTitle(document=document,title = 'Market Fundamentals')
    AddOverviewTable(document=document,number_cols=7,row_data= GetDataForOverviewTable())

def SupplyandDemandSection(document):
    print('Writing Supply and Demand Section')
    AddHeading(document = document, title = 'Supply and Demand', heading_level = 2)

    #Add Overview langauge
    AddDocumentParagraph(document = document, language_variable = supply_and_demand_language)

    AddDocumentPicture(document = document, image_path=(os.path.join(report_folder,'sales_volume.png')),citation='RedFin.com')

    AddDocumentPicture(document = document, image_path=(os.path.join(report_folder,'days_on_market.png')),citation='RedFin.com')

def ValuesSection(document):
    print('Writing Values Section')
    AddHeading(document = document, title = 'Values', heading_level = 2)

    #Add Overview langauge
    AddDocumentParagraph(document = document, language_variable = values_language)

    AddDocumentPicture(document = document, image_path=(os.path.join(report_folder,'sales_price.png')),citation='RedFin.com')

def ConclusionSection(document):
    print('Writing Conclusion Section')
    AddHeading(document = document, title = 'Conclusion', heading_level = 2)

    #Add Overview langauge
    AddDocumentParagraph(document = document, language_variable = conclusion_language)

def WriteReport():
    print('Writing Report')
    #Create Document
    document = Document()
    SetPageMargins(       document = document, margin_size = 1)
    SetDocumentStyle(     document = document)
    AddTitle(             document = document)
    AddMap(               document=    document)
    OverviewSection(      document = document)
    SupplyandDemandSection(document = document)
    ValuesSection(document = document)
    ConclusionSection(document = document)
    
    
    #Save report
    document.save(report_path)  





#Set formatting paramaters for reports
primary_font                  = 'Avenir Next LT Pro Light' 
primary_space_after_paragraph = 8

#Set graph size and format variables
tickangle                     = 0
bowery_grey                   = "#D7DEEA"
bowery_dark_grey              = "#A6B0BF"
bowery_dark_blue              = "#4160D3"
bowery_light_blue             = "#B3C3FF"
bowery_black                  = "#404858"


#Set Graph Size
marginInches      = 1/18
ppi               = 96.85 
width_inches      = 6.5
height_inches     = 3.3

graph_width       = (width_inches  - marginInches)   * ppi 
graph_height      = (height_inches - marginInches)   * ppi

#Set scale for resolution 1 = no change, > 1 increases resolution. Very important for run time of main script. 
scale             = 6

#Set tick font size (also controls legend font size)
tickfont_size     = 8 

#Set Margin parameters/legend location
left_margin       = 0
right_margin      = 0
top_margin        = 75
bottom_margin     = 10
legend_position   = 1.05
title_position    = .95
font_size         = 10.5
backgroundcolor   = 'white'
bowery_grey       = "#D7DEEA"
bowery_dark_grey  = "#A6B0BF"
bowery_dark_blue  = "#4160D3"
bowery_light_blue = "#B3C3FF"
bowery_black      = "#404858"
font_family       = "Avenir Next LT Pro"
font_color        = '#262626'

#Heart of script starts here
DetermineSubjectAndComp()

#Use our 2 dataframes to create key variables 
subject_name             = (df_subject['Region'].iloc[-1])
subject_geo_level        = df_subject['Region Type'].iloc[-1]
subject_property_type    = df_subject['Type'].iloc[-1]
subject_latest_period    = df_subject['Month of Period End'].iloc[-1]

comparison_name          = df_comparison['Region'].iloc[-1]
comparison_geo_level     = df_comparison['Region Type'].iloc[-1]
comparison_property_type = df_comparison['Type'].iloc[-1]
comparison_latest_period = df_comparison['Month of Period End'].iloc[-1]

#Declare output folder
# output_location                =  os.path.join(dropbox_root,'Research', 'Market Analysis',subject_property_type) #Production Output
output_location                =  os.path.join(project_location,'Output\Residential Reports',subject_property_type)      #Testing Output


#Make sure the subject and comparison area have the same last period
assert subject_latest_period == comparison_latest_period

CreateDirectory()

#After making our directory, we can manipulate our subject name
if subject_geo_level == 'Metro':
    subject_name           =subject_name.replace('metro area', 'Metro Area')

CreateLanguage()
CreateGraphs()
WriteReport()
CleanUpPNGs()

