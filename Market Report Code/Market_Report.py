#By Mike Leahy April 29, 2021:
#Summary:
    #Imports 4 clean csv files with summary statistics from CoStar.com on commerical real estate for the 4 main sectors
    #It also imports a coresponding "slices" csv/excel file for each sector which breaks down summary statistics by property sub-type
    #A GUI lauches which prompts the user to select which sector they want to create reports for and which markets/submarkets
    #Loops through these 4 files, loops through each of the markets and submarkets (geographic areas) and creates a directory and word document
    #The word document is a report that reports tables and graphs generated from the data files

import os
import pandas as pd
import numpy as np
from tkinter import *
from tkinter import ttk

from datetime import datetime
import requests
from requests.exceptions import HTTPError 
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry

import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

#These functions are part of the script, writen by us and stored in seperate files
from Graph_Functions import *  
from Language_Functions import *  
from Table_Functions import * 


#Define file pre-paths
dropbox_root                   = os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
project_location               = os.path.join(dropbox_root,'Research','Projects','Research Report Automation Project')                        #Main Folder that stores all output, code, and documentation
# output_location                = os.path.join(project_location,'Output','Market')                                                             #The folder where we store our current reports, testing folder
output_location                = os.path.join(dropbox_root,'Research','Market Analysis','Market')                                             #The folder where we store our current reports, production
map_location                   = os.path.join(project_location,'Data','Market Reports Data','CoStar Maps')                                    #Folders with maps png files  
general_data_location          = os.path.join(project_location,'Data','General Data')                                                         #Folder with data for all report types
costar_data_location           = os.path.join(project_location,'Data','Market Reports Data','CoStar Data')                                    #Folder with clean CoStar CSV files
costar_writeup_location        = os.path.join(project_location,'Data','Market Reports Data','CoStar Writeups')                                #Folder with html files downloaded from CoStar.com
costar_construction_location   = os.path.join(project_location,'Data','Market Reports Data','CoStar Data', 'Construction Data')                                
costar_transaction_location    = os.path.join(project_location,'Data','Market Reports Data','CoStar Data', 'Transaction Data' )                                

#Define functions

def user_selects_reports_or_not():
    #Launches GUI for user to select if they want to write reports, or update the database/CoStar Markets CSV file

    ws = Tk()
    ws.title('Research Automation Project - Market Reports')
    ws.geometry('400x300')
    ws.config(bg='#404858')

    def select_sector(choice):
        global write_reports_yes_or_no
        write_reports_yes_or_no = variable.get()

    #setting variable for Integers
    variable = StringVar()
    variable.set('Write Reports?')

    #creating widget
    dropdown = OptionMenu(ws, variable, *['yes', 'no'], command=select_sector)

    #positioning widget
    dropdown.pack(expand=True)

    #infinite loop 
    ws.mainloop()

def user_selects_sector():
    #Launches GUI for user to select sector
    global df_list, df_slices_list, sector_name_list, selected_sector
    global df_multifamily, df_office, df_retail, df_industrial
    global df_multifamily_slices, df_office_slices, df_retail_slices, df_industrial_slices


    #If we have any custom data, read it in as a dataframe so we can append it to our primary data
    custom_data_file_location      = os.path.join(costar_data_location, 'Clean Data', 'Clean Custom CoStar Data.xlsx')
    if os.path.exists(custom_data_file_location):
        df_custom                  = pd.read_excel(custom_data_file_location)

    #Don't make the user select a sector if they are not trying to write reports
    if write_reports_yes_or_no == 'no':
        selected_sector = 'All'
        
        #Import cleaned data from 1.) Clean Costar Data.py
        df_multifamily                 = pd.read_csv(os.path.join(costar_data_location,'Clean Data','mf_clean.csv')) 
        df_office                      = pd.read_csv(os.path.join(costar_data_location,'Clean Data','office_clean.csv'))
        df_retail                      = pd.read_csv(os.path.join(costar_data_location,'Clean Data','retail_clean.csv'))
        df_industrial                  = pd.read_csv(os.path.join(costar_data_location,'Clean Data','industrial_clean.csv')) 
        
        if os.path.exists(os.path.join(costar_data_location,'Clean Data','mf_slices_clean.csv')):
            df_multifamily_slices          = pd.read_csv(os.path.join(costar_data_location,'Clean Data','mf_slices_clean.csv')) 
        
        if os.path.exists(os.path.join(costar_data_location,'Clean Data','office_slices_clean.csv')):
            df_office_slices               = pd.read_csv(os.path.join(costar_data_location,'Clean Data','office_slices_clean.csv'))
        
        if os.path.exists(os.path.join(costar_data_location,'Clean Data','retail_slices_clean.csv')):
            df_retail_slices               = pd.read_csv(os.path.join(costar_data_location,'Clean Data','retail_slices_clean.csv'))
    
        if os.path.exists(os.path.join(costar_data_location,'Clean Data','industrial_slices_clean.csv')):
            df_industrial_slices           = pd.read_csv(os.path.join(costar_data_location,'Clean Data','industrial_slices_clean.csv')) 

        df_list                        = [df_multifamily, df_office, df_retail, df_industrial]
        sector_name_list               = ['Multifamily','Office','Retail','Industrial']
        try:
            df_slices_list                 = [df_multifamily_slices, df_office_slices, df_retail_slices, df_industrial_slices]
        except:
            df_slices_list                 = [0, 0, 0, 0]

        return('')

    #GUI that lets user specify which sectors they want to run
    ws = Tk()
    ws.title('Research Automation Project - Market Reports')
    ws.geometry('400x300')
    ws.config(bg = '#404858')

    def select_sector(choice):
        global selected_sector
        selected_sector = variable.get()
        
    #setting variable for Integers
    variable = StringVar()
    variable.set('Select a sector')

    #creating widget
    dropdown = OptionMenu(ws, variable, *['Multifamily', 'Office', 'Retail', 'Industrial', 'All'], command = select_sector)

    #positioning widget
    dropdown.pack(expand=True)

    #infinite loop 
    ws.mainloop()

    #Everything below this comment in this function is after the user has selected the sector and tkinter loop has ended
    if selected_sector == 'All':
        #Import cleaned data from 1.) Clean Costar Data.py
        df_multifamily                 = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'mf_clean.csv')) 
        df_office                      = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'office_clean.csv'))
        df_retail                      = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'retail_clean.csv'))
        df_industrial                  = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'industrial_clean.csv')) 

        if os.path.exists(os.path.join(costar_data_location, 'Clean Data', 'mf_slices_clean.csv')):
            df_multifamily_slices          = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'mf_slices_clean.csv')) 
        
        if os.path.exists(os.path.join(costar_data_location, 'Clean Data', 'office_slices_clean.csv')):
            df_office_slices               = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'office_slices_clean.csv'))
        
        if os.path.exists(os.path.join(costar_data_location, 'Clean Data', 'retail_slices_clean.csv')):
            df_retail_slices               = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'retail_slices_clean.csv'))

        if os.path.exists(os.path.join(costar_data_location, 'Clean Data', 'industrial_slices_clean.csv')):    
            df_industrial_slices           = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'industrial_slices_clean.csv')) 

        #Import supplemental data as pandas data frames. This is data we store for ourselves on the differnet markets and submarkets (we will merge into our main data dfs)
        df_multifamily_supplemental   = pd.read_csv(os.path.join(costar_data_location,'Supplemental Data', 'mf_supplemental.csv'),         dtype={'Town': object,}) 
        df_office_supplemental        = pd.read_csv(os.path.join(costar_data_location,'Supplemental Data', 'office_supplemental.csv') ,    dtype={'Town': object,})      
        df_retail_supplemental        = pd.read_csv(os.path.join(costar_data_location,'Supplemental Data', 'retail_supplemental.csv') ,    dtype={'Town': object,})
        df_industrial_supplemental    = pd.read_csv(os.path.join(costar_data_location,'Supplemental Data', 'industrial_supplemental.csv'), dtype={'Town': object,})  	


        #Merge in our supplemental data into our main data frames
        df_multifamily                = pd.merge(df_multifamily, df_multifamily_supplemental,      on=['Geography Name', 'Geography Type'], how = 'left')
        df_office                     = pd.merge(df_office,      df_office_supplemental,           on=['Geography Name', 'Geography Type'], how = 'left')
        df_retail                     = pd.merge(df_retail,      df_retail_supplemental,           on=['Geography Name', 'Geography Type'], how = 'left')
        df_industrial                 = pd.merge(df_industrial,  df_industrial_supplemental,       on=['Geography Name', 'Geography Type'], how = 'left')

        #Do this because we don't have the towns for most of the market so this prevents errors
        df_multifamily['Town']        = df_multifamily['Town'].fillna('')
        df_office['Town']             = df_office['Town'].fillna('')
        df_retail['Town']             = df_retail['Town'].fillna('')
        df_industrial['Town']         = df_industrial['Town'].fillna('')


        df_list                       = [df_multifamily, df_office, df_retail, df_industrial]
        sector_name_list              = ['Multifamily', 'Office', 'Retail', 'Industrial']
        
        #Need this to deal with situations when do not have the slices data
        try:
            df_slices_list                = [df_multifamily_slices, df_office_slices, df_retail_slices, df_industrial_slices]
        except:
            df_slices_list                = [0, 0, 0, 0]

    elif selected_sector == 'Office':

        #Import cleaned data from 1.) Clean Costar Data.py
        df_office                      = pd.read_csv(os.path.join(costar_data_location,'Clean Data', 'office_clean.csv'))
        
        if os.path.exists(os.path.join(costar_data_location,'Clean Data', 'office_slices_clean.csv')):
            df_office_slices               = pd.read_csv(os.path.join(costar_data_location,'Clean Data', 'office_slices_clean.csv'))

        #Import supplemental data as pandas data frames. This is data we store for ourselves on the differnet markets and submarkets (we will merge into our main data dfs)
        df_office_supplemental        = pd.read_csv(os.path.join(costar_data_location, 'Supplemental Data', 'office_supplemental.csv'), dtype = {'Town': object,})      

        #Merge in our supplemental data into our main data frames
        df_office                     = pd.merge(df_office,      df_office_supplemental,           on=['Geography Name','Geography Type'], how = 'left')

        #Do this because we don't have the towns for most of the market so this prevents errors
        df_office['Town']             = df_office['Town'].fillna('')

        try:
            df_office  = df_office.append(df_custom) #Add the custom data to the main data file
        except:
            pass
        df_list                        = [df_office]
        sector_name_list               = ['Office']
        
        try:
            df_slices_list                 = [df_office_slices]
        except:
            df_slices_list                 = [0]

    elif selected_sector == 'Retail':
        #Import cleaned data from 1.) Clean Costar Data.py
        df_retail                      = pd.read_csv(os.path.join(costar_data_location,'Clean Data','retail_clean.csv'))
        
        if os.path.exists(os.path.join(costar_data_location,'Clean Data','retail_slices_clean.csv')):
            df_retail_slices               = pd.read_csv(os.path.join(costar_data_location,'Clean Data','retail_slices_clean.csv'))

        #Import supplemental data as pandas data frames. This is data we store for ourselves on the differnet markets and submarkets (we will merge into our main data dfs)
        df_retail_supplemental        = pd.read_csv(os.path.join(costar_data_location,'Supplemental Data','retail_supplemental.csv') ,dtype={'Town': object,})

        #Merge in our supplemental data into our main data frames
        df_retail                     = pd.merge(df_retail,      df_retail_supplemental, on = ['Geography Name', 'Geography Type'], how = 'left')

        #Do this because we don't have the towns for most of the market so this prevents errors
        df_retail['Town']             = df_retail['Town'].fillna('')

        try:
            df_retail                  = df_retail.append(df_custom) #Add the custom data to the main data file
        except:
            pass
        
        df_list                        = [df_retail]
        sector_name_list               = ['Retail']
        try:
            df_slices_list                 = [df_retail_slices]
        except:
            df_slices_list                 = [0]

    elif selected_sector == 'Multifamily':

        #Import cleaned data from 1.) Clean Costar Data.py
        df_multifamily                 = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'mf_clean.csv')) 
        
        if os.path.exists(os.path.join(costar_data_location, 'Clean Data', 'mf_slices_clean.csv')):
            df_multifamily_slices          = pd.read_csv(os.path.join(costar_data_location, 'Clean Data', 'mf_slices_clean.csv')) 

        #Import supplemental data as pandas data frames. This is data we store for ourselves on the differnet markets and submarkets (we will merge into our main data dfs)
        df_multifamily_supplemental   = pd.read_csv(os.path.join(costar_data_location, 'Supplemental Data', 'mf_supplemental.csv'), dtype={'Town': object,}) 

        #Merge in our supplemental data into our main data frames
        df_multifamily                = pd.merge(df_multifamily, df_multifamily_supplemental,      on=['Geography Name', 'Geography Type'], how = 'left')

        #Do this because we don't have the towns for most of the market so this prevents errors
        df_multifamily['Town']        = df_multifamily['Town'].fillna('')

        try:
            df_multifamily            = df_multifamily.append(df_custom) #Add the custom data to the main data file
        except:
            pass
        
        df_list                       = [df_multifamily]
        sector_name_list              = ['Multifamily']
        try:
            df_slices_list                = [df_multifamily_slices]
        except:
            df_slices_list                = [0]

    elif selected_sector == 'Industrial':
        #Import cleaned data from 1.) Clean Costar Data.py
        df_industrial                  = pd.read_csv(os.path.join(costar_data_location,'Clean Data','industrial_clean.csv')) 
        
        if os.path.exists(os.path.join(costar_data_location,'Clean Data','industrial_slices_clean.csv')):
            df_industrial_slices           = pd.read_csv(os.path.join(costar_data_location,'Clean Data','industrial_slices_clean.csv')) 

        #Import supplemental data as pandas data frames. This is data we store for ourselves on the differnet markets and submarkets (we will merge into our main data dfs)
        df_industrial_supplemental    = pd.read_csv(os.path.join(costar_data_location,'Supplemental Data','industrial_supplemental.csv'),dtype={'Town': object,})  	

        #Merge in our supplemental data into our main data frames
        df_industrial                 = pd.merge(df_industrial,  df_industrial_supplemental,       on=['Geography Name','Geography Type'], how = 'left')

        #Do this because we don't have the towns for most of the market so this prevents errors
        df_industrial['Town']         = df_industrial['Town'].fillna('')
       
        try:
            df_industrial             = df_industrial.append(df_custom) #Add the custom data to the main data file
        except:
            pass
        
        df_list                       = [df_industrial]
        sector_name_list              = ['Industrial']
        try:
            df_slices_list                = [df_industrial_slices]
        except:
            df_slices_list                = [0]
            
#Define functions used to handle the clean CoStar data and help write our repots
def CreateMarketDictionary(df): 
    #Creates a dictionary where each key is a market and the items are lists of its submarkets
     df_markets             = df.loc[df['Geography Type'] == 'Metro'] 
     df_submarkets          = df.loc[df['Geography Type'] == 'Submarket']
     unique_markets_list    = df_markets['Geography Name'].unique()
     unique_submarkets_list = df_submarkets['Geography Name'].unique()
    
     #Now create dictionary to track which submarkets belong to each market
     market_dictionary = {}
     for market in unique_markets_list:
         submarkets = [submarket for submarket in unique_submarkets_list if market in submarket ] #list of sumarkets within current market
         market_dictionary.update({market:submarkets}) 

     return(market_dictionary)

def CleanMarketName(market_name):
    clean_market_name = market_name.replace("""/""",' ')
    clean_market_name = clean_market_name.replace(""":""",'')
    clean_market_name = clean_market_name.replace("""'""",'')
    
    if clean_market_name[-1] == '.':
        clean_market_name = clean_market_name.replace(""".""",'')
    
    clean_market_name = clean_market_name.strip()
    return(clean_market_name)

def UniqueZipCodes(zip_code_list):
    #Converts a list of zip codes with possible duplicates into a unqiue list with no duplicates
    
    #convert the set to the list
    unique_list = list(set(zip_code_list))
    
    #convert from string to int
    for i in range(0, len(unique_list)):
        unique_list[i] = int(unique_list[i])

    return(unique_list)

def AppendAllExcelFilesInDirectory(directory):
    #Takes a directory as input and appends all excel files in it togetehr into dataframe
    #Returns the dataframe or None if no files exist

    i = 0
    for file in os.listdir(directory):

        #Skip the non excel files
        if file.endswith('.xlsx') == False and file.endswith('.csv') == False :
            continue
        
        if file.endswith('.csv'):
            reader_object = pd.read_csv
        elif file.endswith('.xlsx'):
            reader_object = pd.read_excel

        #The first file is our master df we append each new file to
        if i == 0:
            df                 = reader_object(os.path.join(directory, file)) 
        elif i > 0:
            df_alt             = reader_object(os.path.join(directory, file)) 
            df                 = df.append(df_alt)
        i += 1
    if 'df' in locals():
        return(df)
    else:
        return(None)

def CleanTransactionData(df_transactions):
    #Clean Transaction dataframe
    if isinstance(df_transactions,pd.DataFrame) == False:
        return(None)
    
    df_transactions['PropertyType']      = df_transactions['PropertyType'].str.replace('Multi-Family','Multifamily')
    df_transactions['Price/Unit']        = df_transactions['Last Sale Price']/df_transactions['Number Of Units']
    df_transactions['Year Built']        = df_transactions['Year Built'].astype(str)
    df_transactions['Year Built']        = df_transactions['Year Built'].str[0:4]
    df_transactions['Property Address']  =  df_transactions['Property Address'] + ', ' + df_transactions['City'] 
    
    #Convert last sale date to Q# YYYY Format
    df_transactions['Quarter']           = df_transactions['Last Sale Date'].dt.quarter
    df_transactions['Year']              = df_transactions['Last Sale Date'].dt.year
    df_transactions['Quarter']           = df_transactions['Quarter'].astype(str)
    df_transactions['Year']              = df_transactions['Year'].astype(str)
    df_transactions['Last Sale Date']    = 'Q' + df_transactions['Quarter']  + ' ' +  df_transactions['Year']  

    #Create transactions dataframe
    if market == primary_market: #market
        df_submarkets_transactions = df_transactions[   (df_transactions['Market Name'] == primary_market_title)     &
                                                        (df_transactions['PropertyType'] == sector)].copy() 
        #Keep the 20 largest transactions for markets
        df_submarkets_transactions = df_submarkets_transactions.sort_values(by=['Last Sale Price'],ascending=False )

    elif market != primary_market: #submarket
        if sector == 'Multifamily':
            df_submarkets_transactions = df_transactions[(df_transactions['Submarket Cluster'] == market_title) &
                                                        (df_transactions['Market Name'] == primary_market_title)     &
                                                        (df_transactions['PropertyType'] == sector)].copy() 
        else:
            df_submarkets_transactions = df_transactions[(df_transactions['Submarket Name'] == market_title) &
                                                        (df_transactions['Market Name'] == primary_market_title)     &
                                                        (df_transactions['PropertyType'] == sector)].copy() 

    
    df_submarkets_transactions                   = df_submarkets_transactions.sort_values(by=['Last Sale Price'], ascending=False )
    df_submarkets_transactions                   = df_submarkets_transactions.iloc[0:5]
    return(df_submarkets_transactions)

def CleanConstructionData(df_construction):
    #Clean construction dataframe
    if isinstance(df_construction,pd.DataFrame) == False:
        return(None)
    
    df_construction['PropertyType']       = df_construction['PropertyType'].str.replace('Multi-Family','Multifamily')
    df_construction['PropertyType']       = df_construction['PropertyType'].str.replace('Retail (Strip Center)','Retail',regex=False)
    df_construction                       = df_construction.rename(columns={"Number Of Units": "Number of Units"})
    df_construction['Property Address']   = df_construction['Property Address'] + ', ' + df_construction['City'] 
    

    #Create construction dataframe
    if market == primary_market: #market
        df_submarkets_construction = df_construction[   (df_construction['Market Name'] == primary_market_title)     &
                                                        (df_construction['PropertyType'] == sector)].copy() 


    elif market != primary_market: #submarket
        if sector == 'Multifamily':
            df_submarkets_construction = df_construction[(df_construction['Submarket Cluster'] == market_title) &
                                                        (df_construction['Market Name'] == primary_market_title)     &
                                                        (df_construction['PropertyType'] == sector)].copy() 
        else:
            df_submarkets_construction = df_construction[(df_construction['Submarket Name'] == market_title) &
                                                        (df_construction['Market Name'] == primary_market_title)     &
                                                        (df_construction['PropertyType'] == sector)].copy() 
    #Create a dummy variable for if a buidling exists or not
    df_submarkets_construction['Under Construction']  = 0
    df_submarkets_construction.loc[df_submarkets_construction['Building Status']     == 'Under Construction', 'PropertyType'] = 1

    if sector == 'Multifamily':
        sort_var = 'Number of Units'
    else:
        sort_var = 'RBA'
    df_submarkets_construction                   = df_submarkets_construction.sort_values(by=['Under Construction',sort_var], ascending=False, )
    df_submarkets_construction                   = df_submarkets_construction.iloc[0:5]

    return(df_submarkets_construction)

def CreateEmptySalesforceLists():
    global dropbox_primary_markets, dropbox_markets, dropbox_sectors, dropbox_sectors_codes
    global dropbox_links, dropbox_research_names, dropbox_analysis_types, dropbox_states, dropbox_versions, dropbox_statuses, dropbox_document_names
    global dropbox_cbsa_codes
    dropbox_primary_markets        = []
    dropbox_markets                = []
    dropbox_sectors                = []
    dropbox_sectors_codes          = []
    dropbox_links                  = []
    dropbox_research_names         = []
    dropbox_analysis_types         = []
    dropbox_states                 = []
    dropbox_cbsa_codes             = []
    dropbox_versions               = []
    dropbox_statuses               = []
    dropbox_document_names         = []

def UpdateSalesforceMarketList(markets_list, submarkets_list, sector_list, sector_code_list, dropbox_links_list):
    #Add to lists that track our markets and submarkets for salesforce
    markets_list.append(state + '-' + primary_market_name_for_file)
    
    if market == primary_market:
        submarkets_list.append('')
        dropbox_analysis_types.append('Market')
    else:
        submarkets_list.append(state + '-' + market_title)
        dropbox_analysis_types.append('Submarket')

    sector_list.append(sector)
    if sector == 'Multifamily':
        sector_code_list.append('MF')
    else:
        sector_code_list.append(sector[0])

    #Use the output directory to back into the right dropbox link 
    dropbox_link = output_directory.replace(dropbox_root,r'https://www.dropbox.com/home')
    dropbox_link = dropbox_link.replace("\\",r'/')
    dropbox_links_list.append(dropbox_link)

    dropbox_states.append(state)
    dropbox_cbsa_codes.append(df_market_cut.iloc[-1]['CBSA Code'])

    latest_quarter = df_market_cut.iloc[-1]['Period']
    dropbox_versions.append(latest_quarter)

    if market == primary_market:
        dropbox_research_names.append(state + ' - ' + primary_market_name_for_file + ' - ' + sector )
    else:
        if market_title == primary_market_name_for_file : #if market and submarket have same name
            dropbox_research_names.append(state + ' - ' + market_title + ' SUB' + ' - ' + sector )
        else:
            dropbox_research_names.append(state + ' - ' + market_title + ' - ' + sector )

    #If the report is already marked final, record that in our list
    if  os.path.exists(report_path.replace('_draft','_FINAL',1)) or os.path.exists(report_path.replace('_draft','- FINAL',1)) or os.path.exists(report_path.replace('_draft', ' - FINAL', 1)):
        dropbox_document_names.append(report_file_title.replace('_draft', '_FINAL', 1))
        dropbox_statuses.append('Final')
    
    else:
        dropbox_document_names.append(report_file_title)
        dropbox_statuses.append('Draft')

def CreateOutputDirectory():
    #Creates a folder where we put the report document inside of and returns the file path
    sector_folder        = os.path.join(output_location,sector)
    state_folder         = os.path.join(output_location,sector,state)
    market_folder        = os.path.join(state_folder,primary_market_name_for_file)

    if market == primary_market:
        output_directory     = market_folder #Folder where we write report to
    else:
        output_directory     = os.path.join(market_folder,str(market_title)) 

    #Check if output,map, and summary folder already exists, and if it doesnt, make it
    for folder in [sector_folder, state_folder, market_folder, output_directory]:
        if os.path.exists(folder) == False:
            os.mkdir(folder) 

    return(output_directory)

def CreateMapDirectory():
    #Creates a folder where we store png files with map images that we can insert into our reports
    state_folder_map         = os.path.join(map_location,sector,state)
    market_folder_map        = os.path.join(state_folder_map,primary_market_name_for_file)

    if market == primary_market:
        map_directory        = market_folder_map                #Folder where we store map for market or submarket
    else:
        map_directory        = os.path.join(market_folder_map,str(market_title))
    
    #Check if these folders exist and if not, make them
    for folder in [state_folder_map, market_folder_map, map_directory]:
        if os.path.exists(folder) == False:
            os.mkdir(folder) 
    
    return(map_directory)

def CreateWriteupDirectory():
    #Creates a folder where we can save the CoStar HTML file from the website in order to scrape the writeups from CoStar analysts
    sector_folder                   = os.path.join(costar_writeup_location,sector)
    state_folder_writeup            = os.path.join(costar_writeup_location,sector,state)
    market_folder_writeup           = os.path.join(state_folder_writeup,primary_market_name_for_file)

    if market == primary_market:
        writeup_directory           = market_folder_writeup                
    else:
        writeup_directory           = os.path.join(market_folder_writeup,str(market_title))
    
    #Check if folder already exists, and if it doesnt, make it
    for folder in [sector_folder,state_folder_writeup,market_folder_writeup,writeup_directory]:
        if os.path.exists(folder) == False:
            os.mkdir(folder) #Create new folder for market or submarket
    
    return(writeup_directory)

def CreateConstructionDirectory():
    #Creates a folder where we can save the CoStar HTML file from the website in order to scrape the writeups from CoStar analysts
    sector_folder                   = os.path.join(costar_construction_location, sector)
    state_folder                    = os.path.join(costar_construction_location, sector, state)
    market_folder                   = os.path.join(state_folder, primary_market_name_for_file)

    construction_directory           = market_folder               
    
    #Check if folder already exists, and if it doesnt, make it
    for folder in [sector_folder,state_folder, market_folder, construction_directory]:
        if os.path.exists(folder) == False:
            os.mkdir(folder) #Create new folder for market or submarket
    
    return(construction_directory)

def CreateTransactionDirectory():
    #Creates a folder where we can save the CoStar HTML file from the website in order to scrape the writeups from CoStar analysts
    sector_folder                   = os.path.join(costar_transaction_location, sector)
    state_folder                    = os.path.join(costar_transaction_location, sector, state)
    market_folder                   = os.path.join(state_folder, primary_market_name_for_file)
    transaction_directory           = market_folder               
    
    #Check if folder already exists, and if it doesnt, make it
    for folder in [sector_folder,state_folder, market_folder, transaction_directory]:
        if os.path.exists(folder) == False:
            os.mkdir(folder) #Create new folder for market or submarket
    
    return(transaction_directory)

def CreateReportFilePath():
    global report_file_title

    if market == primary_market:
        market_file_name = primary_market_name_for_file
        macro_or_sub     = 'Market'
    else:
        market_file_name = market_title
        macro_or_sub     = 'Submarket'
    
    report_file_title =   latest_quarter  + ' ' +  state + ' - '   + market_file_name + ' - ' + sector + ' ' + macro_or_sub  + '_draft' + '.docx'

    #Make sure we don't hit the 255 max file path limit
    if len(os.path.join(output_directory,report_file_title)) <= 257:
        report_path = os.path.join(output_directory,report_file_title)
    else:
        report_path = os.path.join(output_directory,(latest_quarter + '-' + market_file_name + '_draft' + '.docx')  )

    assert os.path.exists(output_directory)
    return(report_path)

def SetPageMargins():
    #Page Margins
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

def SetStyle():
    style     = document.styles['Normal']
    font      = style.font
    font.name = 'Avenir Next LT Pro Light'
    font.size = Pt(9)

def MakeReportTitle():
    #Write title Heading
    if market == primary_market:
        title = document.add_heading(market_title + ': ' + sector + ' Market Analysis'   ,level=1)
    else:
        title = document.add_heading(market_title + ': ' + sector + ' Submarket Analysis' ,level=1)
    
    title.style                           = document.styles['Heading 2']
    title.paragraph_format.space_after    = Pt(6)
    title.paragraph_format.space_before   = Pt(12)
    title.paragraph_format.keep_with_next = True
    title_style                           = title.style
    title_style.font.name                 = "Avenir Next LT Pro Light"
    title_style.font.size                 = Pt(14)
    title_style.font.bold                 = False
    title_style.font.color.rgb            = RGBColor.from_string('3F65AB')
    title_style.element.xml
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

def MakeCoStarDisclaimer():
    #Write Costar disclaimer: "as of '  + latest_quarter + '"" 
    if market == primary_market:
        disclaimer = document.add_paragraph('The information contained in this report was provided using ' +
                                             latest_quarter                                                + 
                                            ' CoStar data for the '                                        + 
                                            market_title                                                   + 
                                            ' '                                                            + 
                                            sector                                                         + 
                                            """ Market ("Market")."""
                                            )
    #Submarket disclaimer
    else:
        disclaimer = document.add_paragraph('The information contained in this report was provided using ' +
                                            latest_quarter                                                 + 
                                            ' CoStar data for the '                                        + 
                                            market_title                                                   + 
                                            ' '                                                            + 
                                            sector                                                         + 
                                            """ Submarket ("Submarket") """                                +
                                            'located in the '                                              +
                                            primary_market_title                                           + 
                                            """ Market ("Market"). """
                                            )                
    
    disclaimer.style.font.name                = primary_font
    disclaimer.style.font.size                = Pt(9)
    disclaimer.paragraph_format.space_after   = Pt(6)
    disclaimer.paragraph_format.space_before  = Pt(0)
    disclaimer.alignment                      = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer.paragraph_format.keep_together = True

def CleanUpPNGs():
    #Report writing done, delete figures
    files = os.listdir(output_directory)
    for image in files:
        if image.endswith(".png"):
            while os.path.exists(os.path.join(output_directory, image)):
                try:
                    os.remove(os.path.join(output_directory, image))
                except Exception as e: 
                    print(e)

def AddMap():
    
    #Add image of map if there is one in the appropriate map folder
    if os.path.exists(os.path.join(map_directory,'map.png')):
        document.add_picture(os.path.join(map_directory,'map.png'), width=Inches(6.5))
    else:
        document.add_paragraph('')

    last_paragraph           = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def AddDocumentPicture(document, image_path):
    if os.path.exists(image_path):
        fig                                         = document.add_picture(os.path.join(image_path),width = Inches(6.5))
        last_paragraph                              = document.paragraphs[-1] 
        last_paragraph.paragraph_format.space_after = Pt(0)
        last_paragraph.alignment                    = WD_ALIGN_PARAGRAPH.CENTER

def AddDocumentParagraph(document, language_variable):
    assert type(language_variable) == list

    for paragraph in language_variable:
        
        #Skip blank paragraphs
        if paragraph == '':
            continue
        par                                               = document.add_paragraph(str(paragraph))
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        par.paragraph_format.space_before                 = Pt(primary_space_after_paragraph)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
        style                                             = document.styles['Normal']
        font                                              = style.font
        font.name                                         = 'Avenir Next LT Pro Light'
        par.style                                         = document.styles['Normal']

def AddTableTitle(document,title):
    table_title_paragraph                               = document.add_paragraph(title)
    table_title_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.CENTER
    table_title_paragraph.paragraph_format.space_after  = Pt(6)
    table_title_paragraph.paragraph_format.space_before = Pt(12)
    table_title_paragraph.keep_with_next                = True
    table_title_paragraph.keep_together                 = True
    
    #Set font
    for run in table_title_paragraph.runs:
                    font      = run.font
                    font.name = 'Avenir Next LT Pro Medium'

def OverviewSection():

    #Overview Heading
    AddHeading(document, 'Overview', 2)
    
    #Overview Paragraph
    AddDocumentParagraph(document = document, language_variable = overview_language)

    #Overview table title
    AddTableTitle(document = document, title = 'Sector Fundamentals')
    
    #Overview table
    if sector == 'Multifamily':
        if market == 'United States of America':
            AddOverviewTable(document, 8, 4, data_for_overview_table, 1.2)

        else:
            AddOverviewTable(document, 8, 7, data_for_overview_table, 1.2)
    else:
        if market == 'United States of America':
            AddOverviewTable(document, 9, 4, data_for_overview_table, 1.2)

        else:
            AddOverviewTable(document, 9, 7, data_for_overview_table, 1.2)


    #Preamble to historical performance table
    if market == primary_market:
        market_or_submarket = 'Market'
    else:
        market_or_submarket = 'Submarket'
    
    preamble_language = ('Supply and demand indicators, including inventory levels, absorption, vacancy, and rental rates for ' +
                         sector.lower()                                                                                         +
                         ' space in the '                                                                                       +
                         market_or_submarket                                                                                    +
                         ' are presented in the following table.'
                         )


    table_preamble                               = document.add_paragraph(preamble_language)
    table_preamble.alignment                     = WD_ALIGN_PARAGRAPH.JUSTIFY
    table_preamble.paragraph_format.space_after  = Pt(primary_space_after_paragraph)
    table_preamble.paragraph_format.space_before = Pt(6)

    #Market performance table for primary markets
    if market == primary_market:
        AddTableTitle(document = document,title =('Historical ' + sector  + ' Performance: ' +  market_title + ' Market' ))
        AddMarketPerformanceTable(document = document,market_data_frame = df_primary_market,col_width = 1.2,sector=sector)
        document.add_paragraph('')
        
        #For Manhatan submarkets, add a table for each quality slice
        if primary_market == 'Manhattan - NY':
            if len(df_slices) > 0:
                for slice in df_slices['Slice'].unique():
                    df_slices_temp = df_slices.loc[df_slices['Slice'] == slice]

                    if market == primary_market:
                        AddTableTitle(document = document, title =('Historical ' + slice + ' ' + sector  + ' Performance: ' +  market_title + ' Market' ))
                    else:
                        AddTableTitle(document = document, title =('Historical ' + slice + ' ' + sector  + ' Performance: ' +  market_title + ' Submarket'))

                    AddMarketPerformanceTable(document = document, market_data_frame = df_slices_temp, col_width = 1.2, sector=sector)
                    document.add_paragraph('')

    #Submarket market performance table
    else:
        AddTableTitle(document = document, title =('Historical ' + sector  + ' Performance: ' +  market_title + ' Submarket'))
        AddMarketPerformanceTable(document = document, market_data_frame = df_market_cut ,col_width = 1.2,sector=sector)
        document.add_paragraph('')

def SupplyDemandSection():
    #Supply and Demand Section
    AddHeading(document, 'Supply & Demand', 2)
    
    AddDocumentParagraph(document = document, language_variable = demand_language)

    #Vacancy Table
    AddTableTitle(document = document,title ='Vacancy Rates')

    AddTable(document, data_for_vacancy_table, 1.2)

    #Absorption rate Graph
    AddDocumentPicture(document=document, image_path=os.path.join(output_directory,'absorption_rate.png'))
    
def RentSecton():
    AddHeading(document, 'Rents', 3)   

    if len(rent_language) == 2:
        AddDocumentParagraph(document = document, language_variable = [rent_language[0]])
    else:
        AddDocumentParagraph(document = document, language_variable = rent_language)

    #Rent Table
    AddTableTitle(document = document,title = 'Market Rents')
    AddTable(document,data_for_rent_table, col_width = 1.2)
    
    if len(rent_language) == 2:
        document.add_paragraph('')
        AddDocumentParagraph(document = document, language_variable = [rent_language[1]])

    #Insert rent growth graph
    AddDocumentPicture(document=document, image_path = os.path.join(output_directory, 'rent_growth.png'))
    
def ConstructionSection():
    #Construction Section
    AddHeading(document,'Construction & Future Supply',2)

    AddDocumentParagraph(document = document, language_variable = construction_languge)
    
    if 'df_submarkets_construction' in globals() and isinstance(df_submarkets_construction, pd.DataFrame):
        if len(df_submarkets_construction) > 0:
            AddTableTitle(document = document,title = 'Current and Recently Completed Construction Projects')
            AddConstructionTable(document = document, market_data_frame = df_submarkets_construction, col_width = 1.2, sector = sector)

    #Insert construction graph
    AddDocumentPicture(document=document, image_path = os.path.join(output_directory, 'construction_volume.png'))
    
def CapitalMarketsSection():
    #Captial Markets Section
    AddHeading(document,'Capital Markets',2)
    
    if len(sale_language) == 2:
        AddDocumentParagraph(document = document, language_variable = [sale_language[0]])
    else:
        AddDocumentParagraph(document = document, language_variable = sale_language)

    #Sales Volume Graphs
    AddDocumentPicture(document=document, image_path=os.path.join(output_directory,'sales_volume.png'))
    if len(sale_language) == 2:
        AddDocumentParagraph(document = document, language_variable = [sale_language[1]])
    
    if 'df_submarkets_transactions' in globals() and isinstance(df_submarkets_transactions, pd.DataFrame):
        if len(df_submarkets_transactions) > 0:
            AddTableTitle(document = document,title = 'Recent Transactions')
            AddTransactionTable(document = document, market_data_frame = df_submarkets_transactions, col_width = 1.2, sector=sector)

    #Asset Value Graph
    AddDocumentPicture(document = document, image_path=os.path.join(output_directory,'asset_values.png'))
    
def OutlookSection():
    #Outlook Section
    AddHeading(document, 'Outlook', 2)
    AddDocumentParagraph(document = document, language_variable = outlook_language)

def AppendixSection():
    #Adds a table with stats on all submarkets in the market
    if market == primary_market:
        
        if  len(df_submarkets) > 0:
            document.add_paragraph('')
            AddHeading(document, 'Appendix', 2)
            
            AddTableTitle(document = document, title = (market_title + ' ' + sector + ' Market Overview' ))
            AddSubmarketsPerformanceTable(document = document, submarkets_data_frame = df_submarkets, col_width = 1.2, sector = sector)
            document.add_paragraph('')

def GetLanguage(writeup_directory):
    global overview_language, demand_language, sale_language, rent_language, construction_languge, outlook_language
    
    #Overview Language
    try:
        overview_language    = CreateOverviewLanguage(submarket_data_frame = df_market_cut, market_data_frame = df_primary_market, national_data_frame= df_nation, slices_data_frame = df_slices, market_title = market_title, primary_market = primary_market_title, sector = sector, writeup_directory=writeup_directory)
    except Exception as e:
        print(e,'problem creating overview langauge')
        overview_language    = ['']
    
    #Demand Language
    try:
        demand_language      = CreateDemandLanguage(submarket_data_frame = df_market_cut, market_data_frame = df_primary_market, natioanl_data_frame= df_nation, market_title = market_title, primary_market = primary_market, sector = sector, writeup_directory=writeup_directory)
    except Exception as e:
        print(e,'problem creating demand langauge')
        demand_language      = ['']
        
    #Rent Language
    try:
        rent_language        = CreateRentLanguage(submarket_data_frame = df_market_cut, market_data_frame = df_primary_market, natioanl_data_frame= df_nation, market_title = market_title, primary_market = primary_market, sector = sector, writeup_directory=writeup_directory)
    except Exception as e:
        print(e,'problem creating rent langauge')
        rent_language        = ['']

    #Construction Language
    try:
        construction_languge = CreateConstructionLanguage(submarket_data_frame = df_market_cut, market_data_frame = df_primary_market, natioanl_data_frame= df_nation, market_title = market_title, primary_market = primary_market, sector = sector, writeup_directory=writeup_directory)
    except Exception as e:
        print(e,'problem creating construction langauge')
        construction_languge = ['']
    
    #Sale Language
    try:
        sale_language        = CreateSaleLanguage(submarket_data_frame = df_market_cut, market_data_frame = df_primary_market, natioanl_data_frame= df_nation, market_title = market_title, primary_market = primary_market, sector = sector, writeup_directory=writeup_directory)
    except Exception as e:
        print(e,'problem creating sale langauge')
        sale_language        = ['']

    #Outlook Language
    try:
        outlook_language     = CreateOutlookLanguage(submarket_data_frame = df_market_cut, market_data_frame = df_primary_market, natioanl_data_frame= df_nation, market_title = market_title, primary_market = primary_market, sector = sector, writeup_directory=writeup_directory)
    except Exception as e:
        print(e,'problem creating outlook langauge')
        outlook_language     = ['']

def GetOverviewTable():
    #Create Data for overview table
    global market,primary_market,market_title
    #There are 4 possible permuations for this table (market/multifamily, market/non-multifamily, submarket/multifamily, submakert/non-multifamily)
    if sector == 'Multifamily':
        data_for_overview_table = [ [], [], [], [], [], [], [], [] ]
    else:
        data_for_overview_table = [ [], [], [], [], [], [], [], [], [] ]


    #Write Top Row of Report
    if market == 'United States of America':
        data_for_overview_table[0] = ['', market, 'YoY', 'QoQ']


    elif market == primary_market: #market report
        data_for_overview_table[0] = ['', primary_market_title, 'YoY', 'QoQ', df_nation['Geography Name'].iloc[0], 'YoY', 'QoQ']
    else:
        data_for_overview_table[0] = ['', market_title, 'YoY', 'QoQ', primary_market_title, 'YoY', 'QoQ']

    #Rows for non-apt
    if sector != 'Multifamily':
        #Rent Growth Row
        data_for_overview_table[1] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Market Rent/SF',
                                                                'YoY Rent Growth',
                                                                'QoQ Rent Growth',
                                                                '$',
                                                                '%',
                                                                '%',
                                                                'Market Rent/SF'
                                                             )

        #Vacancy Row
        data_for_overview_table[2] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Vacancy Rate',
                                                                'YoY Vacancy Growth',
                                                                'QoQ Vacancy Growth',
                                                                '%',
                                                                'bps',
                                                                'bps',
                                                                'Vacancy Rate'
                                                            )

        #Availability Rate Row
        data_for_overview_table[3] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Availability Rate',
                                                                'YoY Availability Rate Growth',
                                                                'QoQ Availability Rate Growth',
                                                                '%',
                                                                'bps',
                                                                'bps',
                                                                'Availability Rate'
                                                            )
        
        #Absorption Row
        data_for_overview_table[4] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Net Absorption SF',
                                                                'YoY Net Absorption SF Growth',
                                                                'QoQ Net Absorption SF Growth',
                                                                '',
                                                                '%',
                                                                '%',
                                                                'Net Absorption SF'
                                                            )
        
        #Asset Value Row
        data_for_overview_table[5] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Asset Value/Sqft',
                                                                'YoY Asset Value/Sqft Growth',
                                                                'QoQ Asset Value/Sqft Growth',
                                                                '$',
                                                                '%',
                                                                '%',
                                                                'Asset Value/SF'
                                                             )
        
        #Market Cap Rate Row
        data_for_overview_table[6] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Market Cap Rate',
                                                                'YoY Market Cap Rate Growth',
                                                                'QoQ Market Cap Rate Growth',
                                                                '%',
                                                                'bps',
                                                                'bps',
                                                                'Market Cap Rate'
                                                             )
        
        #Transaction Count Row
        data_for_overview_table[7] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Sales Volume Transactions',
                                                                'YoY Transactions Growth',
                                                                'QoQ Transactions Growth',
                                                                '',
                                                                '%',
                                                                '%',
                                                                'Transaction Count'
                                                             )
        #Sales Volume Row
        data_for_overview_table[8] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Total Sales Volume',
                                                                'YoY Total Sales Volume Growth',
                                                                'QoQ Total Sales Volume Growth',
                                                                '$',
                                                                '%',
                                                                '%',
                                                                'Sales Volume'
                                                            )

    #Rows for apt
    if sector == 'Multifamily':
        #Rent row
        data_for_overview_table[1] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Market Effective Rent/Unit',
                                                                'YoY Market Effective Rent/Unit Growth',
                                                                'QoQ Market Effective Rent/Unit Growth',
                                                                '$',
                                                                '%',
                                                                '%',
                                                                'Market Rent/Unit'
                                                             )

        #Vacancy row
        data_for_overview_table[2] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Vacancy Rate',
                                                                'YoY Vacancy Growth',
                                                                'QoQ Vacancy Growth',
                                                                '%',
                                                                'bps',
                                                                'bps',
                                                                'Vacancy Rate'
                                                             )

        #Absorption row
        data_for_overview_table[3] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Absorption Units',
                                                                'YoY Absorption Units Growth',
                                                                'QoQ Absorption Units Growth',
                                                                '',
                                                                '%',
                                                                '%',
                                                                'Net Absorption Units'
                                                            )

        #Asset value row
        data_for_overview_table[4] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Asset Value/Unit',
                                                                'YoY Asset Value/Unit Growth',
                                                                'QoQ Asset Value/Unit Growth',
                                                                '$',
                                                                '%',
                                                                '%',
                                                                'Asset Value/Unit'
                                                            )

        #Market Cap rate row
        data_for_overview_table[5] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Market Cap Rate',
                                                                'YoY Market Cap Rate Growth',
                                                                'QoQ Market Cap Rate Growth',
                                                                '%',
                                                                'bps',
                                                                'bps',
                                                                'Market Cap Rate'
                                                            )

        #Transaction Count row
        data_for_overview_table[6] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Sales Volume Transactions',
                                                                'YoY Transactions Growth',
                                                                'QoQ Transactions Growth',
                                                                '',
                                                                '%',
                                                                '%',
                                                                'Transaction Count')

        #Sales volume row
        data_for_overview_table[7] =    CreateRowDataForTable(  df_market_cut,
                                                                df_primary_market,
                                                                df_nation,
                                                                'Total Sales Volume',
                                                                'YoY Total Sales Volume Growth',
                                                                'QoQ Total Sales Volume Growth',
                                                                '$',
                                                                '%',
                                                                '%',
                                                                'Sales Volume'
                                                             )

    return(data_for_overview_table)

def GetRentTable():
    #Create data for rent Table
    if sector == 'Multifamily':
        return(CreateRowDataForWideTable(data_frame = df_market_cut, data_frame2 = df_primary_market, data_frame3 = df_nation, data_frame4 = df_slices, var1 = 'Market Effective Rent/Unit', modifier = '$', sector = sector))

    else:
        return(CreateRowDataForWideTable(data_frame = df_market_cut, data_frame2 = df_primary_market, data_frame3 = df_nation, data_frame4 = df_slices, var1 = 'Market Rent/SF', modifier = '$', sector = sector))

def CreateMarketReport():
    global market_clean, market_title, output_directory, map_directory
    global df_market_cut, df_primary_market, df_nation, df_submarkets , df_slices
    global latest_quarter, document, data_for_overview_table, data_for_vacancy_table, data_for_rent_table, report_path
    global primary_market, market, primary_market_title
    global df_submarkets_transactions, df_submarkets_construction
    global df_transactions, df_construction
    
    #remove slashes from market names so we can save as folder name
    market_clean        = CleanMarketName(market)
    market_title        = market_clean.replace(primary_market + ' -','').strip()
    
    #Create output, map, and writeup folders for the market of submarket
    output_directory       = CreateOutputDirectory()
    map_directory          = CreateMapDirectory()
    writeup_directory      = CreateWriteupDirectory()
    construction_directory = CreateConstructionDirectory()
    transaction_directory  = CreateTransactionDirectory()

    #Create a dataframe that only has rows for the market or submarket itself
    df_market_cut       = df[df['Geography Name'] == market].copy()                  #df for the market or submarket only
    
    #Get the latest quarter
    latest_quarter      = df_market_cut.iloc[-1]['Period']

    #Get the document name and file path for the report
    report_path         = CreateReportFilePath()

    if write_reports_yes_or_no == 'yes':
        #Create seperate dataframes with only rows from the current (sub)market, the primary market, and the nation 
        
        #Specifiy the different comparsion levels for the culusters and research markets
        if  primary_market == 'Manhattan - NY':
            df_nation         = df[df['Geography Name'] == 'New York - NY'].copy() 
            df_primary_market = df[df['Geography Name'] == 'Manhattan - NY'].copy()          #df for the market only
        
        elif  (primary_market == 'Downtown - NY') or (primary_market == 'Midtown - NY')  or (primary_market == 'Midtown South - NY')  or (primary_market == 'Uptown - NY')  :
            df_nation         = df[df['Geography Name'] == 'Manhattan - NY'].copy() 
            df_primary_market = df[df['Geography Name'] == primary_market].copy()                       
        
        else:
            df_nation         = df[df['Geography Type'] == 'National'].copy()              #df for the USA
            df_primary_market = df[df['Geography Name'] == primary_market].copy()          #df for the market only

        try:
            df_slices             = df2[df2['Geography Name'] == market].copy()        #df for the primary market with the quality/subtype slices
        except:
            column_names = ['Slice', 'Period', 'Vacancy Rate','Market Effective Rent/Unit','Inventory Units','Inventory SF','Market Rent/SF']
            df_slices = pd.DataFrame(columns = column_names)

        #A dataframe that tracks all submarkets in a market at the latest quarter
        df_submarkets         = df.loc[(df['Geography Name'].isin(submarkets) == True) & (df['Period'] == latest_quarter)].copy()
        assert len(df_market_cut)     > 0
        assert len(df_primary_market) > 0
        assert len(df_nation)         == 40

        #Strip the primary market out of the submarket name
        if market != primary_market:
            market_title        = market.replace(primary_market + ' -','').strip()
        elif market == primary_market:
            if market == 'Washington - DC':
                market_title = 'Washington D.C.'
            else:
                market_title        = primary_market.replace(' - ' + state,'').strip()
        
        if primary_market == 'Washington - DC':
            primary_market_title    = 'Washington D.C.'
        else:
            primary_market_title    = primary_market.replace(' - ' + state,'').strip()
        



        #Import transactions data as dataframe that we use regardless of sector selected
        df_transactions            = AppendAllExcelFilesInDirectory(directory=transaction_directory)
        df_submarkets_transactions = CleanTransactionData(df_transactions)  
        
        #Import construction data as dataframe that we use regardless of sector selected
        df_construction             = AppendAllExcelFilesInDirectory(directory=construction_directory)
        df_submarkets_construction  = CleanConstructionData(df_construction)

        #This function calls all the graph functions defined in the Graph_Functions.py file
        CreateAllGraphs(submarket_data_frame = df_market_cut , market_data_frame = df_primary_market, natioanl_data_frame = df_nation , folder = output_directory, market_title = market_title, primary_market = primary_market_title, sector = sector)

        #Create Data for overview table
        #There are 4 possible permuations for this table (market/apt, market/nonapt, submarket/apt, submakert/nonapt)
        data_for_overview_table = GetOverviewTable()

        #Create data for vacancy Table
        data_for_vacancy_table = CreateRowDataForWideTable(data_frame = df_market_cut, data_frame2 = df_primary_market, data_frame3 = df_nation, data_frame4 = df_slices,var1 = 'Vacancy Rate', modifier = '%', sector=sector)
        
        #Create data for rent Table
        data_for_rent_table    = GetRentTable()

        #Get language for paragraphs 
        GetLanguage(writeup_directory = writeup_directory)

        
        #Skip the reports we have already done
        if os.path.exists(report_path.replace('_draft','_FINAL',1)) or os.path.exists(report_path.replace('_draft','- FINAL',1)) or os.path.exists(report_path.replace('_draft',' - FINAL',1)):
            pass
        else:
            #Start writing report
            document = Document()
            SetPageMargins()
            SetStyle()
            MakeReportTitle()
            MakeCoStarDisclaimer()
            AddMap()       
            OverviewSection()
            SupplyDemandSection()
            RentSecton()
            ConstructionSection()
            CapitalMarketsSection()
            OutlookSection()
            AppendixSection()
            
            #Temp fix to replace Manhattan - NY with Manhattan
            for paragraph in document.paragraphs:
                if 'Manhattan - NY' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Manhattan - NY','Manhattan')

            #Save report
            document.save(report_path)

        
        #Report writing done, delete figures
        CleanUpPNGs()

    #Add to lists that track our markets and submarkets for salesforce
    UpdateSalesforceMarketList(markets_list = dropbox_primary_markets, submarkets_list = dropbox_markets, sector_list = dropbox_sectors, sector_code_list = dropbox_sectors_codes, dropbox_links_list = dropbox_links)

def user_selects_market(market_list, market_or_submarket):
    if len(df_list) == 4 or market_list == []:
        return(market_list)
    
    market_list.insert(0, 'All')
    
    def select_market(event):
        global  selected_market
        selected_market = comboExample.get()
        
    app = Tk() 
    app.geometry('600x300')
    app.config(bg='#404858')
    app.title('Research Automation Project - Market Reports') 

    labelTop = Label(app, text = ("Choose your " + market_or_submarket))
    labelTop.grid(column=0, row=0)

    comboExample = ttk.Combobox(app, values = market_list, width=50)

    comboExample.grid(column=0, row=1)
    comboExample.current(0)
    comboExample.bind("<<ComboboxSelected>>", select_market)
    app.mainloop()

    if selected_market == 'All':
        market_list.remove('All')
        return(market_list)
    else:
        return([selected_market])

def CreateDirectoryCSV():
    global dropbox_markets, dropbox_research_names, dropbox_analysis_types, dropbox_states, dropbox_sectors, dropbox_sectors_codes, dropbox_links, dropbox_versions, dropbox_statuses, dropbox_document_names
    global service_api_csv_name, csv_name

    service_api_csv_name = ''

    if write_reports_yes_or_no == 'no':
        #Now create dataframe with list of markets and export to a CSV for Salesforce
        dropbox_df = pd.DataFrame({"Market":            dropbox_primary_markets,
                                "Submarket":            dropbox_markets,
                                'Market Research Name': dropbox_research_names,
                                'Analysis Type':        dropbox_analysis_types,
                                'State':                dropbox_states,
                                'CBSA Code':            dropbox_cbsa_codes,
                                "Property Type":        dropbox_sectors,
                                'Property Type Code':   dropbox_sectors_codes,
                                "Dropbox Links":        dropbox_links,
                                'Version':              dropbox_versions,
                                'Status':               dropbox_statuses,
                                'Document Name':        dropbox_document_names,
                                
                                  }
                                )

        #Create a version of market research name for merge without "SUB" when the submarket name matches the market name
        dropbox_df['Market Research Name Alternative']  = dropbox_df['Market Research Name'].str.replace(' SUB','')

        #We are now going to merge our dataframe with the list of markets and submarkets with the zip codes associated with each submarket
        
        #We first import and clean that zip code level dataset (convert to one row per submarket with a list of zip codes in it)
        df_zipcodes                                     = pd.read_excel(os.path.join(costar_data_location,'Supplemental Data','Zip to Submarket.xlsx'), dtype={'PostalCode': object} ) 
        df_zipcodes                                     = df_zipcodes.loc[(df_zipcodes['PropertyType'] == 'Office') | (df_zipcodes['PropertyType'] == 'Retail') | (df_zipcodes['PropertyType'] == 'Industrial') | (df_zipcodes['PropertyType'] == 'Multi-Family')]
        df_zipcodes.loc[df_zipcodes['PropertyType']     == 'Multi-Family', 'PropertyType'] = 'Multifamily' #Change spelling of multifamily
        df_zipcodes['state']                            = df_zipcodes['MarketName'].str[-2:]
        df_zipcodes['SubmarketName']                    = df_zipcodes['SubmarketName'].apply(CleanMarketName)
        df_zipcodes['Market Research Name Alternative'] = df_zipcodes['state'] + ' - ' + df_zipcodes['SubmarketName'] +  ' - ' +  df_zipcodes['PropertyType'] #form a variable to match on
        df_zipcodes                                     = df_zipcodes.groupby(['Market Research Name Alternative'])['PostalCode'].apply(list)
        df_zipcodes                                     = df_zipcodes.reset_index()

        #Now merge the zip code data with our costar markets csv
        dropbox_df                                      = pd.merge(dropbox_df, df_zipcodes, on='Market Research Name Alternative',how = 'left')
        
        #Replace the Nan zip code values with an empty list
        dropbox_df['PostalCode']                        = dropbox_df['PostalCode'].apply(lambda d: d if isinstance(d, list) else [])
        
        #Aggregate all zip code lists across each market into single list into its own dataframe to merge back in
        dropbox_df_market_zipcodes                      = dropbox_df.groupby(['Market','Property Type']).agg({'PostalCode': 'sum'}).reset_index()
        dropbox_df_market_zipcodes                      = dropbox_df_market_zipcodes.rename(columns={"PostalCode": "Market Zip Codes"})

        #Merge in the market zip codes into the costar csv dataframe 
        dropbox_df                                      = pd.merge(dropbox_df, dropbox_df_market_zipcodes, on = ['Market','Property Type'],how = 'left')
        
        #Replace the empty lists with the aggregate list of zipcodes for markets
        dropbox_df.loc[dropbox_df['Analysis Type']      == 'Market', 'PostalCode'] = dropbox_df['Market Zip Codes']
        dropbox_df                                       = dropbox_df.drop(columns=['Market Zip Codes','Market Research Name Alternative'])
        dropbox_df                                       = dropbox_df.rename(columns={"PostalCode": "Zip Codes"})
        
        #Now clean the zip code variable by keeping only unique values and dropping the quotation marks around each zip code 
        dropbox_df['Zip Codes'] = dropbox_df['Zip Codes'].apply(UniqueZipCodes)
        
        #Change the market research names if there are duplicates. For example, "CBD" is a fairly common market reearch name
        group_df   = dropbox_df.groupby(['Market Research Name']).size().reset_index(name='Market Research Name Count')
        dropbox_df = pd.merge(dropbox_df, group_df, on='Market Research Name',how = 'left')
        
        dropbox_df['Market 1'] = dropbox_df['Market'].str.split('-')

        dropbox_df['Market Research Name 1'] = dropbox_df['Market Research Name'].str.split(' - ')
        dropbox_df['Market Research Name 2'] = dropbox_df['Market Research Name'].str.split(' - ')
        dropbox_df['Market Research Name 3'] = dropbox_df['Market Research Name'].str.split(' - ')

        dropbox_df['Market 1'] = dropbox_df['Market 1'].str[1]
        dropbox_df['Market Research Name 1'] = dropbox_df['Market Research Name 1'].str[0]
        dropbox_df['Market Research Name 2'] = dropbox_df['Market Research Name 2'].str[1]
        dropbox_df['Market Research Name 3'] = dropbox_df['Market Research Name 3'].str[2]

        dropbox_df.loc[(dropbox_df['Market Research Name Count']      > 1) & (dropbox_df['Analysis Type'] == 'Submarket')  , 'Market Research Name'] = dropbox_df['Market Research Name 1'] + ' - ' + dropbox_df['Market Research Name 2'] + ' (' + dropbox_df['Market 1'] + ') - ' + dropbox_df['Market Research Name 3'] 
        dropbox_df = dropbox_df.drop(columns=['Market 1','Market Research Name 1','Market Research Name 2','Market Research Name 3','Market Research Name Count'])

        group_df   = dropbox_df.groupby(['Market Research Name']).size().reset_index(name='Market Research Name Count')
        dropbox_df = pd.merge(dropbox_df, group_df, on='Market Research Name',how = 'left')
        assert dropbox_df['Market Research Name Count'].all() == 1
        dropbox_df = dropbox_df.drop(columns=['Market Research Name Count'])


        #Now we have to add the custom market reports to our dataframe, loop through the directory and get all document names and merge in with our main df (keeping only the ones we didn't already have)
        dropbox_links                  = []
        dropbox_research_names         = []
        dropbox_markets                = []
        dropbox_analysis_types         = []
        dropbox_prop_types             = []
        dropbox_prop_codes             = []
        dropbox_states                 = []
        dropbox_versions               = []
        dropbox_statuses               = []
        dropbox_document_names         = [] 

        for (dirpath, dirnames, filenames) in os.walk(output_location):
            if filenames == []:
                continue
            else:
                for file in filenames:
                    full_path = dirpath + '/' + file
                    
                    if ('.docx' not in file) and ('.dotm' not in file):
                        continue
                    
                    #Skip the folders in the legacy archive outside the quarter folders
                    if ("""Archive""" in full_path):
                        pass
                        if ("""Legacy Archive""" in full_path):
                            if ("""Legacy Archive\\2""" in full_path):
                                pass
                            else:
                                continue


                    #Parse sector and other info from file path string
                    if (os.path.exists(full_path.replace('_draft','_FINAL'))) and ('_draft' in full_path) or ('docx' not in full_path):
                        continue
                    
                    if """\Condo""" in full_path:
                        prop_type ='Condo'
                        prop_code = 'C'

                    elif """\Single Family""" in full_path:
                        prop_type ='Single Family'
                        prop_code = 'SF'

                    elif """\Retail""" in full_path:
                        prop_type ='Retail'
                        prop_code = 'R'

                    elif """\Multifamily""" in full_path:
                        prop_type ='Multifamily'
                        prop_code = 'MF'
                    
                    elif """\Industrial""" in full_path:
                        prop_type ='Industrial'
                        prop_code = 'I'
                    
                    elif """\Office""" in full_path:
                        prop_type ='Office'
                        prop_code = 'O'

                    elif """Market\Other""" in full_path:
                        prop_type = dirpath.split('Other\\')[1]
                        first_slash_position = prop_type.find('\\') 
                        if first_slash_position != -1:
                            prop_type = prop_type[0:first_slash_position]
                        prop_code = prop_type

                    else:
                        prop_type = 'Unknown'
                        prop_code = 'UK'


                    #Get Dropbox link
                    dropbox_link = dirpath.replace(dropbox_root,r'https://www.dropbox.com/home')
                    dropbox_link = dropbox_link.replace("\\",r'/')    
                    
                    #Determine version and state name
                    if """Legacy Archive\\20""" in full_path:
                        full_path_split =  full_path.split("""\\""")
                        version         = full_path_split[9]
                        state_name      = full_path_split[10]

                    else:
                        version       = file[0:7]
                        state_name    = file[8:10]
                    
                    #Determine Status
                    if '_draft' in file:
                        file_status = 'Draft'
                    else:
                        file_status = 'Final'


                    #Get Market Name
                    try:
                        market     = file.split(' - ')[1].strip()
                        research_name = state_name + ' - ' + market + ' - ' + prop_type
                    
                    except:
                        market         = 'FIX FILE FORMAT'
                        research_name  = 'FIX FILE FORMAT'

                    #Add our variables to the lists that will create the dataframe
                    dropbox_prop_types.append(prop_type)
                    dropbox_prop_codes.append(prop_code)
                    dropbox_statuses.append(file_status)
                    dropbox_versions.append(version)
                    dropbox_links.append(dropbox_link)
                    dropbox_document_names.append(file)
                    dropbox_analysis_types.append('Market')
                    dropbox_markets.append(market)
                    dropbox_research_names.append(research_name)
                    dropbox_states.append(state_name)
                
        








        all_files_dropbox_df = pd.DataFrame({'Market Research Name':dropbox_research_names,
                                'Market':dropbox_markets,
                            'Analysis Type': dropbox_analysis_types,
                            'State':         dropbox_states,
                            'Property Type': dropbox_prop_types,
                            'Property Type Code': dropbox_prop_codes,
                            "Dropbox Links":dropbox_links,
                            'Version':dropbox_versions,
                            'Status':dropbox_statuses,
                            'Document Name': dropbox_document_names})
        all_files_dropbox_df = all_files_dropbox_df.sort_values(by=['State','Market Research Name','Version'],ascending = (True, True,False))

        #Drop the rows in this dataframe that are already in our main dropbox df
        all_files_dropbox_df = all_files_dropbox_df.loc[(all_files_dropbox_df['Dropbox Links'].isin(dropbox_df['Dropbox Links'])) == False   ]        

        #Export the CoStar Markets export
        dropbox_df = dropbox_df.append(all_files_dropbox_df)
        
        #Merge the dataframe with a list of states and the inital of who is assigned to complete them
        assigned_to_df                          = pd.read_excel(os.path.join(general_data_location,'Administrative Data','Assigned To States.xlsx')) 
        dropbox_df                              = pd.merge(dropbox_df,assigned_to_df, on=['State'],how = 'left') 

        csv_name = 'CoStar Markets.csv'
        service_api_csv_name = f'CoStar Markets-{datetime.now().timestamp()}.csv'

        dropbox_df.to_csv(os.path.join(output_location, csv_name), index=False)

        if output_location == os.path.join(dropbox_root,'Research','Market Analysis','Market'):
            dropbox_df.to_csv(os.path.join(output_location, service_api_csv_name), index=False)

def UpdateServiceDb(report_type, csv_name, csv_path, dropbox_dir):
    if type == None:
        return
    
    #We only want to update the database when we are in the production folder and the user is not trying to create a report
    if output_location != os.path.join(dropbox_root,'Research','Market Analysis','Market') or write_reports_yes_or_no != 'no':
        return()
    print(f'Updating service database: {report_type}')

    try:
        url = f'http://market-research-service.bowery.link/api/v1/update/{report_type}'
        dropbox_path = f'{dropbox_dir}{csv_name}'
        payload = { 'location': dropbox_path }

        retry_strategy = Retry(
            total=3,
            status_forcelist=[400, 404, 409, 500, 503, 504],
            allowed_methods=["POST"],
            backoff_factor=5,
            raise_on_status=False
                             )

        adapter = HTTPAdapter(max_retries=retry_strategy)
        http    = requests.Session()
        http.mount("https://", adapter)
        http.mount("http://", adapter)

        response = http.post(url, json=payload)
        response.raise_for_status()
        print('Service successfully updated')
    except HTTPError as e:
        print(f'Request failed with status code {response.status_code}')
        json_error = json.loads(response.content)
        print(json.dumps(json_error, indent=2))
        print('Service DB did not successfully update. Please run the script again after fixing the error.')
    finally:
        print(f'Deleting temporary CSV: ', csv_path)
        os.remove(csv_path)           

#Set formatting paramaters for reports
primary_font                    = 'Avenir Next LT Pro Light' 
primary_space_after_paragraph   = 6

#Decide if you want to create report documents or create our csv output and update the database
user_selects_reports_or_not()
user_selects_sector()

#Define these empty lists we will fill during the loops, this is to create a list of markets and submarkets and their dropbox links for Salesforce mapping
CreateEmptySalesforceLists()

#This is the main loop for our program where we loop through the selected sector dataframes, get list of unique markets, loop through those markets creating folders and writing market reports
for df,df2,sector in zip(df_list, df_slices_list, sector_name_list):

    print('--',sector,'--')

    #Create dictionary with each market as key and a list of its submarkets as items
    market_dictionary            = CreateMarketDictionary(df)

    #Use a GUI to let user select a market to create reports for
    selected_market              = user_selects_market(market_list = list(market_dictionary.keys()),market_or_submarket = 'market') 

    #Loop through the market dictionary creating reports for each market and their submarkets
    for primary_market,submarkets in market_dictionary.items():
        state                        = primary_market[-2:] #Get State to make folder that stores markets

        if   primary_market not in selected_market: 
            continue

        print(primary_market)

        primary_market_clean         = CleanMarketName(primary_market)
        primary_market_name_for_file = primary_market_clean.replace(' - ' + state,'' ).strip() #Make a string with just name of market (without the '- STATECODE' portion)

        #"market" is the general variable name used in all functions for the market OR submarket we are doing report for   
        market                       = primary_market 
        CreateMarketReport()
        
        #If the user has selected all markets in a sector, assume they want all submarkets run, otherwise ask them for the submarket they want
        if selected_market == list(market_dictionary.keys()):
            selected_submarket = submarkets
        else:
            selected_submarket = user_selects_market(market_list = submarkets, market_or_submarket = 'submarket') #use a GUI to let user select a market
        
        #Create all the submarket reports for the market
        for submarket in submarkets:
            #If the current submarket is not the one the user selected, skip it
            if submarket not in selected_submarket:
                continue
            market = submarket
            print(submarket)
            CreateMarketReport()

    #Now create national reports
    state                        = 'US'
    market                       = 'United States of America'
    primary_market               = 'United States of America'
    primary_market_clean         = CleanMarketName(primary_market)
    primary_market_name_for_file = primary_market_clean.replace(' - ' + state,'' ).strip() #Make a string with just name of market (without the '- STATECODE' portion)
    market                       = 'United States of America'
    CreateMarketReport() 

#Now call our function that creates a csv with all the current market reports
CreateDirectoryCSV()        

#Post an update request to the Market Research Docs Service to update the database
UpdateServiceDb(report_type = 'markets', 
                csv_name    = service_api_csv_name, 
                csv_path    = os.path.join(output_location, service_api_csv_name),
                dropbox_dir = 'https://www.dropbox.com/home/Research/Market Analysis/Market/')

print('Finished, you rock!')        



