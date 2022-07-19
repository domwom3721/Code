#This file holds the functions used for producing tables in market and submarket reports
import pandas as pd
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import re

def AddOverviewTable(document, number_rows, number_cols, row_data, col_width): #Function we use to insert our overview table into the report document

    #Make sure each row has the same number of items 
    for row in row_data:
        for row2 in row_data:
            assert len(row) == len(row2)


    #create table object
    tab = document.add_table(rows=len(row_data), cols=len(row_data[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True
    #decide if we use standard style or custom

    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,row_data)): 
        assert (len(row_data_list) == number_cols) and (isinstance(row_data_list, list)) #make sure there is an item for each column in this row
        
        #Set height for row
        row.height = Inches(0.17)
        
        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):

            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.23)

            elif current_column == 1:
                cell.width = Inches(1.20)

            elif current_column == 2:
                cell.width = Inches(0.8)

            elif current_column == 3:
                cell.width = Inches(0.8)

            elif current_column == 4:
                cell.width = Inches(1.20)

            elif current_column == 5:
                cell.width = Inches(0.8)

            elif current_column == 6:
                cell.width = Inches(0.8)

            #add border to top row
            if current_row == 1:
                    tcPr      = cell._element.tcPr
                    tcBorders = OxmlElement("w:tcBorders")
                    top       = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')

                    tcBorders.append(top)
                    tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                #Paragaph spacing before and after
                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                
                for run in paragraph.runs:
                    font          = run.font
                    font.size     = Pt(7)
                    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    #make first row bold
                    if current_row == 0: 
                        font      = run.font
                        font.size = Pt(8)
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
    
def AddTable(document, row_data, col_width): #Function we use to insert our wide tables into report document 
    try:
        #Make sure each list in the list of lists have the same number of elements 
        for row in row_data:
            for row2 in row_data:
                assert len(row) == len(row2)
        number_rows = len(row_data)
        number_cols = len(row_data[0])

        #Make sure the data has a list of inputs for each row
        assert number_rows == len(row_data) 

        #create table object
        tab               = document.add_table(rows=number_rows, cols=number_cols)
        tab.alignment     = WD_TABLE_ALIGNMENT.CENTER

        #loop through the rows in the table
        for current_row, (row, row_data_list) in enumerate(zip(tab.rows, row_data)): 
            assert (len(row_data_list) == number_cols) and (isinstance(row_data_list, list)) #make sure there is an item for each column in this row

        
            row.height = Inches(0.17)
            
            #loop through all cells in the current row
            for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):
                
                #We do these replacements so the last 2 column titles dont spill over onto 2 lines
                text_for_cell =  str(cell_data) 
                text_for_cell =  re.sub('202[0-9] Q[0-9]', ("""’""" +  text_for_cell[2:]), text_for_cell) 
                cell.text =  text_for_cell
        
                if current_row == 0:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


                #add border to top row
                if current_row == 1:
                        tcPr        = cell._element.tcPr
                        tcBorders   = OxmlElement("w:tcBorders")
                        top         = OxmlElement('w:top')
                        top.set(qn('w:val'), 'single')
                        tcBorders.append(top)
                        tcPr.append(tcBorders)


                #loop through the paragraphs in the cell and set font and style
                for paragraph in cell.paragraphs:
                    if current_column > 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    paragraph.paragraph_format.space_after  = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)

                    for run in paragraph.runs:
                        font = run.font
                        if current_row == 0:
                            font.size = Pt(8)
                        else:
                            font.size = Pt(7)

                        #make first row bold
                        if current_row == 0: 
                            font.bold = True
                            font.name = 'Avenir Next LT Pro Demi'
        #set column widths
        for i in range(0,number_cols):
            for cell in tab.columns[i].cells:
                cell.width = Inches(col_width)
    except Exception as e:
        print(e,'problem adding wide table')

def AddHeading(document, title, heading_level): #Function we use to insert the table title headers other than the first one in the report for the overview table
            heading                               = document.add_heading(title,level=heading_level)
            heading.style                         = document.styles['Heading 3']
            heading_style                         = heading.style
            heading_style.font.name               = "Avenir Next LT Pro"
            heading_style.font.size               = Pt(11)
            heading_style.font.bold               = False
            heading.paragraph_format.space_after  = Pt(6)
            heading.paragraph_format.space_before = Pt(6)

            #Color
            heading_style.font.color.rgb = RGBColor.from_string('3F65AB')            
            heading_style.element.xml
            rFonts                       = heading_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def CreateRowDataForTable(data_frame, data_frame2, data_frame3, var1, var2, var3, modifier1, modifier2, modifier3, title): #Returns a list which will populate a row in the overview table
    
    #National
    if data_frame.equals(data_frame2) and data_frame.equals(data_frame3):
        row_data_list = [title,
                        data_frame[var1].iloc[-1],
                        data_frame[var2].iloc[-1],
                        data_frame[var3].iloc[-1],
                        ]

    #Market
    elif data_frame.equals(data_frame2):
        row_data_list = [title,
                        data_frame[var1].iloc[-1],
                        data_frame[var2].iloc[-1],
                        data_frame[var3].iloc[-1],
                        data_frame3[var1].iloc[-1],
                        data_frame3[var2].iloc[-1],
                        data_frame3[var3].iloc[-1]
                        ]
    
    
    #Submarket
    else:
        row_data_list = [title,
                        data_frame[var1].iloc[-1],
                        data_frame[var2].iloc[-1],
                        data_frame[var3].iloc[-1],
                        data_frame2[var1].iloc[-1],
                        data_frame2[var2].iloc[-1],
                        data_frame2[var3].iloc[-1]
                        ]

    
    #Add modifiers to variables ($,%,bps)
    first_spot  = 1
    second_spot = 4
    #loop through the 3 modifiers passed into function, going to the places in the row where the modifiers are used ie: (1,4),(2,5),(3,6)
    for modifier in [modifier1,modifier2,modifier3]:
        if modifier == "$":

            if var1 == ('Total Sales Volume') or var1 == ('Asset Value/Unit') or var1 == ('Asset Value/Sqft') or var1 == ('Market Effective Rent/Unit'):
                row_data_list[first_spot] =  str("${:,.0f}".format(row_data_list[first_spot]))

                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = str("${:,.0f}".format(row_data_list[second_spot])) 
                
            else:
                row_data_list[first_spot] = modifier1 + '' + str(row_data_list[first_spot])
                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = modifier1 + '' + str(row_data_list[second_spot])

        elif modifier == '%':
            if var1 == ('Total Sales Volume') or  var1 == ('Sales Volume Transactions'): #only want whole number for percent change in these variables
                row_data_list[first_spot]  = str("{0:,.0f}".format(row_data_list[first_spot]))   + modifier
                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = str("{0:,.0f}".format(row_data_list[second_spot]))  + modifier
            else:
                row_data_list[first_spot]  = str(row_data_list[first_spot]) + modifier
                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = str(row_data_list[second_spot]) + modifier

        elif modifier == 'bps':
            row_data_list[first_spot]  = str("{:,.0f}".format(row_data_list[first_spot]))  + ' ' + modifier
            if second_spot <= len(row_data_list) - 1:
                row_data_list[second_spot] = str("{:,.0f}".format(row_data_list[second_spot])) + ' ' + modifier

        else:
            if var1 == 'Sales Volume Transactions':
                row_data_list[first_spot] =  str("{:,.0f}".format(row_data_list[first_spot]))
                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = str("{:,.0f}".format(row_data_list[second_spot]))
            
            elif var1 == 'Absorption Units' or var1 == 'Net Absorption SF':
                row_data_list[first_spot] =  str("{:,.0f}".format(row_data_list[first_spot]))
                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = str("{:,.0f}".format(row_data_list[second_spot])) 

            else:
                row_data_list[first_spot]  = str(row_data_list[first_spot])  + ' ' + modifier
                if second_spot <= len(row_data_list) - 1:
                    row_data_list[second_spot] = str(row_data_list[second_spot]) + ' ' + modifier
            
        
        first_spot  +=1
        second_spot +=1
        
    for count,cell in enumerate(row_data_list):
        if cell == 'nan' or cell == 'inf' or cell == '-inf%'  or cell == '-inf' or cell == 'nan%' or cell == 'inf%' or cell =='nan bps' or cell =='inf bps':
            row_data_list[count] = 'NA'
    return(row_data_list)

def CreateRowDataForWideTable(data_frame, data_frame2, data_frame3, data_frame4, var1, modifier, sector): #Returns list of lists with data we use to fill rows in the wide table
    #This function takes a variable and returns a list of lists of that variables value over time in the market, submarket, and nation, 
    #and if we are doing a market the different quality slices. Each list in the list represents a row in a table for either rent or vacancy (ADD IN CAP RATES)

    level_1_name = data_frame['Geography Name'].iloc[0]
    level_2_name = data_frame2['Geography Name'].iloc[0]
    level_3_name = data_frame3['Geography Name'].iloc[0] #Typically will be United States, excpet for when doing NYC 


    if level_3_name == 'United States of America':
        level_3_name = 'National'
    elif level_3_name == 'New York - NY':
        level_3_name = 'Metro'

    if data_frame.equals(data_frame2):
        market_or_submarket = 'market'
    else:
        market_or_submarket = 'submarket'


    #Sort the 4 dataframes
    data_frame  = data_frame.sort_values(by=['Period'],ascending = False) 
    data_frame2 = data_frame2.sort_values(by=['Period'],ascending = False) 
    data_frame3 = data_frame3.sort_values(by=['Period'],ascending = False) 
    data_frame4 = data_frame4.sort_values(by=['Slice','Period'],ascending = False) 

    if modifier == '$' and sector == 'Multifamily':
        data_frame[var1]  = data_frame[var1].map('${:,.0f}'.format)
        data_frame2[var1] = data_frame2[var1].map('${:,.0f}'.format)
        data_frame3[var1] = data_frame3[var1].map('${:,.0f}'.format)
        data_frame4[var1] = data_frame4[var1].map('${:,.0f}'.format)

    elif modifier == '$' and sector != 'Multifamily':
        data_frame[var1]  = data_frame[var1].map('${:,.2f}'.format)
        data_frame2[var1] = data_frame2[var1].map('${:,.2f}'.format)
        data_frame3[var1] = data_frame3[var1].map('${:,.2f}'.format)
        data_frame4[var1] = data_frame4[var1].map('${:,.2f}'.format)
   
    else:
        data_frame[var1]  = data_frame[var1].map('{:,.1f}%'.format)
        data_frame2[var1] = data_frame2[var1].map('{:,.1f}%'.format)
        data_frame3[var1] = data_frame3[var1].map('{:,.1f}%'.format)
        data_frame4[var1] = data_frame4[var1].map('{:,.1f}%'.format)

    

    #Keep all rows from the current year and the last quarter from all previous years
    data_frame_last2_quarters             = data_frame.head(2) #dataframe with 2 most recent quarters
    data_frame2_last2_quarters            = data_frame2.head(2) #dataframe with 2 most recent quarters
    data_frame3_last2_quarters            = data_frame3.head(2) #dataframe with 2 most recent quarters
    data_frame4_last2_quarters            = data_frame4.groupby(['Slice']).head(2) #dataframe with 2 most recent quarters

    #Keep all remaining rows in the dataframe
    data_frame                           = data_frame[2:]     
    data_frame                           = data_frame[data_frame['Period'].str.contains("Q4")]
    data_frame['Period']                 = data_frame['Period'].str.replace(r' Q4', '')

    data_frame2                           = data_frame2[2:]     
    data_frame2                           = data_frame2[data_frame2['Period'].str.contains("Q4")]
    data_frame2['Period']                 = data_frame2['Period'].str.replace(r' Q4', '')

    data_frame3                           = data_frame3[2:]     
    data_frame3                           = data_frame3[data_frame3['Period'].str.contains("Q4")]
    data_frame3['Period']                 = data_frame3['Period'].str.replace(r' Q4', '')

    if len(data_frame4) > 0:
        tail_length                           =  int(len(data_frame4))/len(data_frame4['Slice'].unique()) - 2
        data_frame4                           = data_frame4.groupby(['Slice']).tail(tail_length)
        data_frame4                           = data_frame4[data_frame4['Period'].str.contains("Q4")]
        data_frame4['Period']                 = data_frame4['Period'].str.replace(r' Q4', '')
   

    #Append the 2 most recent quarters with the Q4 dataframe
    data_frame = data_frame_last2_quarters.append(data_frame)
    data_frame = data_frame.reset_index()

    data_frame2 = data_frame2_last2_quarters.append(data_frame2)
    data_frame2 = data_frame2.reset_index()

    data_frame3 = data_frame3_last2_quarters.append(data_frame3)
    data_frame3 = data_frame3.reset_index()

    data_frame4 = data_frame4_last2_quarters.append(data_frame4)
    data_frame4 = data_frame4.reset_index()

    #Cut down to the variables we are going to display in the table
    data_frame = data_frame[[var1,'Period']]
    data_frame2 = data_frame2[[var1,'Period']]
    data_frame3 = data_frame3[[var1,'Period']]
    data_frame4 = data_frame4[[var1,'Period','Slice']]

    #Sort again so that most recent quarter is last
    data_frame  = data_frame.sort_values(by=['Period'],ascending = True) 
    data_frame2 = data_frame2.sort_values(by=['Period'],ascending = True) 
    data_frame3 = data_frame3.sort_values(by=['Period'],ascending = True) 
    data_frame4 = data_frame4.sort_values(by=['Slice','Period'],ascending = True) 

    #Add empty rows to the top of the dataframes
    data = []
    data.insert(0, {'Period': '', var1: 'Submarket'})
    data_frame  = pd.concat([pd.DataFrame(data), data_frame], ignore_index=True)
    
    data = []
    data.insert(0, {'Period': '', var1: 'Market'})
    data_frame2  = pd.concat([pd.DataFrame(data), data_frame2], ignore_index=True)

    
    data = []
    data.insert(0, {'Period': '', var1: level_3_name})
    data_frame3  = pd.concat([pd.DataFrame(data), data_frame3], ignore_index=True)

 
    
    #If we are doing a market
    if market_or_submarket == 'market' or len(data_frame4) > 0 :

        try:
            list_of_lists = [data_frame['Period'].tolist(),
                data_frame3[var1].tolist(),
                data_frame2[var1].tolist()]
            
            #Since it's a market, we are going to add a row for each slice we have in our slice data
            for slice in data_frame4['Slice'].unique():
                data_frame_temp     = data_frame4.loc[(data_frame4['Slice'] == slice)]
                data_frame_temp     = data_frame_temp.reset_index()
                slice_list_to_add   = data_frame_temp[var1].tolist()
                slice_list_to_add.insert(0,slice)
                list_of_lists.append(slice_list_to_add)

            #make sure each list (row) in the list of lists (rows) have same number of items
            for list in list_of_lists:
                for list2 in list_of_lists:
                    assert len(list) == len(list2)
    
        #When we don't have full data for the slices
        except:
             list_of_lists = [data_frame['Period'].tolist(),
                data_frame3[var1].tolist(),
                data_frame2[var1].tolist()]

        #remove extra row for national report
        if level_1_name == 'United States of America':
            del list_of_lists[2]


        return(list_of_lists)

    #If we are doing a submarket
    else:
        try:
            list_of_lists = [data_frame['Period'].tolist(),
                    data_frame3[var1].tolist(),
                    data_frame2[var1].tolist(),
                    data_frame[var1].tolist()]
            
            #make sure each list (row) in the list of lists (rows) have same number of items
            for list in list_of_lists:
                for list2 in list_of_lists:
                    assert len(list) == len(list2)

            return(list_of_lists)
        

        except:
            #In some cases, the submarket does not have a full 10 years of data. To handle this, we make sure we display data for quarters we
            #coverage and make the market and nation rows blank
            submarket_period_list = data_frame['Period'].tolist()
            submarket_list        =  data_frame[var1].tolist()
            
            #Create empty rows for market and nation
            market_list           =  ['Market']
            nation_list           =  ['-----National-----']
            for i in range(len(submarket_list) -1):
                market_list.append('') 
                nation_list.append('')

            #Make lists of the data and ensure they are all the same size
            list_of_lists = [submarket_period_list,nation_list,market_list,submarket_list]
            for list in list_of_lists:
                for list2 in list_of_lists:
                    assert len(list) == len(list2)

            return(list_of_lists)

def AddMarketPerformanceTable(document, col_width, market_data_frame, sector): 

    #Convert market dataframe into a (mostly) annual dataset for a handful of variables
    
    #Start by declaring a list of variables we want to display
    if sector == 'Multifamily':
        variables_of_interest = ['Period','Inventory Units','Under Construction Units','Net Delivered Units 12 Mo','Absorption Units 12 Mo',    'Vacancy Rate',                           'Market Effective Rent/Unit']
    elif sector == 'Retail':
        variables_of_interest = ['Period','Inventory SF',   'Under Construction SF',   'Net Delivered SF 12 Mo',    'Net Absorption SF 12 Mo',  'Vacancy Rate',   'Availability Rate',    'Market Rent/SF']
    elif sector == 'Industrial':
        variables_of_interest = ['Period','Inventory SF',   'Under Construction SF',   'Net Delivered SF 12 Mo',    'Net Absorption SF 12 Mo',  'Vacancy Rate',   'Availability Rate',    'Market Rent/SF']
    elif sector == 'Office':
       variables_of_interest = ['Period','Inventory SF',    'Under Construction SF',   'Net Delivered SF 12 Mo',    'Net Absorption SF 12 Mo',  'Vacancy Rate',   'Availability Rate',    'Market Rent/SF']
    
    #Now create an annual dataset where we keep the last period of each year besides the current year
    market_data_frame = market_data_frame.sort_values(by=['Year','Quarter'],ascending = False) 


    for var in variables_of_interest:
        market_data_frame[var] = market_data_frame[var].fillna('NA')



    #Keep all rows from the current year and the last quarter from all previous years
    market_data_frame_last2_quarters     = market_data_frame.head(2) #dataframe with 2 most recent quarters
    market_data_frame                    = market_data_frame[2:]     #dataframe with remaining quarters
    market_data_frame                    = market_data_frame[market_data_frame['Period'].str.contains("Q4")]
    market_data_frame['Period']          = market_data_frame['Period'].str.replace(r' Q4', '')
    
    #Append the 2 most recent quarters with the Q4 dataframe
    market_data_frame = market_data_frame_last2_quarters.append(market_data_frame)
    market_data_frame = market_data_frame.reset_index()

    #Cut down to the variables we are going to display in the table
    market_data_frame = market_data_frame[variables_of_interest]

    #create table object
    number_rows = len(market_data_frame) + 1 #we add extra row for variable names at the top
    number_cols = len(market_data_frame.columns)

    tab = document.add_table(rows=number_rows, cols=number_cols)
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    for current_row,row in enumerate(tab.rows): 
        for current_column,cell in enumerate(row.cells):
            
            if current_row == 0:
                var_name = str(variables_of_interest[current_column])
                cell.text = var_name
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                
            else:
                current_variable = variables_of_interest[current_column]
                data = market_data_frame[current_variable].iloc[current_row-1] #look up the data value for the right period for the current variable
                if type(data) == str:
                    pass
                else:
                    if current_variable == 'Vacancy Rate' or current_variable == 'Availability Rate':
                        data = "{:,.1f}%".format(data) 

                    elif current_variable == 'Market Effective Rent/Unit':
                        data = "${:,.0f}".format(data)
                    
                    elif current_variable == 'Market Rent/SF':
                        data = "${:,.2f}".format(data)
                    
                    elif (current_variable == 'Inventory SF') or (current_variable == 'Inventory Units'):
                        data = "{:,.0f}".format(data)
                    else:
                        data = "{:,.0f}".format(data)

                cell.text = data
                

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)
            else:
                cell.width = Inches(col_width)

            
            #add border to top row
            if current_row == 1:
                tcPr      = cell._element.tcPr
                tcBorders = OxmlElement("w:tcBorders")
                top       = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
            
                tcBorders.append(top)
                tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

                for run in paragraph.runs:
                    font = run.font
                    if current_row == 0:
                        font.size= Pt(8)
                    else:
                        font.size = Pt(7)

                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'

def AddSubmarketsPerformanceTable(document, col_width, submarkets_data_frame, sector): 
    if len(submarkets_data_frame) == 0:
        return() #If there are no submarkets, do nothing

    
    submarkets_data_frame['Submarket'] = submarkets_data_frame['Geography Name']

    #Start by declaring a list of variables we want to display
    if sector == 'Multifamily':
        variables_of_interest = ['Submarket', 'Inventory Units', 'Vacancy Rate', 'Under Construction Units','Market Effective Rent/Unit']
    else:
        variables_of_interest = ['Submarket', 'Inventory SF', 'Vacancy Rate',  'Availability Rate', 'Under Construction SF', 'Market Rent/SF']

    
    for var in variables_of_interest:
        submarkets_data_frame[var] = submarkets_data_frame[var].fillna('NA')

    #Sort from largest to smallest 
    submarkets_data_frame = submarkets_data_frame.sort_values(by=[variables_of_interest[1],'Submarket'],ascending = False) 
    
    #Cut down to the variables we are going to display in the table
    submarkets_data_frame = submarkets_data_frame[variables_of_interest]

    #Drop market preamble from submarket name
    submarkets_data_frame['Submarket'] =  submarkets_data_frame['Submarket'].str.split(' - ').str[2]




    #create table object
    number_rows = len(submarkets_data_frame) + 1 #we add extra row for variable names at the top
    number_cols = len(submarkets_data_frame.columns)

    tab               = document.add_table(rows=number_rows, cols=number_cols)
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    for current_row,row in enumerate(tab.rows): 
        for current_column,cell in enumerate(row.cells):
            
            if current_row == 0:
                var_name                = str(variables_of_interest[current_column])
                cell.text               = var_name
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                
                
                

            else:
                current_variable = variables_of_interest[current_column]
                data = submarkets_data_frame[current_variable].iloc[current_row-1] #look up the data value for the right period for the current variable
                if type(data) == str:
                    pass
                else:
                    if current_variable == 'Vacancy Rate' or current_variable == 'Availability Rate':
                        data = "{:,.1f}%".format(data) 

                    elif current_variable == 'Market Effective Rent/Unit':
                        data = "${:,.0f}".format(data)
                    
                    elif current_variable == 'Market Rent/SF':
                        data = "${:,.2f}".format(data)
                    
                    elif (current_variable == 'Inventory SF') or (current_variable == 'Inventory Units'):
                        data = "{:,.0f}".format(data)
                    else:
                        data = "{:,.0f}".format(data)

                cell.text = data
                

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)
            else:
                cell.width = Inches(col_width)

            
            #add border to top row
            if current_row == 1:
                tcPr      = cell._element.tcPr
                tcBorders = OxmlElement("w:tcBorders")
                top       = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
            
                tcBorders.append(top)
                tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

                for run in paragraph.runs:
                    font = run.font
                    if current_row == 0:
                        font.size= Pt(8)
                    else:
                        font.size = Pt(7)

                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
                    if current_column == 0 and current_row != 0:
                        font.italic = False

def AddTransactionTable(document, col_width, market_data_frame, sector): 

    if len(market_data_frame) == 0:
        return()

    #Start by declaring a list of variables we want to display
    if sector == 'Multifamily':
        variables_of_interest = ['Property Address', 'Number Of Units', 'Building Class', 'Style', 'Year Built', 'Last Sale Date', 'Price/Unit']
    else:
        variables_of_interest = ['Property Address', 'RBA', 'Building Class', 'Year Built', 'Last Sale Date', 'Last Sale Price']

    for var in variables_of_interest:
        market_data_frame[var] = market_data_frame[var].fillna('NA')



    #Cut down to the variables we are going to display in the table
    market_data_frame = market_data_frame[variables_of_interest]

    #create table object
    number_rows = len(market_data_frame) + 1 #we add extra row for variable names at the top
    number_cols = len(market_data_frame.columns)

    tab               = document.add_table(rows=number_rows, cols=number_cols)
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER

    for current_row,row in enumerate(tab.rows): 
        for current_column,cell in enumerate(row.cells):
            
            if current_row == 0:
                var_name                 = str(variables_of_interest[current_column])
                cell.text                = var_name
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                
            else:
                current_variable = variables_of_interest[current_column]
                data             = market_data_frame[current_variable].iloc[current_row-1] #look up the data value for the right period for the current variable
                if type(data) == str:
                    pass
                elif current_variable == 'Last Sale Date':
                    data = data.to_pydatetime()
                    data = data.strftime('%b %d, %Y')
                elif current_variable == 'Year Built':
                    data = str(data)
                elif current_variable == 'Last Sale Price':
                    data = "$" + "{:,.0f}".format(data)
                elif current_variable == 'Price/Unit':
                    data = "$" + "{:,.0f}".format(data)
                else:
                    data = "{:,.0f}".format(data)

                cell.text = data
                

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.35)
            else:
                cell.width = Inches(col_width)

            
            #add border to top row
            if current_row == 1:
                tcPr      = cell._element.tcPr
                tcBorders = OxmlElement("w:tcBorders")
                top       = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
            
                tcBorders.append(top)
                tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

                for run in paragraph.runs:
                    font = run.font
                    if current_row == 0:
                        font.size= Pt(8)
                    else:
                        font.size = Pt(7)

                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'

def AddConstructionTable(document, col_width, market_data_frame, sector): 

    if len(market_data_frame) == 0:
        return()

    #Start by declaring a list of variables we want to display
    if sector == 'Multifamily':
        variables_of_interest = ['Property Address', 'Property Name','Building Status', 'Year Built', 'Building Class', 'Number of Units' ]
    else:
        variables_of_interest = ['Property Address',                 'Building Status', 'Year Built', 'Building Class', 'RBA'             ]

    for var in variables_of_interest:
        market_data_frame[var] = market_data_frame[var].fillna('NA')



    #Cut down to the variables we are going to display in the table
    market_data_frame = market_data_frame[variables_of_interest]

    #create table object
    number_rows = len(market_data_frame) + 1 #we add extra row for variable names at the top
    number_cols = len(market_data_frame.columns)

    tab               = document.add_table(rows=number_rows, cols=number_cols)
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER

    for current_row,row in enumerate(tab.rows): 
        for current_column,cell in enumerate(row.cells):
            
            if current_row == 0:
                var_name                 = str(variables_of_interest[current_column])
                cell.text                = var_name
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                
            else:
                current_variable = variables_of_interest[current_column]
                data             = market_data_frame[current_variable].iloc[current_row-1] #look up the data value for the right period for the current variable
                if type(data) == str:
                    pass
                elif current_variable == 'Last Sale Date':
                    data = data.to_pydatetime()
                    data = data.strftime('%b %d, %Y')
                elif current_variable == 'Year Built':
                    data = str(data)
                elif current_variable == 'Last Sale Price':
                    data = "$" + "{:,.0f}".format(data)
                elif current_variable == 'Price/Unit':
                    data = "$" + "{:,.0f}".format(data)
                else:
                    data = "{:,.0f}".format(data)

                cell.text = data
                

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.35)
            else:
                cell.width = Inches(col_width)

            
            #add border to top row
            if current_row == 1:
                tcPr      = cell._element.tcPr
                tcBorders = OxmlElement("w:tcBorders")
                top       = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
            
                tcBorders.append(top)
                tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

                for run in paragraph.runs:
                    font = run.font
                    if current_row == 0:
                        font.size= Pt(8)
                    else:
                        font.size = Pt(7)

                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
