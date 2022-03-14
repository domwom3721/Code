#By Mike Leahy, April 30 2021
    #Summary: Takes a United States County and fips code as input, prepares economic report for it using data accessed from different APIs

import json
import math
import msvcrt
import os
import re
import sys
import time
from datetime import date, datetime
from random import choice

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import pyautogui
import requests
import shapefile
import wikipedia
from bls_datasets import qcew
from blsconnect import RequestBLS
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fredapi import Fred
from numpy import NaN
from plotly.subplots import make_subplots
from requests.adapters import HTTPAdapter
from requests.exceptions import HTTPError
from requests.packages.urllib3.util.retry import Retry
from selenium import webdriver
from shapely.geometry import LineString, Point, Polygon

#Define file pre-paths
dropbox_root                   =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
project_location               =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)','Research','Projects','Research Report Automation Project') 
main_output_location           =  os.path.join(project_location,'Output', 'Area')                 #Testing
main_output_location           =  os.path.join(dropbox_root,'Research', 'Market Analysis','Area') #Production
general_data_location          =  os.path.join(project_location,'Data', 'General Data')
data_location                  =  os.path.join(project_location,'Data', 'Area Reports Data')
graphics_location              =  os.path.join(project_location,'Data', 'General Data','Graphics')
map_location                   =  os.path.join(project_location,'Data', 'Area Reports Data', 'County Maps')

def DeclareAPIKeys():
    global fred, bls

    #Declare API Key for FRED and BLS
    fred                = Fred(api_key = choice(['7ab383546af7583fae8a058915edc868', '9875b149440961806f0df696105fe12c', '21843961fd781317a7674a5d23d2c7e8', ] ) )   
    bls                = RequestBLS(key = choice(['2b8d15c77bda4527b101a2b1c98551cf', '9f0492293ac04ade8f2e72576d3822db', '708e9d690b604a7ebda9ff55fe634bc3', 'c993e3b3877845b3a60c8bce507acec6', ]))   
    
def CreateEmptySalesforceLists():
    global  dropbox_counties, dropbox_links, dropbox_fips, dropbox_states, dropbox_analysistypes, dropbox_cbsa_codes
    global dropbox_document_names, dropbox_market_research_names, dropbox_statuses, dropbox_versions
    dropbox_counties              = []
    dropbox_fips                  = []
    dropbox_links                 = []
    dropbox_states                = []
    dropbox_analysistypes         = []
    dropbox_cbsa_codes            = []
    dropbox_document_names        = []
    dropbox_market_research_names = []
    dropbox_statuses              = []
    dropbox_versions              = []

def UpdateSalesforceList():
    #Add to lists that track our markets and submarkets for salesforce
    dropbox_counties.append(county)
    dropbox_fips.append(fips)
    dropbox_states.append(state)
    dropbox_analysistypes.append('Area')
    dropbox_cbsa_codes.append(cbsa)
    dropbox_document_names.append(document_name)
    dropbox_market_research_names.append(state + ' - ' + county)
    dropbox_statuses.append('Outdated')
    dropbox_versions.append(current_quarter)
    #Use the output directory to back into the right dropbox link 
    dropbox_link = county_folder.replace(dropbox_root,r'https://www.dropbox.com/home')
    dropbox_link = dropbox_link.replace("\\",r'/')
    dropbox_links.append(dropbox_link)   
   
def CreateDirectory(state, county):
    global county_folder, county_folder_map, report_path, document_name
    state_folder             = os.path.join(main_output_location, state)
    county_folder            = os.path.join(main_output_location, state, county)
    
    state_folder_map         = os.path.join(map_location, state)
    county_folder_map        = os.path.join(map_location, state, county)

    for folder in [state_folder,county_folder,state_folder_map,county_folder_map]:
         if os.path.exists(folder) == False:
            os.mkdir(folder) 

    document_name            = current_quarter + ' ' + state + ' - ' + county + '_draft.docx'
    report_path              = os.path.join(county_folder, document_name)
    return(county_folder)

def UpdateServiceDb(report_type, csv_name, csv_path, dropbox_dir):
    if type == None:
        return
    #Terminate early if output directory is not set to production    
    if main_output_location != os.path.join(dropbox_root,'Research','Market Analysis','Area'):
        return()

    print(f'Updating service database: {report_type}')

    try:
        url = f'http://market-research-service.bowery.link/api/v1/update/{report_type}'
        dropbox_path = f'{dropbox_dir}{csv_name}'
        payload = { 'location': dropbox_path }

        retry_strategy = Retry(
            total=3,
            status_forcelist=[400, 404, 409, 500, 503, 504],
            allowed_methods=["POST"],
            backoff_factor=5,
            raise_on_status=False
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        http = requests.Session()
        http.mount("https://", adapter)
        http.mount("http://", adapter)

        response = http.post(url, json=payload)
        response.raise_for_status()
        print('Service successfully updated')
    except HTTPError as e:
            print(f'Request failed with status code {response.status_code}')
            json_error = json.loads(response.content)
            print(json.dumps(json_error, indent=2))
            print('Service DB did not successfully update. Please run the script again after fixing the error.')
    finally:
        print(f'Deleting temporary CSV: ', csv_path)
        os.remove(csv_path)

def GetStateName(state_code):
    state_names_df    = pd.read_csv(os.path.join(data_location,'State Names.csv'))
    state_names_df    = state_names_df.loc[state_names_df['State Code'] == state_code]
    state_name        = state_names_df['State Name'].iloc[0]
    return(state_name)

def GetCurrentQuarterAndYear():
    #Pulls unemployment for Nassau County, NY as a way to see the most current available month for county level unemployment, 
    #from this we get the quarter and year for our report version
    df                = fred.get_series(series_id = 'NYNASS9URN')
    df                = df.to_frame().reset_index()
    most_recent_month = str(df['index'].iloc[-1])[5:7] #cuts down to just month value eg: 08
    
    if most_recent_month == '12'  or most_recent_month == '01' or most_recent_month == '02':
        quarter = '4'
    elif most_recent_month == '11' or most_recent_month == '10' or most_recent_month == '09':
        quarter = '3'
    elif most_recent_month == '08' or most_recent_month == '07' or most_recent_month == '06':
       quarter  = '2'
    elif most_recent_month == '05' or most_recent_month == '04' or most_recent_month == '03':
        quarter = '1'

    most_recent_year = str(df['index'].iloc[-1])[0:4] #cuts down to just year value eg: 2021

    if  quarter == '1' or quarter == '2'  or quarter == '3':
        year = most_recent_year 
    elif quarter == '4':
        if most_recent_month == '12':
            year = most_recent_year
        else:
            year = str(int(most_recent_year) -1 )

    return[year, quarter]

class TimeoutExpired(Exception):
        pass

def input_with_timeout(prompt, timeout, timer=time.monotonic):

    sys.stdout.write(prompt)
    sys.stdout.flush()
    endtime = timer() + timeout
    result = []
    while timer() < endtime:
        if msvcrt.kbhit():
            result.append(msvcrt.getwche()) #XXX can it block on multibyte characters?
            if result[-1] == '\r':
                return ''.join(result[:-1])
        time.sleep(0.04) # just to yield to other processes/threads
    raise TimeoutExpired

def FindBLSEndYear():
    #The "BLS.Series" method we use to gather employment data uses a start and end year paramter
    #To avoid extra requests for years in the future, we pass a 5 year range. In February, the latest data will be from December 
    #of the previous year. Therefore, we must pass in the previous year as the end year. Every other month, we should use the current year as end year
    current_year   = todays_date.year                      
    current_month  = todays_date.month
    current_day    = todays_date.day
    previous_year  = current_year - 1
    
    #Once we are in March, the latest employment data will be from January so we can use current year as "end_year"
    if current_month >= 3 and current_day >=14:
        return(current_year)
    else:
        return(previous_year)

#####################################################Data Related Functions####################################
#County Data
def GetCountyGDP(fips, observation_start):
    print('Getting County GDP')

    county_gdp_series_code                        = 'REALGDPALL' + fips
    county_gdp_df                                 = fred.get_series(series_id = county_gdp_series_code,observation_start = observation_start)
    county_gdp_df                                 = county_gdp_df.to_frame().reset_index()
    county_gdp_df.columns                         = ['Period','GDP']
    county_gdp_df['GDP']                          = county_gdp_df['GDP'] * 1000
    county_gdp_df['Lagged GDP']                   = county_gdp_df['GDP'].shift(1)
    county_gdp_df['GDP Growth']                   = ((county_gdp_df['GDP'] / county_gdp_df['Lagged GDP']) - 1)  * 100

    if data_export == True:
        county_gdp_df.to_csv(os.path.join(county_folder, 'County GDP.csv'))

    return(county_gdp_df)

def GetCountyPCI(fips, observation_start):
    print('Getting County Per Capita Personal Income')
    #Per Capita Personal Income
    county_pci_series_code = 'PCPI' + fips
    
    #Different series ID for Washington DC
    if fips == '11001':
        county_pci_series_code = 'DCPCPI'

    county_pci_df         = fred.get_series(series_id = county_pci_series_code, observation_start=observation_start)
    county_pci_df         = county_pci_df.to_frame().reset_index()
    county_pci_df.columns = ['Period','Per Capita Personal Income']
    
    if data_export == True:
        county_pci_df.to_csv(os.path.join(county_folder, 'County Per Capita Personal Income.csv'))
    
    return(county_pci_df)

def GetCountyResidentPopulation(fips, observation_start):
    print('Getting County Population')
    #Resident Population 
    resident_population_series_names     = pd.read_excel(os.path.join(data_location,'FRED Series Names','GeoFRED_Resident_Population_by_County_Thousands_of_Persons.xls'),
                dtype={'Region Code': object
                      }
                                                         )
    resident_population_series_names     = resident_population_series_names.loc[resident_population_series_names['Region Code'] == fips]
    county_pop_series_code               = resident_population_series_names['Series ID'].iloc[0]
    
    county_pop_df                        = fred.get_series(series_id = county_pop_series_code, observation_start = observation_start)
    county_pop_df                        = county_pop_df.to_frame().reset_index()
    county_pop_df.columns                = ['Period','Resident Population']
    county_pop_df['Resident Population'] = county_pop_df['Resident Population'] * 1000

    if data_export == True:
        county_pop_df.to_csv(os.path.join(county_folder, 'County Resident Population.csv') )

    return(county_pop_df)

def GetCountyUnemploymentRate(fips, start_year, end_year): 
    print('Getting County Unemployment Rate')

    #Seasonally-adjusted unemployment rate
    series_name            = 'LAUCN' + fips + '0000000003'
    county_ur_df           = bls.series(series_name,start_year=start_year, end_year= end_year) 
    county_ur_df['year']   = county_ur_df['year'].astype(str)
    county_ur_df['period'] = county_ur_df['period'].str[1:3] + '/' +  county_ur_df['year'].str[2:4]      
    county_ur_df           = county_ur_df.rename(columns={series_name: "unemployment_rate"})

    if data_export == True:
        county_ur_df.to_csv(os.path.join(county_folder, 'County Unemployment Rate.csv'))

    return(county_ur_df)

def GetCountyEmployment(fips, start_year, end_year): 
    print('Getting County Total Employment')

    #Total Employment
    series_name                              = 'LAUCN' + fips + '0000000005'
    county_emp_df                            = bls.series(series_name, start_year = (start_year - 1), end_year = end_year)
    county_emp_df['year']                    = county_emp_df['year'].astype(str)
    county_emp_df['period']                  = county_emp_df['period'].str[1:3] + '/' +  county_emp_df['year'].str[2:4] 
    county_emp_df                            = county_emp_df.rename(columns={series_name: "Employment"})
    county_emp_df['Lagged Employment']       = county_emp_df['Employment'].shift(12)
    county_emp_df['Employment Growth']       = round(((county_emp_df['Employment']/county_emp_df['Lagged Employment']) - 1 ) * 100,2 )

    assert len(county_emp_df) > 60
    
    if data_export == True:
        county_emp_df.to_csv(os.path.join(county_folder, 'County Total Employment.csv'))

    return(county_emp_df)

def GetCountyIndustryBreakdown(fips, year, qtr):
    print('Getting County QCEW Employment Breakdown')

    #Pulls employment data from Quarterly Census of Employment and Wages
    df_qcew                  = qcew.get_data('area', rtype='dataframe', year=year,qtr=qtr, area=fips)
    
    if data_export == True:
        df_qcew.to_csv(os.path.join(county_folder,'qcew_raw.csv'))

    #Restrict to county-ownership level (fed,state,local,private), supersector employment
    df_qcew                 = df_qcew.loc[df_qcew['agglvl_code'] == 73] 
    
    #Drop suppresed employment rows
    df_qcew                 = df_qcew.loc[df_qcew['disclosure_code'] != 'N'] 

    #Drop the rows where employment is 0 
    df_qcew                 = df_qcew.loc[(df_qcew['month3_emplvl'] > 0) ] 

    #Create a seperate dataframe with just the weekly wages by industry
    wtavg                   = lambda x: np.average(x.avg_wkly_wage, weights = x.month3_emplvl,axis = 0) #define function to calcuate weighted average wage
    df_qcew_wages           = df_qcew.groupby('industry_code').apply(wtavg).reset_index()
    df_qcew_wages.columns   = ['industry_code','avg_wkly_wage']

    #Create a seperate dataframe with just the location quotient by industry (averaged across sectors)
    wtavg                   = lambda x: np.average(x.lq_month3_emplvl, weights = x.month3_emplvl,axis = 0) #define function to calcuate weighted average wage
    df_qcew_lq              = df_qcew.groupby('industry_code').apply(wtavg).reset_index()
    df_qcew_lq.columns      = ['industry_code','lq_month3_emplvl']

    #Collapse down to total employment across the 3 ownership codes
    df_qcew                 = df_qcew.groupby('industry_code').agg(month3_emplvl=('month3_emplvl', 'sum'),)
    
    #Merge in the wage and location quotient dataframes
    df_qcew                 = pd.merge(df_qcew, df_qcew_wages, on=('industry_code'),how='outer')
    df_qcew                 = pd.merge(df_qcew, df_qcew_lq, on=('industry_code'),how='outer')

    #Change the industry codes to names
    replacements            = {'1011':'Natural Resources & Mining', 
                               '1012':'Construction', 
                               '1013':'Manufacturing', 
                               '1021':'Trade, Transportation, & Utilities', 
                               '1022':'Information', 
                               '1023':'Financial Activities', 
                               '1024':'Professional & Business Services', 
                               '1025':'Education & Health Services', 
                               '1026':'Leisure & Hospitality', 
                               '1027':'Other Services', 
                               '1028':'Public Administration', 
                               '1029':'Unclassified'
                               }

    df_qcew['industry_code'].replace(replacements, inplace=True)

   
    #Sort by total employement
    df_qcew['employment_fraction'] = round(((df_qcew['month3_emplvl']/(df_qcew['month3_emplvl'].sum())) * 100),2)
    df_qcew['county']              = county
    df_qcew                        = df_qcew.loc[df_qcew['industry_code'] != 'Unclassified']
    df_qcew                        = df_qcew.sort_values(by=['month3_emplvl'])

    #Export final data
    if data_export == True:
        df_qcew.to_csv(os.path.join(county_folder, 'County Industry Breakdown.csv'))

    return(df_qcew)

def GetCountyIndustryGrowthBreakdown(fips, year,qtr):
    print('Getting County QCEW Employment Growth Breakdown')


    #Pulls employment data (and the lagged data) from Quarterly Census of Employment and Wages
    df_qcew          = qcew.get_data('area', rtype='dataframe', year=year,qtr=qtr, area=fips)
    df_qcew_lagged   = qcew.get_data('area', rtype='dataframe', year=(str(int(year) - growth_period )),qtr=qtr, area=fips)
    df_qcew_lagged1  = qcew.get_data('area', rtype='dataframe', year=(str(int(year) - 1 )),qtr=qtr, area=fips)

    #Restrict to county-ownership level (fed,state,local,private), supersector employment
    df_qcew          = df_qcew.loc[df_qcew['agglvl_code'] == 73] 
    df_qcew_lagged   = df_qcew_lagged.loc[df_qcew_lagged['agglvl_code'] == 73] 
    df_qcew_lagged1  = df_qcew_lagged1.loc[df_qcew_lagged1['agglvl_code'] == 73] 

    #Restrict to private ownership 
    df_qcew          = df_qcew.loc[df_qcew['own_code'] == 5] 
    df_qcew_lagged   = df_qcew_lagged.loc[df_qcew_lagged['own_code'] == 5] 
    df_qcew_lagged1  = df_qcew_lagged1.loc[df_qcew_lagged1['own_code'] == 5] 

    
    if data_export == True:
        df_qcew.to_csv(os.path.join(county_folder,'qcew_raw.csv'))
        df_qcew_lagged.to_csv(os.path.join(county_folder,'qcew_raw_lagged.csv'))
        df_qcew_lagged1.to_csv(os.path.join(county_folder,'qcew_raw_lagged1.csv'))

    #Add "lagged" and "lagged1" to the column names for the lagged data
    df_qcew_lagged   = df_qcew_lagged.add_prefix('lagged_')
    df_qcew_lagged1  = df_qcew_lagged1.add_prefix('lagged1_')

    #Remove the "lagged" and "lagged1" prefix for the industry and ownership code columns so we can merge on them
    df_qcew_lagged   = df_qcew_lagged.rename(columns={"lagged_own_code": "own_code", "lagged_industry_code": "industry_code"})
    df_qcew_lagged1  = df_qcew_lagged1.rename(columns={"lagged1_own_code": "own_code", "lagged1_industry_code": "industry_code"})


    #Merge together the current quarters data with the data from 1 year ago and with the data from (5) years from now
    df_joint = pd.merge(df_qcew, df_qcew_lagged, on=('industry_code','own_code'),how='outer')
    df_joint = pd.merge(df_joint, df_qcew_lagged1, on=('industry_code','own_code'),how='outer') #now merge in lagged employment data

    #Flag the industries and ownership type rows where the data was suppresed in the past or present
    filter = (df_joint['disclosure_code'] == 'N') | (df_joint['lagged_disclosure_code'] == 'N')
    df_joint['Employment Growth Invalid'] = ''
    df_joint.loc[filter, ['Employment Growth Invalid']] = 1
    df_joint.loc[df_joint['Employment Growth Invalid'] != 1, ['Employment Growth Invalid']] = 0
    
    one_year_filter = (df_joint['disclosure_code'] == 'N') | (df_joint['lagged1_disclosure_code'] == 'N' )
    df_joint['1Y Employment Growth Invalid'] = ''
    df_joint.loc[one_year_filter, ['1Y Employment Growth Invalid']] = 1
    df_joint.loc[df_joint['1Y Employment Growth Invalid'] != 1, ['1Y Employment Growth Invalid']] = 0
    

    #Replace the Employment Growth Invalid column with the maximum value from each row for a given industry
    df_joint['Employment Growth Invalid'] = df_joint.groupby('industry_code')['Employment Growth Invalid'].transform('max')
    df_joint['1Y Employment Growth Invalid'] = df_joint.groupby('industry_code')['1Y Employment Growth Invalid'].transform('max')

    #Drop the rows where employment is 0 
    df_joint          = df_joint.loc[(df_joint['month3_emplvl'] > 0) ] 

    #Create a seperate dataframe with just the current quarters weekly wages by industry
    wtavg = lambda x: np.average(x.avg_wkly_wage, weights = x.month3_emplvl,axis = 0) #define function to calcuate weighted average wage
    df_qcew_wages           = df_joint.groupby('industry_code').apply(wtavg).reset_index()
    df_qcew_wages.columns = ['industry_code','avg_wkly_wage']

    #Collapse down to total employment across the 3 ownership codes
    df_joint                 = df_joint.groupby('industry_code').agg(month3_emplvl=('month3_emplvl', 'sum'),lagged_month3_emplvl=('lagged_month3_emplvl', 'sum'),lagged1_month3_emplvl=('lagged1_month3_emplvl', 'sum'),emp_growth_invalid=('Employment Growth Invalid', 'max'),one_year_emp_growth_invalid=('1Y Employment Growth Invalid', 'max'))
    df_joint                 = pd.merge(df_joint, df_qcew_wages, on=('industry_code'),how='outer')

    #Change the industry codes to names
    replacements = {'1011':'Natural Resources & Mining', 
                    '1012':'Construction', 
                    '1013':'Manufacturing', 
                    '1021':'Trade, Transportation, & Utilities', 
                    '1022':'Information', 
                    '1023':'Financial Activities', 
                    '1024':'Professional & Business Services', 
                    '1025':'Education & Health Services', 
                    '1026':'Leisure & Hospitality', 
                    '1027':'Other Services', 
                    '1028':'Public Administration', 
                    '1029':'Unclassified'}

    df_joint['industry_code'].replace(replacements, inplace=True)

   

    #Calcualte employment growth rates
    df_joint['Employment Growth'] = round((((df_joint['month3_emplvl'] / df_joint['lagged_month3_emplvl']) - 1 ) * 100 ),2)
    df_joint['1 Year Employment Growth'] = round((((df_joint['month3_emplvl'] / df_joint['lagged1_month3_emplvl']) - 1 ) * 100 ),2)
    
    #Drop the employment growth values when the industry is not valid due to data suppression
    growth_filter          = (df_joint['emp_growth_invalid'] == 1)
    one_year_growth_filter = (df_joint['one_year_emp_growth_invalid'] == 1)

    df_joint.loc[growth_filter, ['Employment Growth']] = NaN
    df_joint.loc[one_year_growth_filter, ['1 Year Employment Growth']] = NaN


    #Sort by 5 year growth rate
    df_joint = df_joint.sort_values(by=['Employment Growth'])
    df_joint['county'] = county


    #Export final data
    if data_export == True:
        df_joint.to_csv(os.path.join(county_folder,'County Industry Growth Breakdown.csv'))


    return(df_joint)

def GetCountyMedianListPrice(fips, observation_start):
    print('Getting County Median List Price')
    try:
        mlp_series_names      = pd.read_excel(os.path.join(data_location,'FRED Series Names','GeoFRED_Market_Hotness__Median_Listing_Price_by_County_U.S._Dollars.xls'),
                    dtype={'Region Code': object
                          }
                                             )
        mlp_series_names      = mlp_series_names.loc[mlp_series_names['Region Code'] == fips]
        county_series_code    = mlp_series_names['Series ID'].iloc[0]
        county_mlp_df         = fred.get_series(series_id = county_series_code,observation_start = observation_start)
        county_mlp_df         = county_mlp_df.to_frame().reset_index()
        county_mlp_df.columns = ['Period','Median List Price']
    except Exception as e:
        try:
            county_series_code    = 'MEDLISPRI' + fips
            county_mlp_df         = fred.get_series(series_id = county_series_code,observation_start = observation_start)
            county_mlp_df         = county_mlp_df.to_frame().reset_index()
            county_mlp_df.columns = ['Period','Median List Price']
        except Exception as e:
            try:
                county_series_code    = 'MELIPRCOUNTY' + fips
                county_mlp_df         = fred.get_series(series_id = county_series_code,observation_start = observation_start)
                county_mlp_df         = county_mlp_df.to_frame().reset_index()
                county_mlp_df.columns = ['Period','Median List Price']
            except Exception:
                county_series_code    = 'MELIPRCOUNTY' + fips[1:] #Sometimes FRED series names drop leading 0s
                county_mlp_df         = fred.get_series(series_id = county_series_code,observation_start = observation_start)
                county_mlp_df         = county_mlp_df.to_frame().reset_index()
                county_mlp_df.columns = ['Period','Median List Price']


    if data_export == True:
        county_mlp_df.to_csv(os.path.join(county_folder, 'County Median Home List Price.csv'))

    return(county_mlp_df)

def PolygonToShapeFile(poly):
        # WRITE TO SHAPEFILE USING PYSHP
        target_file_path = os.path.join(county_folder_map,'my.shp')
        shapewriter = shapefile.Writer(target=target_file_path)
        shapewriter.field("field1")
        # print('created writer object')

        # step1: convert shapely to pyshp using the function above
        converted_shape = shapely_to_pyshp(poly)
        # print('created converted shape')
        # step2: tell the writer to add the converted shape
        
        shapewriter.shape(converted_shape)
        # add a list of attributes to go along with the shape
        shapewriter.record(["empty record"])
        # save it
        shapewriter.close()

def shapely_to_pyshp(shapelygeom):
    # first convert shapely to geojson
    try:
        shapelytogeojson = shapely.geometry.mapping
    except:
        import shapely.geometry
        shapelytogeojson = shapely.geometry.mapping
    geoj = shapelytogeojson(shapelygeom)
    # create empty pyshp shape
    record = shapefile.Shape()
    # set shapetype
    if geoj["type"] == "Null":
        pyshptype = 0
    elif geoj["type"] == "Point":
        pyshptype = 1
    elif geoj["type"] == "LineString":
        pyshptype = 3
    elif geoj["type"] == "Polygon":
        pyshptype = 5
    elif geoj["type"] == "MultiPoint":
        pyshptype = 8
    elif geoj["type"] == "MultiLineString":
        pyshptype = 3
    elif geoj["type"] == "MultiPolygon":
        pyshptype = 5
    record.shapeType = pyshptype
    # set points and parts
    if geoj["type"] == "Point":
        record.points = geoj["coordinates"]
        record.parts = [0]
    elif geoj["type"] in ("MultiPoint","Linestring"):
        record.points = geoj["coordinates"]
        record.parts = [0]
    elif geoj["type"] in ("Polygon"):
        record.points = geoj["coordinates"][0]
        record.parts = [0]
    elif geoj["type"] in ("MultiPolygon","MultiLineString"):
        index = 0
        points = []
        parts = []
        for eachmulti in geoj["coordinates"]:
            points.extend(eachmulti[0])
            parts.append(index)
            index += len(eachmulti[0])
        record.points = points
        record.parts = parts
    return (record)

def GetCountyShape(fips):
    global county_shape_polygon
    try:
        shapefile_location = os.path.join(data_location,'cb_2018_us_county_500k','cb_2018_us_county_500k.shp')
        assert os.path.exists(shapefile_location)

        #Open the shapefile
        county_map = shapefile.Reader(shapefile_location)
        
        #Loop through each place in the shape file
        for i in range(len(county_map)):
            county_record = county_map.shapeRecord(i)
            #Look for the record that corresponds to our subject city
            if (county_record.record['STATEFP']) + (county_record.record['COUNTYFP'])  != fips:
                continue
            
            else:
                county_shape        =  county_map.shape(i)
                county_shape_polygon = Polygon(county_shape.points)
                try:
                    PolygonToShapeFile(poly = county_shape_polygon)
                    print('Successfully created polygon object from county shape')
                except Exception as e:
                    print(e,'unable to export county polygon as shape')


                return(county_shape) 
    except Exception as e:
        print(e,'unable to get county shape')

def FindAirport():
    #Specify the file path to the airports shape file
    airport_map_location = os.path.join(general_data_location,'Geographic Data','Airports','Airports.shp')
    
    #Open the shapefile
    airport_map = shapefile.Reader(airport_map_location)
    
    try:
        airports_in_city_index_list = [] #Create empty list that we will fill with numbers that correspond to airports within the subject area
        
        #Find any airports inside the confines of the city
        for i in range(len(airport_map)):
            airport_coords        =  Point(airport_map.shape(i).points[0][0],airport_map.shape(i).points[0][1])
            
            if county_shape_polygon.contains(airport_coords):
                airports_in_city_index_list.append(i)

        airport_info_list = []    
        for airport_index in airports_in_city_index_list:     
            airport_record        = airport_map.shapeRecord(airport_index)
            airport_name          = airport_record.record['Fac_Name']
            airport_type          = airport_record.record['Fac_Type']
            airport_dict          = {'name':airport_name,'type':airport_type}
            
        #Don't consider heliports, seaplane bases etc
            if airport_record.record['Fac_Type'] != 'AIRPORT':
                continue

            airport_info_list.append(airport_dict)

        if airport_info_list == []:
            return(None)
        airport_sentence = (county + ' is served by the following facilities: ')

        for count,airport in enumerate(airport_info_list):
            if count < len(airport_info_list) -1 :
                airport_sentence = airport_sentence + (airport['name'].title()) + ' ('  + (airport['type'].title())   + '), ' 
            else:
                airport_sentence = airport_sentence + 'and ' + (airport['name'].title()) + ' ('  + (airport['type'].title())   + ').' 

        return(airport_sentence)
    except Exception as e:
        print(e,'Unable to locate airport inside the county area')
        return(None)

def FindNearestHighways():
    
    try:
        #Specify the file path to the  shape file
        road_map_location = os.path.join(general_data_location,'Geographic Data','North_American_Roads','North_American_Roads.shp')

        #Open the shapefile
        road_map = shapefile.Reader(road_map_location)
        highways_in_city_index_list = [] #Create empty list that we will fill with numbers that correspond to airports within the subject area
        
        #Find any airports inside the confines of the city
        for i in range(len(road_map)):
            highway_coords        =  LineString(road_map.shape(i).points)           
            if county_shape_polygon.contains(highway_coords):
                highways_in_city_index_list.append(i)
        i = 0
        highway_info_list = []    
        for highway_index in highways_in_city_index_list:     
            highway_record        = road_map.shapeRecord(highway_index)
            highway_name          = highway_record.record['ROADNAME']
            
            #Clean up abbreviations
            highway_name          = highway_name.replace('Hwy','Highway',1)
            highway_name          = highway_name.replace('Pkwy','Parkway',1) 


            #Don't add unnamed highways to our list
            if highway_name == '':
                continue
            #If not the first highway
            if i > 0:
                #Check against the existing highways and make sure it's not a duplicate
                for d in highway_info_list:
                    existingname = d['name']
                    if highway_name == existingname:
                        repeat = 1
                        break
                    repeat = 0
                if repeat == 1:
                    continue
                    
           
            
            highway_type          = highway_record.record['ADMIN']
            highway_dict          = {'name':highway_name,'type':highway_type}
            highway_info_list.append(highway_dict)
            i+=1



        sentence = (county + ' is served by the following roads: ')

        for count,highway in enumerate(highway_info_list):
            if count < len(highway_info_list) -1 :
                sentence = sentence + (highway['name'].title()) + ' ('  + (highway['type'].title())   + '), ' 
            else:
                sentence = sentence + 'and ' + (highway['name'].title()) + ' ('  + (highway['type'].title())   + ').' 

        return(sentence)
    except Exception as e:
        print(e,'Unable to locate major roads inside the county area')
        return(None)
 
def GetCountyData():
    print('Getting County Data')
    global county_gdp, county_pci
    global county_unemployment_rate,county_employment
    global county_resident_pop,county_industry_breakdown,county_industry_growth_breakdown
    global county_mlp
    global county_shape

    #County GDP
    try:
        county_gdp                    = GetCountyGDP(fips = fips,observation_start = observation_start_less1)
    except Exception as e:
        print(e,' - Unable to Get County GDP Data')
        county_gdp                    = ''
    
    #County Unemployment Rate
    try:
        county_unemployment_rate      = GetCountyUnemploymentRate(fips = fips,start_year=start_year,end_year=end_year)
    except Exception as e:
        print(e,' Unable to Get County Unemployment Rate Data')
        county_unemployment_rate      = ''
    
    #County Total Employment
    try:
        county_employment             = GetCountyEmployment(fips = fips,start_year=start_year,end_year=end_year)
    except Exception as e:
        print(e,' Unable to Get County Employment Data')
        county_employment             = ''

    #County Per Capita Income
    try:
        county_pci                    = GetCountyPCI(fips=fips, observation_start=observation_start_less1)
    except Exception as e:
        print(e,' Unable to Get County Per Capita Income Data')
        county_pci                    = ''
    
    #County Population 
    try:
        county_resident_pop           = GetCountyResidentPopulation(fips = fips,observation_start=('01/01/' + str(end_year -12)))
    except Exception as e:
        print(e,' Unable to Get County Population Data')
        county_resident_pop           = ''
    
    #County Industry Breakdown
    try:    
        county_industry_breakdown     = GetCountyIndustryBreakdown(fips=fips,year=qcew_year,qtr=qcew_qtr)
    except Exception as e:
        print(e, ' Unable to get County Industry Breakdown')
        county_industry_breakdown     = ''
    
    #County Industry Growth Breakdown
    try:    
        county_industry_growth_breakdown     = GetCountyIndustryGrowthBreakdown(fips=fips,year=qcew_year,qtr=qcew_qtr)
    except Exception as e:
        print(e, ' Unable to get County Industry growth Breakdown')
        county_industry_growth_breakdown     = ''

    #Get Median List Price
    try:  
        county_mlp                    = GetCountyMedianListPrice(fips = fips,observation_start = observation_start)
    except Exception as e: 
        print(e,' No median home list data price available')
        county_mlp                    = ''
    
    #Get County Shape
    county_shape = GetCountyShape(fips = fips)

#MSA Data
def GetMSAGDP(cbsa, observation_start):
    print('Getting MSA GDP')
    msa_gdp_series_code = 'RGMP' + cbsa
    msa_gdp_df          = fred.get_series(series_id = msa_gdp_series_code,observation_start = observation_start)
    msa_gdp_df          = msa_gdp_df.to_frame().reset_index()
    msa_gdp_df.columns  = ['Period','GDP']
    msa_gdp_df['GDP']   = msa_gdp_df['GDP'] * 1000000
    
    if data_export == True:
        msa_gdp_df.to_csv(os.path.join(county_folder,'MSA GDP.csv'))

    return(msa_gdp_df)

def GetMSAResidentPopulation(cbsa, observation_start):
    print('Getting MSA Population')
    try:
        #Resident Population 
        resident_population_series_names                = pd.read_excel(os.path.join(data_location,'FRED Series Names','GeoFRED_Resident_Population_by_Metropolitan_Statistical_Area_Thousands_of_Persons.xls'),
                    dtype={'Region Code': object
                        })
        resident_population_series_names['Region Code'] = resident_population_series_names['Region Code'].astype(str)
        resident_population_series_names                = resident_population_series_names.loc[resident_population_series_names['Region Code'] == cbsa]
        msa_pop_series_code                             = resident_population_series_names['Series ID'].iloc[0]
    except:
        msa_pop_series_code                             = input('enter the FRED series code for resident population for MSA')

    
    msa_pop_df                        = fred.get_series(series_id = msa_pop_series_code, observation_start = observation_start)
    msa_pop_df                        = msa_pop_df.to_frame().reset_index()
    msa_pop_df.columns                = ['Period','Resident Population']
    msa_pop_df['Resident Population'] = msa_pop_df['Resident Population'] * 1000

    if data_export == True:
        msa_pop_df.to_csv(os.path.join(county_folder,'MSA Resident Population.csv'))
    
    return(msa_pop_df)

def GetMSAPCI(cbsa, observation_start):
    print('Getting MSA Per Capita Personal Income')
    #Per Capita Personal Income
    pci_series_names                = pd.read_excel(os.path.join(data_location,'FRED Series Names','GeoFRED_Per_Capita_Personal_Income_by_Metropolitan_Statistical_Area_Dollars.xls'),
                dtype={'Region Code': object
                      })
    pci_series_names['Region Code'] = pci_series_names['Region Code'].astype(str)
    pci_series_names                = pci_series_names.loc[pci_series_names['Region Code'] == cbsa]


    if len(pci_series_names) == 1: #if the cbsa has a series
        msa_pci_series_code = pci_series_names['Series ID'].iloc[0]
    else:
        msa_pci_series_code = 'RPIPC' + cbsa #use a slightly different seires for the New England MSAs


    msa_pci_df         = fred.get_series(series_id = msa_pci_series_code, observation_start=observation_start)
    msa_pci_df         = msa_pci_df.to_frame().reset_index()
    msa_pci_df.columns = ['Period', 'Per Capita Personal Income']

    if data_export == True:
        msa_pci_df.to_csv(os.path.join(county_folder,'MSA Per Capita Personal Income.csv'))

    return(msa_pci_df)

def GetMSAUnemploymentRate(cbsa, start_year, end_year): 
    print('Getting MSA UR')
    #Seasonally-adjusted unemployment rate

    try:
        #For Non-New England States
        if state not in new_england_states:
            series_name = 'LAUMT' + cbsa_main_state_fips + cbsa + '00000003'
        else:
            series_name = 'LAUMT' + cbsa_main_state_fips + necta_code + '00000003'
        

        msa_ur_df           = bls.series(series_name,start_year=start_year, end_year= end_year) 
        msa_ur_df['year']   = msa_ur_df['year'].astype(str)
        msa_ur_df['period'] = msa_ur_df['period'].str[1:3] + '/' +  msa_ur_df['year'].str[2:4]
        msa_ur_df           = msa_ur_df.rename(columns={series_name: "unemployment_rate"})

        if data_export == True:
            msa_ur_df.to_csv(os.path.join(county_folder,'MSA Unemployment Rate.csv'))
       
        return(msa_ur_df)

    #Sometiems they don't use the primary state's fips in the series
    except Exception as e:
        print(e)
        for state_fips_code in cbsa_all_state_fips:
            try:
                #For Non-New England States
                if state not in new_england_states:
                    series_name = 'LAUMT' + state_fips_code + cbsa + '00000003'
                else:
                    series_name = 'LAUMT' + state_fips_code + necta_code + '00000003'
                

                msa_ur_df = bls.series(series_name,start_year=start_year,end_year= end_year) 
                msa_ur_df['year']   = msa_ur_df['year'].astype(str)
                msa_ur_df['period'] =    msa_ur_df['period'].str[1:3] + '/' +  msa_ur_df['year'].str[2:4]
                msa_ur_df = msa_ur_df.rename(columns={series_name: "unemployment_rate"})

                if data_export == True:
                    msa_ur_df.to_csv(os.path.join(county_folder,'MSA Unemployment Rate.csv'))
                return(msa_ur_df)
            except Exception as e:
                print(e)

def GetMSAEmployment(cbsa, start_year, end_year): 
    print('Getting MSA Total Employment')
    #Total Employment
    try:
        if state not in new_england_states:
            series_name = 'LAUMT' + cbsa_main_state_fips + cbsa + '00000005'
        else:
            series_name = 'LAUMT' + cbsa_main_state_fips + necta_code + '00000005'

        msa_emp_df                            = bls.series(series_name,start_year=(start_year-1), end_year= end_year) 
        msa_emp_df['year']                    = msa_emp_df['year'].astype(str)
        msa_emp_df['period']                  = msa_emp_df['period'].str[1:3] + '/' +  msa_emp_df['year'].str[2:4]      
        msa_emp_df                            = msa_emp_df.rename(columns={series_name: "Employment"})
        msa_emp_df['Lagged Employment']       = msa_emp_df['Employment'].shift(12)
        msa_emp_df['Employment Growth']       = round(((msa_emp_df['Employment']/msa_emp_df['Lagged Employment']) - 1 ) * 100,2 )


        if data_export == True:
            msa_emp_df.to_csv(os.path.join(county_folder,'MSA Total Employment.csv'))
        
        return(msa_emp_df)

    except Exception as e:
        print(e)
        for state_fips_code in cbsa_all_state_fips:
            try:
                if state not in new_england_states:
                    series_name = 'LAUMT' + state_fips_code + cbsa + '00000005'
                else:
                    series_name = 'LAUMT' + state_fips_code + necta_code + '00000005'

                msa_emp_df                            = bls.series(series_name,start_year=(start_year-1), end_year= end_year) 
                msa_emp_df['year']                    = msa_emp_df['year'].astype(str)
                msa_emp_df['period']                  = msa_emp_df['period'].str[1:3] + '/' +  msa_emp_df['year'].str[2:4]      
                msa_emp_df                            = msa_emp_df.rename(columns={series_name: "Employment"})
                msa_emp_df['Lagged Employment']       = msa_emp_df['Employment'].shift(12)
                msa_emp_df['Employment Growth']       = round(((msa_emp_df['Employment']/msa_emp_df['Lagged Employment']) - 1 ) * 100,2 )

                #Drop the extra year we needed to calculate growth rates
                msa_emp_df                            = msa_emp_df.loc[msa_emp_df['year'] != str(start_year-1)]

                if data_export == True:
                    msa_emp_df.to_csv(os.path.join(county_folder,'MSA Total Employment.csv'))
                
                return(msa_emp_df)
            except Exception as e:
                pass

def GetMSAMedianListPrice(cbsa, observation_start):
    print('Getting MSA MLP')
    msa_mlp_series_code = 'MEDLISPRI' + cbsa
    msa_mlp_df          = fred.get_series(series_id = msa_mlp_series_code,observation_start = observation_start)
    msa_mlp_df          = msa_mlp_df.to_frame().reset_index()
    msa_mlp_df.columns  = ['Period','Median List Price']

    if data_export == True:
        msa_mlp_df.to_csv(os.path.join(county_folder,'MSA Median Home List Price.csv'))
        
    return(msa_mlp_df)

def GetMSAIndustryBreakdown(cbsa, year, qtr):
    print('Getting MSA QCEW Employment Breakdown')

    #Pulls employment data from Quarterly Census of Employment and Wages
    series_code      = ('C' + cbsa[0:4])
    df_qcew          = qcew.get_data('area', rtype='dataframe', year=year, qtr=qtr, area = series_code )
    
    if data_export == True:
        df_qcew.to_csv(os.path.join(county_folder,'msa_qcew_raw.csv'))

    #Restrict to county-ownership level (fed,state,local,private), supersector employment
    df_qcew          = df_qcew.loc[df_qcew['agglvl_code'] == 43] 
    
    #Drop suppresed employment rows
    df_qcew          = df_qcew.loc[df_qcew['disclosure_code'] != 'N'] 

    #Drop the rows where employment is 0 
    df_qcew          = df_qcew.loc[(df_qcew['month3_emplvl'] > 0) ] 


    #Create a seperate dataframe with just the weekly wages by industry
    wtavg = lambda x: np.average(x.avg_wkly_wage, weights = x.month3_emplvl,axis = 0) #define function to calcuate weighted average wage
    df_qcew_wages           = df_qcew.groupby('industry_code').apply(wtavg).reset_index()
    df_qcew_wages.columns = ['industry_code','avg_wkly_wage']

    #Create a seperate dataframe with just the location quotient by industry (averaged across sectors)
    wtavg = lambda x: np.average(x.lq_month3_emplvl, weights = x.month3_emplvl,axis = 0) #define function to calcuate weighted average wage
    df_qcew_lq           = df_qcew.groupby('industry_code').apply(wtavg).reset_index()
    df_qcew_lq.columns = ['industry_code','lq_month3_emplvl']

    #Collapse down to total employment across the 3 ownership codes
    df_qcew                 = df_qcew.groupby('industry_code').agg(month3_emplvl=('month3_emplvl', 'sum'),)
    
    #Merge in the wage and location quotient dataframes
    df_qcew                 = pd.merge(df_qcew, df_qcew_wages, on=('industry_code'),how='outer')
    df_qcew                 = pd.merge(df_qcew, df_qcew_lq, on=('industry_code'),how='outer')

    #Change the industry codes to names
    replacements = {'1011':'Natural Resources & Mining', 
                    '1012':'Construction', 
                    '1013':'Manufacturing', 
                    '1021':'Trade, Transportation, & Utilities', 
                    '1022':'Information', 
                    '1023':'Financial Activities', 
                    '1024':'Professional & Business Services', 
                    '1025':'Education & Health Services', 
                    '1026':'Leisure & Hospitality', 
                    '1027':'Other Services', 
                    '1028':'Public Administration', 
                    '1029':'Unclassified'}

    df_qcew['industry_code'].replace(replacements, inplace=True)

   
    #Sort by total employement
    df_qcew['employment_fraction'] = round(((df_qcew['month3_emplvl']/(df_qcew['month3_emplvl'].sum())) * 100),2)
    df_qcew['msa'] = cbsa_name
    df_qcew      = df_qcew.loc[df_qcew['industry_code'] != 'Unclassified']
    df_qcew = df_qcew.sort_values(by=['month3_emplvl'])

    #Export final data
    if data_export == True:
        df_qcew.to_csv(os.path.join(county_folder,'MSA Industry Breakdown.csv'))
    return(df_qcew)

def GetMSAIndustryGrowthBreakdown(cbsa, year, qtr):
    print('Getting MSA QCEW Employment Growth Breakdown')

    #Pulls employment data (and the lagged data) from Quarterly Census of Employment and Wages
    series_code      = ('C' + cbsa[0:4])
    df_qcew          = qcew.get_data('area', rtype='dataframe', year=year,qtr=qtr, area=series_code)
    df_qcew_lagged   = qcew.get_data('area', rtype='dataframe', year=(str(int(year) - growth_period )),qtr=qtr, area=series_code)
    df_qcew_lagged1  = qcew.get_data('area', rtype='dataframe', year=(str(int(year) - 1 )),qtr=qtr, area=series_code)




    #Restrict to county-ownership level (fed,state,local,private), supersector employment
    df_qcew          = df_qcew.loc[df_qcew['agglvl_code'] == 43] 
    df_qcew_lagged   = df_qcew_lagged.loc[df_qcew_lagged['agglvl_code'] == 43] 
    df_qcew_lagged1  = df_qcew_lagged1.loc[df_qcew_lagged1['agglvl_code'] == 43] 

    #Restrict to private ownership 
    df_qcew          = df_qcew.loc[df_qcew['own_code'] == 5] 
    df_qcew_lagged   = df_qcew_lagged.loc[df_qcew_lagged['own_code'] == 5] 
    df_qcew_lagged1  = df_qcew_lagged1.loc[df_qcew_lagged1['own_code'] == 5] 

    
    if data_export == True:
        df_qcew.to_csv(os.path.join(county_folder,'msa_qcew_raw.csv'))
        df_qcew_lagged.to_csv(os.path.join(county_folder,'msa_qcew_raw_lagged.csv'))
        df_qcew_lagged1.to_csv(os.path.join(county_folder,'msa_qcew_raw_lagged1.csv'))

    #Add "lagged" and "lagged1" to the column names for the lagged data
    df_qcew_lagged   = df_qcew_lagged.add_prefix('lagged_')
    df_qcew_lagged1  = df_qcew_lagged1.add_prefix('lagged1_')

    #Remove the "lagged" and "lagged1" prefix for the industry and ownership code columns so we can merge on them
    df_qcew_lagged   = df_qcew_lagged.rename(columns={"lagged_own_code": "own_code", "lagged_industry_code": "industry_code"})
    df_qcew_lagged1  = df_qcew_lagged1.rename(columns={"lagged1_own_code": "own_code", "lagged1_industry_code": "industry_code"})


    #Merge together the current quarters data with the data from 1 year ago and with the data from (5) years from now
    df_joint = pd.merge(df_qcew, df_qcew_lagged, on=('industry_code','own_code'),how='outer')
    df_joint = pd.merge(df_joint, df_qcew_lagged1, on=('industry_code','own_code'),how='outer') #now merge in lagged employment data

    #Flag the industries and ownership type rows where the data was suppresed in the past or present
    filter = (df_joint['disclosure_code'] == 'N') | (df_joint['lagged_disclosure_code'] == 'N')
    df_joint['Employment Growth Invalid'] = ''
    df_joint.loc[filter, ['Employment Growth Invalid']] = 1
    df_joint.loc[df_joint['Employment Growth Invalid'] != 1, ['Employment Growth Invalid']] = 0
    
    one_year_filter = (df_joint['disclosure_code'] == 'N') | (df_joint['lagged1_disclosure_code'] == 'N' )
    df_joint['1Y Employment Growth Invalid'] = ''
    df_joint.loc[one_year_filter, ['1Y Employment Growth Invalid']] = 1
    df_joint.loc[df_joint['1Y Employment Growth Invalid'] != 1, ['1Y Employment Growth Invalid']] = 0
    

    #Replace the Employment Growth Invalid column with the maximum value from each row for a given industry
    df_joint['Employment Growth Invalid'] = df_joint.groupby('industry_code')['Employment Growth Invalid'].transform('max')
    df_joint['1Y Employment Growth Invalid'] = df_joint.groupby('industry_code')['1Y Employment Growth Invalid'].transform('max')

    #Drop the rows where employment is 0 
    df_joint          = df_joint.loc[(df_joint['month3_emplvl'] > 0) ] 

    #Create a seperate dataframe with just the current quarters weekly wages by industry
    wtavg = lambda x: np.average(x.avg_wkly_wage, weights = x.month3_emplvl,axis = 0) #define function to calcuate weighted average wage
    df_qcew_wages           = df_joint.groupby('industry_code').apply(wtavg).reset_index()
    df_qcew_wages.columns = ['industry_code','avg_wkly_wage']

    #Collapse down to total employment across the 3 ownership codes
    df_joint                 = df_joint.groupby('industry_code').agg(month3_emplvl=('month3_emplvl', 'sum'),lagged_month3_emplvl=('lagged_month3_emplvl', 'sum'),lagged1_month3_emplvl=('lagged1_month3_emplvl', 'sum'),emp_growth_invalid=('Employment Growth Invalid', 'max'),one_year_emp_growth_invalid=('1Y Employment Growth Invalid', 'max'))
    df_joint                 = pd.merge(df_joint, df_qcew_wages, on=('industry_code'),how='outer')

    #Change the industry codes to names
    replacements = {'1011':'Natural Resources & Mining', 
                    '1012':'Construction', 
                    '1013':'Manufacturing', 
                    '1021':'Trade, Transportation, & Utilities', 
                    '1022':'Information', 
                    '1023':'Financial Activities', 
                    '1024':'Professional & Business Services', 
                    '1025':'Education & Health Services', 
                    '1026':'Leisure & Hospitality', 
                    '1027':'Other Services', 
                    '1028':'Public Administration', 
                    '1029':'Unclassified'}

    df_joint['industry_code'].replace(replacements, inplace=True)

   

    #Calcualte employment growth rates
    df_joint['Employment Growth'] = round((((df_joint['month3_emplvl'] / df_joint['lagged_month3_emplvl']) - 1 ) * 100 ),2)
    df_joint['1 Year Employment Growth'] = round((((df_joint['month3_emplvl'] / df_joint['lagged1_month3_emplvl']) - 1 ) * 100 ),2)
    
    #Drop the employment growth values when the industry is not valid due to data suppression
    growth_filter          = (df_joint['emp_growth_invalid'] == 1)
    one_year_growth_filter = (df_joint['one_year_emp_growth_invalid'] == 1)

    df_joint.loc[growth_filter,          ['Employment Growth']] = NaN
    df_joint.loc[one_year_growth_filter, ['1 Year Employment Growth']] = NaN


    #Sort by 5 year growth rate
    df_joint        = df_joint.sort_values(by=['Employment Growth'])
    df_joint['msa'] = cbsa_name


    #Export final data
    if data_export == True:
        df_joint.to_csv(os.path.join(county_folder,'MSA Industry Growth Breakdown.csv'))


    return(df_joint)

def GetMSAData():
    global msa_gdp
    global msa_pci
    global msa_unemployment_rate, msa_employment, msa_unemployment
    global msa_resident_pop
    global msa_mlp
    global msa_industry_breakdown, msa_industry_growth_breakdown
    #We create these blank variables so we can use them as function inputs for the graph functions when there is no MSA
    if cbsa == '':
            msa_gdp                         = ''
            msa_pci                         = ''
            msa_unemployment_rate           = ''
            msa_unemployment                = ''
            msa_employment                  = ''
            msa_resident_pop                = ''
            msa_mlp                         = ''
            msa_industry_breakdown          = ''
            msa_industry_growth_breakdown   = ''
    else:
        print('Getting MSA Data')
        msa_gdp                         = GetMSAGDP(cbsa = cbsa, observation_start = observation_start_less1)
        msa_pci                         = GetMSAPCI(cbsa = cbsa, observation_start = observation_start_less1)
        msa_unemployment_rate           = GetMSAUnemploymentRate(cbsa = cbsa, start_year = start_year, end_year = end_year)
        msa_employment                  = GetMSAEmployment(cbsa = cbsa, start_year = start_year, end_year = end_year)
        msa_resident_pop                = GetMSAResidentPopulation(cbsa = cbsa, observation_start = ('01/01/' + str(end_year -12)))

        #Median list price
        try:
            msa_mlp                     = GetMSAMedianListPrice(cbsa = cbsa,observation_start=observation_start)

        except:
            msa_mlp                     = ''


        
        msa_industry_breakdown            = GetMSAIndustryBreakdown(      cbsa = cbsa, year = qcew_year, qtr = qcew_qtr)    
        msa_industry_growth_breakdown     = GetMSAIndustryGrowthBreakdown(cbsa = cbsa, year = qcew_year, qtr = qcew_qtr)

#State Data
def GetStateGDP(state, observation_start):
    print('Getting State GDP')
    state_gdp_series_code = state + 'RGSP'
    state_gdp_df          = fred.get_series(series_id = state_gdp_series_code,observation_start = observation_start)
    state_gdp_df          = state_gdp_df.to_frame().reset_index()
    state_gdp_df.columns  = ['Period','GDP']
    state_gdp_df['GDP']   = state_gdp_df['GDP'] * 1000000

    if data_export == True:
        state_gdp_df.to_csv(os.path.join(county_folder,'State GDP.csv'))
    
    return(state_gdp_df)

def GetStatePCI(state, observation_start):
    print('Getting State Per Capita Personal Income')
    #Per Capita Personal Income
    state_pci_series_code = state + 'PCPI' 
    state_pci_df          = fred.get_series(series_id = state_pci_series_code, observation_start = observation_start)
    state_pci_df          = state_pci_df.to_frame().reset_index()
    state_pci_df.columns  = ['Period','Per Capita Personal Income']

    if data_export == True:
        state_pci_df.to_csv(os.path.join(county_folder, 'State Per Capita Personal Income.csv'))
    
    return(state_pci_df)

def GetStateResidentPopulation(state, observation_start):
    print('Getting State Population')
    #Resident Population 
    state_pop_series_code               = state + 'POP'
    state_pop_df                        = fred.get_series(series_id = state_pop_series_code, observation_start = observation_start)
    state_pop_df                        = state_pop_df.to_frame().reset_index()
    state_pop_df.columns                = ['Period', 'Resident Population']
    state_pop_df['Resident Population'] = state_pop_df['Resident Population'] * 1000

    if data_export == True:
        state_pop_df.to_csv(os.path.join(county_folder,'State Resident Population.csv'))
    
    return(state_pop_df)

def GetStateUnemploymentRate(fips, start_year, end_year): 
    print('Getting State Unemployment Rate')
    #Seasonally-adjusted unemployment rate
    series_name             = 'LASST' + fips[0:2] + '0000000000003'
    state_ur_df             = bls.series(series_name,start_year=start_year, end_year= end_year)

    state_ur_df['year']     = state_ur_df['year'].astype(str)
    state_ur_df['period']   = state_ur_df['period'].str[1:3] + '/' +  state_ur_df['year'].str[2:4]      
    state_ur_df             = state_ur_df.rename(columns={series_name: "unemployment_rate"})

    if data_export == True:
        state_ur_df.to_csv(os.path.join(county_folder,'State Unemployment Rate.csv'))
    
    return(state_ur_df)

def GetStateEmployment(fips, start_year, end_year): 
    print('Getting State Total Employment')
    #Total Employment
    series_name                             = 'LASST' + fips[0:2] + '0000000000005'
    state_emp_df                            = bls.series(series_name,start_year=(start_year-1), end_year= end_year) 
    state_emp_df['year']                    = state_emp_df['year'].astype(str)
    state_emp_df['period']                  = state_emp_df['period'].str[1:3] + '/' +  state_emp_df['year'].str[2:4]      
    state_emp_df                            = state_emp_df.rename(columns={series_name: "Employment"})
    state_emp_df['Lagged Employment']       = state_emp_df['Employment'].shift(12)
    state_emp_df['Employment Growth']       = round(((state_emp_df['Employment']/state_emp_df['Lagged Employment']) - 1 ) * 100,2 )
    
    if data_export == True:
        state_emp_df.to_csv(os.path.join(county_folder, 'State Total Employment.csv'))

    return(state_emp_df)

def GetStateData():
    print('Getting State Data')
    global state_gdp
    global state_pci
    global state_unemployment_rate, state_employment
    global state_resident_pop
    
    state_gdp                        = GetStateGDP(state = state, observation_start = observation_start_less1)
    state_unemployment_rate          = GetStateUnemploymentRate(fips = fips, start_year=start_year,end_year = end_year)
    state_employment                 = GetStateEmployment(fips = fips, start_year=start_year,end_year = end_year)
    state_pci                        = GetStatePCI(state = state, observation_start = observation_start_less1)
    state_resident_pop               = GetStateResidentPopulation(state = state, observation_start = ('01/01/' + str(end_year -12)))

#National Data
def GetNationalPCI(observation_start):
    print('Getting National Per Capita Personal Income')
    #Per Capita Personal Income
    usa_pci_series_code = 'A792RC0Q052SBEA' 
    usa_pci_df          = fred.get_series(series_id = usa_pci_series_code,observation_start = observation_start,frequency = 'a')
    usa_pci_df          = usa_pci_df.to_frame().reset_index()
    usa_pci_df.columns  = ['Period','Per Capita Personal Income']
    usa_pci_df          = usa_pci_df.loc[usa_pci_df['Per Capita Personal Income'] >= 0]
   

    if data_export == True:
        usa_pci_df.to_csv(os.path.join(county_folder,'National Per Capita Personal Income.csv'))

    return(usa_pci_df)

def GetNationalResidentPopulation(observation_start):
    print('Getting National Population')
    #Resident Population 
    usa_pop_series_code               = 'POP'
    usa_pop_df                        = fred.get_series(series_id = usa_pop_series_code, observation_start = observation_start,frequency = 'a')
    usa_pop_df                        = usa_pop_df.to_frame().reset_index()
    usa_pop_df.columns                = ['Period','Resident Population']
    usa_pop_df['Resident Population'] = usa_pop_df['Resident Population'] * 1000

    if data_export == True:
        usa_pop_df.to_csv(os.path.join(county_folder,'National Population.csv'))
    
    return(usa_pop_df)

def GetNationalMedianListPrice(observation_start):
    print('Getting National Median List Price')
    usa_mlp_series_code = 'MEDLISPRIUS'
    usa_mlp_df          = fred.get_series(series_id = usa_mlp_series_code, observation_start = observation_start)
    usa_mlp_df          = usa_mlp_df.to_frame().reset_index()
    usa_mlp_df.columns  = ['Period', 'Median List Price']

    if data_export == True:
        usa_mlp_df.to_csv(os.path.join(county_folder,'National Median Home List Price.csv'))
    
    return(usa_mlp_df)

def GetNationalUnemploymentRate(start_year, end_year):
    print('Getting National Unemployment Rate')
    
    #Seasonally-adjusted unemployment rate
    series_name              = 'LNS14000000'
    national_ur_df           = bls.series(series_name,start_year=start_year,end_year=end_year) 
    national_ur_df['year']   = national_ur_df['year'].astype(str)
    national_ur_df['period'] = national_ur_df['period'].str[1:3] + '/' +  national_ur_df['year'].str[2:4]      
    national_ur_df           = national_ur_df.rename(columns={series_name: "unemployment_rate"})

    if data_export == True:
        national_ur_df.to_csv(os.path.join(county_folder,'National Unemployment Rate.csv'))
    
    return(national_ur_df)

def GetNationalEmployment(start_year, end_year):
    print('Getting National Total Employment')
    #Total Employment
    series_name = 'LNS12000000'
    national_emp_df = bls.series(series_name,start_year=(start_year-1), end_year= end_year)

    national_emp_df['year']   =    national_emp_df['year'].astype(str)
    national_emp_df['period'] =    national_emp_df['period'].str[1:3] + '/' +  national_emp_df['year'].str[2:4] 

    national_emp_df = national_emp_df.rename(columns={series_name: "Employment"})

    national_emp_df['Lagged Employment']       = national_emp_df['Employment'].shift(12)
    national_emp_df['Employment Growth']       =  round(((national_emp_df['Employment']/national_emp_df['Lagged Employment']) - 1 ) * 100,2 )


    if data_export == True:
        national_emp_df.to_csv(os.path.join(county_folder,'National Total Employment.csv'))
    return(national_emp_df)

def GetNationalGDP(observation_start):
    print('Getting National GDP')
    national_gdp_series_code = 'GDP'
    national_gdp_df          = fred.get_series(series_id = national_gdp_series_code,observation_start = observation_start,frequency = 'q')
    national_gdp_df          = national_gdp_df.to_frame().reset_index()
    national_gdp_df.columns  = ['Period','GDP']
    national_gdp_df['GDP']   = national_gdp_df['GDP'] * 1000000000
    
    if data_export == True:
        national_gdp_df.to_csv(os.path.join(county_folder,'National GDP.csv'))
    
    return(national_gdp_df)

def GetNationalData():
    print('Getting National Data')
    global national_pci
    global national_resident_pop
    global national_mlp
    global national_unemployment
    global national_employment
    global national_gdp
    national_pci                       = GetNationalPCI(observation_start = observation_start_less1)
    national_resident_pop              = GetNationalResidentPopulation(observation_start = ('01/01/' + str(end_year - 12)))
    national_mlp                       = GetNationalMedianListPrice(observation_start = observation_start)
    national_unemployment              = GetNationalUnemploymentRate(start_year = start_year, end_year=end_year)
    national_employment                = GetNationalEmployment(start_year = start_year, end_year=end_year)
    national_gdp                       = GetNationalGDP(observation_start = observation_start_less1)

#####################################################Graph Related Functions####################################

def CreateUnemploymentRateEmploymentGrowthGraph(folder):
    print('Creating Unemployment Rate and Employment Growth Graph')
    fig = make_subplots(rows=1, cols=2,subplot_titles=("Unemployment Rate", "Annual Employment Growth"),horizontal_spacing = horizontal_spacing)

    #County unemployment rate
    fig.add_trace(
            go.Scatter(x = county_unemployment_rate['period'].iloc[-60:],
                    y = county_unemployment_rate['unemployment_rate'].iloc[-60:],
                    name = county,
                    line = dict(color = bowery_dark_blue, width = 1,dash = 'dash')
                    ),
            secondary_y = False,
            row         = 1,
            col         = 1
                 )

    #MSA unemployment rate if applicable
    if (cbsa != '') and (msa_unemployment_rate.equals(county_unemployment_rate) == False):
        fig.add_trace(
            go.Scatter(x    = msa_unemployment_rate['period'].iloc[-60:],
                       y    = msa_unemployment_rate['unemployment_rate'].iloc[-60:],
                    name    = cbsa_name + ' (MSA)',
                    line    = dict(color = bowery_light_blue, width = 1)
                    ),
            secondary_y=False,row=1, col=1)

    #State unemployment rate
    if state != 'DC':
        fig.add_trace(
            go.Scatter(x = state_unemployment_rate['period'].iloc[-60:],
                       y = state_unemployment_rate['unemployment_rate'].iloc[-60:],
                    name = state_name,
                    line = dict(color = bowery_dark_grey, width = 1)
                    ),
            secondary_y = False,
            row         = 1, 
            col         = 1
                      )

    #County employment growth 
    fig.add_trace(
        go.Scatter(x = county_employment['period'].iloc[-60:],
                y = county_employment['Employment Growth'].iloc[-60:],
                name = county,
                line = dict(color = bowery_dark_blue, width = 1,dash = 'dash'),showlegend=False
                  ),
        secondary_y=False, 
        row=1, 
        col=2,
                 )

    #MSA employment growth
    if (cbsa != '') and (msa_employment.equals(county_employment) == False):
        fig.add_trace(
            go.Scatter(x = msa_employment['period'].iloc[-60:],
                    y    = msa_employment['Employment Growth'].iloc[-60:],
                    name = cbsa_name + ' (MSA)',
                    line = dict(color=bowery_light_blue, width = 1),
                    showlegend=False
                      ),
            secondary_y=False,row=1, col=2,
                      )

    #State employment growth
    if state != 'DC':
        fig.add_trace(
            go.Scatter(x = state_employment['period'].iloc[-60:],
                      y  = state_employment['Employment Growth'].iloc[-60:],
                    name = state_name,
                    line = dict(color = bowery_dark_grey, width = 1),
                    showlegend=False
                      ),
            secondary_y  = False,row=1, col=2,)

    #Set formatting 
    fig.update_layout(

        title={
            'y':       title_position,
            'x':       0.5,
            'xanchor': 'center',
            'yanchor': 'top'
              },

        legend=dict(
            orientation = "h",
            yanchor     = "bottom",
            y           = legend_position + 0.1,
            xanchor     = "center",
            x           = 0.5,
            font_size   = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",    
                    )

    #Add % to left axis ticks
    fig.update_yaxes(
        tickfont    = dict(size=tickfont_size), 
        ticksuffix  = '%',  
        title       = None ,
        tickmode    = 'auto',
        nticks      = 6,
        secondary_y = False
                    )                 
                    
    
    #Set x axis ticks
    quarter_list      = [i for i in range(len(county_unemployment_rate['period']))]
    quarter_list      = quarter_list[::-12]

    quarter_list_text = [period for period in county_unemployment_rate['period']]
    quarter_list_text = quarter_list_text[::-12]

    fig.update_xaxes(
        tickmode = 'array',
        tickvals = quarter_list,
        ticktext = quarter_list_text,
        tickfont = dict(size=tickfont_size),
        tickangle = 0,
                    )

    #Set size
    fig.update_layout(
    autosize  = False,
    height    = graph_height,
    width     = graph_width,
    margin    = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin,pad=0,autoexpand = True),
                     )
    
    fig.write_image(os.path.join(folder,'unemployment_rate_employment_growth.png'), engine='kaleido', scale = scale)

def CreateUnemploymentRateGraph(folder):

    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #County unemployment rate
    fig.add_trace(
        go.Scatter(x=county_unemployment_rate['period'],
                y=county_unemployment_rate['unemployment_rate'],
                name=county,
                line=dict(color = bowery_dark_blue, width = 1,dash = 'dash'))
        ,secondary_y=False)

    #MSA unemployment rate if applicable
    if (cbsa != '') and (msa_unemployment_rate.equals(county_unemployment_rate) == False):
        fig.add_trace(
            go.Scatter(x=msa_unemployment_rate['period'],
                    y=msa_unemployment_rate['unemployment_rate'],
                    name=cbsa_name + ' (MSA)',
                    line=dict(color=bowery_light_blue, width = 1))
            ,secondary_y=False)

    #State unemployment rate
    fig.add_trace(
        go.Scatter(x=state_unemployment_rate['period'],
                y=state_unemployment_rate['unemployment_rate'],
                name=state_name,
                line=dict(color=bowery_dark_grey, width = 1))
        ,secondary_y=False)

    #Set formatting 
    fig.update_layout(
        title_text="Unemployment Rate",    
        
        title = {
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
              },

        legend = dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",    
        autosize      = False,
        height        = graph_height,
        width         = graph_width,
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin,pad=0,autoexpand = True),

                    )

    #Add % to left axis ticks
    fig.update_yaxes(
        tickfont   = dict(size=tickfont_size), 
        ticksuffix = '%',  
        title      = None ,
        tickmode   = 'auto',
        nticks     = 6,
        secondary_y= False)                 
                    
    
    #Set x axis ticks
    quarter_list      = [i for i in range(len(county_unemployment_rate['period']))]
    quarter_list      = quarter_list[::-12]
    quarter_list_text = [period for period in county_unemployment_rate['period']]
    quarter_list_text = quarter_list_text[::-12]

    fig.update_xaxes(
        tickmode = 'array',
        tickvals = quarter_list,
        ticktext = quarter_list_text,
        tickfont = dict(size=tickfont_size),
        tickangle = 0,
                     )

    fig.write_image(os.path.join(folder,'unemployment_rate.png'), engine = 'kaleido', scale = scale)

def CreateEmploymentGrowthGraph(folder):

    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #County unemployment rate
    fig.add_trace(
        go.Scatter(x = county_employment['period'],
                   y = county_employment['Employment Growth'],
                name = county,
                line = dict(color = bowery_dark_blue)),
        secondary_y  = False
                )

    #MSA unemployment rate if applicable
    if cbsa != '':
        fig.add_trace(
            go.Scatter(x = msa_employment['period'],
                       y = msa_employment['Employment Growth'],
                    name = cbsa_name + ' (MSA)',
                    line = dict(color=bowery_light_blue)),
            secondary_y=False)

    #State unemployment rate
    fig.add_trace(
        go.Scatter(x = state_employment['period'],
                y    = state_employment['Employment Growth'],
                name = state_name,
                line = dict(color=bowery_dark_grey),
                ),
        secondary_y=False)

    #Set formatting 
    fig.update_layout(
        title_text="Annual Employment Growth",    
        
        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
              },

        legend = dict(
                orientation="h",
                yanchor="bottom",
                y=legend_position,
                xanchor="center",
                x=0.5,
                font_size = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",   
        
        #Set size
        autosize  = False,
        height    = graph_height,
        width     = graph_width,
        margin    = dict(l=left_margin, r=right_margin, t=top_margin, b = bottom_margin,pad=0, autoexpand = True) 
                    )

    #Add % to left axis ticks
    fig.update_yaxes(
        tickfont    = dict(size=tickfont_size), 
        ticksuffix  = '%',  
        title       = None ,
        tickmode    = 'auto',
        nticks      = 6,
        secondary_y = False)                
   
    
    #Set x axis ticks
    quarter_list = [i for i in range(len(county_employment['period']))]
    quarter_list = quarter_list[::-12]

    quarter_list_text = [period for period in county_employment['period']]
    quarter_list_text = quarter_list_text[::-12]

    fig.update_xaxes(tickmode = 'array',
        tickvals = quarter_list,
        ticktext = quarter_list_text,
        tickfont = dict(size=tickfont_size),
        tickangle = 0,
                    )


    
    fig.write_image(os.path.join(folder,'employment_growth.png'),engine='kaleido',scale=scale)

def CreatePCIGraph(county_data_frame,msa_data_frame,state_data_frame,national_data_frame,folder):
    print("Creating PCI graph")
    fig = make_subplots(specs=[[{"secondary_y": True}, {"secondary_y": False}]],rows=1, cols=2,subplot_titles=("Per Capita Personal Income", "Annualized Income Growth"),horizontal_spacing = horizontal_spacing)

    #Add county PCI
    if (isinstance(county_data_frame, pd.DataFrame) == True):
        fig.add_trace(
            go.Scatter(x = county_data_frame['Period'],
                    y    = county_data_frame['Per Capita Personal Income'],
                    name = county,
                    line = dict(color = bowery_dark_blue,width = 1,dash = 'dash'),
                    showlegend=False),
                    secondary_y=False,
                    row = 1,
                    col = 1)
        

   #Add MSA PCI if applicable
    if (isinstance(msa_data_frame, pd.DataFrame) == True) and (msa_data_frame.equals(county_data_frame) == False):
        fig.add_trace(
            go.Scatter(x = msa_data_frame['Period'],
                    y    = msa_data_frame['Per Capita Personal Income'],
                    name = cbsa_name + ' (MSA)',
                    line = dict(color=bowery_light_blue,width = 1),
                    showlegend=False)
                    ,secondary_y=False, 
                    row = 1,
                    col = 1)
        
    else:
        #Add state PCI
        fig.add_trace(
            go.Scatter(x = state_data_frame['Period'],
                       y = state_data_frame['Per Capita Personal Income'],
                    name = state_name,
                    line = dict(color=bowery_dark_grey),
                    showlegend=False),
            secondary_y=False,
            row = 1,
            col = 1)
    
    #Add Growth Subfigure
    
    #Calculate annualized growth rates for the county, msa (if available), and state dataframes
    if (isinstance(county_data_frame, pd.DataFrame) == True):
        county_data_frame['Per Capita Personal Income_1year_growth'] =  (((county_data_frame['Per Capita Personal Income']/county_data_frame['Per Capita Personal Income'].shift(1))  - 1) * 100)/1
        county_data_frame['Per Capita Personal Income_3year_growth'] =  (((county_data_frame['Per Capita Personal Income']/county_data_frame['Per Capita Personal Income'].shift(3))   - 1) * 100)/3
        county_data_frame['Per Capita Personal Income_5year_growth'] =  (((county_data_frame['Per Capita Personal Income']/county_data_frame['Per Capita Personal Income'].shift(5))   - 1) * 100)/5

        county_1y_growth                                             = county_data_frame.iloc[-1]['Per Capita Personal Income_1year_growth'] 
        county_3y_growth                                             = county_data_frame.iloc[-1]['Per Capita Personal Income_3year_growth'] 
        county_5y_growth                                             = county_data_frame.iloc[-1]['Per Capita Personal Income_5year_growth']
    else:
        county_1y_growth                                             = 0
        county_3y_growth                                             = 0
        county_5y_growth                                             = 0 


    if (isinstance(msa_data_frame, pd.DataFrame) == True):
        #Make sure we are comparing same years for calculating growth rates for county and msa
        if (isinstance(county_data_frame, pd.DataFrame) == True):
            msa_data_frame = msa_data_frame.loc[msa_data_frame['Period'] <= (county_data_frame['Period'].max())]
        msa_data_frame['Per Capita Personal Income_1year_growth'] =  (((msa_data_frame['Per Capita Personal Income']/msa_data_frame['Per Capita Personal Income'].shift(1))  - 1) * 100)/1
        msa_data_frame['Per Capita Personal Income_3year_growth'] =  (((msa_data_frame['Per Capita Personal Income']/msa_data_frame['Per Capita Personal Income'].shift(3))   - 1) * 100)/3
        msa_data_frame['Per Capita Personal Income_5year_growth'] =  (((msa_data_frame['Per Capita Personal Income']/msa_data_frame['Per Capita Personal Income'].shift(5))   - 1) * 100)/5

        msa_1y_growth  = msa_data_frame.iloc[-1]['Per Capita Personal Income_1year_growth'] 
        msa_3y_growth  = msa_data_frame.iloc[-1]['Per Capita Personal Income_3year_growth'] 
        msa_5y_growth  = msa_data_frame.iloc[-1]['Per Capita Personal Income_5year_growth'] 

    #Make sure we are comparing same years for calculating growth rates for county and state
    if (isinstance(county_data_frame, pd.DataFrame) == True):
        state_data_frame = state_data_frame.loc[state_data_frame['Period'] <= (county_data_frame['Period'].max())]
    state_data_frame['Per Capita Personal Income_1year_growth'] =  (((state_data_frame['Per Capita Personal Income']/state_data_frame['Per Capita Personal Income'].shift(1))  - 1) * 100)/1
    state_data_frame['Per Capita Personal Income_3year_growth'] =  (((state_data_frame['Per Capita Personal Income']/state_data_frame['Per Capita Personal Income'].shift(3))   - 1) * 100)/3
    state_data_frame['Per Capita Personal Income_5year_growth'] =  (((state_data_frame['Per Capita Personal Income']/state_data_frame['Per Capita Personal Income'].shift(5))   - 1) * 100)/5

    state_1y_growth                                             = state_data_frame.iloc[-1]['Per Capita Personal Income_1year_growth'] 
    state_3y_growth                                             = state_data_frame.iloc[-1]['Per Capita Personal Income_3year_growth'] 
    state_5y_growth                                             = state_data_frame.iloc[-1]['Per Capita Personal Income_5year_growth'] 

    #Make sure we are comparing same years for calculating growth rates for county and state
    if  (isinstance(county_data_frame, pd.DataFrame) == True):
        national_data_frame = national_data_frame.loc[national_data_frame['Period'] <= (county_data_frame['Period'].max())]
    national_data_frame['Per Capita Personal Income_1year_growth'] =  (((national_data_frame['Per Capita Personal Income']/national_data_frame['Per Capita Personal Income'].shift(1))  - 1) * 100)/1
    national_data_frame['Per Capita Personal Income_3year_growth'] =  (((national_data_frame['Per Capita Personal Income']/national_data_frame['Per Capita Personal Income'].shift(3))   - 1) * 100)/3
    national_data_frame['Per Capita Personal Income_5year_growth'] =  (((national_data_frame['Per Capita Personal Income']/national_data_frame['Per Capita Personal Income'].shift(5))   - 1) * 100)/5

    national_1y_growth                                             = national_data_frame.iloc[-1]['Per Capita Personal Income_1year_growth'] 
    national_3y_growth                                             = national_data_frame.iloc[-1]['Per Capita Personal Income_3year_growth'] 
    national_5y_growth                                             = national_data_frame.iloc[-1]['Per Capita Personal Income_5year_growth'] 
    

    #Now that we've calculated growth rates, create our plot
    years               = ['5 Years', '3 Years', '1 Year']
    annotation_position = 'outside'
    
    
    #MSA PCI is available, but county is not
    if (isinstance(msa_data_frame, pd.DataFrame) == True)  and (isinstance(county_data_frame, pd.DataFrame) == False):
        print('MSA PCI Available, County PCI is NOT')
        #Add National Growth 
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [national_5y_growth, national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [national_5y_growth, national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                   ),
                row = 1,
                col = 2)

        #Add MSA Growth
        fig.add_trace( 
            go.Bar(
                name = cbsa_name + ' (MSA)',  
                x=years, 
                y=[msa_5y_growth, msa_3y_growth, msa_1y_growth],
                marker_color =bowery_light_blue,
                text = [msa_5y_growth, msa_3y_growth, msa_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis =  False
                ),
            row = 1,
            col = 2)


    
    #MSA PCI is unavailable, but county is (or county is equal to msa)
    elif ((isinstance(msa_data_frame, pd.DataFrame) == False)  and (isinstance(county_data_frame, pd.DataFrame) == True)):
        print('MSA PCI Unavailable, County PCI is available')
        #Add National Growth
        fig.add_trace( 
            go.Bar(
                name = 'United States',  
                x=years, 
                y=[national_5y_growth, national_3y_growth, national_1y_growth],
                marker_color ="#000F44",
                text = [national_5y_growth, national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis =  False
                 ),
            row = 1,
            col = 2
                    )

        #Add State Growth
        fig.add_trace(
            go.Bar(
                name         = state_name,  
                x            = years, 
                y            = [state_5y_growth, state_3y_growth, state_1y_growth],
                marker_color = bowery_dark_grey,
                text         = [state_5y_growth, state_3y_growth, state_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                  ),
            row = 1,
            col = 2)

        #Add County Growth
        fig.add_trace(
            go.Bar(
                name         = county,      
                x            = years, 
                y            = [county_5y_growth,county_3y_growth,county_1y_growth],
                marker_color = bowery_dark_blue,
                text         = [county_5y_growth,county_3y_growth,county_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                ),
            row = 1,
            col = 2)


    #MSA and County are available
    elif (isinstance(msa_data_frame, pd.DataFrame) == True)  and (isinstance(county_data_frame, pd.DataFrame) == True):

        #Add National Growth
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [national_5y_growth, national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [national_5y_growth, national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                  ),
            row = 1,
            col = 2)


        #Add MSA Growth
        fig.add_trace( 
            go.Bar(
                name         = cbsa_name + ' (MSA)',  
                x            = years, 
                y            = [msa_5y_growth, msa_3y_growth, msa_1y_growth],
                marker_color = bowery_light_blue,
                text         = [msa_5y_growth, msa_3y_growth, msa_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                  ),
            row = 1,
            col = 2)

        #Add County Growth
        fig.add_trace(
            go.Bar(
                    name=county,      
                    x=years, 
                    y=[county_5y_growth,county_3y_growth,county_1y_growth],
                    marker_color = bowery_dark_blue,
                    text = [county_5y_growth,county_3y_growth,county_1y_growth],
                    texttemplate = "%{value:.2f}%",
                    textposition = annotation_position,
                    cliponaxis =  False
                ),
            row = 1,
            col = 2)
    else:
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [national_5y_growth, national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [national_5y_growth, national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   = False
                ),
            row = 1,
            col = 2)

    #Change the bar mode
    fig.update_layout(barmode='group')


    #Set Y-Axes format
    fig.update_yaxes(
        ticksuffix = '%',
        tickfont   = dict(size=tickfont_size),
        visible    = False,
        row        = 1,
        col        = 2,
                    )                 

    #Set X-axes format
    fig.update_xaxes(
        tickfont = dict(size=tickfont_size),
        row      = 1,
        col      = 2
                    )

    #Set X-Axis Format
    fig.update_xaxes(
        type = 'date',
        dtick="M12",
        tickformat="%Y",
        tickangle = tickangle,
        tickfont = dict(size=tickfont_size),
        row = 1,
        col = 1,
                    )

    #Set Y-Axis format
    fig.update_yaxes( tickfont = dict(size=tickfont_size),
                      row      = 1,
                      col      = 1,
                    )

    fig.update_yaxes(tickprefix = '$', tickfont = dict(size=tickfont_size), secondary_y=False, row = 1,col = 1)



    
    #Set Legend Layout
    fig.update_layout(
        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position + 0.1 ,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )

    fig.write_image(os.path.join(folder,'per_capita_income_and_growth.png'), engine='kaleido', scale=scale)

def CreateGDPGraph(county_data_frame,msa_data_frame,state_data_frame,folder):
    print('Creating GDP Graph')
    fig = make_subplots(specs=[[{"secondary_y": True}]])

    #Add county GDP
    if (isinstance(county_data_frame, pd.DataFrame) == True):
        fig.add_trace(
            go.Scatter(x = county_data_frame['Period'],
                       y  = county_data_frame['GDP'],
                    name = county+' (L)',
                    line = dict(color = bowery_dark_blue, width = 1,dash = 'dash')
                    ),
            secondary_y=False
                        )

    #Add MSA GDP if applicable
        if (isinstance(msa_data_frame, pd.DataFrame) == True) and ( (msa_data_frame['GDP'].equals(county_data_frame['GDP']))  == False  ):
            fig.add_trace(
            go.Scatter(x=msa_data_frame['Period'],
                    y=msa_data_frame['GDP'],
                    name=cbsa_name + ' (MSA)' + ' (R)',
                    line = dict(color=bowery_light_blue, width = 1),
                    )
            ,secondary_y=True)
        else:
            #Add state GDP
            fig.add_trace(
            go.Scatter(x=state_data_frame['Period'],
                    y=state_data_frame['GDP'],
                    name=state_name+' (R)',
                    line=dict(color=bowery_dark_grey, width = 1),
                    )
            ,secondary_y=True)
    else:
        if (isinstance(msa_data_frame, pd.DataFrame) == True):
            fig.add_trace(
                go.Scatter(x = msa_data_frame['Period'],
                           y = msa_data_frame['GDP'],
                        name = cbsa_name + ' (MSA)' + ' (L)',
                        line = dict(color=bowery_light_blue),
                          ),
                secondary_y=False
            )
        
        #Add state GDP
        fig.add_trace(
            go.Scatter(x = state_data_frame['Period'],
                    y    = state_data_frame['GDP'],
                    name = state_name+' (R)',
                    line = dict(color=bowery_dark_grey),
                    ),
            secondary_y = True
                    )


    #Set X-Axis Format
    fig.update_xaxes(
        type       = 'date',
        dtick      = "M12",
        tickformat = "%Y",
        tickangle  = 0,
        tickfont   = dict(size=tickfont_size),
                    )

    #Set Y-Axis format
    fig.update_yaxes(tickprefix = '$', tickfont = dict(size=tickfont_size), secondary_y=False)
    fig.update_yaxes(tickprefix = '$', tickfont = dict(size=tickfont_size), secondary_y=True)


    #Set Title
    fig.update_layout(
    title_text="Gross Domestic Product",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    #Set Legend Layout
    fig.update_layout(
        legend=dict(
            orientation = "h",
            yanchor     = "bottom",
            y           = legend_position ,
            xanchor     = "center",
            x           = 0.5,
            font_size   = tickfont_size
                   ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )

    fig.write_image(os.path.join(folder,'gdp.png'), engine='kaleido', scale=scale)

def CreatePopulationOverTimeWithGrowthGraph(county_resident_pop, state_resident_pop, msa_resident_pop, national_resident_pop, folder):
    print('Creating population over time with growth bar graph')
    
    #Create figure with secondary y-axis
    fig = make_subplots(specs=[[{"secondary_y": True}, {"secondary_y": False}]],rows=1, cols=2,subplot_titles=("Population", "Annualized Population Growth"),horizontal_spacing = horizontal_spacing)

    #County Population
    fig.add_trace(
        go.Scatter(x = county_resident_pop['Period'],
                y = county_resident_pop['Resident Population'],
                name = county + ' (L)',
                line = dict(color = bowery_dark_blue, width = 1,dash = 'dash'),
                showlegend= False
                ),      
        secondary_y = False, row = 1, col = 1
                 )
    
    #MSA Population if applicable
    if (cbsa != '') and (msa_resident_pop.equals(county_resident_pop) == False):
        fig.add_trace(
            go.Scatter(x = msa_resident_pop['Period'],
                    y = msa_resident_pop['Resident Population'],
                    name = cbsa_name + ' (MSA)' + ' (R)',
                    line = dict(color =bowery_light_blue, width = 1),
                    showlegend = False
                      ),
            secondary_y = True, row = 1, col = 1
                     )
    else:
        #State Population
        fig.add_trace(
        go.Scatter(x=state_resident_pop['Period'],
                y=state_resident_pop['Resident Population'],
                name=state_name + ' (R)',
                line = dict(color=bowery_dark_grey, width = 1),
                showlegend=False
                ),
        secondary_y=True,row=1, col=1,)   


    #Calculate annualized growth rates for the county, msa (if available), and state dataframes
    assert len(county_resident_pop) >= 11
    county_resident_pop['Resident Population_1year_growth'] =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
    county_resident_pop['Resident Population_5year_growth'] =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
    county_resident_pop['Resident Population_10year_growth'] =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10

    county_1y_growth  = county_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
    county_5y_growth  = county_resident_pop.iloc[-1]['Resident Population_5year_growth'] 
    county_10y_growth = county_resident_pop.iloc[-1]['Resident Population_10year_growth']

    if cbsa != '':
        #Make sure we are comparing same years for calculating growth rates for county and msa
        assert len(msa_resident_pop) >= 11
        msa_resident_pop = msa_resident_pop.loc[msa_resident_pop['Period'] <= (county_resident_pop['Period'].max())]
        msa_resident_pop['Resident Population_1year_growth'] =  (((msa_resident_pop['Resident Population']/msa_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
        msa_resident_pop['Resident Population_5year_growth'] =  (((msa_resident_pop['Resident Population']/msa_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
        msa_resident_pop['Resident Population_10year_growth'] =  (((msa_resident_pop['Resident Population']/msa_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10

        msa_1y_growth  = msa_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
        msa_5y_growth  = msa_resident_pop.iloc[-1]['Resident Population_5year_growth'] 
        msa_10y_growth = msa_resident_pop.iloc[-1]['Resident Population_10year_growth']

    #Make sure we are comparing same years for calculating growth rates for county and state
    assert len(state_resident_pop) >= 11
    state_resident_pop = state_resident_pop.loc[state_resident_pop['Period'] <= (county_resident_pop['Period'].max())]
    state_resident_pop['Resident Population_1year_growth'] =  (((state_resident_pop['Resident Population']/state_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
    state_resident_pop['Resident Population_5year_growth'] =  (((state_resident_pop['Resident Population']/state_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
    state_resident_pop['Resident Population_10year_growth'] =  (((state_resident_pop['Resident Population']/state_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10

    state_1y_growth  = state_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
    state_5y_growth  = state_resident_pop.iloc[-1]['Resident Population_5year_growth'] 
    state_10y_growth = state_resident_pop.iloc[-1]['Resident Population_10year_growth']

    #Make sure we are comparing same years for calculating growth rates for county and USA
    national_resident_pop = national_resident_pop.loc[national_resident_pop['Period'] <= (county_resident_pop['Period'].max())]
    national_resident_pop['Resident Population_1year_growth'] =  (((national_resident_pop['Resident Population']/national_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
    national_resident_pop['Resident Population_5year_growth'] =  (((national_resident_pop['Resident Population']/national_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
    national_resident_pop['Resident Population_10year_growth'] =  (((national_resident_pop['Resident Population']/national_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10

    national_1y_growth  = national_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
    national_5y_growth  = national_resident_pop.iloc[-1]['Resident Population_5year_growth'] 
    national_10y_growth = national_resident_pop.iloc[-1]['Resident Population_10year_growth']

    #Now that we've calculated growth rates, create our plot
    years=['10 Years', '5 Years', '1 Year']
    annotation_position = 'outside'


    #If there's a MSA/CBSA include it, otherwise just use county, state, and USA
    if (cbsa != '') and (county_resident_pop.equals(msa_resident_pop) == False):
            
        fig.add_trace(        
        go.Bar(
            name='United States',   
            x=years, 
            y=[national_10y_growth, national_5y_growth, national_1y_growth],
            marker_color ="#000F44",
            text = [national_10y_growth, national_5y_growth, national_1y_growth],
            texttemplate = "%{value:.2f}%",
            textposition = annotation_position ,
            cliponaxis =  False,
            ),
            row = 1,
            col = 2
        )
        

        fig.add_trace(        
        go.Bar(
            name=cbsa_name + ' (MSA)' + ' (R)',   
            x=years, 
            y=[msa_10y_growth, msa_5y_growth, msa_1y_growth],
            marker_color =bowery_light_blue,
            text = [msa_10y_growth, msa_5y_growth, msa_1y_growth],
            texttemplate = "%{value:.2f}%",
            textposition = annotation_position ,
            cliponaxis =  False,
            ),
            row = 1,
            col = 2
        )
        
        fig.add_trace( 
        go.Bar(
            name=county + ' (L)',      
            x=years, 
            y=[county_10y_growth,county_5y_growth,county_1y_growth],
            marker_color = bowery_dark_blue,
            text = [county_10y_growth,county_5y_growth,county_1y_growth],
            texttemplate = "%{value:.2f}%",
            textposition = annotation_position,
            cliponaxis =  False),
            row=1,
            col=2
        )
    
    else:
        fig.add_trace(        
        go.Bar(
            name='United States',   
            x=years, 
            y=[national_10y_growth, national_5y_growth, national_1y_growth],
            marker_color ="#000F44",
            text = [national_10y_growth, national_5y_growth, national_1y_growth],
            texttemplate = "%{value:.2f}%",
            textposition = annotation_position ,
            cliponaxis =  False,
            ),
            row = 1,
            col = 2
        )

        fig.add_trace( 
        go.Bar(
            name=state_name + ' (R)',  
            x=years, 
            y=[state_10y_growth, state_5y_growth, state_1y_growth],
            marker_color =bowery_dark_grey,
            text = [state_10y_growth, state_5y_growth, state_1y_growth],
            texttemplate = "%{value:.2f}%",
            textposition = annotation_position,
            cliponaxis =  False,
            ),
            row = 1,
            col = 2
        )

        fig.add_trace( 
        go.Bar(
                name=county + ' (L)',      
                x=years, 
                y=[county_10y_growth,county_5y_growth,county_1y_growth],
                marker_color = bowery_dark_blue,
                text = [county_10y_growth,county_5y_growth,county_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis =  False,
                ),
                  row = 1,
                  col = 2
                 )
        
        

    #Change the bar mode
    fig.update_layout(barmode='group')

    # fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

     #Set X-axes format
    fig.update_xaxes(
        tickfont = dict(size=tickfont_size),
        row = 1,
        col = 2
        )

    #Set Y-Axes format For Growth Subfigure
    fig.update_yaxes(
        ticksuffix = '%',
        tickfont = dict(size=tickfont_size),row =1,col =2,
        visible = False)                 


    #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position + 0.1 ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=(top_margin + .2), b = (bottom_margin + .2)),
    height    = graph_height,
    width     = graph_width,
    )

    #Set X-Axis Format
    fig.update_xaxes(
        type = 'date',
        dtick="M12",
        tickformat="%Y",
        tickangle = tickangle,
        tickfont = dict(size=tickfont_size),
        # linecolor = 'black',
        row = 1,
        col = 1
        )
    # fig.show()

    # #Set Y-Axis format
    fig.update_yaxes( tickfont = dict(size=tickfont_size),
                    #   linecolor='black',
                      row =1,
                      col =1  
                      )
    #Set Y-Axes format
    fig.update_yaxes(
        ticksuffix = '%',
        tickfont = dict(size=tickfont_size),
        visible = False,
        row = 1,
        col = 2)     

    fig.write_image(os.path.join(folder,'resident_population_and_growth.png'),engine='kaleido',scale=scale)

def CreateMLPWithGrowthGraph(county_data_frame,msa_data_frame,national_data_frame,folder):
    print("Creating MLP graph")
    fig = make_subplots(specs=[[{"secondary_y": True}, {"secondary_y": False}]],rows=1, cols=2,subplot_titles=("Median Home List Price", "Annualized MLP Growth"),horizontal_spacing = horizontal_spacing)

    #Add county MLP
    if  (isinstance(county_data_frame, pd.DataFrame) == True):
        fig.add_trace(
            go.Scatter(x = county_data_frame['Period'],
                      y  = county_data_frame['Median List Price'],
                    name = county,
                    line = dict(color = bowery_dark_blue,width = 1,dash = 'dash'),
                    showlegend=False),
            secondary_y=False,
            row = 1,
            col = 1
                      )
    

    #Add MSA MLP if applicable
    if   (isinstance(msa_data_frame, pd.DataFrame) == True):
        fig.add_trace(
            go.Scatter(x = msa_data_frame['Period'],
                    y    = msa_data_frame['Median List Price'],
                    name = cbsa_name + ' (MSA)',
                    line = dict(color=bowery_light_blue,width = 1),
                    showlegend=False
                      ),
            secondary_y=False, 
            row = 1,
            col = 1)
    

    #Add national MLP    
    if   (isinstance(national_data_frame, pd.DataFrame) == True):
        fig.add_trace(
            go.Scatter(x = national_data_frame['Period'],
                    y    = national_data_frame['Median List Price'],
                    name = 'United States',
                    line = dict(color='#000F44'),
                    showlegend=False),
            secondary_y=False,
            row = 1,
            col = 1)
    
    #Add Growth Subfigure
    #Calculate annualized growth rates for the county, msa (if available), and national dataframes
    if (isinstance(county_data_frame, pd.DataFrame) == True):
        county_data_frame['Median List Price_1year_growth'] =  (((county_data_frame['Median List Price']/county_data_frame['Median List Price'].shift(1 * 12))  - 1) * 100)/1
        county_data_frame['Median List Price_3year_growth'] =  (((county_data_frame['Median List Price']/county_data_frame['Median List Price'].shift(3 * 12))   - 1) * 100)/3
        county_data_frame['Median List Price_5year_growth'] =  (((county_data_frame['Median List Price']/county_data_frame['Median List Price'].shift(5* 12))   - 1) * 100)/5

        county_1y_growth                                    = county_data_frame.iloc[-1]['Median List Price_1year_growth'] 
        county_3y_growth                                    = county_data_frame.iloc[-1]['Median List Price_3year_growth'] 
        county_5y_growth                                    = county_data_frame.iloc[-1]['Median List Price_5year_growth']
    else:
        county_1y_growth                                    = 0
        county_3y_growth                                    = 0
        county_5y_growth                                    = 0 

        
    if (isinstance(msa_data_frame, pd.DataFrame) == True):
        #Make sure we are comparing same years for calculating growth rates for county and msa
        if (isinstance(county_data_frame, pd.DataFrame) == True):
            msa_data_frame = msa_data_frame.loc[msa_data_frame['Period'] <= (county_data_frame['Period'].max())]
        msa_data_frame['Median List Price_1year_growth'] =  (((msa_data_frame['Median List Price']/msa_data_frame['Median List Price'].shift(1 * 12))  - 1) * 100)/1
        msa_data_frame['Median List Price_3year_growth'] =  (((msa_data_frame['Median List Price']/msa_data_frame['Median List Price'].shift(3 * 12))   - 1) * 100)/3
        msa_data_frame['Median List Price_5year_growth'] =  (((msa_data_frame['Median List Price']/msa_data_frame['Median List Price'].shift(5 * 12))   - 1) * 100)/5

        msa_1y_growth                                    = msa_data_frame.iloc[-1]['Median List Price_1year_growth'] 
        msa_3y_growth                                    = msa_data_frame.iloc[-1]['Median List Price_3year_growth'] 
        msa_5y_growth                                    = msa_data_frame.iloc[-1]['Median List Price_5year_growth'] 


    #Make sure we are comparing same years for calculating growth rates for county and state
    if  (isinstance(county_data_frame, pd.DataFrame) == True):
        national_data_frame = national_data_frame.loc[national_data_frame['Period'] <= (county_data_frame['Period'].max())]
    national_data_frame['Median List Price_1year_growth'] =  (((national_data_frame['Median List Price']/national_data_frame['Median List Price'].shift(1 * 12))  - 1) * 100)/1
    national_data_frame['Median List Price_3year_growth'] =  (((national_data_frame['Median List Price']/national_data_frame['Median List Price'].shift(3 * 12))   - 1) * 100)/3
    national_data_frame['Median List Price_5year_growth'] =  (((national_data_frame['Median List Price']/national_data_frame['Median List Price'].shift(5 * 12))   - 1) * 100)/5

    national_1y_growth                                    = national_data_frame.iloc[-1]['Median List Price_1year_growth'] 
    national_3y_growth                                    = national_data_frame.iloc[-1]['Median List Price_3year_growth'] 
    national_5y_growth                                    = national_data_frame.iloc[-1]['Median List Price_5year_growth'] 
    

    #Now that we've calculated growth rates, create our plot
    years               = ['3 Years','1 Year']
    annotation_position = 'outside'
    
   

    #MSA MLP is available
    if (isinstance(msa_data_frame, pd.DataFrame) == True)  and (isinstance(county_data_frame, pd.DataFrame) == False):
        print('MSA MLP Data is available, county is not')
        #Add National Growth 
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                   ),
            row = 1,
            col = 2)

        #Add MSA Growth
        fig.add_trace( 
            go.Bar(
                name         = cbsa_name + ' (MSA)',  
                x            = years, 
                y            = [ msa_3y_growth, msa_1y_growth],
                marker_color = bowery_light_blue,
                text         = [ msa_3y_growth, msa_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                  ),
            row = 1,
            col = 2
                    )

    
    
    #MSA is unavailable, but county is available
    elif (isinstance(msa_data_frame, pd.DataFrame) == False)  and (isinstance(county_data_frame, pd.DataFrame) == True):
        print('MSA MLP Unavailable, County MLP is available')
        #Add National Growth
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                   ),
        row = 1,
        col = 2)

        #Add County Growth
        fig.add_trace(
            go.Bar(
                name         = county,      
                x            = years, 
                y            = [county_3y_growth,county_1y_growth],
                marker_color = bowery_dark_blue,
                text         = [county_3y_growth,county_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                   ),
            row = 1,
            col = 2
                     )

    

    #MSA and County are UNavailable
    elif (isinstance(msa_data_frame, pd.DataFrame) == False)  and (isinstance(county_data_frame, pd.DataFrame) == False):
        print('County and MSA MLP data are unavailable')
        #Add National Growth
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [ national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [ national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                  ),
            row = 1,
            col = 2)
    
    #County AND MSA are available
    elif (isinstance(msa_data_frame, pd.DataFrame) == True)  and (isinstance(county_data_frame, pd.DataFrame) == True):
        print('MSA and County are available')

        #Add National Growth
        fig.add_trace( 
            go.Bar(
                name         = 'United States',  
                x            = years, 
                y            = [ national_3y_growth, national_1y_growth],
                marker_color = "#000F44",
                text         = [ national_3y_growth, national_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                ),
            row = 1,
            col = 2)

        #Add MSA Growth
        fig.add_trace( 
            go.Bar(
                name         = cbsa_name + ' (MSA)',  
                x            = years, 
                y            = [msa_3y_growth, msa_1y_growth],
                marker_color = "#B3C3FF",
                text         = [msa_3y_growth, msa_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                 ),
            row = 1,
            col = 2)

        #Add County Growth
        fig.add_trace(
            go.Bar(
                name         = county,      
                x            = years, 
                y            = [county_3y_growth,county_1y_growth],
                marker_color = bowery_dark_blue,
                text         = [county_3y_growth,county_1y_growth],
                texttemplate = "%{value:.2f}%",
                textposition = annotation_position,
                cliponaxis   =  False
                ),
            row = 1,
            col = 2)

    #Change the bar mode
    fig.update_layout(barmode='group')


    #Set Y-Axes format
    fig.update_yaxes(
        ticksuffix = '%',
        tickfont = dict(size=tickfont_size),
        visible = False,
        row = 1,
        col = 2)                 

    #Set X-axes format
    fig.update_xaxes(
        tickfont = dict(size=tickfont_size),
        row = 1,
        col = 2
        )

    #Set X-Axis Format
    fig.update_xaxes(
        type = 'date',
        dtick="M12",
        tickformat="%Y",
        tickangle = tickangle,
        tickfont = dict(size=tickfont_size),
        row = 1,
        col = 1
        )

    #Set Y-Axis format
    fig.update_yaxes(  tickfont = dict(size=tickfont_size),
                       row = 1,
                       col = 1  
                    )

    fig.update_yaxes(tickprefix = '$', tickfont = dict(size=tickfont_size),secondary_y=False, row = 1,col = 1)


    #Set Legend Layout
    fig.update_layout(
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position + .1 ,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )

    # fig.update_yaxes(automargin = True)  
    fig.write_image(os.path.join(folder,'mlp.png'),engine='kaleido',scale=scale)

def CreateEmploymentByIndustryGraph(county_data_frame,folder):
    print('Creating County Employment by Industry Breakdown Graph')
    
    def format(x):
        return "Weekly Wage: ${:,.0f}".format(x)
    county_data_frame['avg_wkly_wage_string'] = county_data_frame['avg_wkly_wage'].apply(format)
    

    #Employment By Supersector Treemap
    fig = go.Figure(
            go.Treemap(
                values       = county_data_frame['month3_emplvl'],
                labels       = county_data_frame['industry_code'],
                parents      = county_data_frame['county'],
                text         = county_data_frame['avg_wkly_wage_string'],
                textinfo     = "label+text",
                textposition = 'top left',
                marker=dict(
                        colors     = county_data_frame['avg_wkly_wage'],
                        colorscale = 'Blues',),

                              
                     )
                  )


    #Set font               
    #Set Title
    fig.update_layout(
        title={
            'text':    'County Employment Composition & Wages by Industry' + ' (' + qcew_year + ' Q' + qcew_qtr + ')',
            'y':       .985 ,
            'x':       0.5,
            'xanchor': 'center',
            'yanchor': 'top'
            },  

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
    
        margin        = dict(l=left_margin, r=right_margin, t=20, b= 0),
        height        = graph_height,
        width         = graph_width,
                    )

    fig.write_image(os.path.join(folder,'employment_by_industry.png'), engine='kaleido', scale=scale)

def CreateEmploymentGrowthByIndustryGraph(county_data_frame,folder):
    print('Creating County Employment Growth by Industry Graph')
    annotation_position = 'outside'
    county_data_frame   = county_data_frame.loc[county_data_frame['industry_code'] != 'Unclassified'] 

    #Drop industreis where we are missing 5 and 1 year growth
    county_data_frame  = county_data_frame.loc[(county_data_frame['emp_growth_invalid'] != 1) | (county_data_frame['one_year_emp_growth_invalid'] != 1)] 

    fig = go.Figure(
            data=[
            #Add 5 Year Growth Bars
            go.Bar(
                name=str(growth_period) + ' Year Growth',      
                x=county_data_frame['industry_code'], 
                y=county_data_frame['Employment Growth'],
                marker_color= bowery_grey,

                  ),

            #Add 1 year growth circles       
            go.Scatter(
                    name   = '1 Year Growth',      
                    x      = county_data_frame['industry_code'], 
                    y      = county_data_frame['1 Year Employment Growth'],
                    marker = dict(color = bowery_dark_blue, size=9),
                    mode   = 'markers',
                    ),
                    ]
                        )

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    #Set X-axes format
    fig.update_xaxes(
        tickfont = dict(size=tickfont_size),
        title_standoff = 0.10
        )

    #Set Y-Axes format
    fig.update_yaxes(
        ticksuffix = '%',
        tickfont   = dict(size=tickfont_size),
        )                 

    #Set Title
    fig.update_layout(
        title_text= "Private Employment Growth by Industry (County)" + ' (' + qcew_year + ' Q' + qcew_qtr + ')',    

        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
              },
                    
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position ,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),


        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin        = dict(l=left_margin, r=right_margin, t=(top_margin + .2), b = (bottom_margin + .2)),
        height        = graph_height,
        width         = graph_width,
                    )


    fig.write_image(os.path.join(folder,'employment_growth_by_industry.png'),engine='kaleido',scale=scale)

def CreateMSAEmploymentByIndustryGraph(msa_data_frame,folder):
    print('Creating MSA Employment by Industry Breakdown Graph')
    
    def format(x):
        return "Weekly Wage: ${:,.0f}".format(x)
    msa_data_frame['avg_wkly_wage_string'] = msa_data_frame['avg_wkly_wage'].apply(format)
    

    #Employment By Supersector Treemap
    fig = go.Figure(
            go.Treemap(
                values       = msa_data_frame['month3_emplvl'],
                labels       = msa_data_frame['industry_code'],
                parents      = msa_data_frame['msa'],
                text         = msa_data_frame['avg_wkly_wage_string'],
                textinfo     = "label+text",
                textposition = 'top left',
                marker       = dict(
                                colors     = msa_data_frame['avg_wkly_wage'],
                                colorscale ='Blues',
                          ),

                              
                      )
                    )

 
    fig.update_layout(

        title={
            'text': 'MSA Private Employment Composition & Wages by Industry' + ' (' + qcew_year + ' Q' + qcew_qtr + ')',
            'y':.985 ,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
              },  

        font_family="Avenir Next LT Pro",
        font_color='#262626',
        font_size = 10.5,
        paper_bgcolor=paper_backgroundcolor,
        plot_bgcolor ="White",

        margin    =dict(l=left_margin, r=right_margin, t=20, b= 0),
        height    = graph_height,
        width     = graph_width,
                    )

    fig.write_image(os.path.join(folder,'msa_employment_by_industry.png'),engine='kaleido',scale=scale)

def CreateMSAEmploymentGrowthByIndustryGraph(msa_data_frame,folder):
    print('Creating MSA Employment Growth by Industry Graph')
    annotation_position = 'outside'
    msa_data_frame  = msa_data_frame.loc[msa_data_frame['industry_code'] != 'Unclassified'] 

    #Drop industreis where we are missing 5 and 1 year growth
    msa_data_frame  = msa_data_frame.loc[(msa_data_frame['emp_growth_invalid'] != 1) | (msa_data_frame['one_year_emp_growth_invalid'] != 1)] 

    fig = go.Figure(data=[
        #Add 5 Year Growth Bars
        go.Bar(
                name=str(growth_period) + ' Year Growth',      
                x=msa_data_frame['industry_code'], 
                y=msa_data_frame['Employment Growth'],
                marker_color = bowery_grey,
                ),

        #Add 1 year growth circles       
        go.Scatter(
                name='1 Year Growth',      
                x=msa_data_frame['industry_code'], 
                y=msa_data_frame['1 Year Employment Growth'],
                marker=dict(color = bowery_dark_blue, size=9),
                mode = 'markers',
                ),
                            ]
                     )

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    #Set X-axes format
    fig.update_xaxes(
        tickfont = dict(size=tickfont_size),
        title_standoff = 0.10
        )

    #Set Y-Axes format
    fig.update_yaxes(
        ticksuffix = '%',
        tickfont = dict(size=tickfont_size),
        )                 

    #Set Title
    fig.update_layout(
        title_text= "Private Employment Growth by Industry (MSA) " + ' (' + qcew_year + ' Q' + qcew_qtr + ')',    

        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'},
                    
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position ,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),


        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin        = dict(l=left_margin, r=right_margin, t=(top_margin + .2), b = (bottom_margin + .2)),
        height        = graph_height,
        width         = graph_width,
                    )

    fig.write_image(os.path.join(folder,'msa_employment_growth_by_industry.png'), engine='kaleido', scale=scale)

def CreateMLPGraph(county_data_frame,msa_data_frame,folder):
    print('Creating Median List Price Graph')
    if (isinstance(county_data_frame, pd.DataFrame) == False): 
        return()

    #Plot Median Home List price for metro and county
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #Add county median list price
    fig.add_trace(
        go.Scatter(x=county_data_frame['Period'],
                y=county_data_frame['Median List Price'],
                name=county,
                mode='lines',
                line = dict(color = bowery_dark_blue,width = 1,dash = 'dash')
                ),
        secondary_y=False
                 )



    #Add msa median list price if applicable
    if cbsa_name != '':
        fig.add_trace(
            go.Scatter(x=msa_data_frame['Period'],
                    y=msa_data_frame['Median List Price'],
                    name = cbsa_name + ' (MSA)',
                    mode = 'lines',
                    line = dict(color = bowery_light_blue,width = 1)
                    ),
            secondary_y=False
        )
    

    #Add National median list price
    fig.add_trace(
        go.Scatter(x=national_mlp['Period'],
                y=national_mlp['Median List Price'],
                name='United States',
                mode='lines',
                line = dict(color="#000F44",width = 1)
                ),
        secondary_y=False
                 )

    #Set x-axis format
    fig.update_xaxes(
        type = 'date',
        dtick="M4",
        tickformat="%m/%y",
        tickangle = 0,
        tickfont = dict(size=tickfont_size),
        )

    #Set y-axis format
    fig.update_yaxes(
        tickfont = dict(size=tickfont_size),
        tickprefix = '$'
        )


    #Set Title
    fig.update_layout(
        title_text = "Median Home List Price",    

        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
               },

        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position ,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

    
        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
                    )

    fig.write_image(os.path.join(folder,'mlp.png'), engine='kaleido', scale=scale)

def CreateNationalUnemploymentGraph(folder):
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #County unemployment rate
    fig.add_trace(
        go.Scatter(x = national_unemployment['period'],
                  y  = national_unemployment['unemployment_rate'],
                name = 'United States of America',
                line = dict(color = bowery_dark_blue,)),
        secondary_y  = False
                 )

    #Set formatting 
    fig.update_layout(
        
        title_text="National Unemployment Rate",    
        
        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
            },

        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),
        
        font_family="Avenir Next LT Pro",
        font_color='#262626',
        font_size = 10.5,
        paper_bgcolor=paper_backgroundcolor,
        plot_bgcolor ="White",    
                    )

    #Add % to left axis ticks
    fig.update_yaxes(
        tickfont = dict(size=tickfont_size), 
        ticksuffix = '%',  
        title = None ,
        linecolor = 'black',   
        tickmode  = 'auto',
        nticks    = 6,
        showgrid = True,
        gridcolor = 'black',
        secondary_y=False,)   
    
                    
    #Set x axis ticks
    quarter_list      = [i for i in range(len(county_unemployment_rate['period']))]
    quarter_list      = quarter_list[::-12]
    quarter_list_text = [period for period in county_unemployment_rate['period']]
    quarter_list_text = quarter_list_text[::-12]

    fig.update_xaxes(
        tickmode = 'array',
        tickvals = quarter_list,
        ticktext = quarter_list_text,
        tickfont = dict(size=tickfont_size),
        tickangle = 0,
        linecolor = 'black' 
        )

    #Set size
    fig.update_layout(
        autosize  = False,
        height    = graph_height,
        width     = graph_width,
        margin    = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin,pad=0,autoexpand = True),
                     )
    
    fig.write_image(os.path.join(folder,'national_unemployment_rate.png'), engine='kaleido', scale=scale)

def CreateNationalEmploymentGrowthGraph(folder):
    
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    
    fig.add_trace(
        go.Scatter(x = national_employment['period'],
                  y  = national_employment['Employment Growth'],
                name = county,
                line = dict(color = bowery_dark_blue)
                ),
        secondary_y=False
    )



    #Set formatting 
    fig.update_layout(
        title_text="National Annual Employment Growth",   

        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
               },

        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),


        font_family="Avenir Next LT Pro",
        font_color='#262626',
        font_size = 10.5,
        paper_bgcolor=paper_backgroundcolor,
        plot_bgcolor ="White",    
                    )

    #Add % to left axis ticks
    fig.update_yaxes(
        tickfont = dict(size=tickfont_size), 
        ticksuffix = '%',  
        title = None ,
        linecolor = 'black',
        tickmode = 'auto',
        nticks   = 6,
        showgrid = True,
        gridcolor = 'black',
        secondary_y=False)                
    fig.add_hline(y=0, line_width=1, line_color="black")              
    
    
    #Set x axis ticks
    quarter_list = [i for i in range(len(county_employment['period']))]
    quarter_list = quarter_list[::-12]

    quarter_list_text = [period for period in county_employment['period']]
    quarter_list_text = quarter_list_text[::-12]

    fig.update_xaxes(tickmode = 'array',
        tickvals = quarter_list,
        ticktext = quarter_list_text,
        tickfont = dict(size=tickfont_size),
        tickangle = 0,
        linecolor = 'black'
        )

    #Set size
    fig.update_layout(
    autosize=False,
    height    = graph_height,
    width     = graph_width,
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin,pad=0,autoexpand = True),)
    
    fig.write_image(os.path.join(folder,'national_employment_growth.png'),engine='kaleido',scale=scale)

def CreateNationalGDPGraph(folder):
    print('Creating National GDP Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    
    fig.add_trace(
        go.Scatter(x=national_gdp['Period'],
                y=national_gdp['GDP'],
                name='United States of America',
                line=dict(color = bowery_dark_blue)
                ),
        secondary_y=False)

    #Set formatting 
    fig.update_layout(
        title_text="National Gross Domestic Product",    
        
        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
               },

        legend=dict(
                orientation="h",
                yanchor="bottom",
                y=legend_position,
                xanchor="center",
                x=0.5,
                font_size = tickfont_size
                    ),
        
        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 10.5,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",    
                    )

    #Add $ to left axis ticks
    fig.update_yaxes(
        tickfont    = dict(size=tickfont_size), 
        tickprefix  = '$',  
        title       = None ,
        linecolor   = 'black',   
        tickmode    = 'auto',
        nticks      = 6,
        showgrid    = True,
        gridcolor   = 'black',
        secondary_y = False)                 
                    
    fig.update_xaxes(
        tickmode = 'array',
        tickfont = dict(size=tickfont_size),
        tickangle = 0,
        linecolor = 'black' 
        )

    #Set size
    fig.update_layout(
    autosize  = False,
    height    = graph_height,
    width     = graph_width,
    margin    = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin,pad=0,autoexpand = True),
                    )
    
    fig.write_image(os.path.join(folder,'national_gdp_rate.png'), engine='kaleido', scale=scale)

def CreateGraphs():
    print('Creating Graphs')
    CreateUnemploymentRateEmploymentGrowthGraph(                                                                                                                                                                    folder = county_folder)
    CreateGDPGraph(                        county_data_frame = county_gdp, msa_data_frame = msa_gdp,state_data_frame=state_gdp,                                                                                    folder = county_folder)
    CreatePCIGraph(                        county_data_frame = county_pci ,msa_data_frame = msa_pci,state_data_frame=state_pci, national_data_frame = national_pci,                                                  folder = county_folder)
    CreateEmploymentByIndustryGraph(       county_data_frame = county_industry_breakdown,                                                                                                                            folder = county_folder)
    CreateEmploymentGrowthByIndustryGraph( county_data_frame = county_industry_growth_breakdown,                                                                                                                     folder = county_folder)
    CreatePopulationOverTimeWithGrowthGraph(county_resident_pop = county_resident_pop, state_resident_pop = state_resident_pop, msa_resident_pop = msa_resident_pop, national_resident_pop = national_resident_pop, folder = county_folder)
    
    #Median list price graph
    try:
        CreateMLPWithGrowthGraph(county_data_frame = county_mlp, msa_data_frame = msa_mlp, national_data_frame = national_mlp, folder = county_folder)
    except Exception as e:
        print(e)
    
    #MSA Graphs
    if cbsa != '':
        CreateMSAEmploymentByIndustryGraph(      msa_data_frame = msa_industry_breakdown,              folder = county_folder )
        CreateMSAEmploymentGrowthByIndustryGraph(msa_data_frame = msa_industry_growth_breakdown,       folder = county_folder )

#####################################################Language Related Functions####################################
def millify(n):
    millnames = ['',' thousand',' million',' billion',' trillion']
    n         = float(n)
    millidx   = max(0,min(len(millnames)-1,
                        int(math.floor(0 if n == 0 else math.log10(abs(n))/3))))

    return '{:.1f}{}'.format(n / 10**(3 * millidx), millnames[millidx])

def OverviewLanguage():
    print('Writing Overview Langauge')

    #Section 1A: Identify county and state of subject property
    subject_property_location_language = ('The subject property is located in ' + county + ', ' + state_name + '. ')


    #Section 1B: Grab overview summary from Metro_language CSV using the CBSA Code from IdentifyMSA
    try:
        metro_area_language_df = pd.read_csv(os.path.join(data_location,'Metro_Language.csv'),
                                                    dtype={
                                                        'CBSA_Code':             object,
                                                        'metro_name':            object,
                                                        'Overview':              object,
                                                        'LaborMarketConditions': object,
                                                        'growth':                object,
                                                        'UniqueAspects':         object,
                                                        'Infrastructure':        object,
                                                        'Education':             object
                                                        }
                                            )

        #restrict data to only rows with the Subject County CBSA Code
        metro_area_language_df = metro_area_language_df.loc[metro_area_language_df['CBSA_Code'] == cbsa]

        #If we have a cbsa/msa for the county, then use the excel file text
        if len(metro_area_language_df) == 1:   
            CBSA_overview_language       = metro_area_language_df['Overview'].iloc[-1]
        else:
            CBSA_overview_language       = ''
    
    except Exception as e:
        print(e,'trouble getting cbsa overview language')
        CBSA_overview_language       = ''



    #Section 1c: Grab summary text from Wikipedia

    try:
        wikipeida_summary      =  wikipedia.summary((county + ',' + state))
        wikipeida_summary      = wikipeida_summary.replace('the U.S. state of ','')
    
    except Exception as e:
        print(e)
        wikipeida_summary      = ''

    try:
        wikipeida_economy_summary                 = wikipedia.page((county + ',' + state)).section('Economy')
        assert wikipeida_economy_summary != None
    
    except Exception as e:
        print(e)
        wikipeida_economy_summary               = ''


    #Section 2: Create an economic overview paragraph using the data we have on the county, MSA, and state
    try:
        current_period                                    = str(county_employment['period'].iloc[-1])
        current_unemployment                              = county_unemployment_rate['unemployment_rate'].iloc[-1]
        historical_average_unemployment                   = county_unemployment_rate['unemployment_rate'].mean()
        current_gdp_growth                                = ((county_gdp['GDP'].iloc[-1]/county_gdp['GDP'].iloc[-2]) - 1 ) * 100
        current_state_unemployment                        = state_unemployment_rate['unemployment_rate'].iloc[-1]
        largest_industry                                  = county_industry_breakdown['industry_code'].iloc[-1]
        largest_industry_employment_fraction              = county_industry_breakdown['employment_fraction'].iloc[-1]

        #Compare current county unemployment rate to hisorical average
        if current_unemployment > historical_average_unemployment:
            unemployment_above_below_hist_avg = 'above'
        elif current_unemployment < historical_average_unemployment:
            unemployment_above_below_hist_avg = 'below'
        elif current_unemployment == historical_average_unemployment:
            unemployment_above_below_hist_avg = 'equal to'
            
        #Compare current county unemployment rate to state unemployment
        if current_unemployment > current_state_unemployment:
            unemployment_above_below_state = 'above the state level of ' + "{:,.1f}%".format(current_state_unemployment)  
        elif current_unemployment < current_state_unemployment:
            unemployment_above_below_state= 'below the state level of ' +"{:,.1f}%".format(current_state_unemployment)
        elif current_unemployment == current_state_unemployment:
            unemployment_above_below_state = 'equal to the state level'
            

        economic_overview_paragraph = (
                    #GDP Sentence
                    'As of '                  +
                    current_period            +
                    ', '                      +
                    county                    + 
                    """'s"""                  +
                    ' economic output is '    +
                    "{growing_or_contracting}".format(growing_or_contracting = "growing" if (current_gdp_growth >= 0)  else   ('contracting')) +
                    ' at '                    +
                    "{:,.1f}%".format(abs(current_gdp_growth))                                                                                 +
                    ' per year. '             +
                    
                    #Unemployment sentence
                    'The unemployment rate currently sits at '   +
                    "{:,.1f}%".format(current_unemployment)      +
                    ', '                                         +
                    unemployment_above_below_hist_avg            +
                    ' its five-year average '                    +
                    'of '                                        +
                    "{:,.1f}%".format(historical_average_unemployment)                                                                          +
                    ' and '                                      +
                    unemployment_above_below_state               +
                    '. '                                         +

                    #Employment growth and breakdown
                    'The largest industry in terms of employment in '    +
                    county                                               +
                    ' is '                                               +
                    largest_industry                                     +
                    ', which employs '                                   +
                "{:,.1f}%".format(largest_industry_employment_fraction)  +
                    ' of all workers in the County.'

                
                                )

    except Exception as e:
        print(e,'unable to create economic overview paragarph')
        economic_overview_paragraph = ''
        
    #Section 3: Put together a paragraph on the extent to which the county faced job losses during the pandemic
    try:
        february_2020_employment_level = county_employment.loc[(county_employment['year'] == '2020') & (county_employment['periodName'] == 'February')]['Employment'].iloc[-1]
        april_2020_employment_level    = county_employment.loc[(county_employment['year'] == '2020') & (county_employment['periodName'] == 'April')]['Employment'].iloc[-1]    
        pandemic_job_losses            = february_2020_employment_level - april_2020_employment_level 
        pandemic_job_losses_pct        = (pandemic_job_losses/february_2020_employment_level) * 100


        covid_context_pargaraph = (
            'The COVID-19 pandemic slowed economic growth throughout the country, including here in '  + 
             county                                                                                    +
             '. Between February 2020 and April, '                                                     +
             county                                                                                    +
             ' employers shed over '                                                                   + 
             "{:,.0f}".format(pandemic_job_losses)                                                     +
             ' jobs ('                                                                                 +
             "{:,.1f}%".format(pandemic_job_losses_pct)                                                +
             ' of the labor market), as'                                                               +
            ' social distancing protocols were put in place and operating restrictions were imposed. ' + 
            'With the availability of vaccines in early 2021, restrictions eased, and growth returned.'
                                )
    except Exception as e:
        print(e,'unable to create covid context paragraph')
        covid_context_pargaraph = ''

    

    
    #Section 4: Put together our 3 sections and return it
    overview_language = [subject_property_location_language, CBSA_overview_language, wikipeida_summary, wikipeida_economy_summary, covid_context_pargaraph, economic_overview_paragraph]
    return(overview_language)

def CountyEmploymentBreakdownLanguage(county_industry_breakdown):
    print('Writing County Employment Breakdown Langauge')
    
    #Get the largest industries
    largest_industry                                  = county_industry_breakdown['industry_code'].iloc[-1]
    largest_industry_employment                       = county_industry_breakdown['month3_emplvl'].iloc[-1]
    largest_industry_employment_fraction              = county_industry_breakdown['employment_fraction'].iloc[-1]
    
    second_largest_industry                           = county_industry_breakdown['industry_code'].iloc[-2]
    second_largest_industry_employment                = county_industry_breakdown['month3_emplvl'].iloc[-2]
    second_largest_industry_employment_fraction       = county_industry_breakdown['employment_fraction'].iloc[-2]
    

    if len(county_industry_breakdown) > 2:
        third_largest_industry                        = county_industry_breakdown['industry_code'].iloc[-3]
        third_largest_industry_employment             = county_industry_breakdown['month3_emplvl'].iloc[-3]
        third_largest_industry_employment_fraction    = county_industry_breakdown['employment_fraction'].iloc[-3]
        
    else:
        third_largest_industry                        = ''
        third_largest_industry_employment             = ''
        third_largest_industry_employment_fraction    = ''
    

    #Now sort by location quotient to find the highest realative concentration industries
    county_industry_breakdown                           = county_industry_breakdown.sort_values(by=['lq_month3_emplvl'])
    highest_relative_concentration_industry             = county_industry_breakdown['industry_code'].iloc[-1]
    highest_relative_concentration_industry_lq          = county_industry_breakdown['lq_month3_emplvl'].iloc[-1]
    highest_relative_concentration_employment_fraction  = county_industry_breakdown['employment_fraction'].iloc[-1]

    

    #Format Variables
    largest_industry_employment_fraction           = "{:,.1f}%".format(largest_industry_employment_fraction) 
    largest_industry_employment                    = "{:,.0f}".format(largest_industry_employment)

    second_largest_industry_employment_fraction    = "{:,.1f}%".format(second_largest_industry_employment_fraction) 
    second_largest_industry_employment             = "{:,.0f}".format(second_largest_industry_employment)

    third_largest_industry_employment_fraction     = "{:,.1f}%".format(third_largest_industry_employment_fraction) 
    third_largest_industry_employment              = "{:,.0f}".format(third_largest_industry_employment)
    

    county_breakdown_language = (
           #Sentence 1
          'According to the Q'                                                 +
            qcew_qtr                                                           +
            ' '                                                                +
            qcew_year                                                          +
            ' Quarterly Census of Employment and Wages, '                      +
            county                                                             +
            ' employed '                                                       +
            "{:,.0f}".format(county_industry_breakdown['month3_emplvl'].sum()) +
            ' employees, with establishments in the '                          + 
           largest_industry                                                    +
           ', '                                                                +
           second_largest_industry                                             +
           ', and '                                                            +
          third_largest_industry                                               +
           ' industries accounting for the top three employers. '              +

           #Sentence 2
           'These industries employ '                                          +
           largest_industry_employment                                         +
           ' ('                                                                +
            largest_industry_employment_fraction                               +
           '), '                                                               +
            second_largest_industry_employment                                 +
           ' ('                                                                +
           second_largest_industry_employment_fraction                         +
           '), and '                                                           +           
            third_largest_industry_employment                                  +
           ' ('                                                                +
           third_largest_industry_employment_fraction                          +
           ') '                                                                +
           'workers in the County, respectively. '                             +

           #Sentence 3: High concentration industry sentence (only populates if meets threshold)    
            "{high_concentration_sentence}".format(high_concentration_sentence = (county + ' has an especially large share of workers in the ' + highest_relative_concentration_industry + 
             """ industry. In fact, its """                                       + 
            "{:,.1f}%".format(highest_relative_concentration_employment_fraction) +
            ' fraction of workers is ' +  "{:,.1f}".format(highest_relative_concentration_industry_lq) + ' times higher than the National average.'   ) if highest_relative_concentration_industry_lq >= 1.75  else "") 

        )
    return(county_breakdown_language)

def MSAEmploymentBreakdownLanguage(msa_industry_breakdown):
    print('Writing MSA Employment Breakdown Langauge')
    #Get the largest industries
    largest_industry                                  = msa_industry_breakdown['industry_code'].iloc[-1]
    largest_industry_employment                       = msa_industry_breakdown['month3_emplvl'].iloc[-1]
    largest_industry_employment_fraction              = msa_industry_breakdown['employment_fraction'].iloc[-1]
    
    second_largest_industry                           = msa_industry_breakdown['industry_code'].iloc[-2]
    second_largest_industry_employment                = msa_industry_breakdown['month3_emplvl'].iloc[-2]
    second_largest_industry_employment_fraction       = msa_industry_breakdown['employment_fraction'].iloc[-2]
    

    if len(msa_industry_breakdown) > 2:
        third_largest_industry                        = msa_industry_breakdown['industry_code'].iloc[-3]
        third_largest_industry_employment             = msa_industry_breakdown['month3_emplvl'].iloc[-3]
        third_largest_industry_employment_fraction    = msa_industry_breakdown['employment_fraction'].iloc[-3]
        
    else:
        third_largest_industry                        = ''
        third_largest_industry_employment             = ''
        third_largest_industry_employment_fraction    = ''
    

    #Now sort by location quotient to find the highest realative concentration industries
    msa_industry_breakdown                              = msa_industry_breakdown.sort_values(by=['lq_month3_emplvl'])
    highest_relative_concentration_industry             = msa_industry_breakdown['industry_code'].iloc[-1]
    highest_relative_concentration_industry_lq          = msa_industry_breakdown['lq_month3_emplvl'].iloc[-1]
    highest_relative_concentration_employment_fraction  = msa_industry_breakdown['employment_fraction'].iloc[-1]

    

    #Format Variables
    largest_industry_employment_fraction           = "{:,.1f}%".format(largest_industry_employment_fraction) 
    largest_industry_employment                    = "{:,.0f}".format(largest_industry_employment)

    second_largest_industry_employment_fraction    = "{:,.1f}%".format(second_largest_industry_employment_fraction) 
    second_largest_industry_employment             = "{:,.0f}".format(second_largest_industry_employment)

    third_largest_industry_employment_fraction     = "{:,.1f}%".format(third_largest_industry_employment_fraction) 
    third_largest_industry_employment              = "{:,.0f}".format(third_largest_industry_employment)
    

    msa_breakdown_language = (
           #Sentence 1
          'According to the Q'                                              +
            qcew_qtr                                                        +
            ' '                                                             +
            qcew_year                                                       +
            ' Quarterly Census of Employment and Wages, '                   +
            cbsa_name                                                       +
            ' employed '                                                    +
            "{:,.0f}".format(msa_industry_breakdown['month3_emplvl'].sum()) +
            ' private employees, with establishments in the '               +  
           largest_industry                                                 +
           ', '                                                             +
           second_largest_industry                                          +
           ', and '                                                         +
          third_largest_industry                                            +
           ' industries accounting for the top three employers. '           +
           
           #Sentence 2
           'These industries employ '                                       +
           largest_industry_employment                                      +
           ' ('                                                             +
            largest_industry_employment_fraction                            +
           '), '                                                            +
            second_largest_industry_employment                              +
           ' ('                                                             +
           second_largest_industry_employment_fraction                      +
           '), and '                                                        + 
            third_largest_industry_employment                               +
           ' ('                                                             +
           third_largest_industry_employment_fraction                       +
           ') '                                                             +
           'workers in the Metro, respectively. '                           +
            
           #Sentence 3: High concentration industry sentence (only populates if meets threshold)    
            "{high_concentration_sentence}".format(high_concentration_sentence = (cbsa_name + ' has an especially large share of workers in the ' + highest_relative_concentration_industry + """ industry. In fact, its """ +  "{:,.1f}%".format(highest_relative_concentration_employment_fraction) + ' fraction of workers is ' +  "{:,.1f}".format(highest_relative_concentration_industry_lq) + ' times higher than the National average.'   ) if highest_relative_concentration_industry_lq >= 1.75  else "") 

        )
    return(msa_breakdown_language)
  
def UnemploymentLanguage():
    print('Writing Unemployment Langauge')
    try:
        latest_period                      = county_employment['period'].iloc[-1]
        latest_county_employment           = county_employment['Employment'].iloc[-1]
        one_year_ago_county_employment     = county_employment['Employment'].iloc[-13]
        one_year_percent_employment_change = ((latest_county_employment/one_year_ago_county_employment) -1 ) * 100
       
        if one_year_percent_employment_change > 0:
            up_or_down = 'up ' + "{:,.0f}%".format(abs(one_year_percent_employment_change))

            if  one_year_percent_employment_change < 1:
                up_or_down = 'up ' + "{:,.1f}%".format(abs(one_year_percent_employment_change))


        elif one_year_percent_employment_change < 0:
            up_or_down = 'down ' + "{:,.0f}%".format(abs(one_year_percent_employment_change)) 

            if one_year_percent_employment_change < -1:
                up_or_down = 'down ' + "{:,.1f}%".format(abs(one_year_percent_employment_change)) 

        elif one_year_percent_employment_change == 0:
            up_or_down = 'unchanged'

        
        #Get latest state and county unemployment rate
        latest_county_unemployment          = county_unemployment_rate['unemployment_rate'].iloc[-1]

        if cbsa != '':
            latest_msa_unemployment         = msa_unemployment_rate['unemployment_rate'].iloc[-1]
            
            #See if county unemployment rate is above or metro  rate
            if latest_msa_unemployment > latest_county_unemployment:
                msa_county_unemployment_above_or_below = 'below'
            elif latest_msa_unemployment < latest_county_unemployment:
                msa_county_unemployment_above_or_below = 'above'
            elif latest_msa_unemployment == latest_county_unemployment:
                msa_county_unemployment_above_or_below = 'equal to'
            
            
            #Check how far apart msa and county unemployment rates are
            if abs(latest_msa_unemployment - latest_county_unemployment) > 1.5:
                msa_county_unemployment_difference = 'considerably '
            elif abs(latest_msa_unemployment - latest_county_unemployment) > 0:
                msa_county_unemployment_difference = 'just slightly '
            else:
                msa_county_unemployment_difference = ''


        pre_pandemic_unemployment_df        = county_unemployment_rate.loc[(county_unemployment_rate['periodName'] =='February') & (county_unemployment_rate['year'] == '2020')]
        pre_pandemic_unemployment           = pre_pandemic_unemployment_df['unemployment_rate'].iloc[-1]
        one_year_ago_county_unemployment    = county_unemployment_rate['unemployment_rate'].iloc[-13]
        latest_state_unemployment           = state_unemployment_rate['unemployment_rate'].iloc[-1]

        #Change in unemployment rate over past year
        if latest_county_unemployment == one_year_ago_county_unemployment:
            unemployment_change = 'remained stable over the past year at '
        elif latest_county_unemployment > one_year_ago_county_unemployment:
            unemployment_change = 'expanded over the past year to the current rate of '
        elif latest_county_unemployment < one_year_ago_county_unemployment:
            unemployment_change = 'compressed over the past year to the current rate of '

        #See if county unemployment rate is above or below state rate
        if latest_state_unemployment > latest_county_unemployment:
            state_county_unemployment_above_or_below = 'below'
        elif latest_state_unemployment < latest_county_unemployment:
            state_county_unemployment_above_or_below = 'above'
        elif latest_state_unemployment == latest_county_unemployment:
            state_county_unemployment_above_or_below = 'equal to'
        
        #See if county unemployment rate is above pre-pandemic levels
        if pre_pandemic_unemployment > latest_county_unemployment:
            pre_pandemic_unemp_above_or_below = 'has moved below'
        elif pre_pandemic_unemployment < latest_county_unemployment:
            pre_pandemic_unemp_above_or_below = 'remains above'
        elif pre_pandemic_unemployment == latest_county_unemployment:
            pre_pandemic_unemp_above_or_below = 'is equal to'

        #Check how far apart state and county unemployment rates are
        if abs(latest_state_unemployment - latest_county_unemployment) > 1.5:
            state_county_unemployment_difference = 'considerably '
        elif abs(latest_state_unemployment - latest_county_unemployment) > 0:
            state_county_unemployment_difference = 'just slightly '
        else:
            state_county_unemployment_difference = ''

        #Get unemployment rate and emp level in Feb 2020 and Apirl 2020 
        february_2020_employment_level = county_employment.loc[(county_employment['year'] == '2020') & (county_employment['periodName'] == 'February')]['Employment'].iloc[-1]
        april_2020_employment_level    = county_employment.loc[(county_employment['year'] == '2020') & (county_employment['periodName'] == 'April')]['Employment'].iloc[-1]    

        february_2020_unemployment_rate = county_unemployment_rate.loc[(county_unemployment_rate['year'] == '2020') & (county_unemployment_rate['periodName'] == 'February')]['unemployment_rate'].iloc[-1]
        april_2020_unemployment_rate    = county_unemployment_rate.loc[(county_unemployment_rate['year'] == '2020') & (county_unemployment_rate['periodName'] == 'April')]['unemployment_rate'].iloc[-1]    

        pandemic_job_losses             = february_2020_employment_level - april_2020_employment_level 
        pandemic_job_losses_pct         = (pandemic_job_losses/february_2020_employment_level) * 100
        
        #Format variables
        latest_county_employment        = "{:,}".format(latest_county_employment)
        latest_county_unemployment      = "{:,.1f}%".format(latest_county_unemployment)
        latest_state_unemployment       = "{:,.1f}%".format(latest_state_unemployment)


        if cbsa == '':

            unemployment_language=( 
                    #Sentence 1: Discuss covid job losses
                    'At the onset of the pandemic last spring, ' + 
                    county                                       + 
                    ' area employers shed '                      +
                     "{:,.1f}%".format(pandemic_job_losses_pct)  +
                    ' of its workforce, expanding the unemployment rate from '                   +
                    "{:,.1f}%".format(february_2020_unemployment_rate) + ' in February 2020 to ' + 
                    "{:,.1f}%".format(april_2020_unemployment_rate) + ' just two months later. ' + 
                    
                    
                    #Sentence 2: Discuss current unemployment
                    'The unemployment rate in '              +
                    county                                   + 
                    ' has '                                  +
                    unemployment_change                      +
                    latest_county_unemployment               +
                    ', '                                     +        
                    state_county_unemployment_difference     +
                    state_county_unemployment_above_or_below +
                    ' the '                                  +
                    state_name                               + 
                    ' rate'                                  +
                    "{state_unemployment}".format(state_unemployment =(' of '      + latest_state_unemployment ) if (latest_county_unemployment != latest_state_unemployment)  else   ('')) +           
                    '. '                                     +


                    #Sentence 3: Discuss growth in total employment
                    'As of '                        +
                    latest_period                   +
                    ', total employment is '        +
                    up_or_down                      +
                    ' on a year-over-year basis. '  +

                    #Sentence 4: Is unemployment above or below pre-pandemic levels?
                    'The unemployment rate '                      +
                    pre_pandemic_unemp_above_or_below             +
                    ' its pre-pandemic level (Feb 2020) of '      +
                    "{:,.1f}%".format(pre_pandemic_unemployment)  +
                    '.'         
                                )
        else:
            latest_msa_unemployment        = "{:,.1f}%".format(latest_msa_unemployment)
            
            unemployment_language = (#Sentence 1: Discuss covid job losses
                                    'At the onset of the pandemic last spring, ' +
                                    county                                       + 
                                    ' area employers shed '                      + 
                                    "{:,.1f}%".format(pandemic_job_losses_pct)   + 
                                    ' of its workforce, expanding the unemployment rate from ' +
                                    "{:,.1f}%".format(february_2020_unemployment_rate)         + 
                                    ' in February 2020 to '                                    + 
                                    "{:,.1f}%".format(april_2020_unemployment_rate)            + 
                                    ' just two months later. '                                 + 
                                    
                                    #Sentence 2: Discuss current unemployment
                                    'The unemployment rate in '                                +
                                    county                                                     +     
                                    ' has '                                                    +
                                    unemployment_change                                        +
                                    latest_county_unemployment                                 +
                                    ', '                                                       +        
                                    msa_county_unemployment_difference                         +
                                    msa_county_unemployment_above_or_below                     +
                                    ' the '                                                    +
                                    cbsa_name                                                  + 
                                    ' rate'                                                    +
                                    "{msa_unemployment}".format(msa_unemployment =(' of '      + latest_msa_unemployment ) if (latest_county_unemployment != latest_msa_unemployment)  else   ('')) +           
                                    '. '                                                       +


                                    #Sentence 3: Discuss growth in total employment
                                    'As of '                                                    +
                                    latest_period                                               +
                                    ', total employment is '                                    +
                                    up_or_down                                                  +
                                    ' on a year-over-year basis. '                              +

                                    #Sentence 4: Is unemployment above or below pre-pandemic levels?
                                    'The unemployment rate '                                    +
                                    pre_pandemic_unemp_above_or_below                           +
                                    ' its pre-pandemic level (Feb 2020) of '                    +
                                    "{:,.1f}%".format(pre_pandemic_unemployment)                +
                                    '.'         
                                    )
    except Exception as e:
        print(e,'unable to create unemployment language')
        unemployment_language = ''
    
    return([unemployment_language])

def CountyEmploymentGrowthLanguage(county_industry_breakdown):
    print('Writing County Employment Growth Langauge')
    
    #Track employment growth over the past 5 years
    latest_county_employment               = county_industry_breakdown['month3_emplvl'].sum()
    five_years_ago_county_employment       = county_industry_breakdown['lagged_month3_emplvl'].sum()
    five_year_county_employment_growth_pct = ((latest_county_employment/five_years_ago_county_employment) - 1 ) * 100
    five_year_county_employment_growth     = (latest_county_employment - five_years_ago_county_employment) 


    #See if 5 year county employment expaded or contracted
    if five_year_county_employment_growth > 0:
        five_year_county_employment_expand_or_contract = 'expand'
    elif five_year_county_employment_growth < 0:
        five_year_county_employment_expand_or_contract = 'compress'
    else:
        five_year_county_employment_expand_or_contract = 'remained stable'
    


    #Format 5 year county employment growth variables
    five_year_county_employment_growth_pct = "{:,.1f}%".format(abs(five_year_county_employment_growth_pct))
    five_year_county_employment_growth     = "{:,.0f}".format(five_year_county_employment_growth)

    #Drop the industries where employment growth cant be measured properly and the unclassified industry
    county_industry_breakdown              = county_industry_breakdown.loc[county_industry_breakdown['industry_code'] != 'Unclassified']


    #Get the fastest and slowest growing industries
    county_industry_breakdown5y                = county_industry_breakdown.loc[(county_industry_breakdown['emp_growth_invalid'] != 1) ] 
    county_industry_breakdown5y                = county_industry_breakdown5y.sort_values(by=['Employment Growth'])
    fastest_growing_industry_5y                = county_industry_breakdown5y['industry_code'].iloc[-1]
    second_fastest_growing_industry_5y         = county_industry_breakdown5y['industry_code'].iloc[-2]
    
    if len(county_industry_breakdown5y) > 2:
        third_fastest_growing_industry_5y         = county_industry_breakdown5y['industry_code'].iloc[-3]
    else:
        third_fastest_growing_industry_5y         = ''

    slowest_growing_industry_5y              = county_industry_breakdown5y['industry_code'].iloc[0]
    fastest_growth_industry_5y               = county_industry_breakdown5y['Employment Growth'].iloc[-1]
    second_fastest_growth_industry_5y        = county_industry_breakdown5y['Employment Growth'].iloc[-2]
    
    if len(county_industry_breakdown5y) > 2:
        third_fastest_growth_industry_5y         = county_industry_breakdown5y['Employment Growth'].iloc[-3]
    else:
        third_fastest_growth_industry_5y         = ''


    slowest_growth_industry_5y               = county_industry_breakdown5y['Employment Growth'].iloc[0]
    
    
    #Describe the growth of the slowest growing industry
    if slowest_growth_industry_5y < 0:
        slowest_growth_industry_5y_description = 'collapse'
    elif     slowest_growth_industry_5y >= 0:
        slowest_growth_industry_5y_description = 'grow'

    #Format Variables
    fastest_growth_industry_5y               = "{:,.1f}%".format(fastest_growth_industry_5y)
    second_fastest_growth_industry_5y        = "{:,.1f}%".format(second_fastest_growth_industry_5y)
    slowest_growth_industry_5y               = "{:,.1f}%".format(abs(slowest_growth_industry_5y))

    if len(county_industry_breakdown) > 2:
        third_fastest_growth_industry_5y         = "{:,.1f}%".format(third_fastest_growth_industry_5y)

    county_industry_breakdown                = county_industry_breakdown.loc[(county_industry_breakdown['one_year_emp_growth_invalid'] != 1)] 

    #See if all industries lost employment over the past year (or most or some)
    county_industry_breakdown_employment_lossers           = county_industry_breakdown.loc[county_industry_breakdown['1 Year Employment Growth'] < 0]  #Cut down to industries that lost employees
    county_industry_breakdown_employment_winners           = county_industry_breakdown.loc[county_industry_breakdown['1 Year Employment Growth'] >= 0] #Cut down to industries that gained employees

    if len(county_industry_breakdown_employment_lossers) == len(county_industry_breakdown): #all industries lose employment over last year
        employment_loss_1year_all_most_some                     = 'all'
    elif len(county_industry_breakdown_employment_winners) == len(county_industry_breakdown): #no industries lose employment over last year
        employment_loss_1year_all_most_some                     = 'no'
    elif len(county_industry_breakdown_employment_lossers)/len(county_industry_breakdown) >= 0.5:#most industries lose employment over last year
        employment_loss_1year_all_most_some                     = 'most'
    elif len(county_industry_breakdown_employment_lossers) > 0:
        employment_loss_1year_all_most_some                     = 'some'                        #some industries lose employment over last year
    else:
        employment_loss_1year_all_most_some                     = '[all/most/some]'


    #Get industry that has grown the slowest over the last year in the county
    county_industry_breakdown                 = county_industry_breakdown.sort_values(by=['1 Year Employment Growth'])
    slowest_growing_industry_1y               = county_industry_breakdown['industry_code'].iloc[0]
    slowest_growth_industry_1y                = county_industry_breakdown['1 Year Employment Growth'].iloc[0]
    slowest_growth_industry_1y                = "{:,.1f}%".format(slowest_growth_industry_1y)


    county_employment_growth_language = (
                                        'According to the Q'                           +
                                        qcew_qtr                                       +
                                        ' '                                            +
                                        qcew_year                                      +
                                        ' Quarterly Census of Employment and Wages, '  +
                                        county                                         +
                                        ' has experienced private employment '         +
                                        five_year_county_employment_expand_or_contract +
                                        ' '                                            +
                                        five_year_county_employment_growth_pct         +
                                        ' ('                                           +
                                        five_year_county_employment_growth             +
                                        ') '                                           +
                                        'in total over the last five years. '          +
                                        'During that time, the '                       +
                                        fastest_growing_industry_5y                    +
                                        ', '                                           +
                                        second_fastest_growing_industry_5y             +
                                        ', and '                                       +
                                        third_fastest_growing_industry_5y              +
                                        ' industries saw the strongest growth, expanding ' +
                                        fastest_growth_industry_5y                     +
                                        ', '                                           +
                                        second_fastest_growth_industry_5y              + 
                                        ', and '                                       +
                                        third_fastest_growth_industry_5y               +
                                        ', respectively.'                              +
                                        ' Meanwhile, the '                             +
                                        slowest_growing_industry_5y                    +
                                        ' Industry has experienced employment '        +
                                        slowest_growth_industry_5y_description         +
                                        ' '                                            +
                                        slowest_growth_industry_5y                     +
                                        ' over the previous five years.'              
                                        
                                        )


    return(county_employment_growth_language)

def MSAEmploymentGrowthLanguage(msa_industry_breakdown):
    print('Writing MSA Employment Growth Langauge')
    
    #Create MSA paragarph if applicable
    if cbsa != '':
        #Track employment growth over the past 5 years
        latest_msa_employment               = msa_industry_breakdown['month3_emplvl'].sum()
        five_years_ago_msa_employment       = msa_industry_breakdown['lagged_month3_emplvl'].sum()
        five_year_msa_employment_growth_pct = ((latest_msa_employment/five_years_ago_msa_employment) - 1 ) * 100
        five_year_msa_employment_growth     = (latest_msa_employment - five_years_ago_msa_employment) 


        #See if 5 year county employment expaded or contracted
        if five_year_msa_employment_growth > 0:
            five_year_msa_employment_expand_or_contract = 'expand'
        elif five_year_msa_employment_growth < 0:
            five_year_msa_employment_expand_or_contract = 'compress'
        else:
            five_year_msa_employment_expand_or_contract = 'remained stable'
        


        #Format 5 year county employment growth variables
        five_year_msa_employment_growth_pct = "{:,.1f}%".format(abs(five_year_msa_employment_growth_pct))
        five_year_msa_employment_growth     = "{:,.0f}".format(five_year_msa_employment_growth)

        #Drop the industries where employment growth cant be measured properly and the unclassified industry
        msa_industry_breakdown              = msa_industry_breakdown.loc[msa_industry_breakdown['industry_code'] != 'Unclassified']


        #Get the fastest and slowest growing industries
        msa_industry_breakdown5y                     = msa_industry_breakdown.loc[(msa_industry_breakdown['emp_growth_invalid'] != 1) ] 
        msa_industry_breakdown5y                     = msa_industry_breakdown5y.sort_values(by=['Employment Growth'])
        fastest_growing_industry_5y                  = msa_industry_breakdown5y['industry_code'].iloc[-1]
        second_fastest_growing_industry_5y           = msa_industry_breakdown5y['industry_code'].iloc[-2]
        
        if len(msa_industry_breakdown5y) > 2:
            third_fastest_growing_industry_5y         = msa_industry_breakdown5y['industry_code'].iloc[-3]
        else:
            third_fastest_growing_industry_5y         = ''

        slowest_growing_industry_5y                   = msa_industry_breakdown5y['industry_code'].iloc[0]

        fastest_growth_industry_5y                   = msa_industry_breakdown5y['Employment Growth'].iloc[-1]
        second_fastest_growth_industry_5y            = msa_industry_breakdown5y['Employment Growth'].iloc[-2]
        
        if len(msa_industry_breakdown5y) > 2:
            third_fastest_growth_industry_5y         = msa_industry_breakdown5y['Employment Growth'].iloc[-3]
        else:
            third_fastest_growth_industry_5y         = ''


        slowest_growth_industry_5y                   = msa_industry_breakdown5y['Employment Growth'].iloc[0]
        
        
        #Describe the growth of the slowest growing industry
        if slowest_growth_industry_5y < 0:
            slowest_growth_industry_5y_description = 'collapse'
        elif     slowest_growth_industry_5y >= 0:
            slowest_growth_industry_5y_description = 'grow'

        #Format Variables
        fastest_growth_industry_5y               = "{:,.1f}%".format(fastest_growth_industry_5y)
        second_fastest_growth_industry_5y        = "{:,.1f}%".format(second_fastest_growth_industry_5y)
        slowest_growth_industry_5y               = "{:,.1f}%".format(abs(slowest_growth_industry_5y))

        if len(msa_industry_breakdown) > 2:
            third_fastest_growth_industry_5y    = "{:,.1f}%".format(third_fastest_growth_industry_5y)

        msa_industry_breakdown                  = msa_industry_breakdown.loc[(msa_industry_breakdown['one_year_emp_growth_invalid'] != 1)] 

        #See if all industries lost employment over the past year (or most or some)
        msa_industry_breakdown_employment_lossers           = msa_industry_breakdown.loc[msa_industry_breakdown['1 Year Employment Growth'] < 0]  #Cut down to industries that lost employees
        msa_industry_breakdown_employment_winners           = msa_industry_breakdown.loc[msa_industry_breakdown['1 Year Employment Growth'] >= 0] #Cut down to industries that gained employees

        if len(msa_industry_breakdown_employment_lossers) == len(msa_industry_breakdown): #all industries lose employment over last year
            employment_loss_1year_all_most_some                     = 'all'
        elif len(msa_industry_breakdown_employment_winners) == len(msa_industry_breakdown): #no industries lose employment over last year
            employment_loss_1year_all_most_some                     = 'no'
        elif len(msa_industry_breakdown_employment_lossers)/len(msa_industry_breakdown) >= 0.5:#most industries lose employment over last year
            employment_loss_1year_all_most_some                     = 'most'
        elif len(msa_industry_breakdown_employment_lossers) > 0:
            employment_loss_1year_all_most_some                     = 'some'                        #some industries lose employment over last year
        else:
            employment_loss_1year_all_most_some                     = '[all/most/some]'


        #Get industry that has grown the slowest over the last year in the county
        msa_industry_breakdown                    = msa_industry_breakdown.sort_values(by=['1 Year Employment Growth'])
        slowest_growing_industry_1y               = msa_industry_breakdown['industry_code'].iloc[0]
        slowest_growth_industry_1y                = msa_industry_breakdown['1 Year Employment Growth'].iloc[0]
        slowest_growth_industry_1y                = "{:,.1f}%".format(slowest_growth_industry_1y)



        msa_employment_growth_language = (
                                        'According to the Q' +
                                        qcew_qtr                                           +
                                        ' '                                                +
                                        qcew_year                                          +
                                        ' Quarterly Census of Employment and Wages, '      +
                                        cbsa_name + ' Metro '                              +
                                        'has experienced private employment '              +
                                        five_year_msa_employment_expand_or_contract        +
                                        ' '                                                +
                                        five_year_msa_employment_growth_pct                +
                                        ' ('                                               +
                                        five_year_msa_employment_growth                    +
                                        ') '                                               +
                                        'in total over the last five years. '              +
                                        'During that time, the '                           +
                                        fastest_growing_industry_5y                        +
                                        ', '                                               +
                                        second_fastest_growing_industry_5y                 +
                                        ', and '                                           +
                                        third_fastest_growing_industry_5y                  +
                                        ' industries saw the strongest growth, expanding ' +
                                        fastest_growth_industry_5y                         +
                                        ', '                                               +
                                        second_fastest_growth_industry_5y                  + 
                                        ', and '                                           +
                                        third_fastest_growth_industry_5y                   +
                                        ', respectively.'                                  +
                                        ' Meanwhile, the '                                 +
                                        slowest_growing_industry_5y                        +
                                        ' Industry has experienced employment '            +
                                        slowest_growth_industry_5y_description             +
                                        ' '                                                +
                                        slowest_growth_industry_5y                         +
                                        ' over the previous five years.'
                                        )
    else:
        msa_employment_growth_language = ''

    return(msa_employment_growth_language)

def ProductionLanguage(county_data_frame,msa_data_frame,state_data_frame):
    print('Writing Production Langauge')
    county_data_frame['Period'] = county_data_frame['Period'].dt.strftime('%m/%d/%Y')
    latest_period               = county_data_frame['Period'].iloc[-1]
    latest_period               = latest_period[-4:]
    latest_county_gdp           = county_data_frame['GDP'].iloc[-1]
    latest_county_gdp           = millify(latest_county_gdp)
    latest_county_gdp           = "$" + latest_county_gdp
    
    latest_county_gdp_growth    = ((county_data_frame['GDP'].iloc[-1]/county_data_frame['GDP'].iloc[-2]) - 1) * 100
    

    #determine how to describe GDP growth 
    if latest_county_gdp_growth > 3.5:
        gdp_growth_description = 'strong'
    elif latest_county_gdp_growth < 3.5 and (latest_county_gdp_growth >= 1.25):
        gdp_growth_description = 'steady'
    elif latest_county_gdp_growth < 1.25 and (latest_county_gdp_growth >= 0.5):
        gdp_growth_description = 'modest'
    elif latest_county_gdp_growth < 0.5 and (latest_county_gdp_growth >= 0.25):
        gdp_growth_description = 'weak'
    elif latest_county_gdp_growth < 0.25 and (latest_county_gdp_growth >= 0):
        gdp_growth_description = 'stagnant'
    elif latest_county_gdp_growth < 0 :
        gdp_growth_description = 'negative'
    else:
        gdp_growth_description = '[stagnant/steady/strong/weak/negative]'


    if  isinstance(msa_data_frame, pd.DataFrame) == True and msa_data_frame['GDP'].equals(county_data_frame['GDP']) == False:
        msa_data_frame          = msa_data_frame.loc[msa_data_frame['Period'] <= (county_data_frame['Period'].max()) ]
        latest_msa_gdp_growth   = ((msa_data_frame['GDP'].iloc[-1]/msa_data_frame['GDP'].iloc[-2]) - 1) * 100
        latest_msa_gdp_growth   =  "{:,.1f}%".format(latest_msa_gdp_growth)
        msa_or_state_gdp_growth = latest_msa_gdp_growth
        msa_or_state            = 'Metro'
    else:
        state_data_frame        = state_data_frame.loc[state_data_frame['Period'] <= (county_data_frame['Period'].max()) ]
        latest_state_gdp_growth = ((state_data_frame['GDP'].iloc[-1]/state_data_frame['GDP'].iloc[-2]) - 1) * 100
        latest_state_gdp_growth =  "{:,.1f}%".format(latest_state_gdp_growth)
        msa_or_state_gdp_growth = latest_state_gdp_growth
        msa_or_state            = 'State'

    #Fomrmat variables
    latest_county_gdp_growth =  "{:,.1f}%".format(latest_county_gdp_growth)

    production_language = ('While GDP data at the county level is not yet available, '      +
                            latest_period                                                   +
                            ' data from the U.S. Bureau of Economic Analysis points to '    +
                            gdp_growth_description                                          +
                            ' growth for '                                                  +
                            county                                                          +
                            ', which produced ~'                                            +
                            latest_county_gdp                                               +
                            ' of output that year, '                                        +
                            'representing an annual change of '                             +
                            latest_county_gdp_growth                                        +
                            ' compared to '                                                 +
                            msa_or_state_gdp_growth                                         +
                            ' for the '                                                     +
                            msa_or_state                                                    +
                                '.' 
                           )

    boiler_plate_econ_language = ('Economic activity has slowed after historical annual growth of 6.7% in Q2 2021, softening to 2.3% for the third quarter. '                                                           +
                                  'The slowdown in third quarter GDP reflected the continued economic impact of the COVID-19 pandemic. '                                                                                +
                                  'A resurgence of COVID-19 cases resulted in new restrictions and delays in the reopening of establishments in some parts of the country. '                                            +
                                  'Supply-chain disruptions such as delays at U.S. ports and international manufacturing issues contributed to a sharp increase in inflation and pose a risk to the economic outlook. ' +
                                  'Despite supply-side challenges, many economic observers expect that the economy regained momentum in the final quarter and is well positioned for growth in 2022. '
                                 )
    
    return[boiler_plate_econ_language, production_language]

def IncomeLanguage():
    print('Writing Income Langauge')   
    try:
        #Get latest county income level
        latest_county_income                                  = round(county_pci['Per Capita Personal Income'].iloc[-1])
        latest_county_year                                    = str(county_pci['Period'].iloc[-1])[0:4]  

        #Get County growth rates
        county_pci['Per Capita Personal Income_1year_growth'] =  (((county_pci['Per Capita Personal Income']/county_pci['Per Capita Personal Income'].shift(1))  - 1) * 100)/1
        county_pci['Per Capita Personal Income_3year_growth'] =  (((county_pci['Per Capita Personal Income']/county_pci['Per Capita Personal Income'].shift(3))   - 1) * 100)/3
        county_pci['Per Capita Personal Income_5year_growth'] =  (((county_pci['Per Capita Personal Income']/county_pci['Per Capita Personal Income'].shift(5))   - 1) * 100)/5

        county_1y_growth                                      = county_pci.iloc[-1]['Per Capita Personal Income_1year_growth'] 
        county_3y_growth                                      = county_pci.iloc[-1]['Per Capita Personal Income_3year_growth'] 
        county_5y_growth                                      = county_pci.iloc[-1]['Per Capita Personal Income_5year_growth']

        #See if 3 year income growth rate was higher or lower than 5 year growth rate
        if county_3y_growth > county_5y_growth:
            three_five_year_county_declined_or_expanded = 'expanded'
        elif county_3y_growth < county_5y_growth:
            three_five_year_county_declined_or_expanded = 'declined'
        elif county_3y_growth == county_5y_growth:
            three_five_year_county_declined_or_expanded = 'remained stable'

        #Get national growth rates
        national_pci_restricted = national_pci.loc[national_pci['Period'] <= (county_pci['Period'].max())] #Restrict to last year of county data to marke sure we comapre appples to apples
        national_pci_restricted['Per Capita Personal Income_1year_growth'] =  (((national_pci_restricted['Per Capita Personal Income']/national_pci_restricted['Per Capita Personal Income'].shift(1))  - 1) * 100)/1
        national_pci_restricted['Per Capita Personal Income_3year_growth'] =  (((national_pci_restricted['Per Capita Personal Income']/national_pci_restricted['Per Capita Personal Income'].shift(3))   - 1) * 100)/3
        national_pci_restricted['Per Capita Personal Income_5year_growth'] =  (((national_pci_restricted['Per Capita Personal Income']/national_pci_restricted['Per Capita Personal Income'].shift(5))   - 1) * 100)/5

        national_1y_growth  = national_pci_restricted.iloc[-1]['Per Capita Personal Income_1year_growth']
        national_3y_growth  = national_pci_restricted.iloc[-1]['Per Capita Personal Income_3year_growth'] 
        national_5y_growth  = national_pci_restricted.iloc[-1]['Per Capita Personal Income_5year_growth'] 
        
        
        
        #See if 3 year income growth rate was higher or lower than nation
        if county_3y_growth > national_3y_growth:
            county_vs_nation_3y_exceeds_lags = 'exceeds'
        elif county_3y_growth < national_3y_growth:
            county_vs_nation_3y_exceeds_lags = 'lags'
        elif county_3y_growth == national_3y_growth:
            county_vs_nation_3y_exceeds_lags = 'is equal to'



        income_language = (#Sentence 1
                            'Going back five years, '                     +
                            county                                        +
                            """ residents' per capita personal income """ +
                            'has expanded '                               + 
                            "{:,.1f}%".format(county_5y_growth)           + 
                            ' per annum to the '                          +
                            latest_county_year                            + 
                            ' level of '                                  +
                            "${:,}".format(latest_county_income)          +
                            '. '                                          +

                        #Sentence 2    
                        'Over the past three years, growth has '    +
                        three_five_year_county_declined_or_expanded +
                        ', growing '                                +
                            "{:,.1f}%".format(county_3y_growth)     + 
                            ' per annum since '                     + 
                            str(int(latest_county_year) - 3)        + 
                            '. '                                    +
                        
                        #Sentence 3
                        'This growth rate '                         + 
                        county_vs_nation_3y_exceeds_lags            +
                            ' the Nation, which has '               +
                            'expanded'                              +     
                            ' '                                     +   
                            "{:,.1f}%".format(national_3y_growth)   + 
                            ' per year over the last three years. ' 
                )
    except Exception as e:
        print(e,'unable to get income language') 
        income_language = ''

    return([income_language])   

def PopulationLanguage(national_resident_pop):
    print('Writing Demographic Langauge')
    try:
            county_resident_pop['Period'] = county_resident_pop['Period'].dt.strftime('%m/%d/%Y')
            latest_period                 = county_resident_pop['Period'].iloc[-1]
            latest_period                 = latest_period[-4:]
            latest_county_pop             = round(county_resident_pop['Resident Population'].iloc[-1])
            latest_county_pop             = "{:,}".format(latest_county_pop)

            county_resident_pop['Resident Population_1year_growth']  =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
            county_resident_pop['Resident Population_5year_growth']  =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
            county_resident_pop['Resident Population_10year_growth'] =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10

            county_1y_growth  = county_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
            county_5y_growth  = county_resident_pop.iloc[-1]['Resident Population_5year_growth'] 
            county_10y_growth = county_resident_pop.iloc[-1]['Resident Population_10year_growth']

            #Determine how to describe 10 year county population growth
            if county_10y_growth > 0:
                county_10y_expand_or_compress =  'expanded'
            elif county_10y_growth < 0:
                county_10y_expand_or_compress =  'compressed'
            else:
                county_10y_expand_or_compress =  '[remained stagnant with limited growth of ]'
            
            #Determine how to describe 5 year county population growth
            if county_5y_growth > 0:
                county_5y_expand_or_compress =  'growing'
            elif county_5y_growth < 0:
                county_5y_expand_or_compress =  'contracting'
            else:
                county_5y_expand_or_compress =  '[growing/contracting]'

            #Determine if 5 year growth is slower of faster than 10 year growth
            if county_5y_growth > county_10y_growth:
                growth_declined_or_expanded = 'expanded'
            elif county_5y_growth < county_10y_growth:
                growth_declined_or_expanded = 'declined'
            elif county_5y_growth == county_10y_growth:
                growth_declined_or_expanded = 'remained stable'
            else:
                growth_declined_or_expanded = '[declined/expanded]'


            #Make sure we are comparing same years for calculating growth rates for county and USA
            national_resident_pop['Resident Population_1year_growth'] =  (((national_resident_pop['Resident Population']/national_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
            national_resident_pop['Resident Population_5year_growth'] =  (((national_resident_pop['Resident Population']/national_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
            national_resident_pop['Resident Population_10year_growth'] =  (((national_resident_pop['Resident Population']/national_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10
            national_resident_pop = national_resident_pop.loc[national_resident_pop['Period'] <= (county_resident_pop['Period'].max())]

            national_1y_growth  = national_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
            national_5y_growth  = national_resident_pop.iloc[-1]['Resident Population_5year_growth'] 
            national_10y_growth = national_resident_pop.iloc[-1]['Resident Population_10year_growth']

            #Determine if county 5 year growth was slower or faster than national growth
            if county_5y_growth > national_5y_growth:
                county_5y_slower_or_faster_than_national = 'exceeds'
            elif  county_5y_growth < national_5y_growth:
                county_5y_slower_or_faster_than_national = 'falls short of'
            elif  county_5y_growth == national_5y_growth:
                county_5y_slower_or_faster_than_national = 'is equal to'
            else:
                county_5y_slower_or_faster_than_national = '[falls short of/exceeds]'

        
            county_1y_growth  = "{:,.1f}%".format(county_1y_growth)
            county_5y_growth  = "{:,.1f}%".format(abs(county_5y_growth)) 
            county_10y_growth = "{:,.1f}%".format(abs(county_10y_growth))

            national_1y_growth  = "{:,.1f}%".format(national_1y_growth)
            national_5y_growth  = "{:,.1f}%".format(national_5y_growth) 
            national_10y_growth = "{:,.1f}%".format(national_10y_growth)

            population_language = (#Sentence 1:
                                    'Going back ten years, '      +
                                    county                        +
                                    """'s population has """      +
                                    county_10y_expand_or_compress +
                                     ' '                          +
                                    county_10y_growth             +
                                    ' per annum '                 +
                                    'to the '                     +
                                    latest_period                 +      
                                    ' '                           +
                                    'count of '                   +
                                    latest_county_pop             +
                                    '.'                           +

                                    #Sentence 2:
                                    ' Over the past five years, growth has ' +
                                    growth_declined_or_expanded              +
                                    ', '                                     +
                                    county_5y_expand_or_compress             +
                                    ' '                                      +
                                    county_5y_growth                         +
                                    ' per annum since '                      +
                                    str((int(latest_period) - 5))            +
                                    '.' +

                                    #Sentence 3: 
                                    ' This growth rate '                     +
                                    county_5y_slower_or_faster_than_national +
                                    ' the Nation, which has '                +
                                    'expanded'                               +
                                    ' '                                      +
                                    national_5y_growth                       +
                                    ' per year '                             +
                                    'over the last five years.'          
                                    )
    except Exception as e:
        print(e,'unable to create population language')
        population_language = ''
    
    return([population_language])

def InfrastructureLanguage():
    print('Writing Infrastructure Langauge')

    #Section 1: Grab language on infrastructure from Wikipedia API
    page                      =  wikipedia.page((county + ',' + state))
    infrastructure            =  page.section('Infrastructure')
    transportation            =  page.section('Transportation')
    public_transportation     =  page.section('Public transportation')

    infrastructure_language = [] #this is an empty list we will fill with paragraphs and return 
    
    for wikipedia_section in [infrastructure,transportation,public_transportation]:
        if wikipedia_section != None:
            infrastructure_language.append(wikipedia_section)

    #Section 2: Create basic phrase we can insert if there is nothing from Wikipedia
    infrastructure_boiler_plate = 'The '  + county + ' offers a variety of roadways and public transportation infrastructure. With access to multiple interstate systems, travel time to work is about average both within the state and nationally.'    
    if infrastructure_language == []:
        infrastructure_language.append(infrastructure_boiler_plate)

    return(infrastructure_language)

def WikipediaTransitLanguage(category):
    #Searches through a wikipedia page for a number of section titles and returns the text from them (if any)
    try:
        wikipedia_search_terms_df = pd.read_csv(os.path.join(project_location,'Data','General Data','Wikipedia Transit Related Search Terms.csv'))
        wikipedia_search_terms_df = wikipedia_search_terms_df.loc[wikipedia_search_terms_df['category'] == category]
        search_term_list          = wikipedia_search_terms_df['search term']
        page                      = wikipedia.page((county + ',' + state))

        #Create a bs4 html object from the wikipedia page
        soup                      = BeautifulSoup(page.html(),'html.parser')

        #Loop throuhg each search term looking for text, if we find any, return it as a list
        for search_term in search_term_list:
            langauge_paragraphs = []

            #Loop through every heading in the wikipedia page looking for the section we want
            for heading in soup.find_all(re.compile('^h[1-6]$')):

                #If we find the section we're looking for, pull all the text from the paragraphs
                if search_term in heading.text:
                    para      = heading.find_next_sibling('p')
                    para_text = para.text
                    para_text = re.sub('\[\d+\]', '',para_text) #remove wikipedia citations 
                    langauge_paragraphs.append(para_text)       #Once we have found the relevant section, add all the paragraphs into the list of paragraphs

                    #We found the first paragraph and added it to our list with text paragraphs, now keep looking in case there are multiple paragraphs in the section we wanted
                    while True:
                        para = para.find_next_sibling(['p','h1','h2','h3','h4','h5','h6'])            
                        if para == None:
                            break
                        elif para.name != 'p':
                            break
                        else:
                             paratext = para.text
                             paratext = re.sub('\[\d+\]', '',paratext) #remove wikipedia citations 
                             langauge_paragraphs.append(paratext)
            
            if langauge_paragraphs != []:
                return(langauge_paragraphs)

    except Exception as e:
        print(e)
        return(None)

def HousingLanguage():
    print('Writing Housing Langauge')
    
    if isinstance(county_mlp, pd.DataFrame) == False:
        return([''])
    else:
        current_county_mlp        = county_mlp['Median List Price'].iloc[-1]
        current_county_mlp_period = str(county_mlp['Period'].iloc[-1])[0:7]
        current_county_mlp_period = current_county_mlp_period[5:] + '/' + current_county_mlp_period[0:4]
        yoy_county_mlp_growth     = ((county_mlp['Median List Price'].iloc[-1]/county_mlp['Median List Price'].iloc[-13]) - 1 ) * 100   
        yoy_national_mlp_growth   = ((national_mlp['Median List Price'].iloc[-1]/national_mlp['Median List Price'].iloc[-13]) - 1 ) * 100

        #Determine if county year over year growth was positive or negative
        if yoy_county_mlp_growth > 0:
            increase_or_decrease = 'an increase'
        elif yoy_county_mlp_growth < 0 :
            increase_or_decrease = 'a decrease'
        else:
            increase_or_decrease = 'no change'
        
        #Format variables
        current_county_mlp      = "${:,.0f}".format(current_county_mlp)
        
        boiler_plate_housing_language = """Historically low mortgage rates, the desire for more space, and the ability to work from home have led to the highest number of home sales while historically low inventory levels have pushed values to record highs in most counties and metros across the Nation. """ 


        #If we have the metro realtor data
        if isinstance(msa_mlp, pd.DataFrame) == True:
            yoy_msa_mlp_growth = ((msa_mlp['Median List Price'].iloc[-1]/msa_mlp['Median List Price'].iloc[-13]) - 1 ) * 100

            housing_langage =   (#Sentence 1
                                "In "                           +                                           
                                county                          +
                                ', Realtor.com data points to ' +
                                "{growth_description}".format(growth_description = "continued" if  yoy_county_mlp_growth >= 0  else "negative") +                                           
                                ' growth'                       +
                                ' in values. '                 +
                                
                                #Sentence 2
                                'In fact, as of '   +
                                current_county_mlp_period +
                                ', the median home list price sits at ' +
                                current_county_mlp +
                                ', ' +                                        
                                 increase_or_decrease+
                                ' of ' +
                               "{:,.0f}%".format(abs(yoy_county_mlp_growth)) +
                                ' compared to ' +
                                "{msa_growth_description}".format(msa_growth_description = "an increase of " if  yoy_msa_mlp_growth >= 0  else "a decrease of ") +                                           
                                 "{:,.0f}%".format(abs(yoy_msa_mlp_growth)) +
                                ' for the ' +
                                cbsa_name +
                                ' Metro, and ' +
                                "{national_growth_description}".format(national_growth_description = "an increase of " if  yoy_national_mlp_growth >= 0  else "a decrease of ") +      
                                "{:,.0f}%".format(abs(yoy_national_mlp_growth)) +
                                ' across the Nation over the past year.'
                                )

        #If we don't have metro realtor.com data                        
        else: 
            housing_langage = (#Sentence 1
                                'In '                                                                                                            +
                                county                                                                                                           +
                                ', Realtor.com data points to '                                                                                  +
                                "{growth_description}".format(growth_description = "continued " if  yoy_county_mlp_growth >= 0  else "negative") +                                           
                                ' growth'                                                                                                        +
                                ' in values. '                                                                                                   + 
                                
                                #Sentence 2
                                'In fact, as of '                                                                                                +
                                current_county_mlp_period                                                                                        +
                                ', the median home list price sits at '                                                                          +
                                current_county_mlp                                                                                               +
                                ', '                                                                                                             +
                                increase_or_decrease                                                                                             +
                                ' of '                                                                                                           +
                            "{:,.0f}%".format(abs(yoy_county_mlp_growth))                                                                        +
                                ' compared to '                                                                                                  +
                                "{national_growth_description}".format(national_growth_description = "an increase of " if  yoy_national_mlp_growth >= 0  else "a decrease of ") +      
                                "{:,.0f}%".format(abs(yoy_national_mlp_growth))                                                                   +
                                ' across the Nation over the past year.'
                              )
        return([boiler_plate_housing_language,housing_langage])

def CarLanguage():
    print('Creating auto Langauge')
    wikipedia_car_language     = WikipediaTransitLanguage(category = 'car')
    #If we find something on wikipedia, use that, otherwise keep going
    if wikipedia_car_language != None:
        return(wikipedia_car_language)

    else:
        print('No major highway information on wikipedia, using geographic data')
        nearest_highway_language = FindNearestHighways()
        
        if nearest_highway_language != None:
            return(nearest_highway_language)
        else:
            return('Major roads serving ' + county  + ' include .')

def PlaneLanguage():
    print('Creating plane Langauge')

    #First see if any text available on wikipedia, if so use that, if not, use our geographic data
    print('Searching Wikipedia for Airport Info')
    wikipedia_plane_language = WikipediaTransitLanguage(category='air')
    
    if wikipedia_plane_language != None:
        print('Pulled Airport info from Wikipedia')
        return(wikipedia_plane_language)
    
    else:
        #Check to see if there are any airports within the area    
        print('No Airport Information on Wikipedia, using airport shapefile to see if there are any airports within the area')
        airport_language = FindAirport()

        if airport_language != None:
            return(airport_language)

        else:
             return(county + ' is served by  .')

def BusLanguage():
    print('Creating bus Langauge')

    wikipedia_bus_language = WikipediaTransitLanguage(category = 'bus')
    
    if wikipedia_bus_language != None:
        return(wikipedia_bus_language)
    else:
        return(county + ' does not have public bus service.')

def TrainLanguage():
    print('Creating train Langauge')
    wikipedia_train_language = WikipediaTransitLanguage(category='train')
    if wikipedia_train_language != None:
        return(wikipedia_train_language)
    else:
        return(county + ' is not served by any commuter or light rail lines.')

def OutlookLanguage():
    print('Writing Outlook Langauge')
    #First pargarph is the same for every county, second one is specific to the subject county

    #National economy boiler plate
                                
    national_economy_summary = (
                            'The United States economy continues to recover from the aftermath of the Covid-19 pandemic.' + 
                            ' The labor market has restored almost 19 million of the 21 million jobs lost at the beginning of the pandemic, as measured by non-farm employment, bringing the unemployment rate to 3.9% as of December 2021. '+
                            'Employment growth continued in leisure and hospitality, in professional and business services, in retail trade, and in transportation and warehousing. ' + 
                            'After historical growth in Q2, GDP growth slowed to an annual rate of 2.3% in Q3 2021. ' + 
                            'The slowdown in third quarter GDP reflected the continued economic impact of the COVID-19 pandemic. A resurgence of COVID-19 cases resulted in new restrictions and delays in the reopening of establishments in some parts of the country. ' +
                            'Supply-chain disruptions such as delays at U.S. ports and international manufacturing issues contributed to a sharp increase in inflation and pose a risk to the economic outlook. Despite supply-side challenges, many economic observers expect that the economy regained momentum in the final months of the year and is well positioned for continued growth in 2022. ' 
                                 )


    #County GDP/GDP Growth Sentence
    county_gdp_growth              =               ( (county_gdp['GDP'].iloc[-1]) / (county_gdp['GDP'].iloc[0]) - 1 ) * 100
    county_gdp_min_year            =                county_gdp['Period'].min()
    county_gdp_max_year            =                county_gdp['Period'].max()

    #Restrict to years we have for county
    national_gdp_restricted        =               national_gdp.loc[ (national_gdp['Period'] <= county_gdp_max_year) & (national_gdp['Period'] >= county_gdp_min_year)  ]    
    national_gdp_growth            =               ((   (national_gdp_restricted['GDP'].iloc[-1])/(national_gdp_restricted['GDP'].iloc[0])   - 1 ) * 100)
    county_gdp_growth_difference   =                (county_gdp_growth - national_gdp_growth ) * 100

    
    county_gdp_sentence = (#Sentence 1
                           'Between, '                    + 
                            str(county_gdp_min_year)[6:]  + 
                            ' and '                       +
                            str(county_gdp_max_year)[6:]  +
                            ', '                          +
                            county                        +
                            ' GDP grew '                  +
                            "{:,.1f}%".format(county_gdp_growth) +
                            '. '                          +
                             
                             #Sentence 2     
                            'This growth rate '           +
                             "{leads_or_lags}".format(leads_or_lags =('led the national average by ' +  "{:,.0f} bps".format(county_gdp_growth_difference) + ' during this period. ') if (county_gdp_growth_difference > 0)  else   ('lagged the national average by ' + "{:,.0f} bps".format(abs(county_gdp_growth_difference)) + ' during this period. ')) 
                            )

    #Unemployment sentence
    current_unemployment                              = county_unemployment_rate['unemployment_rate'].iloc[-1]
    historical_average_unemployment                   = county_unemployment_rate['unemployment_rate'].mean()
    current_state_unemployment                        = state_unemployment_rate['unemployment_rate'].iloc[-1]
    current_national_unemployment                     = national_unemployment['unemployment_rate'].iloc[-1]

    #Compare current county unemployment rate to hisorical average
    if current_unemployment > historical_average_unemployment:
        unemployment_above_below_hist_avg = 'above'
    elif current_unemployment < historical_average_unemployment:
        unemployment_above_below_hist_avg = 'below'
    elif current_unemployment == historical_average_unemployment:
        unemployment_above_below_hist_avg = 'equal to'

    #Compare current county unemployment rate to state average
    if current_unemployment > current_state_unemployment:
        unemployment_above_below_state = 'above'
    elif current_unemployment < current_state_unemployment:
        unemployment_above_below_state = 'below'
    elif current_unemployment == current_state_unemployment:
        unemployment_above_below_state = 'equal to'

    #Compare current county unemployment rate to natioanl average
    if current_unemployment > current_national_unemployment:
        unemployment_above_below_national = 'above'
    elif current_unemployment < current_national_unemployment:
        unemployment_above_below_national = 'below'
    elif current_unemployment == current_national_unemployment:
        unemployment_above_below_national = 'equal to'
        
        
        
    county_unemployment_sentence = (#Sentence 1
                                    'The current unemployment rate in '     + 
                                    county                                  + 
                                    ' of '                                  + 
                                    "{:,.1f}%".format(current_unemployment) + 
                                    ' is '                                  + 
                                    unemployment_above_below_hist_avg       + 
                                    ' its five-year average. '              +
                                    
                                    #Sentnce 2
                                    'It is ' + unemployment_above_below_state + ' and ' +  unemployment_above_below_national + ' the state ' +  '(' + "{:,.1f}%".format(current_state_unemployment)  + ')'  + ' and national average '  + '(' "{:,.1f}%".format(current_national_unemployment) + ')' ', respectively. '
                                    )

    #Demographics/Population
    county_resident_pop['Resident Population_1year_growth']  =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(1))  - 1) * 100)/1
    county_resident_pop['Resident Population_5year_growth']  =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(5))   - 1) * 100)/5
    county_resident_pop['Resident Population_10year_growth'] =  (((county_resident_pop['Resident Population']/county_resident_pop['Resident Population'].shift(10)) - 1) * 100)/10

    county_1y_growth                                         = county_resident_pop.iloc[-1]['Resident Population_1year_growth'] 
    county_5y_growth                                         = county_resident_pop.iloc[-1]['Resident Population_5year_growth']  
    county_10y_growth                                        = county_resident_pop.iloc[-1]['Resident Population_10year_growth'] 

    if county_5y_growth < 0 and county_1y_growth < 0:
        county_demographic_sentence = (county + ' continues to experience population loss with one- and five-year annual growth rates of ' +  "{:,.1f}%".format(county_1y_growth) + ' and ' + "{:,.1f}%".format(county_5y_growth) + '.'  )
    
    elif county_5y_growth > 0 and county_1y_growth > 0:
        county_demographic_sentence = (county + ' continues to experience population gains with one- and five-year annual growth rates of ' +  "{:,.1f}%".format(county_1y_growth) + ' and ' + "{:,.1f}%".format(county_5y_growth) + '.'  )

    elif  county_5y_growth < 0 and county_1y_growth > 0:
        county_demographic_sentence = ('Although ' + county + ' has experienced population decline of ' +   "{:,.1f}%".format(abs(county_5y_growth)) +' annually over the past five years, growth has returned to positive levels with a most recent one-year growth rate of ' +  "{:,.1f}%".format(county_1y_growth) +'.')
        
    elif county_5y_growth > 0 and county_1y_growth < 0:
        county_demographic_sentence = ('Although ' + county + ' has experienced population growth of ' + "{:,.1f}%".format(county_5y_growth) +  ' over the past five years, it most recently saw a one-year contraction of ' +  "{:,.1f}%".format(county_1y_growth) +'.')

    elif county_5y_growth == 0 and county_1y_growth == 0:
        county_demographic_sentence = (county + """'s""" + ' population has experienced no change over the past five years.') 

    else:
        county_demographic_sentence = ('')

    #County Economy Summary
    county_economy_summary = (
                            county_gdp_sentence + 
                            county_unemployment_sentence + 
                            county_demographic_sentence 
                            )

    return([national_economy_summary,county_economy_summary])

def CreateLanguage():
    global overview_language
    global county_employment_industry_breakdown_language, msa_employment_industry_breakdown_language ,county_employment_growth_language,msa_employment_growth_language,unemployment_language
    global production_language, infrastructure_language,housing_language, outlook_language
    global car_language, train_language, bus_language, plane_language
    global population_language,income_language
    print('Creating Langauge')
    
    #overview language
    overview_language       = OverviewLanguage()

    #County Employment breakdown language
    try:
        county_employment_industry_breakdown_language    = [CountyEmploymentBreakdownLanguage(county_industry_breakdown = county_industry_breakdown)]
    except:
        print('problem with county employment breakdown language')
        county_employment_industry_breakdown_language    = ''

    #MSA Employment breakdown language
    try:
        msa_employment_industry_breakdown_language    = [MSAEmploymentBreakdownLanguage(msa_industry_breakdown = msa_industry_breakdown)]
    except:
        print('problem with MSA employment breakdown language')
        msa_employment_industry_breakdown_language    = ['']

    #Production language
    try:
        production_language     = ProductionLanguage(county_data_frame = county_gdp ,msa_data_frame = msa_gdp,state_data_frame = state_gdp)
    except Exception as e:
        print('problem with production language: ', e)
        production_language = ['']

    #Infrastructure language
    try:
        infrastructure_language = InfrastructureLanguage()
    except:
        print('problem with infrastructure language')
        infrastructure_language = ['']

    #Houing langauge
    try:    
        housing_language        = HousingLanguage()
    except:
        print('problem with housing language')
        housing_language = ['']
   
   #Outloook language
    try:    
        outlook_language        = OutlookLanguage()
    except Exception as e:
        print(e,' problem with outlook language')
        outlook_language = ['']
   

    bus_language                       = BusLanguage() 
    train_language                     = TrainLanguage()
    car_language                       = CarLanguage()
    plane_language                     = PlaneLanguage()


    #Unemployment language
    unemployment_language              = UnemploymentLanguage()

    #County Private Employment Growth language
    try:    
        county_employment_growth_language = [CountyEmploymentGrowthLanguage(county_industry_breakdown=county_industry_growth_breakdown)]
    except Exception as e:
        print(e, ' ---- problem with county emp growth language')
        county_employment_growth_language = ['']
    
    #MSA Private Employment Growth language
    try:    
        msa_employment_growth_language = [MSAEmploymentGrowthLanguage(msa_industry_breakdown=msa_industry_growth_breakdown)]
    except Exception as e:
        print(e, ' ---- problem with msa emp growth language')
        msa_employment_growth_language = ['']

    #Population language 
    population_language = PopulationLanguage(national_resident_pop = national_resident_pop )

    #Income Language
    income_language = IncomeLanguage()

#Table Related functions 
def GetDataAndLanguageForOverviewTable():
    print('Getting Data for overview table')
    
    #Get most recent County values and get county values from 5 years ago
    try:
        assert len(county_employment) > (growth_period * 12) 
        current_county_employment = county_employment['Employment'].iloc[-1]
        lagged_county_employment  = county_employment['Employment'].iloc[-1 - (growth_period * 12)] #the employment data is monthly
        
    except Exception as e:
        print(e,' problem getting county employment for overview table data')
        current_county_employment = 1
        lagged_county_employment  = 1
    
    try:
        current_county_gdp       = county_gdp['GDP'].iloc[-1]
        lagged_county_gdp        = county_gdp['GDP'].iloc[-1 - growth_period]

    except Exception as e:
        print(e,' problem getting county gdp for overview table')
        current_county_gdp        = 1
        lagged_county_gdp         = 1
    
    try:
        current_county_pop        = county_resident_pop['Resident Population'].iloc[-1]
        lagged_county_pop         = county_resident_pop['Resident Population'].iloc[-1- growth_period]
    
    except Exception as e:
        print(e,'problem getting county population for overview table')
        current_county_pop        = 1
        lagged_county_pop         = 1
    
    
    try:
        current_county_pci        = county_pci['Per Capita Personal Income'].iloc[-1]
        lagged_county_pci         = county_pci['Per Capita Personal Income'].iloc[-1- growth_period]
    
    except Exception as e:
        print(e,'problem getting county pci for overview table')
        current_county_pci        = 1
        lagged_county_pci         = 1


    #Calculate county 5-year growth
    try:
        county_employment_growth = ((current_county_employment/lagged_county_employment) - 1 ) * 100
        county_gdp_growth        = ((current_county_gdp/lagged_county_gdp) - 1) * 100
        county_pop_growth        = ((current_county_pop/lagged_county_pop) - 1) * 100
        county_pci_growth        = ((current_county_pci/lagged_county_pci) - 1) * 100

    except Exception as e:
        print(e,'problem calculating growth rates for county in overview table')
        county_employment_growth = 0
        county_gdp_growth        = 0
        county_pop_growth        = 0
        county_pci_growth        = 0


    

    #Make sure we are comparing the same month to month change in values between state and county data
    try:
        state_employment_extra_month_cut_off    = state_employment.loc[state_employment['period']     <= (county_employment['period'].max())]
        state_gdp_extra_month_cut_off           = state_gdp.loc[state_gdp['Period']                   <= (county_gdp['Period'].max())]
        state_resident_pop_extra_month_cut_off  = state_resident_pop.loc[state_resident_pop['Period'] <= (county_resident_pop['Period'].max())]
        state_pci_extra_month_cut_off           = state_pci.loc[state_pci['Period']                   <= (county_pci['Period'].max())]

        current_state_employment                = state_employment_extra_month_cut_off['Employment'].iloc[-1]
        current_state_gdp                       = state_gdp_extra_month_cut_off['GDP'].iloc[-1]
        current_state_pop                       = state_resident_pop_extra_month_cut_off['Resident Population'].iloc[-1]
        current_state_pci                       = state_pci_extra_month_cut_off['Per Capita Personal Income'].iloc[-1]


    except Exception as e:
        print(e,'problem getting state values for overview table data with cut off observations')
        try:
            current_state_employment = state_employment['Employment'].iloc[-1]
            current_state_gdp        = state_gdp['GDP'].iloc[-1]
            current_state_pop        = state_resident_pop['Resident Population'].iloc[-1]
            current_state_pci        = state_pci['Per Capita Personal Income'].iloc[-1]
        except Exception as e:
            print(e,'problem getting state values with most recent data')
            current_state_employment = 1
            current_state_gdp        = 1
            current_state_pop        = 1
            current_state_pci        = 1


    #Get lagged state values from 5 years ago
    try:
        lagged_state_employment = state_employment_extra_month_cut_off['Employment'].iloc[-1 - (growth_period * 12)] #the employment data is monthly
        lagged_state_gdp        = state_gdp_extra_month_cut_off['GDP'].iloc[-1 - growth_period]
        lagged_state_pop        = state_resident_pop_extra_month_cut_off['Resident Population'].iloc[-1- growth_period]
        lagged_state_pci        = state_pci_extra_month_cut_off['Per Capita Personal Income'].iloc[-1- growth_period]

    
    except Exception as e:
        try:
            lagged_state_employment = state_employment['Employment'].iloc[-1 - (growth_period * 12)] #the employment data is monthly
            lagged_state_gdp        = state_gdp['GDP'].iloc[-1 - growth_period]
            lagged_state_pop        = state_resident_pop['Resident Population'].iloc[-1- growth_period]
            lagged_state_pci        = state_pci['Per Capita Personal Income'].iloc[-1- growth_period]

        except Exception as e:
            print(e,'problem getting lagged state values with data without dropped observations')
            lagged_state_employment = 1
            lagged_state_gdp        = 1
            lagged_state_pop        = 1
            lagged_state_pci        = 1

    #Calculate 5-year growth for state
    try:
        state_employment_growth = ((current_state_employment/lagged_state_employment) - 1 ) * 100
        state_gdp_growth        = ((current_state_gdp/lagged_state_gdp) - 1) * 100
        state_pop_growth        = ((current_state_pop/lagged_state_pop) - 1) * 100
        state_pci_growth        = ((current_state_pci/lagged_state_pci) - 1) * 100

    except Exception as e:
        print(e,'problem getting state growth rates for overivew table')
        


    #Get most recent values for MSA if applicable
    if cbsa != '':
        try:
            #Make sure we are comparing the same month to month change in values between msa and county data
            msa_employment_extra_month_cut_off    = msa_employment.loc[msa_employment['period']     <= (county_employment['period'].max())]
            msa_gdp_extra_month_cut_off           = msa_gdp.loc[msa_gdp['Period']                   <= (county_gdp['Period'].max())]
            msa_resident_pop_extra_month_cut_off  = msa_resident_pop.loc[msa_resident_pop['Period'] <= (county_resident_pop['Period'].max())]
            msa_pci_extra_month_cut_off           = msa_pci.loc[msa_pci['Period']                   <= (county_pci['Period'].max())]

            #Now get most recent msa level values
            current_msa_employment                = msa_employment_extra_month_cut_off['Employment'].iloc[-1]
            current_msa_gdp                       = msa_gdp_extra_month_cut_off['GDP'].iloc[-1]
            current_msa_pop                       = msa_resident_pop_extra_month_cut_off['Resident Population'].iloc[-1]
            current_msa_pci                       = msa_pci_extra_month_cut_off['Per Capita Personal Income'].iloc[-1]
        
            #Get lagged msa values from 5 years ago
            lagged_msa_employment               = msa_employment_extra_month_cut_off['Employment'].iloc[-1 - (growth_period * 12)] #the employment data is monthly
            lagged_msa_gdp                      = msa_gdp_extra_month_cut_off['GDP'].iloc[-1 - growth_period]
            lagged_msa_pop                      = msa_resident_pop_extra_month_cut_off['Resident Population'].iloc[-1- growth_period]
            lagged_msa_pci                      = msa_pci_extra_month_cut_off['Per Capita Personal Income'].iloc[-1- growth_period]
        

        
        except Exception as e:
            print(e,'problem getting msa values with observations cut off')
            try:
                #Now get most recent msa level values
                current_msa_employment                = msa_employment['Employment'].iloc[-1]
                current_msa_gdp                       = msa_gdp['GDP'].iloc[-1]
                current_msa_pop                       = msa_resident_pop['Resident Population'].iloc[-1]
                current_msa_pci                       = msa_pci['Per Capita Personal Income'].iloc[-1]
        
                #Get lagged msa values from 5 years ago
                lagged_msa_employment               = msa_employment['Employment'].iloc[-1 - (growth_period * 12)] #the employment data is monthly
                lagged_msa_gdp                      = msa_gdp['GDP'].iloc[-1 - growth_period]
                lagged_msa_pop                      = msa_resident_pop['Resident Population'].iloc[-1- growth_period]
                lagged_msa_pci                      = msa_pci['Per Capita Personal Income'].iloc[-1- growth_period]
            except Exception as e:
                print(e,'problem getting msa values with most recent data')

                #Now get most recent msa level values
                current_msa_employment                = 1
                current_msa_gdp                       = 1
                current_msa_pop                       = 1
                current_msa_pci                       = 1
        
                #Get lagged msa values from 5 years ago
                lagged_msa_employment                 = 1 
                lagged_msa_gdp                        = 1
                lagged_msa_pop                        = 1
                lagged_msa_pci                        = 1

        try:
            #Calculate 5-year growth for msa
            msa_employment_growth               = ((current_msa_employment/lagged_msa_employment) - 1 ) * 100
            msa_gdp_growth                      = ((current_msa_gdp/lagged_msa_gdp) - 1) * 100
            msa_pop_growth                      = ((current_msa_pop/lagged_msa_pop) - 1) * 100
            msa_pci_growth                      = ((current_msa_pci/lagged_msa_pci) - 1) * 100

        except Exception as e:
            print(e,'problem getting msa growth rates for overview table data')


    #Use MSA growth rates for comparison
    try:
        if cbsa != '':
            comparison_emp_growth = msa_employment_growth
            comparison_gdp_growth = msa_gdp_growth
            comparison_pop_growth = msa_pop_growth
            comparison_pci_growth = msa_pci_growth

        #Use State growth rates for comparison
        else:
            comparison_emp_growth = state_employment_growth
            comparison_gdp_growth = state_gdp_growth
            comparison_pop_growth = state_pop_growth
            comparison_pci_growth = state_pci_growth
        
        print('successfully got comparison growth rates for overview table data')

    except Exception as e:
        print(e,'problem getting comparison growth rates')

    try:
        #Determine if county grew faster or slower than state or MSA
        if comparison_emp_growth > county_employment_growth:
            employment_faster_or_slower = 'Slower than'
        elif comparison_emp_growth < county_employment_growth:
            employment_faster_or_slower = 'Faster than'
        elif comparison_emp_growth == comparison_emp_growth:
            employment_faster_or_slower = 'Equal to'
            

        if comparison_gdp_growth > county_gdp_growth:
            gdp_faster_or_slower = 'Slower than'
        elif comparison_gdp_growth < county_gdp_growth:
            gdp_faster_or_slower = 'Faster than' 
        else:
            gdp_faster_or_slower = 'Equal to' 

        if comparison_pop_growth > county_pop_growth:
            pop_faster_or_slower = 'Slower than'
        elif comparison_pop_growth < county_pop_growth:
            pop_faster_or_slower = 'Faster than'
        else:
            pop_faster_or_slower = 'Equal to'

        if comparison_pci_growth > county_pci_growth:
            pci_faster_or_slower = 'Slower than'
        elif comparison_pci_growth < county_pci_growth:
            pci_faster_or_slower = 'Faster than' 
        else:
            pci_faster_or_slower = 'Equal to'
        
    except Exception as e:
        print(e,'problem getting growth comparison descriptions')
        
    #Format Variables
    try:
        current_county_employment = "{:,.0f}".format(current_county_employment)
        current_county_gdp        = '$' + millify(current_county_gdp) 
        current_county_pop        = "{:,.0f}".format(current_county_pop)
        current_county_pci        = "${:,.0f}".format(current_county_pci)

        county_employment_growth  = "{:,.1f}%".format(county_employment_growth)
        county_gdp_growth         = "{:,.1f}%".format(county_gdp_growth)
        county_pop_growth         = "{:,.1f}%".format(county_pop_growth)
        county_pci_growth         = "{:,.1f}%".format(county_pci_growth)

    except Exception as e:
        print(e, 'problem formatting variables for overview table data')
    
    try:
        if cbsa != '':
            overview_table =[ 
                            ['Attribute','County Level Value',str(growth_period) + ' Year Growth Rate','Relative to Baseline ('+ 'MSA' + ')' ], 
                            ['Employment',current_county_employment,county_employment_growth,employment_faster_or_slower + ' MSA' ], 
                            ['GDP',current_county_gdp,county_gdp_growth,gdp_faster_or_slower + ' MSA'],
                            ['Population',current_county_pop,county_pop_growth,pop_faster_or_slower + ' MSA'], 
                            ['Per Capita Personal Income',current_county_pci,county_pci_growth,pci_faster_or_slower + ' MSA'] 
                            ]
        else:
            overview_table =[ 
                            ['Attribute','County Level Value',str(growth_period) + ' Year Growth Rate','Relative to Baseline ('+ state + ')' ], 
                            ['Employment',current_county_employment,county_employment_growth,employment_faster_or_slower + ' State' ], 
                            ['GDP',current_county_gdp,county_gdp_growth,gdp_faster_or_slower + ' State'],
                            ['Population',current_county_pop,county_pop_growth,pop_faster_or_slower + ' State'], 
                            ['Per Capita Personal Income',current_county_pci,county_pci_growth,pci_faster_or_slower + ' State'] 
                            ]
            
        for list in overview_table:
            if list[1] == '$0':
                list[1] = 'NA'
                list[2] = 'NA'
                list[3] = 'NA'
    
    
    except Exception as e:
        print(e,'problem creating list of lists for overview table data')
        overview_table = [
                ['Attribute','County Level Value','5 Year Growth Rate','Relative to Baseline (State Code or MSA)'],
                ['Employment','','','[Faster than/Slower than/Equal to] [State/MSA]'],
                ['GDP','','','[Faster than/Slower than/Equal to] [State/MSA]'],
                ['Population','','','[Faster than/Slower than/Equal to] [State/MSA]'],
                ['Per Capita Personal Income','','','[Faster than/Slower than/Equal to] [State/MSA]']
                ]
    
    try:
        return(overview_table)
    except Exception as e:
        print(e,'failed to return list')

def AddTable(document,data_for_table): #Function we use to insert our overview table into the report document
    #list of list where each list is a row for our table
     
    #make sure each list inside the list of lists has the same number of elements
    for row in data_for_table:
        for row2 in data_for_table:
            assert len(row) == len(row2)


    #create table object
    tab               = document.add_table(rows=len(data_for_table), cols=len(data_for_table[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True

    #loop through the rows in the table
    for current_row, (row,row_data_list) in enumerate(zip(tab.rows,data_for_table)): 

        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):
            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.75)
            elif current_column == 1:
                cell.width = Inches(1.5)
            elif current_column == 2:
                cell.width = Inches(1.5)
            elif current_column == 3:
                cell.width = Inches(2)


            #add border to top row
            if current_row == 1:
                    tcPr      = cell._element.tcPr
                    tcBorders = OxmlElement("w:tcBorders")
                    top       = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                    tcBorders.append(top)
                    tcPr.append(tcBorders)
            
            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                paragraph.paragraph_format.space_after  = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

                for run in paragraph.runs:
                    font              = run.font
                    font.size         = Pt(8)
                    run.alignment     = WD_ALIGN_PARAGRAPH.RIGHT
                    run.space_after   = Pt(0)
                    run.space_before  = Pt(0)
                    
                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
                        font.size = Pt(9)
                    else:
                        font.name = primary_font
             
#####################################################Report document related functions####################################
def SetPageMargins(document, margin_size):
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin   = Inches(margin_size)
        section.right_margin  = Inches(margin_size)

def SetDocumentStyle(document):
    style     = document.styles['Normal']
    font      = style.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(9)

def AddTitle(document):
    title                               = document.add_heading(county + ' Area Analysis',level=1)
    title.style                         = document.styles['Heading 2']
    title.paragraph_format.space_after  = Pt(6)
    title.paragraph_format.space_before = Pt(12)
    title_style                         = title.style
    title_style.font.name               = "Avenir Next LT Pro Light"
    title_style.font.size               = Pt(14)
    title_style.font.bold               = False
    title_style.font.color.rgb          = RGBColor.from_string('3F65AB')
    title_style.element.xml
    rFonts                              = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

    above_map_paragraph = document.add_paragraph("The following analysis includes pertinent aspects of the surrounding region as it pertains to the subject property. " + 
                                                'This report was compiled using data as of ' + current_quarter + ' unless otherwise noted. Data is from a number of sources including the U.S. Bureau of Labor Statistics ("BLS"), the U.S. Bureau of Economic Analysis ("BEA"), and the U.S. Census Bureau.'
                                                )
    above_map_style                                   = above_map_paragraph.style
    above_map_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.JUSTIFY
    above_map_style.font.size                         = Pt(9)
    above_map_paragraph.paragraph_format.space_after  = Pt(primary_space_after_paragraph)

def AddHeading(document, title, heading_level): 
    #Function we use to insert the headers other than the title header
    heading                               = document.add_heading(title, level = heading_level)
    heading.style                         = document.styles['Heading 3']
    heading_style                         = heading.style
    heading_style.font.name               = "Avenir Next LT Pro"
    heading_style.font.size               = Pt(11)
    heading_style.font.bold               = False
    heading.paragraph_format.space_after  = Pt(6)
    heading.paragraph_format.space_before = Pt(12)

    #Color
    heading_style.font.color.rgb          = RGBColor.from_string('3F65AB')            
    heading_style.element.xml
    rFonts                                = heading_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def Citation(document, text):
    citation_paragraph                               = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(6)
    citation_paragraph.paragraph_format.space_before = Pt(0)
    run                                              = citation_paragraph.add_run('Source: ' + text)
    font                                             = run.font
    font.name                                        = primary_font
    font.size                                        = Pt(8)
    font.italic                                      = True
    font.color.rgb                                   = RGBColor.from_string('929292')
    citation_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.RIGHT

def Note(document, text):
    citation_paragraph                               = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(6)
    citation_paragraph.paragraph_format.space_before = Pt(6)
    run                                              = citation_paragraph.add_run('Note: ' + text)
    font                                             = run.font
    font.name                                        = primary_font
    font.size                                        = Pt(8)
    font.italic                                      = True
    font.color.rgb                                   = RGBColor.from_string('929292')
    citation_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.RIGHT

def GetMap():
    try:
        print('Getting map image from Google Maps')
        
        #Search Google Maps for County
        options = webdriver.ChromeOptions()
        options.add_argument("--start-maximized")
        browser = webdriver.Chrome(executable_path=(os.path.join(os.environ['USERPROFILE'], 'Desktop','chromedriver.exe')),options=options)
        browser.get('https:google.com/maps')
        
        #Write county name in box
        Place = browser.find_element_by_class_name("tactile-searchbox-input")
        Place.send_keys((county + ', ' + state))
        
        #Submit county name for search
        Submit = browser.find_element_by_class_name('nhb85d-BIqFsb')
        Submit.click()

        time.sleep(5)
        zoomout = browser.find_element_by_xpath("""//*[@id="widget-zoom-out"]/div""")
        zoomout.click()
        time.sleep(7)

        #Differnent machines have different screen coordinates
        if 'Leahy' in os.environ['USERPROFILE']: 
            print("""Using Mike's coordinates for screenshot""")
            im2 = pyautogui.screenshot(region=(1358, 465, 2142, 1404) ) #left, top, width, and height
        
        elif 'Dominic' in os.environ['USERPROFILE']:
            print("""Using Dom's coordinates for screenshot""")
            im2 = pyautogui.screenshot(region=(3680, 254, 1968, 1231) ) #left, top, width, and height
        
        else:
            return()
        
        time.sleep(.25)
        im2.save(os.path.join(county_folder_map,'map.png'))
        im2.close()
        time.sleep(1)
        browser.quit()
   
   #Upon failure, make sure we close the browser
    except Exception as e:
        print(e)
        try:
            browser.quit()
        except:
            pass
 
def AddMap(document):
    #Add image of map if we already have one
    if os.path.exists(os.path.join(county_folder_map, 'map.png')):
        print('Adding map png to document')
        map = document.add_picture(os.path.join(county_folder_map, 'map.png'), width=Inches(6.5))
    
    #Otherwise, go fetch one from Google Maps
    else:
        GetMap() 
    
        if os.path.exists(os.path.join(county_folder_map, 'map.png')):
            print('Adding map png to document')
            map = document.add_picture(os.path.join(county_folder_map, 'map.png'), width=Inches(6.5))
    
    Citation(document=document,text = 'Google Maps')

def AddTwoColumnTable(document, pic_list, lang_list):
    #Insert the transit graphics(car, bus, plane, train)
    
    #Create table object
    tab = document.add_table(rows=0, cols=2)

    for pic,lang in zip(pic_list,lang_list):
        row_cells      = tab.add_row().cells
        
        left_paragraph = row_cells[0].paragraphs[0]
        run            = left_paragraph.add_run()

        run.add_picture(os.path.join(graphics_location,pic),width=Inches(0.2),height =Inches(0.2))

        right_paragraph             = row_cells[1].paragraphs[0]
        right_paragraph.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
         
        if type(lang) == list:
            for p in lang:
                run             = right_paragraph.add_run()
                run.add_text(str(p))

        else:
            run                 = right_paragraph.add_run()
            run.add_text(str(lang))
     
    #We have now defined our table object,loop through all rows then all cells in each current row
    for row in tab.rows:
        for current_column,cell in enumerate(row.cells):
            #Set Width for cell
            if current_column == 0:
                cell.width = Inches(.3)

            elif current_column == 1:
                cell.width = Inches(6.5)

def AddDocumentParagraph(document, language_variable):
    assert type(language_variable) == list

    for paragraph in language_variable:
        
        #Skip blank paragraphs
        if paragraph == '':
            continue
        
        par                                               = document.add_paragraph(str(paragraph))
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
        style                                             = document.styles['Normal']
        font                                              = style.font
        font.name                                         = primary_font
        par.style                                         = document.styles['Normal']

def AddDocumentPicture(document, image_path, citation):
    if os.path.exists(image_path):
        fig                                         = document.add_picture(os.path.join(image_path),width=Inches(6.5))
        last_paragraph                              = document.paragraphs[-1] 
        last_paragraph.paragraph_format.space_after = Pt(0)
        last_paragraph.alignment                    = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document, citation)

def AddTableTitle(document, title):
    table_title_paragraph                               = document.add_paragraph(title)
    table_title_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.CENTER
    table_title_paragraph.paragraph_format.space_after  = Pt(6)
    table_title_paragraph.paragraph_format.space_before = Pt(12)

    for run in table_title_paragraph.runs:
                    font      = run.font
                    font.name = 'Avenir Next LT Pro Medium'

#Report Section Functions
def OverviewSection(document):
    print('Writing Overview Section')
    AddHeading(document = document, title = 'Overview', heading_level = 2)

    #Add Overview langauge
    AddDocumentParagraph(document = document, language_variable = overview_language)
    
    #Overview table title
    AddTableTitle(document = document, title = 'Area Fundamentals')

    #Creating Overview Table
    try:
        data_for_table = GetDataAndLanguageForOverviewTable()
    except Exception as e:
        print(e,'problem getting data for overview table')
    
    try:
        AddTable(document = document, data_for_table = data_for_table )
    except Exception as e:
        print(e,'problem adding table')
           
def EmploymentSection(document):
    print('Writing Employment Section')
    AddHeading(document = document, title = 'Labor Market Conditions', heading_level = 2)

    #Add MSA Employment Breakdown Language
    AddDocumentParagraph(document = document, language_variable = msa_employment_industry_breakdown_language)

    #Add MSA employment treemap chart
    AddDocumentPicture(document = document, image_path = os.path.join(county_folder,'msa_employment_by_industry.png') ,citation = 'U.S. Bureau of Labor Statistics')

    #Add County Employment Breakdown Language
    AddDocumentParagraph(document = document, language_variable = county_employment_industry_breakdown_language)

    #Add county employment treemap chart
    AddDocumentPicture(document = document, image_path = os.path.join(county_folder,'employment_by_industry.png') ,citation = 'U.S. Bureau of Labor Statistics')
    
    AddDocumentParagraph(document = document, language_variable = unemployment_language)

    #Add combined unemployment rate and employment growth graph
    AddDocumentPicture(document = document, image_path = os.path.join(county_folder,'unemployment_rate_employment_growth.png') ,citation = 'U.S. Bureau of Labor Statistics')

    #MSA Employment growth language
    AddDocumentParagraph(document = document, language_variable = msa_employment_growth_language)

    #Add MSA employment growth by industry bar chart
    AddDocumentPicture(document = document, image_path = os.path.join(county_folder,'msa_employment_growth_by_industry.png') ,citation = 'U.S. Bureau of Labor Statistics')

    #County Employment growth language
    AddDocumentParagraph(document = document, language_variable = county_employment_growth_language)

    #Add county employment growth by industry bar chart
    AddDocumentPicture(document = document, image_path = os.path.join(county_folder,'employment_growth_by_industry.png') ,citation = 'U.S. Bureau of Labor Statistics')

def ProductionSection(document):
    print('Writing Production Section')
    AddHeading(document = document, title = 'Economic Production', heading_level = 2)
    
    #Add GDP Language
    AddDocumentParagraph(document = document, language_variable = production_language)

    #Add GDP Graph
    AddDocumentPicture(document = document, image_path = os.path.join(county_folder,'gdp.png') ,citation = 'U.S. Bureau of Economic Analysis')

def DemographicsSection(document):
    print('Writing Demographic Section')
    AddHeading(document = document, title = 'Demographics', heading_level = 2)

    #Add langugage on population/population growth
    AddDocumentParagraph(document = document, language_variable = population_language)

    #Population graph
    AddDocumentPicture(document = document,image_path = os.path.join(county_folder,'resident_population_and_growth.png'), citation = 'U.S. Census Bureau')

    #Add langugage on per captia income and income growth
    AddDocumentParagraph(document = document, language_variable = income_language)

    #Per Capita Income and Income Growth
    AddDocumentPicture(document = document,image_path = os.path.join(county_folder,'per_capita_income_and_growth.png'), citation = 'U.S. Census Bureau')

def InfrastructureSection(document):
    print('Writing Infrastructure Section')
    AddHeading(document = document, title = 'Infrastructure', heading_level = 2)
    AddDocumentParagraph(document = document, language_variable = infrastructure_language)

    #Insert the transit graphics(car, bus,plane, train)
    AddTableTitle(document = document, title = 'Transportation Methods')
    AddTwoColumnTable(document, pic_list = ['car.png','train.png','bus.png','plane.png'], lang_list = [car_language, train_language, bus_language, plane_language] )

def HousingSection(document):
    print('Writing Housing Section')
    AddHeading(          document = document, title = 'Housing', heading_level = 2)
    AddDocumentParagraph(document = document, language_variable = housing_language)
    AddDocumentPicture(  document = document, image_path = os.path.join(county_folder,'mlp.png'), citation = 'Realtor.com')
        
def OutlookSection(document):
    print('Writing Outlook Section')
    AddHeading(          document = document, title = 'Outlook',            heading_level = 2)
    AddDocumentParagraph(document = document, language_variable = outlook_language)

def WriteReport():
    print('Writing Report')
    #Create Document
    document = Document()
    SetPageMargins(       document = document, margin_size = 1)
    SetDocumentStyle(     document = document)
    AddTitle(             document = document)
    AddMap(               document = document)
    OverviewSection(      document = document)
    EmploymentSection(    document = document)
    ProductionSection(    document = document)
    DemographicsSection(  document = document)
    InfrastructureSection(document = document)
    HousingSection(       document = document)
    OutlookSection(       document = document)

    #Save report
    document.save(report_path)  

def CleanUpPNGs():
    print('Deleting PNG files')
    #Report writing done, delete figures
    files = os.listdir(county_folder)
    for image in files:
        if image.endswith(".png"):
            while os.path.exists(os.path.join(county_folder, image)):
                try:
                    os.remove(os.path.join(county_folder, image))
                except:
                    pass

def CreateDirectoryCSV():
    global service_api_csv_name
    print('Creating CSV with file path information on all existing area reports')
    dropbox_links                  = []
    dropbox_research_names         = []
    dropbox_county_names           = []
    dropbox_analysis_types         = []
    dropbox_states                 = []
    dropbox_versions               = []
    dropbox_statuses               = []
    dropbox_document_names         = []


    for (dirpath, dirnames, filenames) in os.walk(main_output_location):
        for file in filenames:
                
            if ('.docx' not in file):
                continue
            
            full_path = dirpath + '/' + file

            #If there's a draft and final for a county, ignore the draft
            if (os.path.exists(full_path.replace('_draft','_FINAL'))) and ('_draft' in full_path):
                continue
            
            #Create variables on the reports, then add to list
            analyisis_type = 'Area'
            dropbox_link   = dirpath.replace(dropbox_root,r'https://www.dropbox.com/home')
            dropbox_link   = dropbox_link.replace("\\",r'/')    
            
            if '_draft' in file:
                file_status = 'Draft'
            else:
                file_status = 'Final'

            #For our more recent reports, we can parse the file name to grab key info
            if 'Legacy' not in dirpath :
                version                   = file[0:7]
                state_name                = file[8:10]
                first_dash_location       = file.find('- ') 
                first_underscore_location = file.find('_')
                county_name               = file[first_dash_location + 2:first_underscore_location]        
                research_name             = state_name + ' - ' + county_name
            
            elif 'Legacy' in dirpath:
                #Skip these subfolders
                if ('Unknown Period' in dirpath) or ("""Old Templates""" in dirpath) or ("""Source Data""" in dirpath):
                    continue
                
                directory_path_split      = dirpath.split("""\\""")
                version                   = directory_path_split[10]
                state_name                = directory_path_split[11]

                if state_name == 'US':
                    county_name               = 'United States' 
                elif state_name == 'DC':
                    county_name               = 'Washington D.C.'  

                else:
                    county_name               = directory_path_split[12]   
                
                research_name             = state_name + ' - ' + county_name

                

            
            dropbox_document_names.append(file)
            dropbox_analysis_types.append(analyisis_type)
            dropbox_links.append(dropbox_link)
            dropbox_versions.append(version)
            dropbox_statuses.append(file_status)
            dropbox_county_names.append(county_name)
            dropbox_research_names.append(research_name)
            dropbox_states.append(state_name)
            break
        
            

    dropbox_df = pd.DataFrame({'Market Research Name': dropbox_research_names,
                               'County':              dropbox_county_names,
                               'Analysis Type':       dropbox_analysis_types,
                               'State':               dropbox_states,
                               "Dropbox Links":       dropbox_links,
                               'Version':             dropbox_versions,
                               'Status':              dropbox_statuses,
                               'Document Name':       dropbox_document_names
                               }
                             )

    dropbox_df                              = dropbox_df.sort_values(by=['State','Market Research Name','Version'], ascending = (True, True,False))
    assigned_to_df                          = pd.read_excel(os.path.join(general_data_location,'Administrative Data','Assigned To States.xlsx')) 
    dropbox_df                              = pd.merge(dropbox_df,assigned_to_df, on=['State'],how = 'left') 
    
    csv_name                                = 'Dropbox Areas.csv'
    service_api_csv_name                    = f'Dropbox Areas-{datetime.now().timestamp()}.csv'

    dropbox_df.to_csv(os.path.join(main_output_location, csv_name),index=False)

    if main_output_location == os.path.join(dropbox_root,'Research','Market Analysis','Area'):
        dropbox_df.to_csv(os.path.join(main_output_location, service_api_csv_name),index=False)

def Main():
    print('Creating Report for: ', county)
    CreateDirectory(state = state, county = county)
    GetCountyData()
    GetMSAData()
    GetStateData()
    GetNationalData()
    print('Finished Fetching Data')
    
    CreateGraphs()
    print('Finished Creating Graphs')
    
    CreateLanguage()
    print('Finished Creating Langauge')

    WriteReport()
    CleanUpPNGs()
    print('Report Complete')

def IdentifyMSA(fips):
    #Figures out if a county is within a metropolitan statistical area and returns its CBSA code
    cbsa_fips_crosswalk = pd.read_csv(os.path.join(data_location,'cbsa2fipsxw.csv'),
            dtype={'cbsacode':                        object,
                    'metrodivisioncode':              object,
                    'csacode':                        object,
                    'cbsatitle':                      object,
                    'metropolitanmicropolitanstatis': object,
                    'metropolitandivisiontitle':      object,
                    'csatitle':                       object,
                    'countycountyequivalent':         object,
                    'statename':                      object,
                    'fipsstatecode':                  object,
                    'fipscountycode':                 object,
                    'centraloutlyingcounty':          object
                    }
                                    )
    

    #Add missing 0s
    cbsa_fips_crosswalk['fipsstatecode']  =  cbsa_fips_crosswalk['fipsstatecode'].str.zfill(2)
    cbsa_fips_crosswalk['fipscountycode'] =  cbsa_fips_crosswalk['fipscountycode'].str.zfill(3)
    cbsa_fips_crosswalk['FIPS Code']      = cbsa_fips_crosswalk['fipsstatecode'] + cbsa_fips_crosswalk['fipscountycode']
    
    #restrict data to only rows with the subject county fips
    cbsa_fips_crosswalk = cbsa_fips_crosswalk.loc[cbsa_fips_crosswalk['FIPS Code'] == fips] 
    cbsa_fips_crosswalk = cbsa_fips_crosswalk.loc[cbsa_fips_crosswalk['metropolitanmicropolitanstatis'] == 'Metropolitan Statistical Area'] #restrict to msas

    unique_CBSA_list    = cbsa_fips_crosswalk['cbsacode'].unique()
    assert len(unique_CBSA_list) < 2

    if len(cbsa_fips_crosswalk) > 0 :
        cbsa              = cbsa_fips_crosswalk['cbsacode'].iloc[-1]
        cbsa_name         = cbsa_fips_crosswalk['cbsatitle'].iloc[-1]
        
        #Now that we've identified the MSA, we need to know the primary state FIPS code for the BLS API. Many MSAs are in multiple states.
        cbsa_main_state   = cbsa_name.split(', ')[1][0:2] #The 2 character code for the main state of the msa
        all_states_in_msa = cbsa_name.split(', ')[1].split('-')

        state_fips                = pd.read_csv(os.path.join(data_location,'State Names.csv')) #the dataframe with the state fips codes
        state_fips['State FIPS']  = state_fips['State FIPS'].astype(str)
        state_fips['State FIPS']  =  state_fips['State FIPS'].str.zfill(2)  #cleaning the dataframe with the state fips codes
        

        state_fips_restricted           =  state_fips.loc[state_fips['State Code'] == cbsa_main_state] #cutting down dataframe to only the row with the state whose code we are looking up
        cbsa_main_state_fips            =  state_fips_restricted['State FIPS'].iloc[0]
        
        #In case the MSA series on BLS don't use the primary state's FIPS code, collect all of the state codes for the MSA to try
        cbsa_all_state_fips  = []
        for state in all_states_in_msa:
            state_fips_restricted           =  state_fips.loc[state_fips['State Code'] == state] #cutting down dataframe to only the row with the state whose code we are looking up
            state_fip                       =  state_fips_restricted['State FIPS'].iloc[0]
            cbsa_all_state_fips.append(state_fip)

        return([cbsa,cbsa_name, cbsa_main_state_fips, cbsa_all_state_fips])
    else:
        return(['', '', '', ''])

def IdentifyNecta(cbsa):
        cbsa_necta_crosswalk = pd.read_excel(os.path.join(data_location,'cbsa_necta_crosswalk.xls'),
            dtype={'CBSA Code': object,
                   }
                                            )
        
        cbsa_necta_crosswalk['CBSA Code'] = cbsa_necta_crosswalk['CBSA Code'].astype(str)
        cbsa_necta_crosswalk              = cbsa_necta_crosswalk.loc[cbsa_necta_crosswalk['CBSA Code'] == cbsa] 
       
        if len(cbsa_necta_crosswalk) == 1:
            necta_code = cbsa_necta_crosswalk['Necta Code'].iloc[0]
        else:
            necta_code = ''
        return(str(necta_code))

def GetFIPSList():
    #Returns a list of County FIPS codes to do reports for
    fips_list = []
   
    #Give the user 30 seconds to decide if writing reports for metro areas or individual county entries
    try:
        fips_or_cbsa = input_with_timeout('Single county report (c) or reports for all counties in metro (m)?', 30).strip()
    except TimeoutExpired:
        fips_or_cbsa = ''

    #If user enters nothing, return empty list
    if fips_or_cbsa == '':
        return(fips_list)

    #If user enters a metro area, return a list of FIPS for all the counties in that metro area
    elif fips_or_cbsa == 'm': #if user selects whole metro break out of the main loop
        cbsacode = str(input('Enter the MSA CBSA Code')).strip()
        
        cbsa_fips_crosswalk = pd.read_csv(os.path.join(data_location,'cbsa2fipsxw.csv'),
                dtype={'cbsacode': object,
                       'metrodivisioncode': object,
                       'csacode': object,
                       'cbsatitle': object,
                       'metropolitanmicropolitanstatis': object,
                       'metropolitandivisiontitle': object,
                       'csatitle': object,
                       'countycountyequivalent': object,
                       'statename': object,
                       'fipsstatecode': object,
                       'fipscountycode': object,
                       'centraloutlyingcounty': object
                }
                                        )


        #Add missing 0s
        cbsa_fips_crosswalk['fipsstatecode']  =  cbsa_fips_crosswalk['fipsstatecode'].str.zfill(2)
        cbsa_fips_crosswalk['fipscountycode'] =  cbsa_fips_crosswalk['fipscountycode'].str.zfill(3)
        cbsa_fips_crosswalk['FIPS Code']      = cbsa_fips_crosswalk['fipsstatecode'] + cbsa_fips_crosswalk['fipscountycode']

        cbsa_fips_crosswalk                   = cbsa_fips_crosswalk.loc[cbsa_fips_crosswalk['metropolitanmicropolitanstatis'] == 'Metropolitan Statistical Area'] #restrict to MSAs
        cbsa_fips_crosswalk                   = cbsa_fips_crosswalk.loc[cbsa_fips_crosswalk['cbsacode'] == cbsacode] #restrict data to only rows for counties within the MSA
        fips_list                             = cbsa_fips_crosswalk['FIPS Code'].unique()

    #If user enters county, return a list of FIPS that the user can enter
    elif fips_or_cbsa == 'c':#if user selects individual fips, have them add them in manually
        fips = None
        while fips != '':
            fips  =         str(input('What is the 5 digit county FIPS code?')).strip()
            try:
                assert len(fips) == 5
                fips_list.append(fips)
            except Exception as e:
                pass
    
    if fips_list != []:
        print('Preparing Reports for the following fips: ',fips_list)

    return(fips_list)

DeclareAPIKeys()

#Decide if you want to export data in excel files in the county folder
data_export                   = False

#Set formatting paramaters for reports
primary_font                  = 'Avenir Next LT Pro Light' 
primary_space_after_paragraph = 8
#Set graph size and format variables
tickangle                     = 0
bowery_grey                   = "#D7DEEA"
bowery_dark_grey              = "#A6B0BF"
bowery_dark_blue              = "#4160D3"
bowery_light_blue             = "#B3C3FF"
bowery_black                  = "#404858"

marginInches                  = 1/18
ppi                           = 96.85 
width_inches                  = 6.5
height_inches                 = 3.3

graph_width                   = (width_inches - marginInches)   * ppi
graph_height                  = (height_inches  - marginInches) * ppi

#Set scale for resolution 1 = no change, > 1 increases resolution. Very important for run time of main script). 
scale                         = 7

#Set tick font size (also controls legend font size)
tickfont_size                 = 8 

#Set Margin parameters/legend location
left_margin                   = 0
right_margin                  = 0
top_margin                    = 75
bottom_margin                 = 10
legend_position               = 1.10
paper_backgroundcolor         = 'white'
title_position                = .95
horizontal_spacing            = 0.1

todays_date                   = date.today()
current_year_and_quarter      = GetCurrentQuarterAndYear()
current_year                  = current_year_and_quarter[0]
current_quarter_number        = current_year_and_quarter[1]
current_quarter               = current_year + ' Q' + current_quarter_number
new_england_states            = ['MA','VT','RI','ME','NH','CT']


#Set number of years we want to look back to calculate employment growth
growth_period                 = 5
end_year                      = FindBLSEndYear()
start_year                    = end_year - growth_period        #For BLS Series 
observation_start             = '01/01/' + str(start_year -1)   #For FRED
observation_start_less1       = '01/01/' + str(start_year -2)   #For FRED for series 1 year behind the rest
qcew_year                     = current_year                    #for quarterly census of employment and wages
qcew_qtr                      = '3'                             #for quarterly census of employment and wages


#This is the main loop for our program, we loop through a list of county FIPS codes
for i, fips in enumerate(GetFIPSList()):
    
    try:
        assert type(fips) == str
        master_county_list = pd.read_excel(os.path.join(data_location, 'County_Master_List.xls'),
                dtype={'FIPS Code': object
                      }
                                          )
        master_county_list = master_county_list.loc[(master_county_list['FIPS Code'] == fips)]
        assert len(master_county_list) == 1

        state                = master_county_list['State'].iloc[0]    
        state_name           = GetStateName(state_code=state)
        county               = master_county_list['County Name'].iloc[0]
        cbsa                 = IdentifyMSA(fips)[0]
        cbsa_name            = IdentifyMSA(fips)[1]
        cbsa_main_state_fips = IdentifyMSA(fips)[2] #the state fips code of the first state listed for a msa        
        cbsa_all_state_fips  = IdentifyMSA(fips)[3] #a list of 2 digit FIPS codes for each state the MSA is in
        
        if state in new_england_states:
            necta_code       = IdentifyNecta(cbsa = cbsa)
        county               = county.split(",")[0]    

        Main()
        
    except Exception as e:
        print(e)
        print('Report Creation Failed for : ',fips)


#Now that we are done creating reports, crawl through our directory and create CSV file on exisiting reports
CreateDirectoryCSV()
 
#Post an update request to the Market Research Docs Service to update the database
UpdateServiceDb(report_type = 'areas', 
                csv_name    = service_api_csv_name, 
                csv_path    = os.path.join(main_output_location, service_api_csv_name),
                dropbox_dir = 'https://www.dropbox.com/home/Research/Market Analysis/Area/'
                )




