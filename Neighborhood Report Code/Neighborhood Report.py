#By Mike Leahy
#Started 06/30/2021
#Summary: This script creates reports on neighborhoods/cities for Bowery

import json
import msvcrt
import os
import random
import re
import shutil
import sys
import time
from datetime import date, datetime
from operator import ne
from statistics import mean

import googlemaps
import mpu
import pandas as pd
import plotly.graph_objects as go
import pyautogui
import requests
import shapefile
import us
import wikipedia
from bs4 import BeautifulSoup
from census import Census
from census_area import Census as CensusArea
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps
from plotly.subplots import make_subplots
from requests.adapters import HTTPAdapter
from requests.exceptions import HTTPError
from requests.packages.urllib3.util.retry import Retry
from selenium import webdriver
from shapely.geometry import LineString, Point, Polygon
from shapely.ops import nearest_points
from walkscore import WalkScoreAPI
from yelpapi import YelpAPI

#Define file paths
dropbox_root                   =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)')               
project_location               =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)','Research','Projects', 'Research Report Automation Project') 
main_output_location           =  os.path.join(project_location,'Output', 'Neighborhood')                     #testing
main_output_location           =  os.path.join(dropbox_root,'Research', 'Market Analysis','Neighborhood')     #production
data_location                  =  os.path.join(project_location,'Data', 'Neighborhood Reports Data')
general_data_location          =  os.path.join(project_location,'Data', 'General Data')
graphics_location              =  os.path.join(project_location,'Data', 'General Data','Graphics')
map_location                   =  os.path.join(project_location,'Data', 'Neighborhood Reports Data','Neighborhood Maps')
nyc_cd_map_location            =  os.path.join(project_location,'Data', 'Neighborhood Reports Data','NYC_CD Maps')
neighborhood_shapes_location   =  os.path.join(data_location, 'Neighborhood Shapes')
salesforce_report              =  os.path.join(project_location,'Data', 'Neighborhood Reports Data','Salesforce') 

#Data Manipulation functions
def ConvertListElementsToFractionOfTotal(raw_list):
    #Convert list with raw totals into a list where each element is a fraction of the total
    total = sum(raw_list)

    converted_list = []
    for i in raw_list:
        assert i >= 0
        converted_list.append(i/total * 100)
    
    return(converted_list)

def AggregateAcrossDictionaries(neighborhood_tracts_data, fields_list):
    aggregate_dict = {} 
    for field in fields_list:
        total_value = 0

        #Add up all the values from each dictionary
        for d in neighborhood_tracts_data:
            value       = d[field]
            total_value = total_value + int(value)
        
        #Add the current field to the new aggregate_dict
        aggregate_dict[field] = total_value
    
    return(aggregate_dict)

def FindMedianCategory(frequency_list, category_list):
    #Takes a list with a fequency distribution (eg [10%,30%,60%]) and the corresponding cateorgories [Red,Blue,Green]
    #Returns the median category, in this case Green
    assert len(frequency_list) == len(category_list)

    total_value_fraction = 0
    for i,value_category_fraction in enumerate(frequency_list):
        total_value_fraction += value_category_fraction
        if total_value_fraction >= 50:
            median_cat_index = i
            break
    
    median_category     = category_list[median_cat_index]
    return(median_category)

def FilterDictionary(dictionary,desired_keys):
    #This function takes a dictionary and a list of desired keys. 
    #It loops through the dictionary keeping only the keys that are in our list
    #Returns the dictionary without the unwanted items
    assert type(dictionary) == dict

    unwanted_keys = []          
    for key in dictionary.keys():
        if key not in desired_keys:
            unwanted_keys.append(key)
    
    #Go through our list of unwanted keys and delete them
    for unwanted_key in unwanted_keys:
        del dictionary[unwanted_key]
    return(dictionary)
#####################################################Geographic Data Related Functions####################################
def GetLatandLon():
    #Look up lat and lon of area with geocoding using google maps api
    gmaps          = googlemaps.Client(key=google_maps_api_key) 
    
    if neighborhood_level == 'custom':
        geocode_result = gmaps.geocode(address=(neighborhood + ', ' + comparison_area + ',' + comparison_state),)
    else:
        geocode_result = gmaps.geocode(address=(neighborhood + ',' + hood_state),)
    
    latitude       = geocode_result[0]['geometry']['location']['lat']
    longitude      = geocode_result[0]['geometry']['location']['lng']

    return([latitude,longitude]) 

def GetNeighborhoodShape():
    global neighborhood_shape_polygon
    if neighborhood_level == 'custom':
        city_geo_json_file_path = os.path.join(neighborhood_shapes_location,'Custom Hood Shapes',comparison_area + '.geojson')
        
        #If we have the geojson file downloaded for the comparison city
        if os.path.exists(city_geo_json_file_path):
            
            #Method 1: Pull geojson from file with city name
            with open(city_geo_json_file_path) as infile: #Open a geojson file with the city as the name the name of the file with the neighborhood boundries for that city
                my_shape_geojson = json.load(infile)
            
            print('Successfully opened geojson file for ' + comparison_area)

            try:    
                #Iterate through the features in the file (each feature is a negihborhood) and find the boundry of interest
                for i in range(len(my_shape_geojson['features'])):
                    feature_hood_name = my_shape_geojson['features'][i]['properties']['name']
                    if feature_hood_name == neighborhood:
                        neighborhood_shape = my_shape_geojson['features'][i]['geometry']
                        print('Successfully pulled hood shape from stored geojson file')
                        
                        #Now that we have grabbed the coordinates for the area, export it as shapefile
                        try:
                            coordinates                    = neighborhood_shape['coordinates'][0]
                            neighborhood_shape_polygon     = Polygon(coordinates)
                            PolygonToShapeFile(poly        = neighborhood_shape_polygon)

                        except Exception as e:
                            print(e,'unable to export neighborhood polygon as shape on first try')
                            try:
                                coord_tuple_list           = [tuple(l) for l in neighborhood_shape['coordinates'][0][0]]
                                neighborhood_shape_polygon = Polygon(coord_tuple_list)
                                PolygonToShapeFile(poly    = neighborhood_shape_polygon)
                            except Exception as e:
                                print(e,'unable to export neighborhood polygon as shape on second try, bummer ')
                        return(neighborhood_shape) 
                
            except Exception as e:
                print(e,'unable to get geography from the city geojson file even tho it exists')
        
        #If we don't have the geojson file downloaded for the comparison city        
        confirmation = input('Unable to find geography from the ' + comparison_area + ' geojson file. ' + 'Press enter to confirm you are using a hand drawn geography downloaded from online')
        
        #Define file locations
        file_download_location             = os.path.join(os.environ['USERPROFILE'],'Downloads', 'map.geojson') #download from here: http://geojson.io/#map=5/34.071/-72.817
        new_geojson_file_location          = os.path.join(data_location,'Neighborhood Shapes','Custom Hood Shapes', 'map.geojson')
        final_geojson_file_location        = os.path.join(hood_folder_map,'map.geojson')


        #Step 1: Move the exported geojson file from downloads to data folder 
        if os.path.exists(file_download_location):
            print('Moving custom geojson file from downloads folder into data folder') 
            shutil.move(file_download_location, new_geojson_file_location)
            time.sleep(3)
        with open(new_geojson_file_location) as infile: #Open a geojson file with the city as the name the name of the file with the neighborhood boundries for that city
                my_shape_geojson = json.load(infile)


        neighborhood_shape = my_shape_geojson['features'][0]['geometry']
        print('Successfully pulled hood shape from downloaded custom geojson file')

         #Now that we have grabbed the coordinates for the area, export it as shapefile
        try:
            coord_tuple_list = [tuple(l) for l in neighborhood_shape['coordinates'][0]]
            neighborhood_shape_polygon = Polygon(coord_tuple_list)
            PolygonToShapeFile(poly = neighborhood_shape_polygon)
        except Exception as e:
                print(e,'unable to export neighborhood polygon as shape')
        
        #Move the geojson file from the data folder into the subject hoods map folder
        shutil.move(new_geojson_file_location, final_geojson_file_location)

        return(neighborhood_shape) 


    elif neighborhood_level == 'place':
        try:
            shapefile_location = os.path.join(neighborhood_shapes_location,'Census Place Shapes',('tl_2021_' + hood_state_fips + '_place'),('tl_2021_' + hood_state_fips + '_place.shp'))
            assert os.path.exists(shapefile_location)

            #Open the shapefile
            place_map = shapefile.Reader(shapefile_location)
        
            #Loop through each place in the shape file
            for i in range(len(place_map)):
                place_record = place_map.shapeRecord(i)
                #Look for the record that corresponds to our subject city
                if place_record.record['PLACEFP'] != hood_place_fips:
                    continue
                else:
                    neighborhood_shape         =  place_map.shape(i)
                    neighborhood_shape_polygon = Polygon(neighborhood_shape.points)
                    print('Successfully pulled census shape from shapefile')
                    try:
                        PolygonToShapeFile(poly = neighborhood_shape_polygon)
                    except Exception as e:
                        print(e,'unable to export city polygon as shape')

                    print('Successfully created polygon object from census shape')

                    return(neighborhood_shape) 
        except Exception as e:
            print(e,'unable to get shape for census place')
    
    elif neighborhood_level == 'county subdivision':
        try:
            shapefile_location = os.path.join(neighborhood_shapes_location,'Census County Subdivision Shapes',('tl_2021_' + hood_state_fips + '_cousub'),('tl_2021_' + hood_state_fips + '_cousub.shp'))
            assert             os.path.exists(shapefile_location)

            #Open the shapefile
            place_map          = shapefile.Reader(shapefile_location)

        

            #Loop through each subdivision in the shape file
            for i in range(len(place_map)):
                place_record = place_map.shapeRecord(i)
                #Look for the record that corresponds to our subject city
                if place_record.record['COUNTYFP'] != hood_county_fips or  place_record.record['COUSUBFP'] != hood_suvdiv_fips:
                    continue
                else:
                    neighborhood_shape         =  place_map.shape(i)
                    neighborhood_shape_polygon = Polygon(neighborhood_shape.points)
                    print('Successfully pulled census shape from shapefile')
                    try:
                        PolygonToShapeFile(poly = neighborhood_shape_polygon)
                    except Exception as e:
                        print(e,'unable to export subdivision polygon as shape')

                    print('Successfully created polygon object from census shape')

                    return(neighborhood_shape) 
        except Exception as e:
            print(e,'unable to get shape for census place')

def PolygonToShapeFile(poly):
        #WRITE TO SHAPEFILE USING PYSHP
        target_file_path = os.path.join(hood_folder_map, 'my.shp')
        shapewriter      = shapefile.Writer(target = target_file_path)
        shapewriter.field("field1")

        #step1: convert shapely to pyshp using the function above
        converted_shape = shapely_to_pyshp(poly)

        #step2: tell the writer to add the converted shape
        shapewriter.shape(converted_shape)
        
        #add a list of attributes to go along with the shape
        shapewriter.record(["empty record"])
        
        #save it
        shapewriter.close()

def shapely_to_pyshp(shapelygeom):
    #first convert shapely to geojson
    try:
        shapelytogeojson = shapely.geometry.mapping
    except:
        import shapely.geometry
        shapelytogeojson = shapely.geometry.mapping
    
    geoj = shapelytogeojson(shapelygeom)
    
    #create empty pyshp shape
    record = shapefile.Shape()
    
    #set shapetype
    if geoj["type"] == "Null":
        pyshptype = 0
    elif geoj["type"] == "Point":
        pyshptype = 1
    elif geoj["type"] == "LineString":
        pyshptype = 3
    elif geoj["type"] == "Polygon":
        pyshptype = 5
    elif geoj["type"] == "MultiPoint":
        pyshptype = 8
    elif geoj["type"] == "MultiLineString":
        pyshptype = 3
    elif geoj["type"] == "MultiPolygon":
        pyshptype = 5
    record.shapeType = pyshptype
    
    #set points and parts
    if geoj["type"] == "Point":
        record.points = geoj["coordinates"]
        record.parts = [0]
    elif geoj["type"] in ("MultiPoint","Linestring"):
        record.points = geoj["coordinates"]
        record.parts = [0]
    elif geoj["type"] in ("Polygon"):
        record.points = geoj["coordinates"][0]
        record.parts = [0]
    elif geoj["type"] in ("MultiPolygon","MultiLineString"):
        index = 0
        points = []
        parts = []
        for eachmulti in geoj["coordinates"]:
            points.extend(eachmulti[0])
            parts.append(index)
            index += len(eachmulti[0])
        record.points = points
        record.parts = parts
    return (record)

def GetListOfNeighborhoods(city):
    try:
            #Method 1: Pull geojson from file with city name
        with open(os.path.join(data_location,'Neighborhood Shapes','Custom Hood Shapes',city + '.geojson')) as infile: #Open a geojson file with the city as the name the name of the file with the neighborhood boundries for that city
                    my_shape_geojson = json.load(infile)
                
        #Iterate through the features in the file (each feature is a neighborhood) and find the boundry of interest
        feature_hood_names = []
        for i in range(len(my_shape_geojson['features'])):
            feature_hood_name = my_shape_geojson['features'][i]['properties']['name']
            feature_hood_names.append(feature_hood_name)
            
        return(feature_hood_names) 
    except Exception as e:
        print(e)
        return([])

def DetermineNYCCommunityDistrict(lat,lon):
    print('Determining NYC Community District')
    try:
        #Method 1: Pull geojson from file with city name
        with open(os.path.join(neighborhood_shapes_location,'Custom Hood Shapes','NY','nyc_communitydistricts.json')) as infile: #Open a geojson file with the city as the name the name of the file with the neighborhood boundries for that city
            my_shape_geojson = json.load(infile)
        
        #Iterate through the features in the file (each feature is a communtiy district) and find the boundry of interest
        for communtiy_district in range(len(my_shape_geojson['features'])):
            communtiy_district_number = my_shape_geojson['features'][communtiy_district]['properties']["BoroCD"]
            communtiy_district_shape  = my_shape_geojson['features'][communtiy_district]['geometry']['coordinates'][0]
            
            try:
                point   = Point(lon, lat)
                polygon = Polygon([tuple(l) for l in communtiy_district_shape])      
                                
                #Check if lat and lon is inside the communtiy district
                if polygon.contains(point) == True:
                    print('The NYC Community District is ',str(communtiy_district_number))
                    return(str(communtiy_district_number))

            except Exception as e:
                continue
        
        return('x')


    except Exception as e:
        print(e,'Unable to search for NYC communtiy district')
        return('x')

def TransitAgencyIdToName(id):
    #Takes a National Transit Database agency ID number and returns the name of the transit agency
    try:
        #Open the ID to name crosswalk file with agency info
        df           = pd.read_excel(os.path.join(general_data_location,'Geographic Data','GTFS_NTM_Stops','2020 Agency Information.xlsx')) #from here: https://www.transit.dot.gov/ntd/data-product/2020-annual-database-agency-information
        df['NTD ID'] = df['NTD ID'].astype(str)
        df           = df.loc[df['NTD ID'] == id]
        agency_name  = df['Agency Name'].iloc[0]    
        return(agency_name)
    except Exception as e:
        print(e,'Couldnt find Transit Agency name from ID number')
        return('----------')
#####################################################User FIPS input proccessing Functions####################################

def ProcessPlaceFIPS(place_fips):
    #This function takes a user provided 7 digit census place fips code and returns a list of key variables about that fips code
    #eg: the place name, type, state name, state code, etc

    #Process the FIPS code provided    
    place_fips                      = place_fips.replace('-','').strip()
    assert len(place_fips)          == 7
    state_fips                      = place_fips[0:2]
    place_fips                      = place_fips[2:]

    #Get name of the hood using the FIPS code provided
    place_name                      = c.sf1.state_place(fields=['NAME'], state_fips = state_fips, place = place_fips)
    assert                          place_name != []
    place_name                      = place_name[0]['NAME']
    state_full_name                 = place_name.split(',')[1].strip()
    place_name                      = place_name.split(',')[0].strip()
    place_type                      = place_name.split(' ')[len(place_name.split(' '))-1] #eg: village, city, etc
    place_name                      = ' '.join(place_name.split(' ')[0:len(place_name.split(' '))-1]).title()
    place_name                      = place_name.replace("""\\""",'-')
    place_name                      = place_name.replace("""/""",'-')
    

    
    #Name of State
    state                           = us.states.lookup(state_full_name) #convert the full state name to the 2 letter abbreviation
    state                           = state.abbr
    assert len(state)               == 2

    return([place_fips, state_fips, place_name, state_full_name, state,place_type])

def ProcessCountyFIPS(county_fips):
    
    #Process the FIPS code provided by user
    county_fips               = county_fips.replace('-','').strip()
    assert len(county_fips)   == 5
    state_fips                = county_fips[0:2]
    county_fips               = county_fips[2:]

    #Get name of county
    name                      = c.sf1.state_county(fields=['NAME'], state_fips = state_fips, county_fips = county_fips)[0]['NAME']
    state_full_name           = name.split(',')[1].strip()
    name                      = name.split(',')[0].strip().title()
    name                      = name.replace("""'S""","""'s""")


    #Change county name for NYC counties
    if county_fips == '061' and name == 'New York County':
        name = 'Manhattan'
    elif county_fips == '081' and name == 'Queens County':
        name = 'Queens'
    elif county_fips == '047' and name == 'Kings County':
        name = 'Brooklyn'
    elif county_fips == '085' and name == 'Richmond County':
        name = 'Staten Island'
    elif county_fips == '005' and name == 'Bronx County':
        name = 'The Bronx'

    #Name of State
    state                      = us.states.lookup(state_full_name) #convert the full state name to the 2 letter abbreviation
    state                      = state.abbr

    assert len(state)          == 2

    return[county_fips, state_fips, name, state_full_name, state]

def ProcessCountySubdivisionFIPS(county_subdivision_fips):
    #Proccess FIPS code provided
    county_subdivision_fips = county_subdivision_fips.replace('-','').strip()
    assert len(county_subdivision_fips) == 10
    state_fips       = county_subdivision_fips[0:2]
    county_fips      = county_subdivision_fips[2:5]
    suvdiv_fips      = county_subdivision_fips[5:]

    #Get name of hood using the FIPS code provided
    name             = c.sf1.state_county_subdivision(fields=['NAME'],state_fips = state_fips,county_fips = county_fips, subdiv_fips = suvdiv_fips)[0]['NAME']
    state_full_name  = name.split(',')[2].strip()
    name             = name.split(',')[0].strip().title()
    place_type       = name.split(' ')[len(name.split(' '))-1] #eg: village, city, etc
    name             = ' '.join(name.split(' ')[0:len(name.split(' '))-1]).title()
    
    if place_type    == 'Township':
        name         = name + ' ' + place_type
    
    #Name of State
    state            = us.states.lookup(state_full_name) #convert the full state name to the 2 letter abbreviation
    state            = state.abbr
    assert len(state) == 2

    return([suvdiv_fips, county_fips, name, state_fips, state_full_name, state,place_type])
     
def ProcessCountyTract(tract,county_fips):
    #Takes a user provided county fips code and a census tract number and returns a list of key variables
    county_fips               = county_fips.replace('-','').strip()
    assert len(county_fips)   == 5
    state_fips                = county_fips[0:2]
    county_fips               = county_fips[2:]

    tract                     = tract.replace('-','').strip()
    assert len(tract)         == 6

    #Get name of tract
    name                      = c.sf1.state_county_tract(fields=['NAME'],state_fips = state_fips, county_fips = county_fips,tract = tract)[0]['NAME']
    state_full_name           = name.split(',')[2].strip()
    name                      = name.split(',')[0] + ',' +  name.split(',')[1]
    name                      = name.strip().title()

    #Name of State
    state                     = us.states.lookup(state_full_name) #convert the full state name to the 2 letter abbreviation
    state                     = state.abbr
    assert len(state)         == 2

    return([county_fips, tract, name, state_full_name, state, state_fips])

def ProcessZipCode(zip_code):
    #Process the zip code provided
    zip_code                            = str(zip_code).replace('-','').strip()
    assert len(zip_code) == 5
    
    #Get the state FIPS code (eg New York: 36)
    zip_county_crosswalk_df            = pd.read_excel(os.path.join(data_location,'Census Area Codes','ZIP_COUNTY_092021.xlsx')) #read in crosswalk file
    zip_county_crosswalk_df['ZIP']     = zip_county_crosswalk_df['ZIP'].astype(str)
    zip_county_crosswalk_df['ZIP']     = zip_county_crosswalk_df['ZIP'].str.zfill(5)
    zip_county_crosswalk_df['COUNTY']  = zip_county_crosswalk_df['COUNTY'].astype(str)
    zip_county_crosswalk_df['COUNTY']  = zip_county_crosswalk_df['COUNTY'].str.zfill(5)

    zip_county_crosswalk_df            = zip_county_crosswalk_df.loc[zip_county_crosswalk_df['ZIP'] == zip_code]                 #restrict to rows for zip code
    county_fips                        = str(zip_county_crosswalk_df['COUNTY'].iloc[-1])[2:]
    state_fips                         = str(zip_county_crosswalk_df['COUNTY'].iloc[-1])[0:2] #Get state fips from the county fips code (the county the zip code is in)
    assert len(state_fips) == 2

    #Get name of hood
    name                               = c.sf1.state_zipcode(fields=['NAME'],state_fips=state_fips, zcta=zip_code)[0]['NAME']
    state_full_name                    = name.split(',')[1].strip()
    name                               = name.split(',')[0].replace('ZCTA5','').strip().title() + ' (Zip Code)'

    #Name of State
    state                              = us.states.lookup(state_full_name) #convert the full state name to the 2 letter abbreviation
    state                              = state.abbr
    assert                             len(state) == 2


    return([county_fips, zip_code, name, state_full_name, state, state_fips])

def PlaceFIPSToCountyFIPS(place_fips,state_fips):
    print('Looking for county fips code')
    #Takes 7 digit place fips code for a city and returns the 5 digit fips code for that city
    
    #Open file with place fips code and county fips code
    place_county_crosswalk_df                            = pd.read_csv(os.path.join(data_location,'Census Area Codes','national_places.csv'),encoding='latin-1') #read in crosswalk file
    
    place_county_crosswalk_df['PLACEFP']                 = place_county_crosswalk_df['PLACEFP'].astype(str)
    place_county_crosswalk_df['PLACEFP']                 = place_county_crosswalk_df['PLACEFP'].str.zfill(5)

    place_county_crosswalk_df['STATEFP']                 = place_county_crosswalk_df['STATEFP'].astype(str)
    place_county_crosswalk_df['STATEFP']                 = place_county_crosswalk_df['STATEFP'].str.zfill(2)

    place_county_crosswalk_df['County_FIPS']             = place_county_crosswalk_df['County_FIPS'].astype(str)
    place_county_crosswalk_df['County_FIPS']             = place_county_crosswalk_df['County_FIPS'].str.zfill(7)


    #Restrict to observations that include the provieded place fips
    place_county_crosswalk_df                            = place_county_crosswalk_df.loc[(place_county_crosswalk_df['PLACEFP'] == str(place_fips)) & (place_county_crosswalk_df['STATEFP'] == str(state_fips))].reset_index()                 
    place_county_crosswalk_df                            = place_county_crosswalk_df.loc[(place_county_crosswalk_df['TYPE'] != 'County Subdivision')].reset_index()                 
    assert len(place_county_crosswalk_df)                == 1
    county_fips                                          = str(place_county_crosswalk_df['County_FIPS'].iloc[-1])[0:5]
    
    #When the city is in multiple counties, let the user select, if not, return the coutny fips
    if pd.isnull(place_county_crosswalk_df['Second_County'].iloc[-1]):
        return(county_fips)

    else:
        input_county_fips = str(input('The city/place is in more than 1 county, enter the desired comparison county FIPS'))
        assert len(input_county_fips) == 5
        return(input_county_fips)
   
def PlaceNameToPlaceFIPS(place_name,state_code):
    #Takes place name and returns the 7 digit fips code for a city 
    
    #Open file with place fips code and county fips code
    place_county_crosswalk_df                            = pd.read_csv(os.path.join(data_location,'Census Area Codes','national_places.csv'),encoding='latin-1') #read in crosswalk file
    
    place_county_crosswalk_df['PLACEFP']                 = place_county_crosswalk_df['PLACEFP'].astype(str)
    place_county_crosswalk_df['PLACEFP']                 = place_county_crosswalk_df['PLACEFP'].str.zfill(5)

    place_county_crosswalk_df['State_Place_FP']          = place_county_crosswalk_df['State_Place_FP'].astype(str)
    place_county_crosswalk_df['State_Place_FP']          = place_county_crosswalk_df['State_Place_FP'].str.zfill(7)


    #Restrict to observations that include the provieded place fips
    place_county_crosswalk_df                            = place_county_crosswalk_df.loc[(place_county_crosswalk_df['Neighborhood_District'] == str(place_name)) & (place_county_crosswalk_df['STATE'] == str(state_code)) & (place_county_crosswalk_df['TYPE'] != 'County Subdivision') ].reset_index()                 
    
    #Return the last row if that's there's only one, otherwise ask user to choose
    if len(place_county_crosswalk_df) == 1:
        county_fips                                      = str(place_county_crosswalk_df['State_Place_FP'].iloc[-1])[0:7]
    
    elif len(place_county_crosswalk_df) > 1:
        print(place_county_crosswalk_df)
        selected_county = int(input('There are more than 1 counties for this city: enter the number of your choice'))  
        county_fips                                      = str(place_county_crosswalk_df['State_Place_FP'].iloc[selected_county])[0:7]
    else:
        return(None)


    return(county_fips)

def SubdivsionNameToFIPS(subdivision_name,state_code):
    #Takes subdivision name and returns the 10 digit fips code for it
    
    #Open file with place fips code and county fips code
    place_county_crosswalk_df                             = pd.read_csv(os.path.join(data_location,'Census Area Codes','national_cousub.csv'),encoding='latin-1') #read in crosswalk file
    
    place_county_crosswalk_df['COUSUBFP']                 = place_county_crosswalk_df['COUSUBFP'].astype(str)
    place_county_crosswalk_df['COUSUBFP']                 = place_county_crosswalk_df['COUSUBFP'].str.zfill(5)

    place_county_crosswalk_df['COUNTYFP']                 = place_county_crosswalk_df['COUNTYFP'].astype(str)
    place_county_crosswalk_df['COUNTYFP']                 = place_county_crosswalk_df['COUNTYFP'].str.zfill(3)

    place_county_crosswalk_df['STATEFP']                  = place_county_crosswalk_df['STATEFP'].astype(str)
    place_county_crosswalk_df['STATEFP']                  = place_county_crosswalk_df['STATEFP'].str.zfill(2)

    place_county_crosswalk_df['COUSUBNAME']               = place_county_crosswalk_df['COUSUBNAME'].astype(str)
    place_county_crosswalk_df['COUSUBNAME']               = place_county_crosswalk_df['COUSUBNAME'].str.strip()
    place_county_crosswalk_df['COUSUBNAME']               = place_county_crosswalk_df['COUSUBNAME'].str.split(' ',-1)
    place_county_crosswalk_df['COUSUBNAMELEN']            = place_county_crosswalk_df['COUSUBNAME'].str.len()

    
    #Cut off the last word in each county subdivision name
    def drop_last_item(item):
        if len(item) > 1:
            return( ' '.join(item[0:(int(len(item))-1)]))
        else:
            return(item)
        
    place_county_crosswalk_df['COUSUBNAME'] = place_county_crosswalk_df['COUSUBNAME'].apply(drop_last_item)
    
    #Restrict to observations that include the provieded place fips
    place_county_crosswalk_df               = place_county_crosswalk_df.loc[(place_county_crosswalk_df['COUSUBNAME'] == str(subdivision_name)) & (place_county_crosswalk_df['STATE'] == str(state_code))  ].reset_index()                 
    
    #Return the last row if that's there's only one, otherwise ask user to choose
    if len(place_county_crosswalk_df['COUNTYFP'].unique()) == 1:
        subdiv_fips                         = str(place_county_crosswalk_df['STATEFP'].iloc[-1])  +  str(place_county_crosswalk_df['COUNTYFP'].iloc[-1])   +  str(place_county_crosswalk_df['COUSUBFP'].iloc[-1])  
    
    #If there's more than one unique county, let user choose
    elif  len(place_county_crosswalk_df['COUNTYFP'].unique()) > 1:
        print(place_county_crosswalk_df)
    
        try:
            selected_county = int(input_with_timeout('There are more than 1 counties for this subdivision: enter the number of your choice',10))  
        
        except TimeoutExpired:
            selected_county = 0

        subdiv_fips                         = str(place_county_crosswalk_df['STATEFP'].iloc[selected_county])  +  str(place_county_crosswalk_df['COUNTYFP'].iloc[selected_county])   +  str(place_county_crosswalk_df['COUSUBFP'].iloc[selected_county])  


    else:
        return(None)
	
    return(subdiv_fips)

def SalesforcePlaceFIPSList():
    print('Getting list of Place Fips from our salesforce export')
    #Open Salesforce Report
    salesforce_df              =  pd.read_csv(os.path.join(salesforce_report,'report.csv'))
    
    place_fips_list            = []
    city_name_list             = list(salesforce_df['Property: Neighborhood/District'])
    state_code_list            = list(salesforce_df['Property: State'])
        
    for loop_city,sc in zip(city_name_list, state_code_list):
        place_fips_list.append(PlaceNameToPlaceFIPS(place_name = loop_city, state_code = sc))  

    return(place_fips_list)     

def SalesforceSubdivisionFIPSList():
    print('Getting list of subdivision Fips from our salesforce export')

    #Open Salesforce Report
    salesforce_df              =  pd.read_csv(os.path.join(salesforce_report,'report.csv'))
    
    subdiv_fips_list           = []
    subdiv_name_list           = list(salesforce_df['Property: Neighborhood/District']) 
    state_code_list            = list(salesforce_df['Property: State'])
        
    for loop_subdiv,sc in zip(subdiv_name_list,state_code_list):
        subdiv_fips_list.append(SubdivsionNameToFIPS(subdivision_name= loop_subdiv,state_code = sc))  

    return(subdiv_fips_list)

def CountyInputPlaceFIPSList(county_fips):
    #Takes a county fips code and returns a list of place fips code in that county
    print('Getting list of place fips within ' + county_fips)

    #Open file with place fips code and county fips code
    place_county_crosswalk_df                                   = pd.read_csv(os.path.join(data_location,'Census Area Codes','national_places.csv'),encoding='latin-1',dtype={'County_FIPS':str,'State_Place_FP':str}) #read in crosswalk file
    
    place_county_crosswalk_df['State_Place_FP']                 = place_county_crosswalk_df['State_Place_FP'].astype(str)
    place_county_crosswalk_df['State_Place_FP']                 = place_county_crosswalk_df['State_Place_FP'].str.zfill(7)
    
    place_county_crosswalk_df['County_FIPS']                    = place_county_crosswalk_df['County_FIPS'].astype(str)
    place_county_crosswalk_df['County_FIPS']                    = place_county_crosswalk_df['County_FIPS'].str.zfill(5)
  
    #Restrict to observations that fall within the county fips provided
    place_county_crosswalk_df                                   = place_county_crosswalk_df.loc[(place_county_crosswalk_df['County_FIPS'] == str(county_fips)) ].reset_index()                 

    return(list(place_county_crosswalk_df['State_Place_FP']))
  
def CountyInputSubdivisionFIPSList(county_fips):
    #Takes a county fips code and returns a list of subdivision fips code in that county
    print('Getting list of subdivision fips within ' + county_fips)

    #Open file with place fips code and county fips code
    place_county_crosswalk_df                            = pd.read_csv(os.path.join(data_location,'Census Area Codes','national_cousub.csv'),encoding='latin-1',dtype={'STATEFP':str,'COUNTYFP':str,'COUSUBFP':str}) #read in crosswalk file
    
    place_county_crosswalk_df['STATEFP']                 = place_county_crosswalk_df['STATEFP'].astype(str)
    place_county_crosswalk_df['STATEFP']                 = place_county_crosswalk_df['STATEFP'].str.zfill(2)
    
    place_county_crosswalk_df['COUNTYFP']                = place_county_crosswalk_df['COUNTYFP'].astype(str)
    place_county_crosswalk_df['COUNTYFP']                = place_county_crosswalk_df['COUNTYFP'].str.zfill(3)
    place_county_crosswalk_df['COUNTYFP']                = place_county_crosswalk_df['STATEFP'] + place_county_crosswalk_df['COUNTYFP']

    place_county_crosswalk_df['COUSUBFP']                = place_county_crosswalk_df['COUSUBFP'].astype(str)
    place_county_crosswalk_df['COUSUBFP']                = place_county_crosswalk_df['COUSUBFP'].str.zfill(5)

    place_county_crosswalk_df['SUBDIVFIPS']              = place_county_crosswalk_df['COUNTYFP'] + place_county_crosswalk_df['COUSUBFP'] 

    #Restrict to observations that fall within the county fips provided
    place_county_crosswalk_df                            = place_county_crosswalk_df.loc[(place_county_crosswalk_df['COUNTYFP'] == str(county_fips)) ].reset_index()                 

    return(list(place_county_crosswalk_df['SUBDIVFIPS']))

#####################################################Misc Functions####################################
def CreateDirectory():
    print('Creating Directories and file name')
    global report_path,hood_folder_map,hood_folder
    
    state_folder_map         = os.path.join(map_location,hood_state)

    state_folder             = os.path.join(main_output_location,hood_state)

    if neighborhood_level == 'custom':

        if os.path.exists(state_folder) == False:
            os.mkdir(state_folder)  
        
        if os.path.exists(state_folder_map) == False:
            os.mkdir(state_folder_map)  

        if hood_state == ('NY') and ((comparison_area == 'Brooklyn') or (comparison_area == 'Staten Island') or (comparison_area == 'The Bronx') or  (comparison_area == 'Manhattan') or (comparison_area == 'Queens')   ):
            
            if (comparison_area == 'Brooklyn'):
                city_folder     =  os.path.join(main_output_location,hood_state,'NYC','BK')
                city_folder_map =  os.path.join(map_location,hood_state,'NYC','BK')

            elif (comparison_area == 'Staten Island'):
                city_folder     =  os.path.join(main_output_location,hood_state,'NYC','SI')
                city_folder_map =  os.path.join(map_location,hood_state,'NYC','SI')

            elif (comparison_area == 'The Bronx'):
                city_folder     =  os.path.join(main_output_location,hood_state,'NYC','BX')
                city_folder_map =  os.path.join(map_location,hood_state,'NYC','BX')
            
            elif (comparison_area == 'Manhattan'):
                city_folder     =  os.path.join(main_output_location,hood_state,'NYC','MA')
                city_folder_map =  os.path.join(map_location,hood_state,'NYC','MA')
            
            elif (comparison_area == 'Queens'):
                city_folder     =  os.path.join(main_output_location,hood_state,'NYC','QU')
                city_folder_map =  os.path.join(map_location,hood_state,'NYC','QU')
        
        
        else:
            city_folder         =  os.path.join(main_output_location,hood_state,comparison_area)
            city_folder_map     =  os.path.join(map_location,hood_state,comparison_area)

        if os.path.exists(city_folder) == False:
            os.mkdir(city_folder) 
        
        if os.path.exists(city_folder_map) == False:
            os.mkdir(city_folder_map) 


        hood_folder              = os.path.join(main_output_location,hood_state,city_folder,neighborhood).replace('/','-')
        hood_folder_map          = os.path.join(map_location,hood_state,city_folder_map,neighborhood).replace('/','-')


    else:
        hood_folder              = os.path.join(main_output_location,hood_state,neighborhood).replace('/','-')
        hood_folder_map          = os.path.join(map_location,hood_state,neighborhood).replace('/','-')
    

    for folder in [state_folder, hood_folder, state_folder_map, hood_folder_map]:
         if os.path.exists(folder):
            pass 
         else:
            os.mkdir(folder) 
    
    report_path = os.path.join(hood_folder,current_year + ' ' + hood_state + ' - ' + neighborhood.replace('/','-')  + ' - hood' + '_draft')[:255] 
    report_path = report_path + '.docx'


def FindZipCodeDictionary(zip_code_data_dictionary_list,zcta,state_fips):
    #This function takes a list of dictionaries, where each zip code gets its own dictionary. Takes a zip code and state fips code and finds and returns just that dictionary.
    #We need to use this, because the census api is causing an error that requires us to retrive data for all zip codes in the country
    for zcta_dictionary in  zip_code_data_dictionary_list:
        if zcta_dictionary['zip code tabulation area'] == zcta and zcta_dictionary['state'] == state_fips:
            return(zcta_dictionary)
    print('Could not find dictionary for given zip code: ', zcta )

#Data Gathering Related Functions
def DeclareAPIKeys():
    global census_api_key, walkscore_api_key, google_maps_api_key, yelp_api_key, yelp_api, yelp_client_id, location_iq_api_key
    global c, c_area, walkscore_api, zoneomics_api_key
    
    #Declare API Keys
    census_api_key_df             = pd.read_csv(  os.path.join(data_location,'API Keys','CensusAPIKeys.csv') )
    census_api_keys               = list(census_api_key_df['Key'])
    census_api_key                = random.choice(census_api_keys)

    walkscore_api_key             = '057f7c0a590efb7ec06da5a8735e536d'
    google_maps_api_key           = 'AIzaSyBMcoRFOW2rxAGxURCpA4gk10MROVVflLs'
    yelp_client_id                = 'NY9c0_9kvOU4wfzmkkruOQ'
    yelp_api_key                  = 'l1WjEgdgSMpU9PJtXEk0bLs4FJdsVLONqJLhbaA0gZlbFyEFUTTkxgRzBDc-_5234oLw1CLx-iWjr8w4nK_tZ_79qVIOv3yEMQ9aGcSS8xO1gkbfENCBKEl34COVYXYx'
    location_iq_api_key           = 'pk.8937271b8b15004065ca62552e7d06f7'
    zoneomics_api_key             = 'd69b3eee92f8d3cec8c71893b340faa8cb52e1b8'

    c                             = Census(census_api_key)     #Census API wrapper package
    c_area                        = CensusArea(census_api_key) #Census API package, sepearete extension of main package that allows for custom boundries
    yelp_api                      = YelpAPI(yelp_api_key)
    walkscore_api                 = WalkScoreAPI(api_key = walkscore_api_key)

#Data Gathering Related Functions
def DeclareFormattingParameters():
    global primary_font
    global primary_space_after_paragraph
    
    #Set formatting paramaters for reports
    primary_font                  = 'Avenir Next LT Pro Light' 
    primary_space_after_paragraph = 8
class TimeoutExpired(Exception):
    pass

def input_with_timeout(prompt, timeout, timer=time.monotonic):
    sys.stdout.write(prompt)
    sys.stdout.flush()
    endtime = timer() + timeout
    result  = []
    while timer() < endtime:
        if msvcrt.kbhit():
            result.append(msvcrt.getwche()) #XXX can it block on multibyte characters?
            if result[-1] == '\r':
                return ''.join(result[:-1])
        time.sleep(0.04) # just to yield to other processes/threads
    raise TimeoutExpired

#####################################################Census Data Related Functions####################################
def GetCensusFrequencyDistribution(geographic_level,hood_or_comparison_area,fields_list,operator):
    #A general function that takes a list of census variables that represent a set of all possible categories (eg: a list of home value categories)
    #It then creates a list with the number of observations in each cateogry,
    #It then converts the total ammount elements to fractions of the total
    
    #The basic mechanics are this ['men','women'] ----> [20,20] ----> [50,50]

    if operator == c.sf1:
        year = decennial_census_year
    elif operator == c.acs5:
        year = acs_5y_year
    else:
        assert(False)

    #Speicify geographic level specific varaibles
    if geographic_level == 'place':
        if hood_or_comparison_area == 'hood':
            place_fips = hood_place_fips
            state_fips = hood_state_fips
        
        elif hood_or_comparison_area == 'comparison area':
            place_fips = comparison_place_fips
            state_fips = comparison_state_fips

        try:
            neighborhood_household_size_distribution_raw = operator.state_place(fields = fields_list,state_fips = state_fips,place=place_fips,year= year)[0]
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area)
            return()
    
    elif geographic_level == 'county':
        if hood_or_comparison_area == 'hood':
            county_fips = hood_county_fips
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            county_fips = comparison_county_fips
            state_fips  = comparison_state_fips
        
        try:
            neighborhood_household_size_distribution_raw =operator.state_county(fields = fields_list, state_fips = state_fips,county_fips = county_fips,year= year)[0]
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'county subdivision':
        if hood_or_comparison_area == 'hood':
            county_fips = hood_county_fips
            subdiv_fips = hood_suvdiv_fips
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            county_fips = comparison_county_fips
            subdiv_fips = comparison_suvdiv_fips
            state_fips  = comparison_state_fips
        try:
            neighborhood_household_size_distribution_raw = operator.state_county_subdivision(fields=fields_list,state_fips=state_fips,county_fips=county_fips,subdiv_fips=subdiv_fips,year = year)[0]
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'zip':
        if hood_or_comparison_area == 'hood':
            zcta        = hood_zip
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            zcta        = comparison_zip
            state_fips  = comparison_state_fips
        try:
            neighborhood_household_size_distribution_raw = operator.state_zipcode(fields=fields_list,state_fips=state_fips,zcta=zcta,year= year)[0]
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'tract':
        if hood_or_comparison_area == 'hood':
            tract       = hood_tract 
            county_fips = hood_county_fips
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            tract       = comparison_tract
            county_fips = comparison_county_fips
            state_fips  = comparison_state_fips
        
        try:
            neighborhood_household_size_distribution_raw = operator.state_county_tract(fields=fields_list, state_fips = state_fips,county_fips=county_fips,tract=tract,year= year)[0]
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'custom':
        if operator == c.acs5:
            operator  = c_area.acs5
        elif operator == c.sf1:
            operator  = c_area.sf1
        
        try:
            #Create empty list we will fill with dictionaries (one for each census tract within the custom shape/neighborhood)
            neighborhood_tracts_data = []

            #Fetch census data for all relevant census tracts within the neighborhood
            raw_census_data = operator.geo_tract(fields_list, neighborhood_shape,year= year)
            for tract_geojson, tract_data, tract_proportion in raw_census_data:
                neighborhood_tracts_data.append((tract_data))

            #Convert the list of dictionaries into a single dictionary where we aggregate all values across keys
            neighborhood_household_size_distribution_raw = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_tracts_data, fields_list = fields_list )
        except Exception as e:
            print(e,'Unable to get distribtuion for custom area')

    #General data manipulation (same for all geographic levels)
    distribution = []
    for field in fields_list:
            distribution.append(neighborhood_household_size_distribution_raw[field])

    try:    
        distribution = ConvertListElementsToFractionOfTotal(distribution)
        return(distribution)
    except Exception as e:
        print(e)

def GetCensusValue(geographic_level,hood_or_comparison_area,field,operator,aggregation_method):
    #A general function that takes a single census variable (eg: median home value)
    #It pulls the value for that variable from census API and returns it

    if operator == c.sf1:
        year = decennial_census_year
    elif operator == c.acs5:
        year = acs_5y_year
    elif operator == c.pl:
        year = 2020
    else:
        assert(False)

    #Speicify geographic level specific varaibles
    if geographic_level == 'place':
        if hood_or_comparison_area == 'hood':
            place_fips = hood_place_fips
            state_fips = hood_state_fips
        
        elif hood_or_comparison_area == 'comparison area':
            place_fips = comparison_place_fips
            state_fips = comparison_state_fips
        
        try:
            values = operator.state_place(fields = field,state_fips = state_fips,place=place_fips,year= year)[0] 
            values = FilterDictionary(dictionary = values, desired_keys= field)
            values = list(values.values())
            return(values)
        
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area)
    
    elif geographic_level == 'county':
        if hood_or_comparison_area == 'hood':
            county_fips = hood_county_fips
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            county_fips = comparison_county_fips
            state_fips  = comparison_state_fips
        
        try:
            values = operator.state_county(fields = field, state_fips = state_fips,county_fips = county_fips,year= year)[0]
            values = FilterDictionary(dictionary = values, desired_keys= field)
            values = list(values.values())
            return(values)
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )

    elif geographic_level == 'county subdivision':
        if hood_or_comparison_area == 'hood':
            county_fips = hood_county_fips
            subdiv_fips = hood_suvdiv_fips
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            county_fips = comparison_county_fips
            subdiv_fips = comparison_suvdiv_fips
            state_fips  = comparison_state_fips
        
        try:
            values = operator.state_county_subdivision(fields=field,state_fips=state_fips,county_fips=county_fips,subdiv_fips=subdiv_fips,year = year)[0]
            values = FilterDictionary(dictionary = values, desired_keys= field)
            values = list(values.values())
            return(values)
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )

    elif geographic_level == 'zip':
        if hood_or_comparison_area == 'hood':
            zcta = hood_zip
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            zcta       = comparison_zip
            state_fips = comparison_state_fips

        try:
            values = operator.state_zipcode(fields=field,state_fips=state_fips,zcta=zcta,year= year)[0]
            values = FilterDictionary(dictionary = values, desired_keys= field)
            values = list(values.values())
            return(values)
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )

    elif geographic_level == 'tract':
        if hood_or_comparison_area == 'hood':
            tract       = hood_tract 
            county_fips = hood_county_fips
            state_fips  = hood_state_fips

        elif hood_or_comparison_area == 'comparison area':
            tract       = comparison_tract
            county_fips = comparison_county_fips
            state_fips  = comparison_state_fips
        try:
            values = operator.state_county_tract(fields=field, state_fips = state_fips,county_fips=county_fips,tract=tract,year= year)[0]
            values = FilterDictionary(dictionary = values, desired_keys= field)
            values = list(values.values())
            return(values)
        except Exception as e:
            print(e, 'Problem getting data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )

    elif geographic_level == 'custom':
        if operator == c.acs5:
            operator = c_area.acs5
        elif operator == c.sf1:
            operator = c_area.sf1

        #Use this method when you want to average across census tracts
        if aggregation_method == 'mean':
            custom_values = []
            for f in field:
                try:
                    #Create empty list we will fill with values (one for each census tract within the custom shape/neighborhood)
                    neighborhood_tracts_data = []

                    #Fetch census data for all relevant census tracts within the neighborhood
                    raw_census_data = operator.geo_tract(f, neighborhood_shape,year = year)
                    for tract_geojson, tract_data, tract_proportion in raw_census_data:
                        tract_value = int(tract_data[f])
                        if tract_value > 0:
                            neighborhood_tracts_data.append((tract_value))
                    #We take the simple mean of the census tracts in the area
                    assert(len(neighborhood_tracts_data) > 0)
                    custom_value = mean(neighborhood_tracts_data)

                except Exception as e:
                    print(e,'Not able to get mean value using tracts')
                    custom_value = 0
                custom_values.append(custom_value)
            


            return(custom_values)

        #Use this method when adding all the values together
        elif aggregation_method == 'total':
            custom_values = []
            for f in field:
                try:
                    neighborhood_tracts_data = []

                    #Fetch census data for all relevant census tracts within the neighborhood
                    raw_census_data = operator.geo_tract(f, neighborhood_shape,year = decennial_census_year)
                    
                    for tract_geojson, tract_data, tract_proportion in raw_census_data:
                        neighborhood_tracts_data.append(tract_data)
                    #Convert the list of dictionaries into a single dictionary where we aggregate all values across keys
                    value_raw_data          = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_tracts_data, fields_list = [f])
                    value                   = value_raw_data[f]
        
                except Exception as e:
                    print(e,'Not able to get total value using tracts')
                    value = 0
                
                custom_values.append(value)
            return(custom_values)

        #Raise an error if passed an unsupported aggregation method
        else:
            print('Not a supported aggregation method')
            assert(False)

#Households by number of memebrs
def GetHouseholdSizeData(geographic_level,hood_or_comparison_area):
    print('Getting household size data for: ',hood_or_comparison_area)
    neighborhood_household_size_distribution       = GetCensusFrequencyDistribution(geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area, fields_list=['H013002','H013003','H013004','H013005','H013006','H013007','H013008'],operator=c.sf1)          #Neighborhood households by size
    return(neighborhood_household_size_distribution)
    
#Household Tenure (owner-occupied vs renter-occupied)
def GetHousingTenureData(geographic_level,hood_or_comparison_area):
    #Occupied Housing Units by Tenure
    print('Getting tenure data for: ',hood_or_comparison_area)
    neighborhood_tenure_distribution  = GetCensusFrequencyDistribution(     geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area, fields_list = ['H004004','H004003','H004002'],operator=c.sf1) 

    #add together the owned free and clear percentage with the owned with a mortgage percentage to simply an owner-occupied fraction
    neighborhood_tenure_distribution[1] = neighborhood_tenure_distribution[1] +  neighborhood_tenure_distribution[2]
    del neighborhood_tenure_distribution[2]
    return(neighborhood_tenure_distribution)

#Housing related data functions
def GetHousingValues(geographic_level,hood_or_comparison_area):
    print('Getting housing value data for: ',hood_or_comparison_area)
    household_value_data = GetCensusFrequencyDistribution(geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area, fields_list = ["B25075_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,28)],operator=c.acs5)  
    return(household_value_data)

#Year Housing Built Data
def GetHouseYearBuiltData(geographic_level,hood_or_comparison_area):
    print('Getting year built data for: ',hood_or_comparison_area)
    year_built_data = GetCensusFrequencyDistribution(    geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area,  fields_list = ["B25034_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,12)],operator= c.acs5)
    year_built_data.reverse()
    return(year_built_data)

#Travel Method to work
def GetTravelMethodData(geographic_level,hood_or_comparison_area):
    print('Getting travel method to work data for: ' + hood_or_comparison_area)
    neighborhood_method_to_work_distribution = GetCensusFrequencyDistribution(geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area, fields_list =['B08006_003E','B08006_004E','B08006_008E','B08006_015E','B08006_017E','B08006_014E','B08006_016E'],operator=c.acs5)  
    return(neighborhood_method_to_work_distribution) 

#Household Income data functions
def GetHouseholdIncomeValues(geographic_level,hood_or_comparison_area):
    print('Getting household income data for: ',hood_or_comparison_area)
    total_income_breakdown = GetCensusFrequencyDistribution( geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area, fields_list = ["B19001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,18)],operator=c.acs5) 
    return(total_income_breakdown)

#Travel Time to Work
def GetTravelTimeData(geographic_level,hood_or_comparison_area):
    print('Getting travel time data for: ', hood_or_comparison_area)
     #5 Year ACS travel time range:   B08012_003E - B08012_013E
    travel_time_data = GetCensusFrequencyDistribution(        geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area,fields_list = ["B08012_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,14)],operator=c.acs5)   
    return(travel_time_data)

#Age Related Data Functions
def GetAgeData(geographic_level,hood_or_comparison_area):
    print('Getting age breakdown for: ',hood_or_comparison_area)
    #Return a list with the fraction of the population in different age groups 

    #Define 2 lists of variables, 1 for male age groups and another for female
    male_fields_list   =  ["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,26)]  #5 Year ACS age variables for men range:  B01001_003E - B01001_025E
    female_fields_list =  ["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(27,50)] #5 Year ACS age variables for women range:  B01001_027E - B01001_049E

   
    if geographic_level == 'place':
        try:
            if hood_or_comparison_area == 'hood':
                place_fips  = hood_place_fips
                state_fips  = hood_state_fips


            elif hood_or_comparison_area == 'comparison area':
                place_fips  = comparison_place_fips
                state_fips  = comparison_state_fips

            male_age_data   = c.acs5.state_place(fields=male_fields_list, state_fips=state_fips,place=place_fips,year=acs_5y_year)[0]
            female_age_data = c.acs5.state_place(fields=female_fields_list,state_fips=state_fips,place=place_fips,year=acs_5y_year)[0]
        except Exception as e:
            print(e, 'Problem getting age data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()
    
    elif geographic_level == 'county':
        try:
            if hood_or_comparison_area == 'hood':
                county_fips = hood_county_fips
                state_fips  = hood_state_fips

            elif hood_or_comparison_area == 'comparison area':
                county_fips = comparison_county_fips
                state_fips  = comparison_state_fips

            male_age_data   = c.acs5.state_county(fields=male_fields_list,state_fips=state_fips,county_fips=county_fips,year=acs_5y_year)[0]
            female_age_data = c.acs5.state_county(fields=female_fields_list,state_fips=state_fips,county_fips=county_fips,year=acs_5y_year)[0]

        except Exception as e:
            print(e, 'Problem getting age data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()
    
    elif geographic_level == 'county subdivision':
        try:
            if hood_or_comparison_area == 'hood':
                county_fips = hood_county_fips
                subdiv_fips = hood_suvdiv_fips
                state_fips  = hood_state_fips


            elif hood_or_comparison_area == 'comparison area':
                county_fips = comparison_county_fips
                subdiv_fips = comparison_suvdiv_fips
                state_fips  = comparison_state_fips


            male_age_data   = c.acs5.state_county_subdivision(fields = male_fields_list, state_fips = state_fips, county_fips = county_fips, subdiv_fips = subdiv_fips, year = acs_5y_year)[0]
            female_age_data = c.acs5.state_county_subdivision(fields = female_fields_list, state_fips = state_fips, county_fips = county_fips, subdiv_fips = subdiv_fips, year = acs_5y_year)[0]
        except Exception as e:
            print(e, 'Problem getting age data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'zip':
            try:        
                if hood_or_comparison_area == 'hood':
                    zcta = hood_zip

                elif hood_or_comparison_area == 'comparison area':
                    zcta = comparison_zip
            
                male_age_data         = c.acs5.zipcode(fields = male_fields_list, zcta = '*',year=acs_5y_year)
                male_age_data         = FindZipCodeDictionary(zip_code_data_dictionary_list =   male_age_data  , zcta = zcta, state_fips = state_fips )

                female_age_data       = c.acs5.zipcode(fields = female_fields_list, zcta = '*',year=acs_5y_year)
                female_age_data       = FindZipCodeDictionary(zip_code_data_dictionary_list =   female_age_data  , zcta = zcta, state_fips = state_fips )

            except Exception as e:
                print(e, 'Problem getting age data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
                return()

    elif geographic_level == 'tract':
        try:
            if hood_or_comparison_area == 'hood':
                tract       = hood_tract 
                county_fips = hood_county_fips

            elif hood_or_comparison_area == 'comparison area':
                tract       = comparison_tract
                county_fips = comparison_county_fips

            male_age_data   = c.acs5.state_county_tract(fields = male_fields_list, state_fips = state_fips, county_fips = county_fips, tract = tract, year = acs_5y_year)[0]
            female_age_data = c.acs5.state_county_tract(fields = female_fields_list, state_fips = state_fips, county_fips = county_fips, tract = tract, year = acs_5y_year)[0]
        except Exception as e:
            print(e, 'Problem getting age data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()
        
    elif geographic_level == 'custom':
        
        try:
            #Create empty list we will fill with dictionaries (one for each census tract within the custom shape/neighborhood)
            neighborhood_male_tracts_data   = []
            neighborhood_female_tracts_data = []

            #Fetch census data for all relevant census tracts within the neighborhood
            raw_male_census_data   = c_area.acs5.geo_blockgroup(male_fields_list, neighborhood_shape,year=acs_5y_year)
            raw_female_census_data = c_area.acs5.geo_blockgroup(female_fields_list, neighborhood_shape,year=acs_5y_year)
            

            for tract_geojson, tract_data, tract_proportion in raw_male_census_data:
                neighborhood_male_tracts_data.append((tract_data))
            
            for tract_geojson, tract_data, tract_proportion in raw_female_census_data:
                neighborhood_female_tracts_data.append((tract_data))

            #Convert the list of dictionaries into a single dictionary where we aggregate all values across keys
            male_age_data   = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_male_tracts_data, fields_list   = male_fields_list )
            female_age_data = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_female_tracts_data, fields_list = female_fields_list )
        except Exception as e:
            print(e,'unable to use blockgroup level, using tract')
            #Create empty list we will fill with dictionaries (one for each census tract within the custom shape/neighborhood)
            neighborhood_male_tracts_data   = []
            neighborhood_female_tracts_data = []

            #Fetch census data for all relevant census tracts within the neighborhood
            raw_male_census_data   = c_area.acs5.geo_tract(male_fields_list, neighborhood_shape,year=acs_5y_year)
            raw_female_census_data = c_area.acs5.geo_tract(female_fields_list, neighborhood_shape,year=acs_5y_year)
            

            for tract_geojson, tract_data, tract_proportion in raw_male_census_data:
                neighborhood_male_tracts_data.append((tract_data))
            
            for tract_geojson, tract_data, tract_proportion in raw_female_census_data:
                neighborhood_female_tracts_data.append((tract_data))

            #Convert the list of dictionaries into a single dictionary where we aggregate all values across keys
            male_age_data   = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_male_tracts_data, fields_list   = male_fields_list )
            female_age_data = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_female_tracts_data, fields_list = female_fields_list )

    #Create an empty list and place the age values from the dictionary inside of it
    male_age_breakdown = []
    for field in male_fields_list:
        male_age_breakdown.append(male_age_data[field])


    #Create an empty list and place the age values from the dictionary inside of it
    female_age_breakdown = []
    for field in female_fields_list:
        female_age_breakdown.append(female_age_data[field])
    

    total_age_breakdown = []
    for (men, women) in zip(male_age_breakdown, female_age_breakdown):
        total = (men + women)
        total_age_breakdown.append(total)

    
    #Consolidate some of the age groups into larger groups
    total_age_breakdown[0] = sum(total_age_breakdown[0:5])
    total_age_breakdown[1] = sum(total_age_breakdown[5:8])
    total_age_breakdown[2] = sum(total_age_breakdown[8:10])
    total_age_breakdown[3] = sum(total_age_breakdown[10:13])
    total_age_breakdown[4] = sum(total_age_breakdown[13:18])
    total_age_breakdown[5] = sum(total_age_breakdown[18:])
    del[total_age_breakdown[6:]]


    #Convert from raw numbers to fractions of total
    total_age_breakdown = ConvertListElementsToFractionOfTotal(total_age_breakdown)

    return(total_age_breakdown)

#Number of Housing Units based on number of units in building
def GetNumberUnitsData(geographic_level,hood_or_comparison_area):
    print('Getting housing units by number of units data for: ', hood_or_comparison_area)
    
    
    owner_occupied_fields_list  = ["B25032_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,11)]   #5 Year ACS owner occupied number of units variables range:  B25032_003E - B25032_010E
    renter_occupied_fields_list = ["B25032_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(14,22)]  #5 Year ACS renter occupied number of units variables range: B25032_014E - B25032_021E 

    if geographic_level == 'place':
        try:
            if hood_or_comparison_area == 'hood':
                place_fips = hood_place_fips
                state_fips = hood_state_fips

            elif hood_or_comparison_area == 'comparison area':
                place_fips = comparison_place_fips
                state_fips = comparison_state_fips
        
            owner_occupied_units_raw_data = c.acs5.state_place(fields = owner_occupied_fields_list,state_fips=state_fips,place=place_fips,year=acs_5y_year)[0]
            renter_occupied_units_raw_data = c.acs5.state_place(fields = renter_occupied_fields_list,state_fips=state_fips,place=place_fips,year=acs_5y_year)[0]
        
        except Exception as e:
            print(e, 'Problem getting number units data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'county':
        try:
            
            if hood_or_comparison_area == 'hood':
                county_fips = hood_county_fips
                state_fips = hood_state_fips

            elif hood_or_comparison_area == 'comparison area':
                county_fips = comparison_county_fips
                state_fips = comparison_state_fips

            owner_occupied_units_raw_data  = c.acs5.state_county(fields = owner_occupied_fields_list,  state_fips = state_fips, county_fips = county_fips,year=acs_5y_year)[0]
            renter_occupied_units_raw_data = c.acs5.state_county(fields = renter_occupied_fields_list, state_fips = state_fips, county_fips = county_fips,year=acs_5y_year)[0]
        except Exception as e:
            print(e, 'Problem getting number units data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'county subdivision':
        try:
            if hood_or_comparison_area == 'hood':
                county_fips = hood_county_fips
                subdiv_fips = hood_suvdiv_fips
                state_fips  = hood_state_fips


            elif hood_or_comparison_area == 'comparison area':
                county_fips = comparison_county_fips
                subdiv_fips = comparison_suvdiv_fips
                state_fips = comparison_state_fips

        
            owner_occupied_units_raw_data  = c.acs5.state_county_subdivision(fields = owner_occupied_fields_list, state_fips  = state_fips, county_fips=county_fips,  subdiv_fips=subdiv_fips,year=acs_5y_year)[0]
            renter_occupied_units_raw_data = c.acs5.state_county_subdivision(fields = renter_occupied_fields_list, state_fips = state_fips, county_fips=county_fips,  subdiv_fips=subdiv_fips,year=acs_5y_year)[0]
        except Exception as e:
            print(e, 'Problem getting number units data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'zip':
        try:        
            if hood_or_comparison_area == 'hood':
                zcta = hood_zip

            elif hood_or_comparison_area == 'comparison area':
                zcta = comparison_zip
        
            owner_occupied_units_raw_data       = c.acs5.zipcode(fields = owner_occupied_fields_list,  zcta = '*',year=acs_5y_year )
            owner_occupied_units_raw_data       = FindZipCodeDictionary(zip_code_data_dictionary_list =   owner_occupied_units_raw_data  , zcta = zcta, state_fips = state_fips )

            renter_occupied_units_raw_data      = c.acs5.zipcode(fields = renter_occupied_fields_list, zcta = '*',year=acs_5y_year )
            renter_occupied_units_raw_data      = FindZipCodeDictionary(zip_code_data_dictionary_list =   renter_occupied_units_raw_data  , zcta = zcta, state_fips = state_fips )


        except Exception as e:
            print(e, 'Problem getting number units data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'tract':
        try:
            if hood_or_comparison_area == 'hood':
                tract       = hood_tract 
                county_fips = hood_county_fips

            elif hood_or_comparison_area == 'comparison area':
                tract       = comparison_tract
                county_fips = comparison_county_fips

            owner_occupied_units_raw_data  = c.acs5.state_county_tract(fields = owner_occupied_fields_list, state_fips=state_fips, county_fips=county_fips,  tract=tract,year=acs_5y_year)[0]
            renter_occupied_units_raw_data = c.acs5.state_county_tract(fields = renter_occupied_fields_list, state_fips=state_fips, county_fips=county_fips, tract=tract,year=acs_5y_year)[0]
        
        except Exception as e:
            print(e, 'Problem getting number units data for: Geographic Level - ' + geographic_level + ' for ' + hood_or_comparison_area )
            return()

    elif geographic_level == 'custom':
        try:
            #Create empty list we will fill with dictionaries (one for each census tract within the custom shape/neighborhood)
            neighborhood_oo_tracts_data   = []
            neighborhood_ro_tracts_data = []

            #Fetch census data for all relevant census tracts within the neighborhood
            raw_oo_census_data   = c_area.acs5.geo_blockgroup(owner_occupied_fields_list, neighborhood_shape,year=acs_5y_year)
            raw_ro_census_data = c_area.acs5.geo_blockgroup(renter_occupied_fields_list, neighborhood_shape,year=acs_5y_year)
            

            for tract_geojson, tract_data, tract_proportion in raw_oo_census_data:
                neighborhood_oo_tracts_data.append((tract_data))
            
            for tract_geojson, tract_data, tract_proportion in raw_ro_census_data:
                neighborhood_ro_tracts_data.append((tract_data))

            #Convert the list of dictionaries into a single dictionary where we aggregate all values across keys
            owner_occupied_units_raw_data   = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_oo_tracts_data, fields_list   = owner_occupied_fields_list )
            renter_occupied_units_raw_data  = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_ro_tracts_data, fields_list = renter_occupied_fields_list )
        except Exception as e:
            print(e,'unable to block group level, using tracts instead')
            #Create empty list we will fill with dictionaries (one for each census tract within the custom shape/neighborhood)
            neighborhood_oo_tracts_data   = []
            neighborhood_ro_tracts_data = []

            #Fetch census data for all relevant census tracts within the neighborhood
            raw_oo_census_data   = c_area.acs5.geo_tract(owner_occupied_fields_list, neighborhood_shape,year=acs_5y_year)
            raw_ro_census_data = c_area.acs5.geo_tract(renter_occupied_fields_list, neighborhood_shape,year=acs_5y_year)
            

            for tract_geojson, tract_data, tract_proportion in raw_oo_census_data:
                neighborhood_oo_tracts_data.append((tract_data))
            
            for tract_geojson, tract_data, tract_proportion in raw_ro_census_data:
                neighborhood_ro_tracts_data.append((tract_data))

            #Convert the list of dictionaries into a single dictionary where we aggregate all values across keys
            owner_occupied_units_raw_data   = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_oo_tracts_data, fields_list   = owner_occupied_fields_list )
            renter_occupied_units_raw_data  = AggregateAcrossDictionaries(neighborhood_tracts_data = neighborhood_ro_tracts_data, fields_list = renter_occupied_fields_list )
        
        
    #Create an empty list and place the values from the dictionary inside of it
    owner_occupied_units_data = []
    for field in owner_occupied_fields_list:
        owner_occupied_units_data.append(owner_occupied_units_raw_data[field])

    #Now repeat for the renter occupied fields
    #Create an empty list and place the values from the dictionary inside of it
    renter_occupied_units_data = []
    for field in renter_occupied_fields_list:
        renter_occupied_units_data.append(renter_occupied_units_raw_data[field])

    
    total_units = sum(owner_occupied_units_data) + sum(renter_occupied_units_data)

    total_unit_data = []
    for (oo, ro) in zip(owner_occupied_units_data, renter_occupied_units_data):
        total_unit_data.append(( (oo + ro )/total_units) * 100)

    
    return(total_unit_data)

#Occupations Data
def GetTopOccupationsData(geographic_level,hood_or_comparison_area):
    print('Getting occupation data for: ',hood_or_comparison_area)
    top_occupations_data = GetCensusFrequencyDistribution(        geographic_level = geographic_level, hood_or_comparison_area = hood_or_comparison_area,fields_list =  ['B24011_002E','B24011_018E','B24011_026E','B24011_029E','B24011_036E'],operator=c.acs5)   
    return(top_occupations_data)
    
def GetOverviewTable(hood_geographic_level,comparison_geographic_level):
    print('Getting Overview table data')
    global _2010_hood_pop,  _2010_hood_hh
    global current_hood_pop, current_hood_hh
    global _2010_comparison_pop, _2010_comparison_hh
    global current_comparison_pop, current_comparison_hh
    global  hood_pop_growth,comparsion_hh_growth      
    global comparsion_pop_growth,comparsion_hh_growth
        
    #Pulls in our population and hh variables and estiatmes growth rates, returns a list of list that we will use to populate the overview table
    total_pop_field               = 'P001001'
    total_households_field        = 'H003002'
    total_families_field          = 'P035001'

    acs_total_pop_field           = 'B01001_001E'
    acs_total_households_field    = 'B25124_002E'  

    redistricting_total_pop_field = 'P1_001N'
    redistricting_total_hh_field  = 'H1_002N'
    redistricting_total_f_field   = 'P1_035N'

    print('Getting 2010 Population and Total Households Estimate for Hood')
    _2010_hood_pop = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood', field = [total_pop_field],        operator = c.sf1, aggregation_method = 'total')[0]
    _2010_hood_hh  = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood', field = [total_households_field], operator = c.sf1, aggregation_method = 'total')[0]
    print(_2010_hood_pop,_2010_hood_hh)
    # assert int(_2010_hood_pop) > 0
    # assert int(_2010_hood_hh) > 0 

    print('Successfully grabbed 2010 Population and HH count for hood')
    
    print('Getting current or 2020 Population and Total Households Estimate for Hood')
    #calculate table variables for hood
    if hood_geographic_level != 'custom':
        current_estimate_period = '2020 Census'
        current_hood_pop        = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood', field = [redistricting_total_pop_field],   operator = c.pl, aggregation_method = 'total')[0]
        current_hood_hh         = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood', field = [redistricting_total_hh_field],    operator = c.pl, aggregation_method = 'total')[0]

    
    elif hood_geographic_level == 'custom':
        current_estimate_period = 'Current Estimate'
        current_hood_pop        = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood', field = [acs_total_pop_field],   operator = c.acs5, aggregation_method = 'total')[0]
        current_hood_hh         = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood', field = [acs_total_households_field],   operator = c.acs5, aggregation_method = 'total')[0]



    print('Successfully grabbed 2020 Population and HH count for hood')
    # assert current_hood_pop > 0

    #Table variables for comparison area
    print('Getting 2010 Population and Total Households for comparison area')
    _2010_comparison_pop = GetCensusValue(geographic_level = comparison_level,   hood_or_comparison_area = 'comparison area',  field = [total_pop_field],        operator = c.sf1, aggregation_method = 'total')[0]
    _2010_comparison_hh  = GetCensusValue(geographic_level = comparison_level, hood_or_comparison_area =  'comparison area', field = [total_households_field], operator = c.sf1, aggregation_method = 'total')[0]
    
    print('Getting current Population and Total Households for comparison area')
    if comparison_geographic_level != 'custom':
        current_comparison_pop    = GetCensusValue(geographic_level = comparison_level, hood_or_comparison_area = 'comparison area', field = [redistricting_total_pop_field],   operator = c.pl, aggregation_method = 'total')[0]
        current_comparison_hh     = GetCensusValue(geographic_level = comparison_level, hood_or_comparison_area = 'comparison area', field = [redistricting_total_hh_field],    operator = c.pl, aggregation_method = 'total')[0]

    elif comparison_geographic_level == 'custom':
        #Custom vs. custom still not supported
        pass

    print('Successfully grabbed 2020 Population and HH count for comparison area')

    #Set growth periods
    if hood_geographic_level == 'custom':
        pop_growth_period = acs_5y_year - decennial_census_year
        hh_growth_period  = acs_5y_year - decennial_census_year
        assert pop_growth_period == hh_growth_period == 9
    else:
        pop_growth_period = (decennial_census_year + 10) - decennial_census_year
        hh_growth_period  = (decennial_census_year + 10) - decennial_census_year
        assert pop_growth_period == hh_growth_period == 10

    #Calculate growth rates
    hood_pop_growth        = (((int(current_hood_pop)/int(_2010_hood_pop)) - 1) * 100 )/pop_growth_period
    
    #Total Households not available in american community survey
    if hood_geographic_level != 'custom':
        hood_hh_growth         = (((int(current_hood_hh)/int(_2010_hood_hh))   - 1) * 100)/hh_growth_period
        hood_hh_growth          = "{:,.1f}%".format(hood_hh_growth)
        current_hood_hh        = "{:,.0f}".format(int(current_hood_hh))
    else:
        hood_hh_growth         = 'NA'

    comparsion_pop_growth  =  ((int(current_comparison_pop)/int(_2010_comparison_pop) - 1) * 100)/pop_growth_period
    comparsion_hh_growth   =  ((int(current_comparison_hh)/int(_2010_comparison_hh)   - 1) * 100)/hh_growth_period

    #Format pop and hh variables
    _2010_hood_pop         = "{:,.0f}".format(int(_2010_hood_pop))
    _2010_hood_hh          = "{:,.0f}".format(int(_2010_hood_hh))
    _2010_comparison_pop   = "{:,.0f}".format(int(_2010_comparison_pop))
    _2010_comparison_hh    = "{:,.0f}".format(int(_2010_comparison_hh))
    current_hood_pop       = "{:,.0f}".format(int(current_hood_pop))
    current_comparison_pop = "{:,.0f}".format(int(current_comparison_pop))
    current_comparison_hh  = "{:,.0f}".format(int(current_comparison_hh))

    #Format growth variables
    hood_pop_growth         = "{:,.1f}%".format(hood_pop_growth)
    comparsion_pop_growth   = "{:,.1f}%".format(comparsion_pop_growth)
    comparsion_hh_growth    = "{:,.1f}%".format(comparsion_hh_growth)

    #each row represents a row of data for overview table
    row1 = [''          , 'Area',             '2010 Census',            current_estimate_period,                                      'Annual % Change']
    row2 = ['Population', neighborhood,        _2010_hood_pop,          current_hood_pop ,                                 hood_pop_growth ]
    row3 = [''          , comparison_area,     _2010_comparison_pop,    current_comparison_pop,                       comparsion_pop_growth]
    row4 = ['Households', neighborhood,        _2010_hood_hh,           current_hood_hh,                                     hood_hh_growth]
    row5 = [''          , comparison_area,     _2010_comparison_hh,     current_comparison_hh,                        comparsion_hh_growth ]
    
    if neighborhood_level != 'custom': #Don't include household rows for custom neighborhoods
        return( 
                [ 
                row1,
                row2,
                row3,
                row4,
                row5
                    ])
    else:
        return(    
                [ 
                row1,
                row2,
                row3,
                    ])

#####################################################Non Census Sources Data Functions####################################
def GetWikipediaPage():
    global page
    print('Getting Wikipedia page')
    if (neighborhood_level == 'place') or (neighborhood_level == 'county subdivision') or (neighborhood_level == 'county'): #Don't bother looking for wikipedia page if zip code
            wikipedia_page_search_term    = (neighborhood + ', ' + hood_state_full_name)
                       
    elif (neighborhood_level == 'custom'):
            wikipedia_page_search_term    = (neighborhood + ', ' + comparison_area )
    try:
        page                          =  wikipedia.page(wikipedia_page_search_term)   
    except Exception as e:
        print(e,': problem getting wikipedia page')
        page = None

def GetWalkScore(lat,lon):

    lat = str(lat)
    lon = str(lon)
    url = """https://api.walkscore.com/score?format=json&address=None&""" + """lat=""" + lat + """&lon=""" + lon + """&transit=1&bike=1&wsapikey=""" + walkscore_api_key
    print('Getting Walk Score: ', url)
    walkscore_response = requests.get(url).json()
    
    #Get Walk score from response
    try:
        walk_score           = walkscore_response['walkscore']
        walk_description     = walkscore_response['description']
        walk_table_entry     = ('Walk Score: ' + str(walk_score) + ' (' + walk_description + ')')
    except Exception as e:
        print(e,'could not get walk score')
        walk_table_entry     = 'NA'
    
    #Get Transit score from response
    try:
        transit_score        = walkscore_response['transit']['score']
        transit_description  = walkscore_response['transit']['description']
        transit_table_entry  = ('Transit Score: ' + str(transit_score) + ' (' + transit_description + ')')

    except Exception as e:
        print(e,'could not get transit score')
        transit_table_entry  = 'Transit Score: NA'


    #Get Bike score from response
    try:
         bike_score           =  walkscore_response['bike']['score']
         bike_description     =  walkscore_response['bike']['description']
         bike_table_entry     = ('Bike Score: ' + str(bike_score)  + ' (' + bike_description + ')')

    except Exception as e:
        print(e,'could not get bike score')
        bike_table_entry     = 'Bike Score: NA'

    
    #Return a list of the 3 scores
    walk_scores = [walk_table_entry, transit_table_entry, bike_table_entry]
    return(walk_scores)

def GetYelpData(lat,lon,radius):
    print('Getting Yelp Data')
    #Return a dictionary where each key is a business caategory and the values are a list of the 5 most recomended businesseses on Yelp.com
    business_categories = {'retail':[], 'banks, gyms':[], 'parks and recreation':[], 'education':[], 'transportation':[]}

    try:
        for category in business_categories.keys():
            response              = yelp_api.search_query(categories=category, longitude = lon, latitude = lat, radius = radius,sort_by = 'best_match')
            
            #Loop through the results of the yelp search and pull business names
            for i in range(5):
                business_name = response['businesses'][i]['name']
                business_categories[category].append(business_name)
        
            time.sleep(.1)
    except Exception as e:
        print(e)
        
        
    return(business_categories)

def FindAirport():
    #Specify the file path to the airports shape file
    airport_map_location = os.path.join(general_data_location,'Geographic Data','Airports','Airports.shp')
    
    #Open the shapefile
    airport_map = shapefile.Reader(airport_map_location)
    
    try:
        airports_in_city_index_list = [] #Create empty list that we will fill with numbers that correspond to airports within the subject area
        
        #Find any airports inside the confines of the city
        for i in range(len(airport_map)):
            airport_coords        =  Point(airport_map.shape(i).points[0][0],airport_map.shape(i).points[0][1])
            
            if neighborhood_shape_polygon.contains(airport_coords):
                airports_in_city_index_list.append(i)

        airport_info_list = []    
        for airport_index in airports_in_city_index_list:     
            airport_record        = airport_map.shapeRecord(airport_index)
            airport_name          = airport_record.record['Fac_Name']
            airport_type          = airport_record.record['Fac_Type']
            airport_dict          = {'name':airport_name,'type':airport_type}
            
        #Don't consider heliports, seaplane bases etc
            if airport_record.record['Fac_Type'] != 'AIRPORT':
                continue

            airport_info_list.append(airport_dict)

        if airport_info_list == []:
            return(None)
        
        if len(airport_info_list) > 1:
            airport_sentence = (neighborhood + ' is served by the following facilities: ')

            for count,airport in enumerate(airport_info_list):
                if count < len(airport_info_list) -1 :
                    airport_sentence = airport_sentence + (airport['name'].title()) + ' ('  + (airport['type'].title())   + '), ' 
                else:
                    airport_sentence = airport_sentence + 'and ' + (airport['name'].title()) + ' ('  + (airport['type'].title())   + ').' 

        elif len(airport_info_list) == 1:
            airport_sentence = (neighborhood + ' is served by ' + (airport_info_list[0]['name'].title()) + ' '  + (airport_info_list[0]['type'].title()) +'.')

        return(airport_sentence)
    except Exception as e:
        print(e,'Unable to locate airport inside the neighborhood area')
        return(None)

def FindNearestAirport(lat,lon):
    
    #Specify the file path to the airports shape file
    airport_map_location = os.path.join(general_data_location,'Geographic Data','Airports','Airports.shp')
    
    #Open the shapefile
    airport_map = shapefile.Reader(airport_map_location)
       
    #Find the cloeset airport
    #Loop through each feature/point in the shape file
    for i in range(len(airport_map)):
        airport        =  airport_map.shape(i)
        airport_record = airport_map.shapeRecord(i)
        
        #Don't consider heliports, seaplane bases etc
        if airport_record.record['Fac_Type'] != 'AIRPORT':
            continue

        #Don't consider private airports
        if airport_record.record['Fac_Use'] != 'PU':
            continue

        airport_coord = airport.points
        airport_lat_lon = (airport_coord[0][1], airport_coord[0][0])
        dist = mpu.haversine_distance(airport_lat_lon, (lat, lon)) * 0.621371 #measure distance between airport and subject property   

        if i == 0:
            min_dist           = dist
            cloest_airport_num = i
        elif i > 0 and dist < min_dist:
            min_dist           = dist
            cloest_airport_num = i

    closest_airport = airport_map.shapeRecord(cloest_airport_num)
    airport_distance = "{:,.1f} miles".format(min_dist)   
    airport_lang = (neighborhood + ' is roughly ' + airport_distance + ' from ' + closest_airport.record['Fac_Name'].title() + ', a public ' +  closest_airport.record['Fac_Type'].lower() + '.' )
    
    airport_lang = airport_lang.replace('Intl','International')
    airport_lang = airport_lang.replace('Rgnl','Regional')

    return(airport_lang)

def FindNearestHighways():
    
    try:
        #Specify the file path to the shape file
        road_map_location = os.path.join(general_data_location,'Geographic Data','North_American_Roads','North_American_Roads.shp')

        #Open the shapefile
        road_map = shapefile.Reader(road_map_location)
        highways_in_city_index_list = [] #Create empty list that we will fill with numbers that correspond to roads within the subject area
        
        #Loop through the road map and find any roads inside the confines of the city. Once we identify one, add the index number to our list of highways (indexes)
        for i in range(len(road_map)):
            highway_coords        =  LineString(road_map.shape(i).points)           
            if neighborhood_shape_polygon.intersects(highway_coords):
                highways_in_city_index_list.append(i)
        
        #Now loop through our list of index numbers, for each index number, create a dictionary with key info (name, etc), append that dictionary to empty list
        i = 0
        highway_info_list = []    
        for highway_index in highways_in_city_index_list:     
            highway_record        = road_map.shapeRecord(highway_index)
            

            highway_name          = highway_record.record['ROADNAME'].title()
            #Clean up abbreviations
            highway_name          = highway_name.replace('Hwy','Highway')
            highway_name          = highway_name.replace('Pkwy','Parkway') 

            highway_number        = str(highway_record.record['ROADNUM'])

            #Don't add unnamed highways to our list
            if highway_name == '' or highway_name == 'Unknown':
                continue
           
            if i > 0:  #If not the first highway check against the existing highways and make sure it's not a duplicate
                for d in highway_info_list:
                    existingname = d['name']
                    if highway_name == existingname:
                        repeat = 1
                        break
                    repeat = 0
                if repeat == 1:
                    continue
                        
            highway_type          = highway_record.record['ADMIN'].title()
            highway_dict          = {'name':highway_name,'type':highway_type,'number':highway_number}
            highway_info_list.append(highway_dict)
            i+=1

        #When there are more than 1 roads
        if len(highway_info_list) > 2:
            sentence = (neighborhood + ' is served by the following roads: ')

            for count,highway in enumerate(highway_info_list):
               

                if count < len(highway_info_list) -1 :
                    sentence = sentence + (highway['name']) + ' ('  + (highway['type'])   + ') ' + ' (' + highway['number'] + '),'
                else:
                    sentence = sentence + 'and ' + (highway['name']) + ' ('  + (highway['type'])   + ')' + ' (' + highway['number'] + ').' 
        
        
        
        #When we only have 1 major road in our list
        elif len(highway_info_list) == 1:
            sentence = (highway_info_list[0]['name'] + ' (' + highway_info_list[0]['number'] + ')' + ' is the main road connecting ' + neighborhood + '.')
        
        elif len(highway_info_list) == 2:
           
            sentence = ((highway_info_list[0]['name']) + ' (' + highway_info_list[0]['number'] + ')' + ' ('  + (highway_info_list[0]['type'])   + ') ' + 'and '  + 
                        (highway_info_list[1]['name']) + ' (' + highway_info_list[1]['number'] + ')'+ ' ('  + (highway_info_list[1]['type'])   + ') '
                        + ' are the main roads connecting ' + neighborhood + '.')
        elif len(highway_info_list) == 0:
            sentence = None

        return(sentence)
    except Exception as e:
        print(e,'Unable to locate major roads inside the neighborhood area')
        return(None)

def FindTrainLines():
    
    try:
        #Specify the file path to the shape file
        map_location = os.path.join(general_data_location,'Geographic Data','GTFS_NTM_Shapes (Non Bus)','GTFS_NTM_Shapes.shp')

        #Open the shapefile
        map = shapefile.Reader(map_location)
        index_list = [] #Create empty list that we will fill with numbers that correspond to routes within the subject area
        
        #Loop through the road map and find any roads inside the confines of the city. Once we identify one, add the index number to our list of highways (indexes)
        for i in range(len(map)):
            highway_coords        =  LineString(map.shape(i).points)           
            if neighborhood_shape_polygon.intersects(highway_coords):
                index_list.append(i)

        #Now loop through our list of index numbers, for each index number, create a dictionary with key info (name, etc), append that dictionary to empty list
        i = 0
        info_list = []    
        for index in index_list:     
            highway_record        = map.shapeRecord(index)
            
            name          = highway_record.record['route_long'].title()
            type          = highway_record.record['route_ty_1'].title()
            agency        = highway_record.record['agency_id'].title()
            info_dict     = {'name':name,'type':type,'agency':agency}
            
            if i > 0:  #If not the first highway check against the existing highways and make sure it's not a duplicate
                for d in info_list:
                    existingname = d['name']
                    if name == existingname:
                        repeat = 1
                        break
                    repeat = 0
                if repeat == 1:
                    continue
            info_list.append(info_dict)
            i+=1
            
        #When there are more than 1 routes
        if len(info_list) > 2:
            sentence = (neighborhood + ' is served by the following transit routes: ')

            for count,route in enumerate(info_list):
               
                if count < len(info_list) -1 :
                    sentence = sentence + (route['name']) + ' ('  + (route['type'])   + '), ' 
                else:
                    sentence = sentence + 'and ' + (route['name']) + ' ('  + (route['type'])   + ').' 
        
        #When we only have 1 major road in our list
        elif len(info_list) == 1:
            if agency != '':
                sentence = (info_list[0]['agency'] + """'s""" + info_list[0]['name'] + ' line is the main ' + info_list[0]['type'] + ' route connecting ' + neighborhood + '.')
            else:
                sentence = ('The ' + info_list[0]['name'] + ' line is the main ' + info_list[0]['type'].lower() + ' route connecting ' + neighborhood + '.')


        elif len(info_list) == 2:
           
            sentence = ((info_list[0]['name']) + ' ('  + (info_list[0]['type'])   + ') ' + 'and '  + 
                        (info_list[1]['name']) + ' ('  + (info_list[1]['type'])   + ') '
                        + ' are the main transit routes connecting ' + neighborhood + '.')

        elif len(info_list) == 0:
            sentence = None
        
        return(sentence)
    except Exception as e:
        print(e,'Unable to locate transit routes inside the neighborhood area')
        return(None)

def FindBusLines():
    try:
        #Specify the file path to the shape file
        map_location = os.path.join(general_data_location,'Geographic Data','GTFS_NTM_Stops','GTFS_NTM_Stops.shp')

        #Open the shapefile
        map = shapefile.Reader(map_location)
        index_list = [] #Create empty list that we will fill with numbers that correspond to routes within the subject area
        
        #Loop through the road map and find any roads inside the confines of the city. Once we identify one, add the index number to our list of highways (indexes)
        for i in range(len(map)):
            points = map.shape(i).points
            coords        =  Point(points[0])     
            if neighborhood_shape_polygon.contains(coords):
                index_list.append(i)

        #Now loop through our list of index numbers, for each index number, create a dictionary with key info (name, etc), append that dictionary to empty list
        i = 0
        info_list = []    
        for index in index_list:     
            highway_record        = map.shapeRecord(index)
            
            name          = highway_record.record['stop_name']
            agency_id     = highway_record.record['ntd_id']
            info_dict     = {'name':name,'type':type,'agency_id':agency_id}
            
            if i > 0:  #If not the first highway check against the existing highways and make sure it's not a duplicate
                for d in info_list:
                    existingname = d['name']
                    if name == existingname:
                        repeat = 1
                        break
                    repeat = 0
                if repeat == 1:
                    continue
            info_list.append(info_dict)
            i+=1
            
        #When there are more than 1 routes
        if len(info_list) > 0:
            agency_id = str(info_list[0]['agency_id'])
            agency_name = TransitAgencyIdToName(id = agency_id)
            sentence = (agency_name + ' provides public bus service within ' + neighborhood + '.')

        elif len(info_list) == 0:
            sentence = None
        
        return(sentence)
    except Exception as e:
        print(e,'Unable to locate bus Stop inside the neighborhood area')
        return(None)

def SearchGreatSchoolDotOrg():
    print('Getting education data')
    if os.path.exists(os.path.join(hood_folder_map,'education_map.png')): #If we already have a map for this area skip it 
        return()
   
    try:
        if neighborhood_level == 'custom':
            search_term = (neighborhood + ', ' + comparison_area)
        else:
            search_term = (neighborhood + ', ' + hood_state)

        #Search https://www.greatschools.org/ for the area
        options = webdriver.ChromeOptions()
        options.add_argument("--start-maximized")
        browser = webdriver.Chrome(executable_path=(os.path.join(os.environ['USERPROFILE'], 'Desktop','chromedriver.exe')),options=options)
        browser.get('https://www.greatschools.org/')
        
        #Write hood name in box
        Place = browser.find_element_by_class_name("search_form_field")
        Place.send_keys(search_term)
        time.sleep(1)
        
        #Submit hood name for search
        Submit = browser.find_element_by_class_name('search_form_button')
        Submit.click()
        time.sleep(3)
        
        #Zoom out map
        pyautogui.moveTo(3261, y=1045)
        time.sleep(1)
        for i in range(1):
           pyautogui.click()
           time.sleep(1)
        time.sleep(3)

        if 'Leahy' in os.environ['USERPROFILE']: #differnet machines have different screen coordinates
            print('Using Mikes coordinates for screenshot')
            im2 = pyautogui.screenshot(region=(1167,872, 2049, 1316) ) #left, top, width, and height
        
        elif 'Dominic' in os.environ['USERPROFILE']:
            print('Using Doms coordinates for screenshot')
            im2 = pyautogui.screenshot(region=(3680,254,1968 ,1231) ) #left, top, width, and height
        
        else:
            im2 = pyautogui.screenshot(region=(1167,872, 2049, 1316) ) #left, top, width, and height
        
        time.sleep(.25)
        im2.save(os.path.join(hood_folder_map,'education_map.png'))
        im2.close()
        time.sleep(1)
        browser.quit()
    except Exception as e:
        print(e)
        try:
            browser.quit()
        except:
            pass

def ApartmentDotComSearchTerm():
    #Takes the name of the city or neighborhood and creates a url for apartments.com
    if neighborhood_level == 'place':
        search_term = 'https://www.apartments.com/' + '-'.join(neighborhood.lower().split(' ')) + '-' + hood_state.lower() + '/'
    elif neighborhood_level == 'custom':
        search_term = 'https://www.apartments.com/' + '-'.join(neighborhood.lower().split(' ')) + '-' + '-'.join(comparison_area.lower().split(' ')) +  '-' + hood_state.lower() + '/'
    elif neighborhood_level == 'county subdivision':
        search_term = 'https://www.apartments.com/' + '-'.join(neighborhood.lower().split(' ')) + '-' + hood_state.lower() + '/'
    
    return(search_term)

def ApartmentsDotComSearch():
    print('Searching Apartments.com:',ApartmentDotComSearchTerm())
    try:
        search_term = ApartmentDotComSearchTerm() 

        response    = requests.get(search_term,
                                   headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_2) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/47.0.2526.106 Safari/537.36"}
                                   )
        soup_data   = BeautifulSoup(response.text, 'html.parser')
        
        marketing_blurb_section = soup_data.find(id='marketingBlurb')
        marketing_paragraphs    = marketing_blurb_section.find_all('p')
        
        descriptive_paragraphs = []
        for count,paragraph in enumerate(marketing_paragraphs):
            if ('Learn'  in paragraph.text) and ('More' in paragraph.text):
                return(descriptive_paragraphs)
            if count > 3 :
                continue
            if 'Let Apartments.com help you find' in  paragraph.text:
                continue          
            descriptive_paragraphs.append(paragraph.text)
        
        return(descriptive_paragraphs)
    
    
    except Exception as e:
        print(e)
        return([''])

def Zoneomics(address):
    #Searches the Zoneomics API for screenshot of Local Zoning
    print('Getting Zoneomics Zoning Map')
    url = 'https://www.zoneomics.com/api/get_zone_screen_shot?address=address&api_key=api_key&map_zoom_level=17'

    data = {
    'key':     zoneomics_api_key,
    'address': address
        }

    try:
        response = requests.get(url, params=data).json()
    except Exception as e:
        print(e)
        response = [{}]

#Main data function
def GetData():
    #List of 5 Year American Community Survey Variables here: https://api.census.gov/data/2019/acs/acs5/variables.html
    #List of 2010 Census variables here:                      https://api.census.gov/data/2010/dec/sf1/variables.html
    #List of 2020 Redistricting variables here:               https://api.census.gov/data/2020/dec/pl/variables.html
    print('Getting Data')
    # global total_number_households, average_household_size
    global overview_table_data
    global neighborhood_household_size_distribution,comparison_household_size_distribution
    global neighborhood_tenure_distribution, comparison_tenure_distribution
    global neighborhood_time_to_work_distribution, comparison_time_to_work_distribution
    global neighborhood_method_to_work_distribution
    global neighborhood_age_data,comparison_age_data
    global neighborhood_housing_value_data,comparison_housing_value_data
    global neighborhood_number_units_data,comparison_number_units_data
    global neighborhood_household_income_data, comparison_household_income_data
    global neighborhood_year_built_data, comparison_year_built_data   
    global walk_score_data
    global nyc_community_district
    global neighborhood_median_home_value, comparison_median_home_value
    global neighborhood_median_year_built, comparison_median_year_built
    global neighborhood_median_age, comparison_median_age
    global neighborhood_average_hh_size,comparison_average_hh_size
    global neighborhood_median_hh_inc,comparison_median_hh_inc
    print('Getting Data for ' + neighborhood)

    #Start by getting our distributions for our graphs
    print('Getting distributions for hood')
    sleep_time = 1
    
    neighborhood_household_size_distribution          = GetHouseholdSizeData(     geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Neighborhood households by size
    time.sleep(sleep_time)
    
    neighborhood_tenure_distribution                  = GetHousingTenureData(     geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Housing Tenure (owner occupied/renter)
    time.sleep(sleep_time)
    
    neighborhood_housing_value_data                   = GetHousingValues(         geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Owner Occupied housing units by value
    time.sleep(sleep_time)
    
    neighborhood_year_built_data                      = GetHouseYearBuiltData(    geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Housing Units by year structure built
    time.sleep(sleep_time)
    
    neighborhood_method_to_work_distribution          = GetTravelMethodData(      geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Travel Mode to Work
    time.sleep(sleep_time)
    
    neighborhood_household_income_data                = GetHouseholdIncomeValues( geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Households by household income data
    time.sleep(sleep_time)
    
    neighborhood_time_to_work_distribution            = GetTravelTimeData(        geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Travel Time to Work
    time.sleep(sleep_time)
    
    neighborhood_number_units_data                    = GetNumberUnitsData(       geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Housing Units by units in building
    time.sleep(sleep_time)
    
    neighborhood_age_data                             = GetAgeData(               geographic_level = neighborhood_level, hood_or_comparison_area = 'hood')          #Population by age data
    time.sleep(sleep_time)
    
    #Now grab single values for our language
    print('Getting median values for hood')
    neighborhood_values                                 = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood',field = ['B25077_001E','B25035_001E','B01002_001E','B19013_001E'],operator = c.acs5,aggregation_method = 'mean')
    neighborhood_median_home_value                      = neighborhood_values[0]
    neighborhood_median_year_built                      = neighborhood_values[1]
    neighborhood_median_age                             = neighborhood_values[2]
    neighborhood_median_hh_inc                          = neighborhood_values[3]
    time.sleep(sleep_time)
    neighborhood_average_hh_size                        = GetCensusValue(geographic_level = neighborhood_level, hood_or_comparison_area = 'hood',field = ['H012001'],    operator = c.sf1, aggregation_method = 'mean')[0]

    #Handle missing values 
    if neighborhood_median_year_built == '0':
        neighborhood_median_year_built = 'NA'
    

    #Start by getting our distributions for our graphs
    print('Getting Data For ' + comparison_area)
    print('Getting distributions for comparison area')
    comparison_household_size_distribution            = GetHouseholdSizeData(    geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')
    time.sleep(sleep_time)
    comparison_tenure_distribution                    = GetHousingTenureData(    geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')
    time.sleep(sleep_time)
    comparison_housing_value_data                     = GetHousingValues(        geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')    
    time.sleep(sleep_time)
    comparison_year_built_data                        = GetHouseYearBuiltData(   geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')
    time.sleep(sleep_time)
    comparison_household_income_data                  = GetHouseholdIncomeValues(geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')   
    time.sleep(sleep_time)
    comparison_time_to_work_distribution              = GetTravelTimeData(       geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')
    time.sleep(sleep_time)
    comparison_age_data                               = GetAgeData(              geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')
    time.sleep(sleep_time)
    comparison_number_units_data                      = GetNumberUnitsData(      geographic_level  = comparison_level,   hood_or_comparison_area = 'comparison area')    
    time.sleep(sleep_time)
    
    print('Getting median values for comparison area')
    #Now grab single values for our language
    
    comparison_values                                 = GetCensusValue(geographic_level = comparison_level, hood_or_comparison_area = 'comparison area',field = ['B25077_001E','B25035_001E','B01002_001E','B19013_001E'],operator = c.acs5,aggregation_method = 'mean')
    comparison_median_home_value                      = comparison_values[0]
    comparison_median_year_built                      = comparison_values[1]
    comparison_median_age                             = comparison_values[2]
    comparison_median_hh_inc                          = comparison_values[3]
    time.sleep(sleep_time)
    comparison_average_hh_size                        = GetCensusValue(geographic_level = comparison_level, hood_or_comparison_area = 'comparison area',field = ['H012001'],    operator = c.sf1,aggregation_method  = 'mean')[0]
        
    # Handle missing values 
    if comparison_median_year_built == ('0' or 0):
        comparison_median_year_built = 'NA'

    # Walk score
    walk_score_data                                   = GetWalkScore(            lat = latitude, lon = longitude                                                    )

    #Overview Table Data
    overview_table_data                               = GetOverviewTable(hood_geographic_level = neighborhood_level ,comparison_geographic_level = comparison_level)
    nyc_community_district                            = DetermineNYCCommunityDistrict(lat = latitude, lon = longitude )
    
#####################################################Graph Related Functions####################################
def SetGraphFormatVariables():
    
    global graph_width, graph_height, scale,tickfont_size,left_margin,right_margin,top_margin,bottom_margin,legend_position,paper_backgroundcolor,title_position
    global fig_width
    global main_graph_font_size

    #Set graph size and format variables
    marginInches         = 1/18
    ppi                  = 96.85 
    width_inches         = 6.5
    height_inches        = 3.3
    fig_width            = 4.5 #width for the pngs (graph images) we insert into report document
    graph_width          = (width_inches - marginInches)   * ppi
    graph_height         = (height_inches  - marginInches) * ppi

    #Set scale for resolution 1 = no change, > 1 increases resolution. Very important for run time of main script. 
    scale                = 7

    #Set tick font size (also controls legend font size)
    tickfont_size        = 14 
    main_graph_font_size = 14

    #Set Margin parameters/legend location
    left_margin           = 0
    right_margin          = 0
    top_margin            = 75
    bottom_margin         = 10
    legend_position       = 1.05

    #Paper color
    paper_backgroundcolor = 'white'

    #Title Position
    title_position        = .95    

def CreateHouseholdSizeHistogram():
    print('Creating Household size graph')

    household_size_categories = ['1', '2', '3', '4', '5', '6', '7+']
    fig = make_subplots(specs=[[{"secondary_y": False}]])
    
    #Add Bars with neighborhood household size distribution
    fig.add_trace(
        go.Bar(        
            y            = neighborhood_household_size_distribution,
            x            = household_size_categories,
            name         = neighborhood,
            marker_color = "#4160D3"),
        secondary_y=False
                )

    fig.add_trace(
        go.Bar(
            y            = comparison_household_size_distribution,
            x            = household_size_categories,
            name         = comparison_area,
            marker_color = "#B3C3FF"
             ),
        secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout (
    title_text="Households by Household Size",    

    title = {
        'y': title_position,
        'x': 0.5,
        'xanchor': 'center',
        'yanchor': 'top'
          },
                    

    legend = dict(
                orientation = "h",
                yanchor     = "bottom",
                y           = legend_position ,
                xanchor     = "center",
                x           = 0.5,
                font_size   = tickfont_size
                  ),

    margin  = dict(
                l = left_margin, 
                r = right_margin, 
                t = top_margin, 
                b = bottom_margin
                    ),

    height        = graph_height,
    width         = graph_width,
    font_family   = "Avenir Next LT Pro",
    font_color    = '#262626',
    font_size     = main_graph_font_size,
    paper_bgcolor = paper_backgroundcolor,
    plot_bgcolor  = "White"
                     )

    #Add % to y-axis ticks
    fig.update_yaxes(title = None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size = tickfont_size), tickformat = '.0f', secondary_y = False)
    fig.update_xaxes(tickfont = dict(size = tickfont_size))       
    fig.write_image(os.path.join(hood_folder, 'household_size_graph.png'), engine = 'kaleido',scale = scale)

def CreateHouseholdTenureHistogram():
    print('Creating Household tenure graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    tenure_categories = ['Renter Occupied', 'Owner Occupied']
    
    #Add Bars with neighborhood household size distribution
    fig.add_trace(
        go.Bar(
            y             = neighborhood_tenure_distribution,
            x             = tenure_categories,
            name          = neighborhood,
            marker_color  = "#4160D3"),
        secondary_y  = False
            )
    #Add Comparison Bars
    fig.add_trace(
    go.Bar(y             = comparison_tenure_distribution,
           x             = tenure_categories,
           name          = comparison_area,
           marker_color  = "#B3C3FF")
           ,secondary_y  = False
            )
    
    
    #Set Title
    fig.update_layout(
        title_text    = "Occupied Housing Units by Tenure",    
        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
        
        title = {
            'y':        title_position,
            'x':        0.5,
            'xanchor': 'center',
            'yanchor': 'top'
                },
                    

        legend = dict(
            orientation = "h",
            yanchor     = "bottom",
            y           = legend_position ,
            xanchor     = "center",
            x           = 0.5,
            font_size   = tickfont_size
                      ),

 
        margin    = dict(l = left_margin, r = right_margin, t = top_margin, b = bottom_margin),
        height    = graph_height,
        width     = graph_width, 
                  
                  )

    #Add % to  axis ticks
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False)    
    fig.update_xaxes(tickfont = dict(size=tickfont_size))       
    fig.write_image(os.path.join(hood_folder,'household_tenure_graph.png'),engine='kaleido',scale=scale)

def CreateHouseholdNumberUnitsInBuildingHistogram():
    print('Creating Household by number of units in structure graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    number_units_categories = ['Single Family Homes','Townhomes','Duplexes','3-4 Units','5-9 Units','10-19 Units','20-49 Units','50 >= Units']
   
    #Add Bars with neighborhood distribution
    fig.add_trace(
        go.Bar(
            y            = neighborhood_number_units_data,
            x            = number_units_categories,
            name         = neighborhood,
            marker_color = "#4160D3"),
        secondary_y = False
            )

    #Add Bars with comparison distribution
    fig.add_trace(
    go.Bar(
           y            = comparison_number_units_data,
           x            = number_units_categories,
           name         = comparison_area,
           marker_color = "#B3C3FF"),
    secondary_y = False
            )
        
    #Set Title
    fig.update_layout(
    
        title_text = "Housing Units by Units in Structure",    

        title      = {
            'y':        title_position,
            'x':        0.5,
            'xanchor': 'center',
            'yanchor': 'top'
                     },
                        
        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position ,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )

    #Add % to  axis ticks
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False)       
    fig.update_xaxes(tickfont = dict(size=tickfont_size))       
    fig.write_image(os.path.join(hood_folder,'household_units_in_structure_graph.png'),engine='kaleido',scale=scale)

def CreateHouseholdYearBuiltHistogram():
    print('Creating Household Year Built graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    year_built_categories = ['2014 >=','2010-2013','2000-2009','1990-1999','1980-1989','1970-1979','1960-1969','1950-1959','1940-1949','<= 1939']
    year_built_categories.reverse()

    #Add Bars with neighborhood year built data
    fig.add_trace(
        go.Bar(
            y            = neighborhood_year_built_data,
            x            = year_built_categories,
            name         = neighborhood,
            marker_color = "#4160D3"),
        secondary_y = False
               )

    #Add bars for comparison area
    fig.add_trace(
        go.Bar(
            y            = comparison_year_built_data,
            x            = year_built_categories,
            name         = comparison_area,
            marker_color = "#B3C3FF"),
        secondary_y=False
            )
        
    fig.update_layout(
        #Set Title
        title_text="Housing Units by Year Structure Built",    

        title = {
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
                },
                        

        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position ,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )

    #Add % to  axis ticks
    fig.update_yaxes(title = None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size = tickfont_size), tickformat = '.0f', secondary_y = False)
    fig.update_xaxes(tickfont = dict(size = tickfont_size))       
    fig.write_image(os.path.join(hood_folder, 'household_year_built_graph.png'), engine = 'kaleido', scale = scale)

def CreateHouseholdValueHistogram():
    print('Creating Household value graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    housing_value_categories = ['$10,000 <','$10,000-14,999','$15,000-19,999','$20,000-24,999','$25,000-29,999','$30,000-34,000','$35,000-39,999','$40,000-49,000','$50,000-59,9999','$60,000-69,999','$70,000-79,999','$80,000-89,999','$90,000-99,999','$100,000-124,999','$125,000-149,999','$150,000-174,999','$175,000-199,999','$200,000-249,999','$250,000-299,999','$300,000-399,999','$400,000-499,999','$500,000-749,999','$750,000-999,999','$1,000,000-1,499,999','$1,500,000-1,999,999','$2,000,000 >=']
    assert len(neighborhood_housing_value_data) == len(comparison_housing_value_data)
    assert len(housing_value_categories) == len(neighborhood_housing_value_data) == len(comparison_housing_value_data)
    
    #Add Bars with neighborhood house value distribution
    fig.add_trace(
        go.Bar(
                y        = neighborhood_housing_value_data,
                x        = housing_value_categories,
            name         = neighborhood,
            marker_color = "#4160D3"),
            secondary_y = False
            )

    fig.add_trace(
        go.Bar(
            y            = comparison_housing_value_data,
            x            = housing_value_categories,
            name         = comparison_area,
            marker_color = "#B3C3FF"),
        secondary_y=False
                 )
    
    
    #Set Title
    fig.update_layout(
        title_text = "Owner Occupied Housing Units by Value",    

        title={
            'y':        title_position,
            'x':        0.5,
            'xanchor': 'center',
            'yanchor': 'top'
            },
                        
        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position ,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

    
        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = 14,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin    = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height    = graph_height,
        width     = graph_width,
        
                    )


    #Add % to  axis ticks
    fig.update_yaxes(title=None)
    fig.update_xaxes(tickangle = 45, tickfont = dict(size=tickfont_size - 3))       
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'household_value_graph.png'),engine='kaleido',scale=scale)

def CreatePopulationByAgeHistogram():
    print('Creating Population by Age Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    age_ranges = ['0-19','20-24','25-34','35-49','50-66','67+']
    
    assert len(neighborhood_age_data) == len(age_ranges) 

    #Add Bars with neighborhood household size distribution
    fig.add_trace(
        go.Bar(
                y        = neighborhood_age_data,
                x        = age_ranges,
            name         = neighborhood,
            marker_color = "#4160D3"),
        secondary_y=False
            )

    #Add bars with comparison area age distribution
    fig.add_trace(
        go.Bar(
            y            = comparison_age_data,
            x            = age_ranges,
            name         = comparison_area,
            marker_color = "#B3C3FF"),
        secondary_y=False
            )
    
    
    fig.update_layout(
    
    #Set Title
    title_text = "Population by Age",    

    title = {
        'y':        title_position,
        'x':        0.5,
        'xanchor': 'center',
        'yanchor': 'top'
            },
                    

    legend = dict(
        orientation = "h",
        yanchor   = "bottom",
        y         = legend_position ,
        xanchor   = "center",
        x         = 0.5,
        font_size = tickfont_size
                ),

    font_family   = "Avenir Next LT Pro",
    font_color    = '#262626',
    font_size     = main_graph_font_size,
    paper_bgcolor = paper_backgroundcolor,
    plot_bgcolor  = "White",

    margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height        = graph_height,
    width         = graph_width,
        
                    )


    #Add % to  axis ticks
    fig.update_xaxes(tickangle = 0)  
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False) 
    fig.update_xaxes(tickfont = dict(size = tickfont_size))       
    fig.write_image(os.path.join(hood_folder, 'population_by_age_graph.png'), engine = 'kaleido', scale=scale)

def CreatePopulationByIncomeHistogram():
    print('Creating Population by Income Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    income_categories = ['< $10,000',
                         '$10,000-14,999',
                         '$15,000-19,999',
                         '$20,000-24,999',
                         '$25,000-29,999',
                         '$30,000-34,999',
                         '$35,000-39,999',
                         '$40,000-44,999',
                         '$45,000-49,999',
                         '$50,000-59,999',
                         '$60,000-74,999',
                         '$75,000-99,999',
                         '$100,000-124,999',
                         '$125,000-149,999',
                         '$150,000-199,999',
                         '> $200,000'
                        ]

    assert len(income_categories) == len(neighborhood_household_income_data) == len(comparison_household_income_data)
    
    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(
           y            = neighborhood_household_income_data,
           x            = income_categories,
           name         = neighborhood,
           marker_color = "#4160D3")
            ,secondary_y = False
            )

    #Add bars for comparison area        
    fig.add_trace(
        go.Bar(
            y            = comparison_household_income_data,
            x            = income_categories,
            name         = comparison_area,
            marker_color = "#B3C3FF"),
        secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
        title_text="Households by Household Income",    

        title     = {
            'y':        title_position,
            'x':        0.5,
            'xanchor': 'center',
            'yanchor': 'top'
            },
                        

        legend = dict(
                    orientation = "h",
                    yanchor     = "bottom",
                    y           = legend_position ,
                    xanchor     = "center",
                    x           = 0.5,
                    font_size   = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    ='#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )


    #Add % to Y-axis ticks
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False) 
    fig.update_xaxes(tickfont = dict(size=tickfont_size - 3),tickangle = 45)       
    fig.write_image(os.path.join(hood_folder,'population_by_income_graph.png'),engine='kaleido',scale=scale)

def CreateTopOccupationsHistogram():
    print('Creating Top Occupations Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])
    
    occupations_categories       =  ['Management and Business', 'Service', 'Sales and Office', 'Natural Resources', 'Production'] 
    assert                       len(occupations_categories) == len(neighborhood_top_occupations_data)
        
    #We have a list of categories and a list of their respecitive employment shares. Covert to list, then sort from smallest to largest
    top_occ_dict = {occupations_categories[i]: neighborhood_top_occupations_data[i] for i in range(len(occupations_categories))}
    
    #Now sort dict by values
    top_occ_dict = {k: v for k, v in sorted(top_occ_dict.items(), key=lambda item: item[1])}
    
    sorted_occupations_categories = list(top_occ_dict.keys())
    sorted_occupations_shares     = list(top_occ_dict.values())

    #Add Bars with neighborhood household size distribution
    fig.add_trace(
        go.Bar(y=sorted_occupations_shares,
            x=sorted_occupations_categories,
            name=neighborhood,
            marker_color="#4160D3"),
        secondary_y=False
            )
    
    #Set Title
    fig.update_layout(
        title_text = "Top Employment Occupations",    

        title ={
            'y':        title_position,
            'x':        0.5,
            'xanchor': 'center',
            'yanchor': 'top'
               },
                        

        legend=dict(
                orientation = "h",
                yanchor     = "bottom",
                y           = legend_position ,
                xanchor     = "center",
                x           = 0.5,
                font_size   = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
    
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )

    #Add % to  y-axis ticks
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False) 
    fig.update_xaxes(tickfont = dict(size=tickfont_size))       
    fig.write_image(os.path.join(hood_folder,'top_occupations_graph.png'),engine='kaleido',scale=scale)

def CreateTravelTimeHistogram():
    print('Creating Travel Time to work Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    travel_time_categories = ['< 5 Minutes','5-9 Minutes','10-14 Minutes','15-19 Minutes','20-24 Minutes','25-29 Minutes','30-34 Minutes','35-39 Minutes','40-44 Minutes','45-59 Minutes','60-89 Minutes','> 90 Minutes']
    
    #Add Bars with neighborhood household size distribution
    fig.add_trace(
        go.Bar(
            y            = neighborhood_time_to_work_distribution,
            x            = travel_time_categories,
            name         = neighborhood,
            marker_color = "#4160D3"),
        secondary_y=False
            )

    fig.add_trace(
        go.Bar(
            y            = comparison_time_to_work_distribution,
            x            = travel_time_categories,
            name         = comparison_area,
            marker_color = "#B3C3FF"),
        secondary_y=False
            )

    fig.update_layout(
        
        #Set Title
        title_text = "Travel Time to Work",    

        title = {
            'y':        title_position,
            'x':        0.5,
            'xanchor': 'center',
            'yanchor': 'top'
            },
                        

        legend = dict(
                    orientation="h",
                    yanchor="bottom",
                    y=legend_position ,
                    xanchor="center",
                    x=0.5,
                    font_size = tickfont_size
                    ),


        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",

        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
            
                    )


    #Add % to y-axis ticks   
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size), tickformat='.0f', secondary_y = False)   
    fig.update_xaxes(tickfont = dict(size = tickfont_size))       
    fig.write_image(os.path.join(hood_folder,'travel_time_graph.png'), engine = 'kaleido', scale=scale)

def CreateTravelModeHistogram():
    print('Creating Travel Mode to work Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    travel_method_categories = ['Drove Alone','Car Pooled','Public Transportation','Walked','Worked from Home','Biked','Other']
    assert len(neighborhood_method_to_work_distribution) == len(travel_method_categories)
    
    fig.add_trace(
        go.Bar(
                y        = neighborhood_method_to_work_distribution,
                x        = travel_method_categories,
                name     = neighborhood,
            marker_color = "#4160D3"),
        secondary_y=False
            )
    
    #Set Title
    fig.update_layout(
        title_text="Travel Mode to Work",    

        title={
            'y':title_position,
            'x':0.5,
            'xanchor': 'center',
            'yanchor': 'top'
              },
                        
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=legend_position ,
            xanchor="center",
            x=0.5,
            font_size = tickfont_size
                    ),

        font_family   = "Avenir Next LT Pro",
        font_color    = '#262626',
        font_size     = main_graph_font_size,
        paper_bgcolor = paper_backgroundcolor,
        plot_bgcolor  = "White",
        margin        = dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
        height        = graph_height,
        width         = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(title=None)
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.0f',secondary_y=False)    
    fig.update_xaxes(tickfont = dict(size=tickfont_size))       
    fig.write_image(os.path.join(hood_folder,'travel_mode_graph.png'),engine='kaleido',scale=scale)

def CreateGraphs():
    print('Creating Graphs')
    #Household Size Graph
    try:
        CreateHouseholdSizeHistogram()
    except Exception as e:
        print(e,'unable to create household size graph')
    
    #Household Tenure Graph
    try:
        CreateHouseholdTenureHistogram()
    except Exception as e:
        print(e,'unable to create household tenure graph')
    
    #Household Value graph
    try:
        CreateHouseholdValueHistogram()
    
    except Exception as e:
        print(e,'unable to create housing value graph')
    
    #House Year Built Graph
    try:
        CreateHouseholdYearBuiltHistogram()

    except Exception as e:
        print(e,'unable to create year built graph')
    
    #Number Units in Building Graph
    try:
        CreateHouseholdNumberUnitsInBuildingHistogram()
    except Exception as e:
        print(e,'unable to create number units graph')
    
    #Population By Age Histogram
    try:
        CreatePopulationByAgeHistogram()
    except Exception as e:
        print(e,'unable to create population by age graph')

    #Population by Income Graph
    try:
        CreatePopulationByIncomeHistogram()
    except Exception as e:
        print(e,'unable to create population by income graph')

    #Travel time Graph
    try:
        CreateTravelTimeHistogram()
    except Exception as e:
        print(e,'unable to create travel time graph')

    #Travel Mode Graph
    try:
        CreateTravelModeHistogram()
    except Exception as e:
        print(e,'unable to create travel mode graph')

#####################################################Language Related Functions####################################
def WikipediaTransitLanguage(category):
    #Searches through a wikipedia page for a number of section titles and returns the text from them (if any)
    try:
        #Open a file that holds a list of search terms for different transit categories
        wikipedia_search_terms_df = pd.read_csv(os.path.join(project_location,'Data','General Data','Wikipedia Transit Related Search Terms.csv'))
        
        #Restrict to only rows (search terms) that are relevant to the current cateogry we are looking for (eg: roads)
        wikipedia_search_terms_df = wikipedia_search_terms_df.loc[wikipedia_search_terms_df['category'] == category]
        search_term_list = wikipedia_search_terms_df['search term']
        
        #Create a bs4 html object from the wikipedia page
        soup = BeautifulSoup(page.html(),'html.parser')

        #Loop throuhg each search term looking for text, if we find any, return it as a list
        for search_term in search_term_list:
            langauge_paragraphs = []

            #Loop through every heading in the wikipedia page looking for the section we want
            for heading in soup.find_all(re.compile('^h[1-6]$')):

                #If we find the section we're looking for, pull all the text from the paragraphs
                if search_term in heading.text:
                    para = heading.find_next_sibling('p')
                    para_text = para.text
                    para_text = re.sub('\[\d+\]', '',para_text) #remove wikipedia citations
                    langauge_paragraphs.append(para_text) #Once we have found the relevant section, add all the paragraphs into the list of paragraphs

                    #We found the first paragraph and added it to our list with text paragraphs, now keep looking in case there are multiple paragraphs in the section we wanted
                    while True:
                        para = para.find_next_sibling(['p','h1','h2','h3','h4','h5','h6'])            
                        if para == None:
                            break
                        elif para.name != 'p':
                            break
                        else:
                            para_text = para.text
                            para_text = re.sub('\[\d+\]', '',para_text) #remove wikipedia citations
                            langauge_paragraphs.append(para_text)
            
            if langauge_paragraphs != []:
                return(langauge_paragraphs)


    except Exception as e:
        print(e,'problem getting wikipedia language for ' + category)

def SummaryLangauge():
    print('Creating Summary Langauge')
    try:
        print('Getting summary from wikipedia')
        wikipedia_summary = page.summary
    except Exception as e:
        print(e,'trouble getting wikiepdia summary')
        wikipedia_summary = ('')

    try:
        apartmentsdotcomlanguage = ApartmentsDotComSearch() #neighborhood summary pulled from Apartments.com
    except Exception as e:
        print(e,'trouble getting apartments.com summary')
        apartmentsdotcomlanguage = ('')
        
    return[wikipedia_summary,apartmentsdotcomlanguage]

def CommunityAssetsLanguage():
    print('Creating Community Assets Langauge')
    try:
        community_assets_language = (neighborhood + ' offers a number of community assets throughout. Commercial corridors such as [INSERTROADS], provide access to restaurants, retailers, and other small businesses that are sufficient for weekly errands.')
    except Exception as e:
        print(e,'unable to get community assets langauge')
        community_assets_language = ''

    return([community_assets_language])

def CarLanguage():
    print('Creating auto Langauge')
    wikipedia_car_language     = WikipediaTransitLanguage(category='car')
    
    if wikipedia_car_language != None:
        return(wikipedia_car_language)
    else:
        print('No major highway information on wikipedia, using geographic data')
        nearest_highway_language = FindNearestHighways()
        if nearest_highway_language != None:
            return(nearest_highway_language)
        else:
            return(neighborhood + ' does not have immediate access to any major highways or roads. ' + 'Local corridors include ___________. ')

def PlaneLanguage():
    print('Creating plane Langauge')

    #First see if any text available on wikipedia, if so use that, if not, use our geographic data
    print('Searching Wikipedia for Airport Info')
    wikipedia_plane_language = WikipediaTransitLanguage(category='air')
    if wikipedia_plane_language != None:
        print('Pulled Airport info from Wikipedia')
        return(wikipedia_plane_language)
    
    else:
        #Check to see if there are any airports within the area    
        print('No Airport Information on Wikipedia, using airport shapefile to see if there are any airports within the area')
        airport_language = FindAirport()

        if airport_language != None:
            return(airport_language)

        else:
            print('No Airport Information on Wikipedia or inside the area, using airport shapefile to get closest airport')
            nearest_airport_language = FindNearestAirport(lat = latitude, lon = longitude)
            if (nearest_airport_language != None) and (nearest_airport_language != ''):
                return(nearest_airport_language)
            else:
                print('Unable to find airport using airport shapefile')
                return(neighborhood + ' is roughly ' + '[---]' + ' miles from ' + '------' + ', a [-------] ' + '.')

def BusLanguage():
    print('Creating bus Langauge')

    wikipedia_bus_language = WikipediaTransitLanguage(category='bus')
    if wikipedia_bus_language != None:
        return(wikipedia_bus_language)
    
    bus_lang = FindBusLines()
    if bus_lang != None:
        return(bus_lang)
    else:
        return( neighborhood + ' does not have public bus service, but service is accessible in nearby ________.' ) 
                
def TrainLanguage():
    print('Creating train Langauge')
    wikipedia_train_language = WikipediaTransitLanguage(category='train')
    if wikipedia_train_language != None:
        return(wikipedia_train_language)
    print('No train language on wikipedia, using geographic data')
    train_lang = FindTrainLines()
    if train_lang != None:
        return(train_lang)
    
    
    else:
        print('Did not identify any train routes from  geographic data')
        return('[There is limited use of public transit in ' + neighborhood + '. In fact, it is not served by any commuter or light-rail lines. For public transit options, residents and visitors utilize service in ____.]' +     
              ' [---- provides public train service within ' + neighborhood + '.]' )

def OutlookLanguage():
    print('Creating Outlook Langauge')
    pop_growth_description = '[negative/modest/moderate/strong/extreme]'

    outlook_language = (neighborhood + 
                        ' is a '     + 
                        hood_place_type + 
                        ' in '          + 
                        comparison_area + 
                        ', '            + 
                        comparison_state_full_name + 
                        ' with limited/access to interstate highways, public transportation, a diverse housing stock, recreational amenities, and businesses for everyday use. ' +
                        
                        #Growth sentance
                        'Given the lack of access to/Despite access to/ these amenities,' + neighborhood + ' has experienced ' +
                        pop_growth_description +
                        ' population growth over the past decade, a trend that is expected to continue in the near-term.'
                        
                         )
    
    return([outlook_language])

def HousingIntroLanguage():
    print('Creating housing intro Langauge')
    try:
        housing_mix_description = ('[diverse with varying density, tenure status, and price points but largely homogeneous with the majority of the housing stock consisting of ')

        housing_intro_language = ('Housing is one of the most identifiable characteristics of an area. Different factors, such as property type, ' +
            'renter/owner mix, housing age, and household characteristics play roles in how an area is defined. ' +
            'In ' + neighborhood + ', housing is ' + housing_mix_description +  '. ')
    
    except Exception as e:
        print(e,'Unable to create housing intro language')  
	
    return([housing_intro_language])

def HousingTypeTenureLanguage():
    print('Creating housing type and tenure Langauge')
    try:
        number_units_categories = ['Single-family homes','Townhomes','Duplexes','3-4-unit properties','5-9-unit properties','10-19-unit properties','20-49-unit properties','properties with 50+ units']
        assert len(number_units_categories) == len(neighborhood_number_units_data)
        
        #We have a list of categories and a list of their respecitive employment shares. Covert to list, then sort from smallest to largest
        top_size_dict = {number_units_categories[i]: neighborhood_number_units_data[i] for i in range(len(number_units_categories))}
            
        #Now sort dict by values
        top_size_dict =  {k: v for k, v in sorted(top_size_dict.items(), key=lambda item: item[1])}
        
        size_list   = list(top_size_dict.keys())
        most_common_category        = size_list[len(size_list)     - 1]
        second_most_common_category = size_list[len(size_list)     - 2]
        third_most_common_category = size_list[len(size_list)     - 3]


        
        
        hood_owner_occupied_fraction        =  neighborhood_tenure_distribution[1] 
        comparsion_owner_occupied_fraction  =  comparison_tenure_distribution[1]

        if hood_owner_occupied_fraction > comparsion_owner_occupied_fraction:
            hood_owner_ouccupied_higher_lower   =  'higher than'
        elif hood_owner_occupied_fraction < comparsion_owner_occupied_fraction:
            hood_owner_ouccupied_higher_lower   =  'lower than'
        elif hood_owner_occupied_fraction > comparsion_owner_occupied_fraction:
            hood_owner_ouccupied_higher_lower   =  'equal to'
        else:
            hood_owner_ouccupied_higher_lower   =  '[lower than/higher than/equal to]'
        
        single_family_fraction    = float(neighborhood_number_units_data[0])
        fifty_plus_units_fraction = float(neighborhood_number_units_data[len(neighborhood_number_units_data) -1])
        if  single_family_fraction > 0 and fifty_plus_units_fraction > 0:
            available_housing_types_list = ('single-family homes, some smaller multifamily properties, along with larger garden style properties, and even some buildings with 50+ units')
        else:
            available_housing_types_list = ('single-family homes, some smaller multifamily properties, and some larger garden style properties')


        housing_type_tenure_langugage = (most_common_category                                                    +
                                        ', followed by '                                                         +
                                        second_most_common_category.lower()                                              + 
                                        ' then '                                                                 +
                                        third_most_common_category.lower()                                               +
                                        ' account for the most common forms of housing in '                      +
                                        neighborhood                                                             +                                          
                                        '. '                                                                     +                                     
                                        "{:,.0f}%".format(hood_owner_occupied_fraction)                          +   
                                        ' of the housing units in '                                              + 
                                        neighborhood                                                             + 
                                        ' are occupied by their owner. '                                         +
                                        'This percentage of owner-occupation is '                                +
                                        hood_owner_ouccupied_higher_lower                                        + 
                                        ' the '                                                                  +
                                        comparison_area                                                          +
                                        ' level of '                                                             +
                                        "{:,.0f}%".format(comparsion_owner_occupied_fraction)                    +
                                        '.'
                                        )
    except Exception as e:
        print(e,'unable to get housing tenure type language')
        housing_type_tenure_language = ''




    return([housing_type_tenure_langugage])

def HousingValueLanguage():
    print('Creating Household by value Langauge')
    try:
        housing_value_categories = ['under $10k','$10,000-14,999','$15,000-19,999','$20,000-24,999','$25,000-29,999','$30,000-34,000','$35,000-39,999','under $50k','$50,000-59,9999','$60,000-69,999','$70,000-79,999','$80k - $89k','$90k - $99k','$100k - $124k','$125k - $149k','$150k - $174k','$175k - $199k','$200k - $249k','$250k - $299k','$300k - $399k','$400k - $499k','$500k - $749k','$750k - $1 million','$1 million - $1.5m','$1.5 million -$2m','over $2 million']
        assert len(neighborhood_housing_value_data) == len(housing_value_categories) == len(comparison_housing_value_data)
        
        hood_largest_value_category = housing_value_categories[neighborhood_housing_value_data.index(max(neighborhood_housing_value_data))] #get the most common income category
        comp_largest_value_category = housing_value_categories[comparison_housing_value_data.index(max(comparison_housing_value_data))]

        if (hood_largest_value_category != 'under $10,000') and (hood_largest_value_category != 'over $2,000,000'):
            hood_largest_value_category = 'between ' + hood_largest_value_category
        
        if (comp_largest_value_category != 'under $10,000') and (comp_largest_value_category != 'over $2,000,000'):
            comp_largest_value_category = 'between ' + comp_largest_value_category
            


        value_language = (  'Homes in '                                        +
                        neighborhood                                           + 
                        ' have a median value of '                             + 
                            "${:,.0f}".format(neighborhood_median_home_value)  +
                        ', compared to '                                       +
                        "${:,.0f}".format(comparison_median_home_value)        +
                        ' for '                                                +  
                        comparison_area                                        +
                        '. In '                                                + 
                        neighborhood                                           + 
                        ', the largest share of homes have a value '           +
                        hood_largest_value_category                            +
                        ', compared to '                                       +
                        comp_largest_value_category                            +
                        ' for '                                                +
                            comparison_area                                    +
                            '.'
                            )
    except Exception as e:
        print(e,'unable to get housing value langauge')
        value_language = ''
    
    return([value_language])

def HousingYearBuiltLanguage():
    print('Creating House by Year Built Langauge')
    try:
        year_built_categories       = ['since 2014','since 2010','between 2000-2009','between 1990-1999','between 1980-1989','between 1970-1979','between 1960-1969','between 1950-1959','during WWII','before WWII']
        year_built_categories.reverse()

        #Largest cateorgies for hood and comparison area
        hood_largest_yrblt_category = year_built_categories[neighborhood_year_built_data.index(max(neighborhood_year_built_data))] #get the most common income category
        comp_largest_yrblt_category = year_built_categories[comparison_year_built_data.index(max(comparison_year_built_data))]



        yrblt_language = (  'Homes in '                                         +
                        neighborhood                                            + 
                        ' have a median year built of '                         + 
                           str(neighborhood_median_year_built)[0:4]             +
                        ', compared to '                                        +
                            str(comparison_median_year_built)                   +
                            ' for '                                             +
                            comparison_area                                     +
                        '. '                                                    +
                        
                        'In '                                                   + 
                        neighborhood                                            + 
                        ', the largest share of homes were built '              +
                        hood_largest_yrblt_category                             +
                        ', compared to '                                        +
                        comp_largest_yrblt_category                             +
                        ' for '                                                 +
                            comparison_area                                     +
                            '.'
                            )
    except Exception as e:
        print(e,'Unable to create housing year built langauge')    
        yrblt_language = ''
    
    return([yrblt_language])

def EmploymentLanguage():
    print('Creating Employment by Industry langauge')

    try:
        occupations_categories       =  ['Management and Business','Service','Sales and Office','Natural Resources','Production'] 
        assert                       len(occupations_categories) == len(neighborhood_top_occupations_data)
        
        #We have a list of categories and a list of their respecitive employment shares. Covert to list, then sort from smallest to largest
        top_occ_dict = {occupations_categories[i]: neighborhood_top_occupations_data[i] for i in range(len(occupations_categories))}
        
        #Now sort dict by values
        top_occ_dict = {k: v for k, v in sorted(top_occ_dict.items(), key=lambda item: item[1])}
        
        industry_list   = list(top_occ_dict.keys())
        top_industry    = industry_list[len(industry_list) - 1]
        second_industry = industry_list[len(industry_list) - 2]
        third_industry  = industry_list[len(industry_list) - 3]

        employment_language = ('The majority of working age residents are employed in the ' + top_industry + ', ' + second_industry + ', and ' + third_industry + ' industries. ')
    
    
    except:
        employment_language = ('')

    return([employment_language])

def HouseholdSizeLanguage():
    print('Creating Household by Size Langauge')

    household_size_categories = ['1 person','2 person','3 person','4 person','5 person','6 person','7 or more people']

    #Largest cateogy for hood and comparsion area
    hood_largest_size_category = household_size_categories[neighborhood_household_size_distribution.index(max(neighborhood_household_size_distribution))] #get the most common household size category
    comp_largest_size_category = household_size_categories[comparison_household_size_distribution.index(max(comparison_household_size_distribution))]

    #Compare the average household size of the hood and comparison area
    if neighborhood_average_hh_size  > comparison_average_hh_size:
        avg_hh_size_comparison  = 'The household size in ' + neighborhood + ' tends to be larger than those in ' + comparison_area + '. '    
        comp_smaller_larger     = 'smaller'         
    elif neighborhood_average_hh_size < comparison_average_hh_size:
        avg_hh_size_comparison  = 'The household size in ' + neighborhood + ' tends to be smaller than those in ' + comparison_area + '. '  
        comp_smaller_larger     = 'larger'         
    elif neighborhood_average_hh_size == comparison_average_hh_size:
        avg_hh_size_comparison  = 'The average size of a ' + neighborhood + ' household is equal to those in ' + comparison_area + '. In fact, '
        comp_smaller_larger     = 'equal in size on average'         

    family_vs_individual_break_down = 'Given the average age, household size, and age distribution, the majority of households consist of [individuals, couples, and families].'
    
    #Compare the modal categories
    hood_largest_size_category_numeric = int(hood_largest_size_category[0])
    comp_largest_size_category_numeic  = int(comp_largest_size_category[0])
    assert 0 < hood_largest_size_category_numeric < 8
    assert 0 < comp_largest_size_category_numeic  < 8

    if hood_largest_size_category_numeric != comp_largest_size_category_numeic:
        hh_size_category_comparison  = 'The largest share of households consist of ' + hood_largest_size_category + ' compared to ' + comp_largest_size_category + ' for ' +  comparison_area + ', where the average household tends to be ' + comp_smaller_larger +  '. '          
    elif hood_largest_size_category_numeric == comp_largest_size_category_numeic:
        hh_size_category_comparison  = comp_largest_size_category + ' households account for the largest share in both ' + neighborhood + ' and ' + comparison_area + '. '
 
    household_size_language = (avg_hh_size_comparison + 
                             'Households in '                                          +
                               neighborhood                                            + 
                              ' have an average size of '                              + 
                              "{:,.2f} people".format(neighborhood_average_hh_size)    +
                              ', compared to '                                         +
                              "{:,.2f} people".format(comparison_average_hh_size)      +
                              ' in '                                                   +   
                              comparison_area                                          + 
                              '. '                                                     +
                              hh_size_category_comparison                              
                            #   +
                            #   family_vs_individual_break_down                                              
                            )
    
    return([household_size_language])

def PopLanguage():
    print('Creating Population intro Langauge')

    #global _2010_hood_pop,  _2010_hood_hh
    # global current_hood_pop, current_hood_hh
    # global _2010_comparison_pop, _2010_comparison_hh
    # global current_comparison_pop, current_comparison_hh
    # global  hood_pop_growth,comparsion_hh_growth      
    # global comparsion_pop_growth,comparsion_hh_growth
    
    disclaimer_language = ("""The following demographic profile, created with data from the U.S. Census Bureau, reflects the subject's municipality and market. """)
	

    if neighborhood_level != 'custom':
        
        if float(hood_pop_growth.replace('%','')) < 0:
            hood_pop_growth_or_contract =  'contracted'  
        elif float(hood_pop_growth.replace('%','')) >= 0:
            hood_pop_growth_or_contract =  'grown'  
 


        population_description = ( 'As of the 2010 Census, ' + 
                                   neighborhood             +
                                   ' had a population of '  +
                                    _2010_hood_pop          +
                                    ' people and '          +
                                    _2010_hood_hh           +
                                    ' households. '         +
                                    'Preliminary 2020 Census data shows its population has ' +
                                     hood_pop_growth_or_contract                             +
                                    ' by '                                                   +
                                    hood_pop_growth                                          +
                                    ' per year to '                                          +
                                    current_hood_pop                                         +
                                    ' residents.'
                                    )
    #We don't generate this langauge for custom areas becuase we don't have all the needed info yet
    else:
        population_description = (" Population estimates for "        + 
                                  neighborhood                        +
                                  " reflect the sum of population estimates for census tracts that overlap its geographic boundaries. " +
                                  "Current population estimates for " + 
                                  neighborhood                        +
                                  ' and '                             +
                                  comparison_area                     +
                                " reflect "                            +
                                "data from the most recent 5-year American Community Survey (ACS) and the 2020 Census Redistricting Data Program, respectively.")

    pop_intro_language = (disclaimer_language + population_description)

    return([pop_intro_language])

def PopulationAgeLanguage():
    print('Creating Population by Age Langauge')
    try:
        age_ranges = [' children & teens',' young adults',' young professionals',' those between 35 and 49 years',' those between 50 and 66 years',' those over 67 years old']

        hood_largest_age_category  = age_ranges[neighborhood_age_data.index(max(neighborhood_age_data))] #get the most common income category
        comp_largest_age_category  = age_ranges[comparison_age_data.index(max(comparison_age_data))]

        #Median age higher in hood
        if neighborhood_median_age > comparison_median_age:
            median_age_comparison   = 'The median age of ' + neighborhood + ' residents is older than in ' + comparison_area + '. '             
        
        #Median age lower in hood
        elif neighborhood_median_age < comparison_median_age:
            median_age_comparison   = 'The median age of ' + neighborhood + ' residents is younger than in ' + comparison_area + '. '             
        
        #Median age equal
        elif neighborhood_median_age == comparison_median_age:
            median_age_comparison   = 'The median age of ' + neighborhood + ' residents is equal to ' + comparison_area + '. '             
        
        #Largest Cohort Comparison
        if hood_largest_age_category != comp_largest_age_category:
            largest_age_cohort_comparison = ('In '                                                  + 
                        neighborhood                                                                +
                        ','                                                                         + 
                        hood_largest_age_category                                                   +
                        ' account for the largest cohort, compared to '                             +
                        comparison_area                                                             +
                        ' where'                                                                    +
                        comp_largest_age_category                                                   +
                        ' accounts for the largest cohort. '                                         )

        elif hood_largest_age_category == comp_largest_age_category:
            largest_age_cohort_comparison = ('In both '                                             + 
                        neighborhood                                                                + 
                        ' and '                                                                     +
                        comparison_area                                                             +
                        ','                                                                         + 
                        hood_largest_age_category                                                   +
                        ' account for the largest cohort.')

        age_language = (median_age_comparison                                                       + 
                        'Residents of '                                                             +
                        neighborhood                                                                + 
                        ' have a median age of '                                                    + 
                        "{:,.1f}".format(neighborhood_median_age)                                   +
                        ', compared to '                                                            +
                        "{:,.1f}".format(comparison_median_age)                                     +
                        ' in '                                                                      +
                        comparison_area                                                             +
                        '. '                                                                        +
                        largest_age_cohort_comparison              
                        )

    except Exception as e:
        print(e,'unable to create population by age langauge')
        age_language = ''
        
    return([age_language])

def IncomeLanguage():
    print('Creating HH Income Langauge')
    income_categories = ['under $10k',
                         'between $10k-14,999',
                         'between $15k-19,999',
                         'between $20k-24,999',
                         'between $25k-29,999',
                         'between $30k-34,999',
                         'between $35k-39,999',
                         'between $40k-44,999',
                         'between $45k-49,999',
                         'between $50k-59,999',
                         'between $60k-74,999',
                         'between $75k-99,999',
                         'between $100k-124,999',
                         'between $125k-149,999',
                         'between $150k-199,999',
                         '$200k or higher']

    hood_largest_income_category = income_categories[neighborhood_household_income_data.index(max(neighborhood_household_income_data))] #get the most common income category
    comp_largest_income_category = income_categories[comparison_household_income_data.index(max(comparison_household_income_data))]

    #Compare Median HH Income levels
    if neighborhood_median_hh_inc > comparison_median_hh_inc:
        median_hh_comparison  = 'Households in ' + neighborhood + ' have a higher median income than those in ' + comparison_area + '. '             
    elif neighborhood_median_hh_inc < comparison_median_hh_inc:
        median_hh_comparison  = 'Households in ' + neighborhood + ' have a lower median income than those in ' + comparison_area + '. '
    elif neighborhood_median_hh_inc == comparison_median_hh_inc:
        median_hh_comparison  = 'Median household income levels are equal in ' + neighborhood + ' and ' + comparison_area + '. '
    else:
        median_hh_comparison  = 'Household income levels are similar in ' + neighborhood + ' and ' + comparison_area + '. '

    #Compare the largest income cateogry
    if hood_largest_income_category !=  comp_largest_income_category:
        largest_income_cohort_comparison = (
                        'In '                                                          +                                
                       neighborhood                                                    + 
                       ', the largest share of households has a household income of ' +
                       hood_largest_income_category                                    +
                       ', compared to '                                                +
                       comp_largest_income_category                                    +
                       ' for '                                                         +
                        comparison_area                                                +
                        '.' 
                                                )

    elif hood_largest_income_category ==  comp_largest_income_category:
        largest_income_cohort_comparison = (
                        'In both '                                                     +                                
                       neighborhood                                                    +
                       ' and '                                                         +
                       comparison_area                                                 + 
                       ', the largest share of households has a household income of ' +
                       hood_largest_income_category                                    +
                       '.'
                                                )
        

    income_language = (median_hh_comparison                                  +
                      'Households in '                                       +
                       neighborhood                                          + 
                       ' have a median income of '                           + 
                        "${:,.0f}".format(neighborhood_median_hh_inc)        +
                       ', compared to '                                      +
                       "${:,.0f}".format(comparison_median_hh_inc)           +
                       ' for households in '                                 + 
                       comparison_area                                       +
                       '. The chart below indicates the share of households by income brackets. ' +
                       largest_income_cohort_comparison
                        )
    
    return([income_language])

def TravelMethodLanguage():
    print('Creating Travel Method Langauge')
    try:
        travel_method_categories = ['drive alone','carpooling','use public transit','walking','work from home','biking','other']
        assert len(travel_method_categories) == len(neighborhood_method_to_work_distribution)
        
        #We have a list of categories and a list of their respecitive employment shares. Covert to list, then sort from smallest to largest
        top_method_dict = {travel_method_categories[i]: neighborhood_method_to_work_distribution[i] for i in range(len(travel_method_categories))}
            
        #Now sort dict by values
        top_method_dict =  {k: v for k, v in sorted(top_method_dict.items(), key=lambda item: item[1])}
        
        methods_list   = list(top_method_dict.keys())
        frac_list      = list(top_method_dict.values())
        
        hood_largest_travel_category        = methods_list[len(methods_list)     - 1]
        hood_largest_travel_category_frac   = frac_list[len(methods_list)        - 1]

        second_most_common_category         = methods_list[len(methods_list)     - 2]
        second_most_common_frac             = frac_list[len(methods_list)        - 2]





        travel_method_language = ('In ' + neighborhood + ', the majority of residents ' + hood_largest_travel_category.lower()  + ' with ' +  "{:,.0f}%".format(hood_largest_travel_category_frac) + ' choosing to do so.' +
                                 
                                 ' The second most popular method is ' + 
                                 second_most_common_category +
                                 ' with ' +
                                 "{:,.0f}%".format(second_most_common_frac) +
                                 '.'
                                )
    except Exception as e:
        print(e,'problem creating travel method langauge')
        travel_method_language = ''
    
    return([travel_method_language])
    
def TravelTimeLanguage():
    print('Creating Travel Time Langauge')
    travel_time_categories = ['< 5 minutes','5-9 minutes','10-14 minutes','15-19 minutes','20-24 minutes','25-29 minutes','30-34 minutes','35-39 minutes','40-44 minutes','45-59 minutes','60-89 minutes','> 90 minutes']


    #Estimate a median household income from a category freqeuncy distribution
    hood_median_time_range   = FindMedianCategory(frequency_list=neighborhood_time_to_work_distribution, category_list = travel_time_categories) 
    hood_median_time_range   = hood_median_time_range.replace(' minutes','')
    hood_median_time_range   = hood_median_time_range.replace(',','').split('-')
    hood_median_time         = (int(hood_median_time_range[0]) + int(hood_median_time_range[1]))/2

    #Estimate a median household income from a category freqeuncy distribution
    comp_median_time_range   = FindMedianCategory(frequency_list=comparison_time_to_work_distribution, category_list = travel_time_categories) 
    comp_median_time_range   = comp_median_time_range.replace(' minutes','')
    comp_median_time_range   = comp_median_time_range.replace(',','').split('-')
    comp_median_time         = (int(comp_median_time_range[0]) + int(comp_median_time_range[1]))/2
    
    hood_largest_time_category = travel_time_categories[neighborhood_time_to_work_distribution.index(max(neighborhood_time_to_work_distribution))] #get the most common income category
    comp_largest_time_category = travel_time_categories[comparison_time_to_work_distribution.index(max(comparison_time_to_work_distribution))]

    time_language = ('Commuters in ' + neighborhood + 
                       ' have a median commute time of about '                      + 
                        "{:,.0f} minutes".format(hood_median_time)                   +
                       '. '                     +
                       
                       'In '                                                  + 
                       neighborhood                                           + 
                       ', the largest share of commuters have a commute between ' +
                       hood_largest_time_category +
                       ', compared to ' +
                       comp_largest_time_category        +
                       ' for '           +
                        comparison_area +
                        '.'
                        )
    
    return([time_language])

def EducationLanguage():
    #This function returns a string we will place in the community assets table in the education row 
    education_list                         = LocationIQPOIList(lat = latitude, lon = longitude , category = ['school','college'],radius=10000, limit = 5 ) 
    
    education_language                      = (neighborhood + 
                                         
                                         ' has a number of education options available including ' + 

                                         ', '.join(education_list) + 
                                         '.'
                                         
                                         )
    
    return(education_language)

def FoodLanguage():
    #This function returns a string we will place in the community assets table in the food row 
    food_list                          = LocationIQPOIList(lat = latitude, lon = longitude,  category = ['restaurant'],radius=10000, limit = 5 ) 
    
    food_language                      = ('For sit down restaurants, quick bites, and other eating locations, ' + neighborhood + ' offers options such as ' +

                                         ', '.join(food_list) + 
                                         ', and more.'
                                         
                                         )
    
    return(food_language)

def HospitalLanguage():
    #This function returns a string we will place in the community assets table in the hospital row 
    hospital_list                      = LocationIQPOIList(lat = latitude, lon = longitude,  category = ['hospital'],radius=10000, limit = 5 ) 
    #most residents leave the Borough for nearby communities with clinics and hospitals. 
    hospital_language                  = ('For healthcare needs, residents of the community and region have access to a number of ' + 

                                         'medical facilities including ' +
                                         ', '.join(hospital_list) + 
                                         '.'
                                         
                                         )
    
    return(hospital_language)

def ParkLangauge():
    #This function returns a string we will place in the community assets table in the park row 
    park_list                          = LocationIQPOIList(lat = latitude, lon = longitude,  category = ['park','stadium'],radius=10000, limit = 4) 
    
    park_language                      = ('In ' + neighborhood + 
                                         
                                         ' and the surrounding area, there are several public parks and facilities including ' + 

                                         ', '.join(park_list) + 
                                         ', and a few others.'
                                         
                                         )
    
    return(park_language)

def RetailLanguage():
    #This function returns a string we will place in the community assets table in the retail row 
    retail_list                        = LocationIQPOIList(lat = latitude, lon = longitude,  category = ['shop'],radius=10000, limit = 5 ) 
    
    retail_language                    = ('For shopping needs, ' +
                                         neighborhood + 
                                         ' offers ' + 
                                         ', '.join(retail_list) + 
                                         ', and more.'
                                         
                                         ) 
    
    return(retail_language)

def LandUseLanguage():
    lul = (neighborhood + ' is in NYC community district ' + nyc_community_district + '.')
    return([lul])

def LocationIQPOIList(lat,lon,category,radius,limit):
    #Searches the Locate IQ API for points of interest
    print('Searching Location IQ API for: ',category)

    url = "https://us1.locationiq.com/v1/nearby.php"

    data = {
    'key': location_iq_api_key,
    'lat': lat,
    'lon': lon,
    'tag': category,
    'radius': radius,
    'format': 'json',
    'limit': limit
        }

    try:
        response = requests.get(url, params=data).json()
        poi_list = [x['name'] for x in response]
        return(poi_list)

    except Exception as e:
        try:
            print(e,'Unable to get location IQ result on first attempt, sleeping for 5 seconds and trying again')
            time.sleep(5)
            response = requests.get(url, params=data).json()
            poi_list = [x['name'] for x in response]
            return(poi_list)
        except:
            print('Unable to get location IQ results on second try')
            return(([]))

def CreateLanguage():
    
    print('Creating Langauge')

    global summary_langauge, conclusion_langauge
    global bus_language, car_language, plane_language, train_language
    global population_age_language, income_language, population_language
    global travel_method_language, travel_time_language
    global housing_value_language, year_built_language
    global household_size_language, housing_intro_language, housing_type_tenure_language
    global community_assets_language, education_language, food_language, hospital_language, park_language, retail_language
    global land_use_language

    summary_langauge                   =  SummaryLangauge()
    housing_type_tenure_language       =  HousingTypeTenureLanguage()
    housing_intro_language             =  HousingIntroLanguage()
    housing_value_language             =  HousingValueLanguage()
    year_built_language                =  HousingYearBuiltLanguage()
    population_language                =  PopLanguage()

    #Communtiy assets langauge variables
    community_assets_language          = CommunityAssetsLanguage()
    education_language                 = EducationLanguage() 
    food_language                      = FoodLanguage()
    hospital_language                  = HospitalLanguage()
    park_language                      = ParkLangauge()
    retail_language                    = RetailLanguage()

    #Paragraph Language
    population_age_language            = PopulationAgeLanguage()
    income_language                    = IncomeLanguage()
    household_size_language            = HouseholdSizeLanguage()
    travel_method_language             = TravelMethodLanguage() 
    travel_time_language               = TravelTimeLanguage()

    #Transit Table Language
    bus_language                       = BusLanguage() 
    train_language                     = TrainLanguage()
    car_language                       = CarLanguage()
    plane_language                     = PlaneLanguage()

    land_use_language                  = LandUseLanguage()
    conclusion_langauge                = OutlookLanguage()

#####################################################Report document related functions####################################
def SetPageMargins(document,margin_size):
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin   = Inches(margin_size)
        section.right_margin  = Inches(margin_size)

def SetDocumentStyle(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(9)

def AddTitle(document):
    main_title = document.add_heading('Neighborhood & Demographic Overview draftdraftdraft',level=0) 
    main_title.style = document.styles['Heading 1']
    main_title.paragraph_format.space_after  = Pt(6)
    main_title.paragraph_format.space_before = Pt(12)
    main_title_style = main_title.style
    main_title_style.font.name = "Avenir Next LT Pro Light"
    main_title_style.font.size = Pt(18)
    main_title_style.font.bold = False
    main_title_style.font.color.rgb = RGBColor.from_string('3F65AB')
    main_title_style.element.xml
    rFonts = main_title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

def AddHeading(document,title,heading_level,heading_number,font_size): #Function we use to insert the headers other than the title header
            heading = document.add_heading(title,level=heading_level)
            heading.style = document.styles[heading_number]
            heading_style =  heading.style
            heading_style.font.name = "Avenir Next LT Pro Light"
            heading_style.font.size = Pt(font_size)
            heading_style.font.bold = False
            heading.paragraph_format.space_after  = Pt(6)
            heading.paragraph_format.space_before = Pt(12)

            #Color
            heading_style.font.color.rgb = RGBColor.from_string('3F65AB')            
            heading_style.element.xml
            rFonts = heading_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def AddTableTitle(document,title):
    table_title_paragraph = document.add_paragraph(title)
    table_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_title_paragraph.paragraph_format.space_after  = Pt(6)
    table_title_paragraph.paragraph_format.space_before = Pt(12)
    for run in table_title_paragraph.runs:
                    font = run.font
                    font.name = 'Avenir Next LT Pro Medium'

def Citation(document,text):
    citation_paragraph = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(0)
    citation_paragraph.paragraph_format.space_before = Pt(0)
    run = citation_paragraph.add_run('Source: ' + text)
    font = run.font
    font.name = primary_font
    font.size = Pt(8)
    font.italic = True
    font.color.rgb  = RGBColor.from_string('929292')
    citation_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text != 'Google Maps':
        blank_paragraph = document.add_paragraph('')
        blank_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def GetMap():
    print('Getting Map')
    try:
        #Search Google Maps for hood
        options = webdriver.ChromeOptions()
        options.add_argument("--start-maximized")
        browser = webdriver.Chrome(executable_path=(os.path.join(os.environ['USERPROFILE'], 'Desktop','chromedriver.exe')),options=options)
        browser.get('https:google.com/maps')
            
        #Write hood name in box
        Place = browser.find_element_by_class_name("tactile-searchbox-input")

        if neighborhood_level != 'custom':
            Place.send_keys((neighborhood + ', ' + hood_state))
        
        elif neighborhood_level == 'custom':
            Place.send_keys((neighborhood + ', ' + comparison_area))

        #Submit hood name for search
        Submit = browser.find_element_by_class_name('nhb85d-BIqFsb')
        Submit.click()
    
        if 'Leahy' in os.environ['USERPROFILE']:
            #Move to the left button "Collapse Side Pannel" 
            # to make the grey message go away
            collapse_side_pannel_x = 1048
            collapse_side_pannel_y = 1292
            collapse_side_pannel_duration = 0.5
            pyautogui.moveTo(x=collapse_side_pannel_x,y=collapse_side_pannel_y,duration = collapse_side_pannel_duration)
            time.sleep(2)
            pyautogui.moveTo(x=collapse_side_pannel_x - 200,y=collapse_side_pannel_y-100,duration=collapse_side_pannel_duration)
            time.sleep(1)
            print('Using Mikes coordinates for screenshot')
            im2 = pyautogui.screenshot(region=(1358,465, 2142, 1404) ) #left, top, width, and height
        elif 'Dominic' in os.environ['USERPROFILE']:
            time.sleep(12)
            print('Using Doms coordinates for screenshot')
            im2 = pyautogui.screenshot(region=(3680,254,1968 ,1231) ) #left, top, width, and height

        time.sleep(.1)
        im2.save(os.path.join(hood_folder_map,'map.png'))
        time.sleep(.2)

        # second photo, zoomed out
        zoomout = browser.find_element_by_xpath("""//*[@id="widget-zoom-out"]/div""")
        for i in range(3):
            zoomout.click() 
        time.sleep(2.5)


        if 'Leahy' in os.environ['USERPROFILE']: #differnet machines have different screen coordinates
            print('Using Mikes coordinates for screenshot')
            im2 = pyautogui.screenshot(region=(1358,465, 2142, 1404) ) #left, top, width, and height
        
        elif 'Dominic' in os.environ['USERPROFILE']:
            print('Using Doms coordinates for screenshot')
            im2 = pyautogui.screenshot(region=(3680,254,1968 ,1231) ) #left, top, width, and height
        
        else:
            im2 = pyautogui.screenshot(region=(1089,276, 2405, 1754) ) #left, top, width, and height
        time.sleep(1)

        im2.save(os.path.join(hood_folder_map,'map2.png'))
        im2.close()

        #Wait till we have saved both png images before proceeding
        while (os.path.exists(os.path.join(hood_folder_map,'map.png')) == False) or (os.path.exists(os.path.join(hood_folder_map,'map2.png')) == False):
            pass

        browser.quit()
    except Exception as e:
         print(e)
         try:
            browser.quit()
         except:
            pass

def add_border(input_image, output_image, border):
    #adds border around png images
    img = Image.open(input_image)
    if isinstance(border, int) or isinstance(border, tuple):
        bimg = ImageOps.expand(img, border=border)
    else:
        raise RuntimeError('Border is not an image or tuple')
    bimg.save(output_image)

def OverlayMapImages():
    print("Creating overlayed map image")
    map_path  =  os.path.join(hood_folder_map,'map.png')
    map2_path = os.path.join(hood_folder_map,'map2.png')
    map3_path = os.path.join(hood_folder_map,'map3.png')
    
    #Make sure we have map 1 and map 2 in order to create map 3 (the overlayed map image)
    try:
        assert (os.path.exists(map_path)) and (os.path.exists(map2_path))
    except:
        print('Unable to make overlayed map')
        return('\n')

    #Open zommed out map
    img1 = Image.open(map2_path)
    
    add_border(map_path, output_image = map_path, border=5)
    time.sleep(.2)
    #Open zommed in map
    img2 = Image.open(map_path)

    #Reduce size of zommed in image by a constant factor
    image_reduction_scale = 3
    img2 = img2.resize((int(img2.size[0]/image_reduction_scale),int(img2.size[1]/image_reduction_scale)))
    
    #Add the zoomed in map on top of the zoomed out map and save as new png image
    # No transparency mask specified,                                      
    # simulating an raster overlay
    img1.paste(img2, (img1.size[1] - 25,900))
    
    img1.save(map3_path)
    
def AddMap(document):
    map_path        = os.path.join(hood_folder_map,'map.png')
    map2_path       = os.path.join(hood_folder_map,'map2.png')
    map3_path       = os.path.join(hood_folder_map,'map3.png')
    
    if (os.path.exists(map_path) == False) or (os.path.exists(map2_path) == False): #If we don't have a zommed in map image or a zoomed out map, create one
        GetMap()    
    if os.path.exists(map3_path) == False: #If we don't have an image with a zommed in map overlayed on zoomed out map, create one
        OverlayMapImages()
   
    print('Adding Map') 
    if os.path.exists(map3_path):
        paragraph = document.add_paragraph('')
        paragraph.add_run().add_picture(map3_path,width=Inches(6.5))
        paragraph.paragraph_format.space_after         = Pt(0)
        
def PageBreak(document):
    #Add page break
    page_break_paragraph = document.add_paragraph('')
    run = page_break_paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def AddDocumentParagraph(document,language_variable):
    assert type(language_variable) == list
    for paragraph in language_variable:
        if paragraph == '':
            continue
        par                                               = document.add_paragraph(paragraph)
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Next LT Pro Light'
        par.style = document.styles['Normal']

def AddDocumentPicture(document,image_path,citation):
    if os.path.exists(image_path):
        fig = document.add_picture(os.path.join(image_path),width=Inches(fig_width))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.paragraph_format.space_after       = Pt(0)

        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,citation)

def AddTable(document,data_for_table): #Function we use to insert our overview table into the report document
    #list of list where each list is a row for our table
     
    #make sure each list inside the list of lists has the same number of elements
    for row in data_for_table:
        for row2 in data_for_table:
            assert len(row) == len(row2)


    #create table object
    tab = document.add_table(rows=len(data_for_table), cols=len(data_for_table[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True
    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,data_for_table)): 

    
        row.height = Inches(0)

        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):
            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)

            elif current_column == 1:
                cell.width = Inches(1.25)




            #add border to top row
            if current_row == 1:
                    tcPr = cell._element.tcPr
                    tcBorders = OxmlElement("w:tcBorders")
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                    tcBorders.append(top)
                    tcPr.append(tcBorders)
            
            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                paragraph.paragraph_format.space_after   = Pt(0)
                paragraph.paragraph_format.space_before  = Pt(0)


                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(8)
                    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'

def AddTwoColumnTable(document,pic_list,lang_list):
    #Insert the transit graphics(car, bus,plane, train)
    tab = document.add_table(rows=0, cols=2)
    for pic,lang in zip(pic_list,lang_list):
        pic_path = os.path.join(graphics_location,pic)
        if os.path.exists(pic_path) == False:
            continue
        row_cells = tab.add_row().cells
        
        left_paragraph = row_cells[0].paragraphs[0]
        left_paragraph.alignment                                    = WD_ALIGN_PARAGRAPH.JUSTIFY

        run            = left_paragraph.add_run()

        #Add Picture
        run.add_picture(pic_path,width=Inches(0.2),height =Inches(0.2))

        right_paragraph = row_cells[1].paragraphs[0]
        right_paragraph.alignment                                    = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if type(lang) == list:
            for p in lang:
                run             = right_paragraph.add_run()
                run.add_text(str(p))

        else:
            run             = right_paragraph.add_run()
            run.add_text(str(lang))

    #We have now defined our table object,loop through all rows then all cells in each current row
    for row in tab.rows:
        for current_column,cell in enumerate(row.cells):
            #Set Width for cell
            if current_column == 0:
                cell.width = Inches(.2)
            elif current_column == 1:
                cell.width = Inches(6)

def AddPointOfInterestsTable(document,data_for_table): #Function we use to insert our table with Location IQ points of interest into the report document
    print(data_for_table)
    print(type(data_for_table))
    print(type(data_for_table[0]))

    #Convert the data from location IQ from json to list of list where each list is a row for the table
    converted_data_for_table = [ list(data_for_table[0].keys())  ]
    
    
    for i in data_for_table:
        new_list = list(i.values())
        converted_data_for_table.append(new_list)

    print(converted_data_for_table)
    assert type(converted_data_for_table) == list

    #make sure each list inside the list of lists has the same number of elements
    for row in converted_data_for_table:
        for row2 in converted_data_for_table:
            assert len(row) == len(row2)


    #create table object
    tab = document.add_table(rows=len(converted_data_for_table), cols=len(converted_data_for_table[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True
    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,converted_data_for_table)): 

    
        row.height = Inches(0)

        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):
            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)

            elif current_column == 1:
                cell.width = Inches(1.25)




            #add border to top row
            if current_row == 1:
                    tcPr = cell._element.tcPr
                    tcBorders = OxmlElement("w:tcBorders")
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                    tcBorders.append(top)
                    tcPr.append(tcBorders)
            
            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(8)
                    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'

#####################################################Report sections functions####################################
def IntroSection(document):
    print('Writing Intro Section')
    AddTitle(document = document)
    #Add Map Marker
    fig = document.add_picture(os.path.join(graphics_location,'marker.png'))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.paragraph_format.space_after       = Pt(0)
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    AddMap(document = document)
    Citation(document,'Google Maps')
    AddHeading(document = document, title =  (neighborhood + ' at a Glance'),            heading_level = 1,heading_number='Heading 3',font_size=11)
   
    #Add neighborhood overview language
    AddDocumentParagraph(document = document,language_variable =  summary_langauge)

def LandUseandZoningSection(document):
    nyc_cd_map_path = os.path.join(nyc_cd_map_location,nyc_community_district,'map.png')
    if os.path.exists(nyc_cd_map_path):
        print('Writing Land Use Section')
        AddHeading(document = document, title =  ('Land Use and Zoning'),            heading_level = 1,heading_number='Heading 3',font_size=11)
        print('Adding NYC Community District Map') 
        nyc_map = document.add_picture(nyc_cd_map_path,width=Inches(6.5))
        Citation(document,'NYC Zola')
   
        #Add neighborhood overview language
        AddDocumentParagraph(document = document,language_variable =  land_use_language)

def CommunityAssetsSection(document):
    print('Writing Community Assets Section')
    #Community Assets Section
    AddHeading(document = document, title = 'Community Assets',            heading_level = 1,heading_number='Heading 3',font_size=11)

    AddDocumentParagraph(document = document,language_variable =  community_assets_language)

    #Table Title
    AddTableTitle(document = document, title = 'Community Assets')

    #Add Community Assets Table                 
    AddTwoColumnTable(document,pic_list      = ['education.png','food.png','medical.png','park.png','retail.png'],lang_list =[education_language, food_language, hospital_language, park_language, retail_language] )

def HousingSection(document):
    print('Writing Housing Section')
    AddHeading(document = document, title = 'Housing',                  heading_level = 1,heading_number='Heading 3',font_size=11)
    
    AddDocumentParagraph(document = document,language_variable =  housing_intro_language)

    #Add tenure language
    AddDocumentParagraph(document = document,language_variable =  housing_type_tenure_language)

    #Insert household units by units in_structure graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'household_units_in_structure_graph.png'),citation ='U.S. Census Bureau')

    #Insert Household Tenure graph
    if os.path.exists(os.path.join(hood_folder,'household_tenure_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_tenure_graph.png'),width=Inches(fig_width))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Add housing value language
    AddDocumentParagraph(document = document,language_variable =  housing_value_language)

    #Insert Household value graph
    if os.path.exists(os.path.join(hood_folder,'household_value_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_value_graph.png'),width=Inches(fig_width))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')

    #Add language
    AddDocumentParagraph(document = document,language_variable =  year_built_language)

    #Insert household units by year built graph
    if os.path.exists(os.path.join(hood_folder,'household_year_built_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_year_built_graph.png'),width=Inches(fig_width))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')

def DevelopmentSection(document):
    print('Writing Development Section')
    
    #Development subsection
    AddHeading(document = document, title = 'Development',                  heading_level = 1,heading_number='Heading 3',font_size=11)

def EducationSection(document):
    print('Writing Education Section')

    AddHeading(document = document, title = 'Education',                  heading_level = 1,heading_number='Heading 3',font_size=11)

    if os.path.exists(os.path.join(hood_folder_map,'education_map.png')):
        fig = document.add_picture(os.path.join(hood_folder_map,'education_map.png'),width=Inches(fig_width))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'greatschools.org')

def PopulationSection(document):
    print('Writing Population Section')
    
    AddHeading(document = document, title = 'Population and Households',                                     heading_level = 1,heading_number='Heading 3',font_size=11)
    AddDocumentParagraph(document = document,language_variable =  population_language)
    if neighborhood_level == 'custom':
        AddTableTitle(document = document, title = 'Population Growth')
    else:
        AddTableTitle(document = document, title = 'Population and Household Growth')
    
    try:
        #Add Overview Table
        AddTable(document = document,data_for_table = overview_table_data )
    except Exception as e:
        print(e,'Unable to add overview table')
    
    #household size langauge
    for paragraph in household_size_language:
        if paragraph == '':
            continue
        par                                               = document.add_paragraph(paragraph)
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        par.paragraph_format.space_before                 = Pt(12)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
    # AddDocumentParagraph(document = document,language_variable =  household_size_language)

    #Insert Household size graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'household_size_graph.png'),citation ='U.S. Census Bureau')

    #Age langauge
    AddDocumentParagraph(document = document,language_variable =  population_age_language)
    
    #Insert population by age graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'population_by_age_graph.png'),citation ='U.S. Census Bureau')
    
    #Income langauge
    AddDocumentParagraph(document = document,language_variable =  income_language)

    #Insert population by income graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'population_by_income_graph.png'),citation ='U.S. Census Bureau')

def EmploymentSection(document):
    print('Writing Employment Section')

    #Employment and Transportation Section
    AddHeading(document = document, title = 'Employment',                  heading_level = 1,heading_number='Heading 3',font_size=11)

    AddDocumentParagraph(document = document,language_variable =  employment_language)
    
    #Insert top occupations graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'top_occupations_graph.png'),citation ='U.S. Census Bureau')
        
def TransportationSection(document):
    print('Writing Transportation Section')
    #Employment and Transportation Section
    AddHeading(document = document, title = 'Transportation',                  heading_level = 1,heading_number='Heading 3',font_size=11)

    #Travel method Lanaguage
    AddDocumentParagraph(document = document,language_variable =  travel_method_language)

    #Travel time Lanaguage
    AddDocumentParagraph(document = document,language_variable =  travel_time_language)

    #Insert Travel Time to Work graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'travel_time_graph.png'),citation ='U.S. Census Bureau')
    
    #Insert Transport Method to Work graph
    AddDocumentPicture(document = document, image_path = os.path.join(hood_folder,'travel_mode_graph.png'),citation ='U.S. Census Bureau')
    
    #Transportation Methods table
    AddTableTitle(document = document, title = 'Transportation Methods')
    
    #Only include walk score for custom neighborhoods
    if neighborhood_level == 'custom':
        AddTwoColumnTable(document,pic_list      = ['walk.png','car.png','train.png','bus.png','plane.png',],lang_list =[walk_score_data[0],car_language, train_language, bus_language, plane_language] )
    else:
        AddTwoColumnTable(document,pic_list      = ['car.png','train.png','bus.png','plane.png',],lang_list =[car_language, train_language, bus_language, plane_language] )

def OutlookSection(document):
    print('Writing Outlook Section')
    AddHeading(document = document, title = 'Conclusion',            heading_level = 1,heading_number='Heading 3',font_size=11)

    AddDocumentParagraph(document = document,language_variable =  conclusion_langauge)
    
def WriteReport():
    print('Writing Report')
    #Create Document
    document = Document()
    SetPageMargins(           document  = document, margin_size=1)
    SetDocumentStyle(         document = document)
    IntroSection(             document = document)
    LandUseandZoningSection(  document = document)
    PopulationSection(        document = document)
    HousingSection(           document = document)
    CommunityAssetsSection(   document = document)
    TransportationSection(    document = document)
    OutlookSection(           document = document)


    #Save report
    document.save(report_path)  

def CleanUpPNGs():
    print('Deleting PNG files')
    #Report writing done, delete figures
    files = os.listdir(hood_folder)
    for image in files:
        if image.endswith(".png"):
              while os.path.exists(os.path.join(hood_folder, image)):
                try:
                    os.remove(os.path.join(hood_folder, image))
                except:
                    pass
           

def CreateDirectoryCSV():
    global service_api_csv_name
    print('Creating CSV with file path information on all existing hood reports')
    dropbox_links                  = []
    dropbox_research_names         = []
    dropbox_neighborhoods          = []
    dropbox_analysis_types         = []
    dropbox_states                 = []
    dropbox_versions               = []
    dropbox_statuses               = []
    dropbox_document_names         = []


    for (dirpath, dirnames, filenames) in os.walk(main_output_location):
        if filenames == []:
            continue
        else:
            for file in filenames:
                
                full_path = dirpath + '/' + file

                if file == 'Dropbox Neighborhoods.csv' or ('Legacy Archive' in dirpath) or ('~' in file):
                    continue
                
                if (os.path.exists(full_path.replace('_draft','_FINAL'))) and ('_draft' in full_path) or ('docx' not in full_path):
                    continue

                dropbox_document_names.append(file)
                dropbox_analysis_types.append('Neighborhood')
                dropbox_link = dirpath.replace(dropbox_root,r'https://www.dropbox.com/home')
                dropbox_link = dropbox_link.replace("\\",r'/')    
                dropbox_links.append(dropbox_link)
                dropbox_versions.append(file[0:4])
                if '_draft' in file:
                    file_status = 'Draft'
                else:
                    file_status = 'Final'

                dropbox_statuses.append(file_status)

                
                state_name    = file[5:7]
                
                try:
                    hood_name     = file.split(' - ')[1].strip()
                    research_name = state_name + ' - ' + file.split(' - ')[1].strip() + ' - ' + 'Hood'
                
                except:
                    hood_name     = 'FIX FILE FORMAT'
                    research_name = 'FIX FILE FORMAT'
                
                dropbox_neighborhoods.append(hood_name)
                dropbox_research_names.append(research_name)
                dropbox_states.append(state_name)
            
        

    dropbox_df = pd.DataFrame({'Market Research Name':dropbox_research_names,
                            'Neighborhood':dropbox_neighborhoods,
                           'Analysis Type': dropbox_analysis_types,
                           'State':         dropbox_states,
                           "Dropbox Links":dropbox_links,
                           'Version':dropbox_versions,
                           'Status':dropbox_statuses,
                           'Document Name': dropbox_document_names})
    dropbox_df = dropbox_df.sort_values(by=['State','Market Research Name','Version'], ascending = (True, True,False))

    #Merge the dataframe with a list of states and the inital of who is assigned to complete them
    assigned_to_df                          = pd.read_excel(os.path.join(general_data_location,'Administrative Data','Assigned To States.xlsx')) 
    dropbox_df                              = pd.merge(dropbox_df,assigned_to_df, on=['State'],how = 'left') 

    csv_name = 'Dropbox Neighborhoods.csv'
    service_api_csv_name = f'Dropbox Neighborhoods-{datetime.now().timestamp()}.csv'
    dropbox_df.to_csv(os.path.join(main_output_location, csv_name),index=False)

    if main_output_location == os.path.join(dropbox_root,'Research','Market Analysis','Neighborhood'):
        dropbox_df.to_csv(os.path.join(main_output_location, service_api_csv_name),index=False)

def DecideIfWritingReport():
    global report_creation
    #Give the user 10 seconds to decide if writing reports for metro areas or individual county entries
    try:
        if batch_mode == False:
            report_creation = input_with_timeout('Create new report? y/n', 30).strip()
        else:
            report_creation = 'y'

    except TimeoutExpired:
        report_creation = ''

def UserSelectsNeighborhoodLevel(batch_mode):
    global analysis_type_number
    global neighborhood_level,comparison_level

    if batch_mode == True:
        analysis_type_number = batch_type_number
    else:
        analysis_type_number = int(input('What is the geographic level of the neighborhood and comparison area?' + '\n'
    
                                    '1.) = Place  vs. County'+ '\n' #+
                                    '2.) = County Subdivison vs. County' + '\n' +
                                    '3.) = Custom vs. Place'  + '\n' +
                                    '4.) = Place  vs. County Subdivison'+ '\n' +
                                    # '5.) = Zip    vs. Place'+ '\n' #+

                                    # '6.) = Tract vs. Place'   + '\n' +
                                    # '7.) = Tract vs. County ' + '\n' +
                                    # '8.) = Tract vs. Zip'     + '\n' +
                                    # '9.) = Tract vs. County Subdivison'+ '\n' +
                                    # '10.) = Tract vs. Custom'+ '\n' +
                                    # '11.) = Tract vs. None'+ '\n' +

                                    # '12.) = Place  vs. Zip'+ '\n' +
                                    # '13.) = Place  vs. Custom'+ '\n' +
                                    # '14.) = Place  vs. Tract'+ '\n' +
                                    # '15.) = Place  vs. None'+ '\n' +


                                    # '16.) = County  vs. Place'+ '\n' +
                                    # '17.) = County  vs. Tract' + '\n' +
                                    # '18.) = County vs. Zip' + '\n' +
                                    # '19.) = County vs. Custom'+ '\n' +
                                    # '20.) = County vs. County Subdivison'+ '\n' +
                                    # '21.) = County  vs. None'+ '\n' +

                                    # '22.) = Zip vs. Tract '+ '\n' +
                                    # '23.) = Zip vs. Custom'+ '\n' +
                                    # '24.) = Zip vs. County Subdivison'+ '\n' +
                                    # '25.) = Zip vs. County'+ '\n' +
                                    # '26.) = Zip vs. None'+ '\n' +

                                    # '27.) = County Subdivison vs. Place'  + '\n' +
                                    # '28.) = County Subdivison vs. Custom' + '\n' +
                                    # '29.) = County Subdivison vs. Zip'+ '\n' +
                                    # '30.) = County Subdivison vs. Tract'+ '\n' +
                                    # '31.) = County Subdivison vs. None'  + '\n' +
                                
                                    # '32.) = Custom vs. Tract'  + '\n' +
                                    # '33.) = Custom vs. County Subdivison' + '\n' +
                                    '34.) = Custom vs. County' + '\n' +
                                    # '35.) = Custom vs. Zip'  + '\n' +
                                    # '36.) = Custom  vs. None'  
                                    ''
                                    ))

    

    
    #Each number corresponds to a different analysis level pair eg: place vs county, zip vs. place, etc
    if analysis_type_number == 1:       #Place  vs. County
        neighborhood_level              = 'place'
        comparison_level                = 'county'
    elif analysis_type_number == 2:     #County Subdivison vs. County
        neighborhood_level              = 'county subdivision'
        comparison_level                = 'county'
    elif analysis_type_number == 3:     #Custom vs. Place
        neighborhood_level              = 'custom'
        comparison_level                = 'place'
    elif analysis_type_number == 4:     #Place vs. County Subdivison
        neighborhood_level              = 'place'
        comparison_level                = 'county subdivision'
    elif analysis_type_number == 5:     #Zip vs. Place
        neighborhood_level              = 'zip'
        comparison_level                = 'place'


    elif analysis_type_number == 6:     #Tract vs. Place
        neighborhood_level              = 'tract'
        comparison_level                = 'place'
    elif analysis_type_number == 7:     #Tract vs. County
        neighborhood_level              = 'tract'
        comparison_level                = 'county'
    elif analysis_type_number == 8:     #Tract vs. Zip
        neighborhood_level              = 'tract'
        comparison_level                = 'zip'
    elif analysis_type_number == 9:     #Tract vs. County Subdivison
        neighborhood_level              = 'tract'
        comparison_level                = 'county subdivision'
    elif analysis_type_number == 10:    #Tract vs. Custom
        neighborhood_level              = 'tract'
        comparison_level                = 'custom'
    elif analysis_type_number == 11:    #Tract vs. None
        neighborhood_level              = 'tract'
        comparison_level                = 'None'


    elif analysis_type_number == 12:    #Place  vs. Zip
        neighborhood_level              = 'place'
        comparison_level                = 'zip'
    elif analysis_type_number == 13:    #Place  vs. Custom
        neighborhood_level              = 'place'
        comparison_level                = 'custom'
    elif analysis_type_number == 14:    #Place  vs. Tract
        neighborhood_level              = 'place'
        comparison_level                = 'tract'
    elif analysis_type_number == 15:    #Place  vs. None
        neighborhood_level              = 'place'
        comparison_level                = 'None'


    elif analysis_type_number == 16:    #County  vs. Place
        neighborhood_level              = 'county'
        comparison_level                = 'place'
    elif analysis_type_number == 17:    #County  vs. Tract
        neighborhood_level              = 'county'
        comparison_level                = 'tract'
    elif analysis_type_number == 18:    #County vs. Zip
        neighborhood_level              = 'county'
        comparison_level                = 'zip'
    elif analysis_type_number == 19:    #County vs. Custom
        neighborhood_level              = 'county'
        comparison_level                = 'custom'
    elif analysis_type_number == 20:    #County vs. County Subdivison
        neighborhood_level              = 'county'
        comparison_level                = 'county subdivision'
    elif analysis_type_number == 21:    #County  vs. None
        neighborhood_level              = 'county'
        comparison_level                = 'None'


    elif analysis_type_number == 22:    #Zip vs. Tract
        neighborhood_level              = 'zip'
        comparison_level                = 'tract'
    elif analysis_type_number == 23:    #Zip vs. Custom
        neighborhood_level              = 'zip'
        comparison_level                = 'custom'
    elif analysis_type_number == 24:    #Zip vs. County Subdivison
        neighborhood_level              = 'zip'
        comparison_level                = 'county subdivision'
    elif analysis_type_number == 25:    #Zip vs. County
        neighborhood_level              = 'zip'
        comparison_level                = 'county'
    elif analysis_type_number == 26:    #Zip vs. None
        neighborhood_level              = 'zip'
        comparison_level                = 'None'


    elif analysis_type_number == 27:    #County Subdivison vs. Place
        neighborhood_level              = 'county subdivision'
        comparison_level                = 'place'
    elif analysis_type_number == 28:    #County Subdivison vs. Custom
        neighborhood_level              = 'county subdivision'
        comparison_level                = 'custom'
    elif analysis_type_number == 29:    #County Subdivison vs. Zip
        neighborhood_level              = 'county subdivision'
        comparison_level                = 'zip'
    elif analysis_type_number == 30:    #County Subdivison vs. Tract
        neighborhood_level              = 'county subdivision'
        comparison_level                = 'tract'
    elif analysis_type_number == 31:    #County Subdivison vs. None
        neighborhood_level              = 'county subdivision'
        comparison_level                = 'None'



    elif analysis_type_number == 32:    #Custom vs. Tract
        neighborhood_level              = 'custom'
        comparison_level                = 'tract'
    elif analysis_type_number == 33 :   #Custom vs. County Subdivison
        neighborhood_level              = 'custom'
        comparison_level                = 'county subdivision'
    elif analysis_type_number == 34 :   #Custom vs. County
        neighborhood_level              = 'custom'
        comparison_level                = 'county'
    elif analysis_type_number == 35:    #Custom vs. Zip
        neighborhood_level              = 'custom'
        comparison_level                = 'zip'
    elif analysis_type_number == 36:    #Custom  vs. None
        neighborhood_level              = 'custom'
        comparison_level                = 'None'
    else:
        print('Not a supported level currently')
    return(int(analysis_type_number))

def GetUserInputs():
    global neighborhood, hood_tract, hood_zip, hood_place_fips, hood_place_type, hood_suvdiv_fips, hood_county_fips
    global hood_state, hood_state_fips, hood_state_full_name

    #Get User input on neighborhood/subject area
    if neighborhood_level == 'place':        #when our neighborhood is a town or city eg: East Rockaway Village, New York
       
        if batch_mode == False:
            place_fips_info                 = ProcessPlaceFIPS(input('Enter the 7 digit Census Place FIPS Code'))
        else:
            place_fips_info                 = ProcessPlaceFIPS(place_fips=place_fips)


        hood_place_fips                 = place_fips_info[0]
        hood_state_fips                 = place_fips_info[1]
        neighborhood                    = place_fips_info[2]
        hood_state_full_name            = place_fips_info[3]
        hood_state                      = place_fips_info[4]
        hood_place_type                 = place_fips_info[5]

    elif neighborhood_level == 'county subdivision':     #when our neighborhood is county subdivison eg: Town of Hempstead, New York (A large town in Nassau County with several villages within it)
        
        if batch_mode == False:
            subdivision_fips_info           = ProcessCountySubdivisionFIPS(county_subdivision_fips=input('Enter the 10 digit county subdivision FIPS Code'))

        else:
            subdivision_fips_info           = ProcessCountySubdivisionFIPS(county_subdivision_fips=subdiv_fips)

        hood_suvdiv_fips                = subdivision_fips_info[0]
        hood_county_fips                = subdivision_fips_info[1]
        neighborhood                    = subdivision_fips_info[2]
        hood_state_fips                 = subdivision_fips_info[3]
        hood_state_full_name            = subdivision_fips_info[4]
        hood_state                      = subdivision_fips_info[5]
        hood_place_type                 = subdivision_fips_info[6]

    elif neighborhood_level == 'tract':      #when our neighborhood is a census tract eg: Tract 106.01 in Manhattan
        tract_info                      = ProcessCountyTract(tract = input('Enter the 6 digit tract code for hood'), county_fips =  input('Enter the 5 digit County FIPS Code for the county the hood tract is in'))
        hood_county_fips                = tract_info[0]
        hood_tract                      = tract_info[1]
        neighborhood                    = tract_info[2]
        hood_state_full_name            = tract_info[3]
        hood_state                      = tract_info[4]
        hood_state_fips                 = tract_info[5]
        hood_place_type                 = 'census tract'
                    
    elif neighborhood_level == 'zip':      #When our neighborhood is a zip code eg: 11563
        zip_info                         = ProcessZipCode(zip_code=input('Enter the 5 digit zip code for hood'))
        hood_county_fips                 = zip_info[0]
        hood_zip                         = zip_info[1]
        neighborhood                     = zip_info[2]
        hood_state_full_name             = zip_info[3]
        hood_state                       = zip_info[4]
        hood_state_fips                  = zip_info[5]
        hood_place_type                  = 'zip code'

    elif neighborhood_level == 'county':      #When our neighborhood is a county eg Nassau County, New York
        county_fips_info                = ProcessCountyFIPS(county_fips =   input('Enter the 5 digit county FIPS Code for the hood'))
        hood_county_fips                = county_fips_info[0]
        hood_state_fips                 = county_fips_info[1]
        neighborhood                    = county_fips_info[2]
        hood_state_full_name            = county_fips_info[3]
        hood_state                      = county_fips_info[4]
        hood_place_type                 = 'county'

    elif neighborhood_level == 'custom': #When our neighborhood is a neighboorhood within a city (eg: Financial District, New York City)
        #Get name of hood
        if batch_mode == False:
            neighborhood        = input('Enter the name of the custom neighborhood').strip()
        hood_place_type         = 'neighborhood'

def GetComparsionInfo():
    global comparison_area, comparison_tract ,comparison_zip, comparison_place_fips, comparison_suvdiv_fips, comparison_county_fips
    global comparison_state, comparison_state_fips, comparison_state_full_name
    global comparison_place_type
    global hood_state

    #Get user input on comparison area
    if comparison_level == 'county':          #When our comparison area is a county eg Nassau County, New York
        if neighborhood_level == 'place':
            county_fips_info                      = ProcessCountyFIPS(PlaceFIPSToCountyFIPS(hood_place_fips,hood_state_fips))
        elif neighborhood_level == 'county subdivision':
             county_fips_info                      = ProcessCountyFIPS(county_fips =   hood_state_fips + hood_county_fips)
        elif neighborhood_level == 'custom':
            if batch_mode == False:
                county_fips_info                      = ProcessCountyFIPS(county_fips =   input('Enter the 5 digit county FIPS Code for the hood'))
            elif batch_mode == True:
                county_fips_info                      = ProcessCountyFIPS(county_fips =   place_fips)


        else:
            county_fips_info                      = ProcessCountyFIPS(county_fips =   input('Enter the 5 digit county FIPS Code for the hood'))

        comparison_county_fips                = county_fips_info[0]
        comparison_state_fips                 = county_fips_info[1]
        comparison_area                       = county_fips_info[2]
        comparison_state_full_name            = county_fips_info[3]
        comparison_state                      = county_fips_info[4]
        comparison_place_type                 = 'county'

    elif comparison_level == 'place':        #when our comparison area is a town or city eg: East Rockaway Village, New York
        if batch_mode == True:
            place_fips_info                      = ProcessPlaceFIPS(place_fips) #use this for batches of  all the hoods in a city
        elif batch_mode == False:
            place_fips_info                      = ProcessPlaceFIPS(place_fips = input('Enter the 7 digit Census Place FIPS Code') )
        comparison_place_fips                = place_fips_info[0]
        comparison_state_fips                = place_fips_info[1]
        comparison_area                      = place_fips_info[2]
        comparison_state_full_name           = place_fips_info[3]
        comparison_state                     = place_fips_info[4]
        comparison_place_type                = place_fips_info[5]

    elif comparison_level == 'county subdivision':       #when our comparison area is county subdivison eg: Town of Hempstead, New York (A large town in Nassau County with several villages within it)
        subdivision_fips_info                 = ProcessCountySubdivisionFIPS(county_subdivision_fips=input('Enter the 10 digit county subdivision FIPS Code'))
        comparison_suvdiv_fips                = subdivision_fips_info[0]
        comparison_county_fips                = subdivision_fips_info[1]
        comparison_area                       = subdivision_fips_info[2]
        comparison_state_fips                 = subdivision_fips_info[3]
        comparison_state_full_name            = subdivision_fips_info[4]
        comparison_state                      = subdivision_fips_info[5]
        comparison_place_type                 = subdivision_fips_info[6]

    elif comparison_level == 'zip':        #When our comparison area is a zip code eg: 11563
        zip_info                               = ProcessZipCode(zip_code=input('Enter the 5 digit zip code for comparison area'))
        comparison_county_fips                 = zip_info[0]
        comparison_zip                         = zip_info[1]
        comparison_area                        = zip_info[2]
        comparison_state_full_name             = zip_info[3]
        comparison_state                       = zip_info[4]
        comparison_state_fips                  = zip_info[5]
        comparison_place_type                  = 'zip code'
      
    elif comparison_level == 'tract':        #when our comparison area is a census tract eg: Tract 106.01 in Manhattan
        tract_info                            = ProcessCountyTract(tract = input('Enter the 6 digit tract code'), county_fips =  input('Enter the 5 digit County FIPS Code for the county the hood tract is in'))
        comparison_county_fips                = tract_info[0]
        comparison_tract                      = tract_info[1]
        comparison_area                       = tract_info[2]
        comparison_state_full_name            = tract_info[3]
        comparison_state                      = tract_info[4]
        comparison_state_fips                 = tract_info[5]
        comparison_place_type                 = 'census tract'

    elif comparison_level == 'custom':   #When our comparison area is a neighboorhood within a city (eg: Financial District, New York City)
        comparison_area                       = input('Enter the name of the custom comparison area').strip()
        comparison_place_type                 = 'neighborhood'

    #Use comparison area state when doing a custom report
    if neighborhood_level == 'custom':
        hood_state                      = comparison_state

def UpdateServiceDb(report_type, csv_name, csv_path, dropbox_dir):
    if type == None:
        return
    print(f'Updating service database: {report_type}')

    try:
        url = f'http://market-research-service.bowery.link/api/v1/update/{report_type}'
        dropbox_path = f'{dropbox_dir}{csv_name}'
        payload = { 'location': dropbox_path }

        retry_strategy = Retry(
            total=3,
            status_forcelist=[400, 404, 409, 500, 503, 504],
            allowed_methods=["POST"],
            backoff_factor=5,
            raise_on_status=False
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        http = requests.Session()
        http.mount("https://", adapter)
        http.mount("http://", adapter)

        response = http.post(url, json=payload)
        response.raise_for_status()
        print('Service successfully updated')
    except HTTPError as e:
        print(f'Request failed with status code {response.status_code}')
        json_error = json.loads(response.content)
        print(json.dumps(json_error, indent=2))
        print('Service DB did not successfully update. Please run the script again after fixing the error.')
    finally:
        print(f'Deleting temporary CSV: ', csv_path)
        os.remove(csv_path)           

def Main():
    DecideIfWritingReport()
   
    if report_creation == 'y':
        UserSelectsNeighborhoodLevel(batch_mode)
        GetUserInputs() #user selects if they want to run report and gives input for report subject
        GetComparsionInfo()
        print('Preparing report for: ' + neighborhood + ' compared to ' + comparison_area)
        global latitude
        global longitude
        global current_year
        global neighborhood_shape
        current_year       = str(date.today().year)
        CreateDirectory()
        coordinates        = GetLatandLon()
        latitude           = coordinates[0] 
        longitude          = coordinates[1]
        neighborhood_shape = GetNeighborhoodShape()
    
        #Skip places we have already done
        if os.path.exists(report_path) == False and os.path.exists(report_path.replace('_draft','_FINAL')) == False:
            GetWikipediaPage()
            GetData()
            CreateGraphs()
            CreateLanguage()
            WriteReport()
            CleanUpPNGs()
        print('Report for: ---------' + neighborhood + ' compared to ' + comparison_area + ' Complete ----------------')

SetGraphFormatVariables()
DeclareAPIKeys()
DeclareFormattingParameters()

decennial_census_year = 2010 #The year of the cenus for our variables we pull from sf1 10 year census files
acs_5y_year           = 2019 #The year of the american community survey for our variables we pull from 5 year acs

#Let User Decide (on timer) if they want to use batch mode or not
try:
    batch_mode_input = int(input_with_timeout('Batch Mode: True (1) or False (0)',30))
    if batch_mode_input == 1:
        batch_mode = True
    elif batch_mode_input == 0:
        batch_mode = False
except:
    batch_mode = True


if batch_mode == True:
    
    #Let user select batch number
    try:
        batch_type_number =  int(input_with_timeout('Select batch type (1 = places, 2 = subdivisions, 3 = neighborhoods vs city, 34 = neighborhoods vs county )',30))  #controls what report type we are doing batches of
    except:
        batch_type_number = 1 #controls what report type we are doing batches of
    
    #Let user select if doing batches from salesforce or user inputs
    try:
        salesforce_batch_mode =  int(input_with_timeout('Select batch type (1 = salesforce list, 2 = user proivded county fips',30))  
    except:
        salesforce_batch_mode = 1 #controls what report type we are doing batches of
    
    if salesforce_batch_mode == 1:
        salesforce_batch_mode = True
    else:
        salesforce_batch_mode = False



    #When we are doing a batch of different custom neighborhoods within a single city
    if batch_type_number == (3) or batch_type_number == (34):
        place_fips = input('Enter the 7 digit fips code of the city you want to do all the neighborhoods of')
        for  neighborhood in GetListOfNeighborhoods(input('Enter the name of the city you want to do all the neighborhoods of (Must match with geojson file name')):
            try:
                Main()
            except Exception as e:
                print(e,'REPORT CREATION FAILED')

    #When we are doing a batch of different cities
    elif batch_type_number == 1:
        if salesforce_batch_mode == True:
            place_fips_list             = SalesforcePlaceFIPSList() #Retrieve a list of place fips based on the place names in our salesforce export
        else:
            place_fips_list             = CountyInputPlaceFIPSList(county_fips = input('Enter the 5 digit county fips code')) 
        
        
        for place_fips in place_fips_list:
            if place_fips != None:
                try:
                    Main() #This is our main function that calls all other functions we will use
                except Exception as e:
                    print(e,'REORT CREATION FAILED for',place_fips)
                
    
    
    
    #When we are doing a batch of different county subdivisions
    elif batch_type_number == 2:
        if salesforce_batch_mode == True:
            subdiv_fips_list             = SalesforceSubdivisionFIPSList() #Retrieve a list of place fips based on the place names in our salesforce export
        else:
            subdiv_fips_list             = CountyInputSubdivisionFIPSList(county_fips = input('Enter the 5 digit county fips code'))         
        
        for subdiv_fips in subdiv_fips_list:
            try:
                Main() #This is our main function that calls all other functions we will use
            except Exception as e:
                print(e,'REORT CREATION FAILED for',subdiv_fips)


#When we are doing single reports
else:
    Main() #This is our main function that calls all other functions we will use


#Crawl through directory and create CSV with all current neighborhood report documents
CreateDirectoryCSV()

#Post an update request to the Market Research Docs Service to update the database
if main_output_location == os.path.join(dropbox_root,'Research','Market Analysis','Neighborhood'): 
    UpdateServiceDb(report_type='neighborhoods', 
                csv_name=service_api_csv_name, 
                csv_path=os.path.join(main_output_location, service_api_csv_name),
                dropbox_dir='https://www.dropbox.com/home/Research/Market Analysis/Neighborhood/')
print('Finished!')
