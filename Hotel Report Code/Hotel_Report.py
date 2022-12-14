#Summary: This script injests data on key metrics for different hotel markets across the country from CoStar.com.
          #It creates a report document for each market and saves it in a corresponding folder

import os
import pandas as pd
from us import states
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Inches, Pt, RGBColor

#Define file pre paths
dropbox_root                   =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
project_location               =  os.path.join(dropbox_root,'Research','Projects','Research Report Automation Project')      #Main Folder that stores all output, code, and documentation
#output_location                = os.path.join(dropbox_root,'Research','Market Analysis','Market','Other','Hotel')           #The folder where we store our current reports, production
output_location                = os.path.join(project_location,'Output','Hotel')                                              #The folder where we store our current reports, testing folder
map_location                   = os.path.join(project_location,'Data','Hotel Reports Data','Hotel Maps')                     #Folders with maps png files  
general_data_location          =  os.path.join(project_location,'Data','General Data')                                        #Folder with data used in multiple report types
hotel_data_location            = os.path.join(project_location,'Data','Hotel Reports Data')                                   #Folder with data for hotel reports only


#Import hotel data
hotel_df                       = pd.read_csv(os.path.join(hotel_data_location,'Clean Hotel Data','clean_hotel_data.csv')) 
 

def CreateMarketDictionary(df): #Creates a dictionary where each key is a market and the items are lists of its submarkets
     df_markets             = df.loc[df['Geography Type'] == 'Market'] 
     df_submarkets          = df.loc[df['Geography Type'] == 'Submarket']
     unique_markets_list    = df_markets['Geography Name'].unique()
     unique_submarkets_list = df_submarkets['Geography Name'].unique()
    
     #Now create dictionary to track which submarkets belong to each market
     market_dictionary = {}
     for market in unique_markets_list:
         submarkets = [submarket for submarket in unique_submarkets_list if market in submarket ] #list of sumarkets within current market
         market_dictionary.update({market:submarkets}) 

     return(market_dictionary)

def CleanMarketName(name):
    #Takes a market or submarket name and returns a clean version, also makes windows operating system compliant version for files wihthout slashes
    global market_for_report
    global state

    if name != primary_market: #submarkets
        name  = market.replace(primary_market + ' - ','')
    elif name == primary_market: #markets
        if name == 'Hawaii/Kauai Islands':
            state = 'HI'
        elif name == 'Grand Rapids & Michigan West':
            state = 'MI'
        elif name == 'Central New Jersey':
            state = 'NJ'
        elif name == 'West Virginia':
            state = 'WV'
        elif name == 'Rhode Island':
            state = 'RI'
        elif name == 'Long Island':
            state = 'NY'
        elif name == 'New Hampshire':
            state = 'NH'
        elif name == 'New Jersey Shore':
            state = 'NJ'
        elif name == 'New Mexico North':
            state = 'NM'
        elif name == 'New Mexico South':
            state = 'NM'
        elif name == 'New York State':
            state = 'NY'
        elif name == 'North Carolina East':
            state = 'NC'
        elif name == 'North Carolina West':
            state = 'NC'
        elif name == 'North Dakota':
            state = 'ND'
        elif name == 'South Carolina Area':
            state = 'SC'
        elif name == 'South Dakota':
            state = 'SD'
        elif ' - ' in name:
            state = name[-2:] 
        else:
            state = name.split(' ')[0]

        #Convert long state name to 2 letter state code
        if len(state) > 2:
            state = states.lookup(state)
            state = state.abbr
        assert len(state) == 2
        name = name.replace((' - ' + state),'')
        

    #Create variable without the characrters removed so we can use it in the report document
    market_for_report = name


    for char in ['/',"""\\"""]:
        name = name.replace(char,' ')

    return(name)
    
def CreateReportFilePath():
    return(os.join(output_directory,market_clean + ''))

def CreateOutputDirectory():
    global report_path
    state_folder         = os.path.join(output_location,state)
    market_folder        = os.path.join(output_location,state,primary_market_clean)

    if market == primary_market:
        output_directory     = market_folder                    #Folder where we write report to
        market_or_submarket = 'Market'
    else:
        output_directory     = os.path.join(state_folder,primary_market_clean,market_clean)
        market_or_submarket = 'Submarket'

    document_name = current_quarter + ' ' + state + ' - ' + market_clean + ' - ' + 'Hotel ' + market_or_submarket  + '_draft.docx'
    report_path = os.path.join(output_directory,document_name)

    #Check if output,map, and summary folder already exists, and if it doesnt, make it
    for folder in [state_folder,market_folder,output_directory]:
       
        if os.path.exists(folder) == False:
            os.mkdir(folder) #Create new folder for market or submarket
    return(output_directory)

def CreateMapDirectory():
    global map_path
    state_folder         = os.path.join(map_location,state)
    market_folder        = os.path.join(map_location,state,primary_market_clean)

    if market == primary_market:
        output_directory     = market_folder                    #Folder where we write report to
    else:
        output_directory     = os.path.join(state_folder,primary_market_clean,market_clean)
    map_path = os.path.join(output_directory,'map.png')
    #Check if output,map, and summary folder already exists, and if it doesnt, make it
    for folder in [state_folder,market_folder,output_directory]:
       
        if os.path.exists(folder) == False:
            os.mkdir(folder) #Create new folder for market or submarket
    return(output_directory)

def CreateMarketReport():
    global market,market_clean,output_directory,map_directory,primary_market_clean

    #remove slashes from market names so we can save as folder name
    primary_market_clean = CleanMarketName(name = primary_market)
    market_clean         = CleanMarketName(name = market)

    #Create output, map, and writeup folders for the market of submarket
    output_directory    = CreateOutputDirectory()
    map_directory       = CreateMapDirectory()

    WriteReport()

###############################Report Related Functions###############################
def SetPageMargins(document,margin_size):
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin   = Inches(margin_size)
        section.right_margin  = Inches(margin_size)

def SetDocumentStyle(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(9)

def AddTitle(document):
    title = document.add_heading(market_clean + ' Hotel Analysis',level=1)
    title.style = document.styles['Heading 2']
    title.paragraph_format.space_after  = Pt(6)
    title.paragraph_format.space_before = Pt(12)
    title_style = title.style
    title_style.font.name = "Avenir Next LT Pro Light"
    title_style.font.size = Pt(14)
    title_style.font.bold = False
    title_style.font.color.rgb = RGBColor.from_string('3F65AB')
    title_style.element.xml
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")


    if market == primary_market:
        above_map_paragraph = document.add_paragraph('The information contained in this report was provided using ' +
                                             current_quarter + 
                                            ' CoStar data for the ' + 
                                            market + 
                                            ' ' + 
                                            'Hotel' + 
                                            """ Market ("Market").""")
    #Submarket disclaimer
    else:

        above_map_paragraph = document.add_paragraph('The information contained in this report was provided using ' +
                                            current_quarter  + 
                                            ' CoStar data for the ' + 
                                            market_clean + 
                                            ' ' + 
                                            'Hotel' + 
                                            """ Submarket ("Submarket") """ +
                                            'located in the ' +
                                            primary_market_clean +
                                            """ Market ("Market"). """ 
                                            )     



    above_map_style = above_map_paragraph.style
    above_map_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    above_map_style.font.size = Pt(9)
    above_map_paragraph.paragraph_format.space_after  = Pt(primary_space_after_paragraph)

def AddHeading(document,title,heading_level): #Function we use to insert the headers other than the title header
            heading = document.add_heading(title,level=heading_level)
            heading.style = document.styles['Heading 3']
            heading_style =  heading.style
            heading_style.font.name = "Avenir Next LT Pro"
            heading_style.font.size = Pt(11)
            heading_style.font.bold = False
            heading.paragraph_format.space_after  = Pt(6)
            heading.paragraph_format.space_before = Pt(12)

            #Color
            heading_style.font.color.rgb = RGBColor.from_string('3F65AB')            
            heading_style.element.xml
            rFonts = heading_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def Citation(document,text):
    citation_paragraph = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(6)
    citation_paragraph.paragraph_format.space_before = Pt(6)
    run = citation_paragraph.add_run('Source: ' + text)
    font = run.font
    font.name = primary_font
    font.size = Pt(8)
    font.italic = True
    font.color.rgb  = RGBColor.from_string('929292')
    citation_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text != 'Google Maps':
        pass

def Note(document,text):
    citation_paragraph = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(6)
    citation_paragraph.paragraph_format.space_before = Pt(6)
    run = citation_paragraph.add_run('Note: ' + text)
    font = run.font
    font.name = primary_font
    font.size = Pt(8)
    font.italic = True
    font.color.rgb  = RGBColor.from_string('929292')
    citation_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def AddDocumentParagraph(document,language_variable):
    assert type(language_variable) == list
    for paragraph in language_variable:
        if paragraph == '':
            continue
        par                                               = document.add_paragraph(str(paragraph))
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
        style = document.styles['Normal']
        font = style.font
        font.name = 'Avenir Next LT Pro Light'
        par.style = document.styles['Normal']

def AddDocumentPicture(document,image_path,citation):
    if os.path.exists(image_path):
        fig = document.add_picture(os.path.join(image_path),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.paragraph_format.space_after       = Pt(0)

        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,citation)

def AddTableTitle(document,title):
    table_title_paragraph = document.add_paragraph(title)
    table_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_title_paragraph.paragraph_format.space_after  = Pt(6)
    table_title_paragraph.paragraph_format.space_before = Pt(12)
    for run in table_title_paragraph.runs:
                    font = run.font
                    font.name = 'Avenir Next LT Pro Medium'


def OverviewSection():

    #Overview Heading
    AddHeading(document, 'Overview', 2)
    
    #Overview Paragraph
    AddDocumentParagraph(document = document, language_variable = overview_language)

    #Overview table title
    AddTableTitle(document = document, title = 'Sector Fundamentals')
    
    #Overview table

def SupplyDemandSection():
    #Supply and Demand Section
    AddHeading(document, 'Supply & Demand', 2)
    
    AddDocumentParagraph(document = document, language_variable = supply_and_demand_language)

#Language Related functions
def OverviewLanguage():
    try:
        overview_paragraph = 'The subject property is located in the ' + market_clean + '.'
    except Exception as e:
        print(e, 'Unable to create overview language')
        overview_paragraph = ''
    return([overview_paragraph])

def SupplyDemandLanguage():
    try:
        supply_demand_paragraph =  ('The current inventory level in ' + market_clean + ' is ' + "{:,.0f}".format(hotel_df['Inventory'].iloc[-1]) + '. ' +  
                                    'There were ' + "{:,.0f}".format(hotel_df['Homes Sold'].iloc[-1]) + ' homes sold in ' + subject_latest_period.strftime('%B') +
                                    ' with an average sale to list price ratio of ' + "{:,.1f}".format(hotel_df['Average Sale To List'].iloc[-1])  + '%.'
                                    )

    except Exception as e:
        print(e, 'Unable to create supply and demand language')
        supply_demand_paragraph = ''
    return([supply_demand_paragraph])

def ADRREVPARLanguage():
    try:
        ADRREVPAR_paragraph = ('The current median price/SF in ' + market_clean + ' is ' + "${:,.0f}".format(hotel_df['Median Price/SF'].iloc[-1]) + 
                            '/SF compared to ' +  "${:,.0f}".format(df_comparison['Median Price/SF'].iloc[-1])  + '/SF in ' + comparison_name.strip() + '.'    
                           )
         
    except Exception as e:
        print(e, 'Unable to create values language')
        ADRREVPAR_paragraph = ''
    return([ADRREVPAR_paragraph])

def ConstructionLanguage():
    try:
        Construction_paragraph = ('The current median price/SF in ' + market_clean + ' is ' + "${:,.0f}".format(hotel_df['Median Price/SF'].iloc[-1]) + 
                            '/SF compared to ' +  "${:,.0f}".format(df_comparison['Median Price/SF'].iloc[-1])  + '/SF in ' + comparison_name.strip() + '.'    
                           )
         
    except Exception as e:
        print(e, 'Unable to create values language')
        Construction_paragraph = ''
    return([Construction_paragraph])

def ValuesLanguage():
    try:
        values_paragraph = ('The current median price/SF in ' + market_clean + ' is ' + "${:,.0f}".format(hotel_df['Median Price/SF'].iloc[-1]) + 
                            '/SF compared to ' +  "${:,.0f}".format(df_comparison['Median Price/SF'].iloc[-1])  + '/SF in ' + comparison_name.strip() + '.'    
                           )
         
    except Exception as e:
        print(e, 'Unable to create values language')
        values_paragraph = ''
    return([values_paragraph])

def ConclusionLanguage():
    try:
        outlook_paragraph = ''
    except Exception as e:
        print(e, 'Unable to create conclusion language')
        outlook_paragraph = ''

    return([outlook_paragraph])

def CreateLanguage():
    print('Creating Language')
    global overview_language, supply_and_demand_language, ADRREVPAR_language, Construction_language, values_language, conclusion_language
    overview_language          = OverviewLanguage()
    supply_and_demand_language = SupplyDemandLanguage()
    ADRREVPAR_language         = ADRREVPARLanguage()
    Construction_language      = ConstructionLanguage()
    values_language            = ValuesLanguage()
    conclusion_language        = ConclusionLanguage()



def WriteReport():
    global document
    #Create Document
    document = Document()
    SetPageMargins(document   = document, margin_size=1)
    SetDocumentStyle(document = document)
    AddTitle(document = document)
    AddMap()
    AddDocumentParagraph
    AddTableTitle

    #Save report
    document.save(report_path)  

def AddMap():
    #Add image of map if there is one in the appropriate map folder
    if os.path.exists(map_path):
        map = document.add_picture(map_path, width=Inches(6.5) )
    else:
        map = document.add_paragraph('')

    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


#Set global paramaters
current_quarter               = '2022 Q2'
primary_font                  = 'Avenir Next LT Pro Light' 
primary_space_after_paragraph = 8

#Create dictionary with each market as key and a list of its submarkets as items
market_dictionary            = CreateMarketDictionary(hotel_df)

#Our main loop where we loop through each item in the market dictionary
for primary_market,submarkets in market_dictionary.items():


    market = primary_market
    print(primary_market)
    CreateMarketReport()
    
    for submarket in submarkets:
        market = submarket
        print(submarket)
        CreateMarketReport()