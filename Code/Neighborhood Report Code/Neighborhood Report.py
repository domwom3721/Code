#By Mike Leahy
#Started 06/30/2021
#Summary: This script creates reports on neighborhoods/cities for Bowery

import math
import os
import time
from datetime import date
from pprint import pprint
from random import randrange

import docx
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import pyautogui
import requests
import wikipedia
from bls_datasets import oes, qcew
from blsconnect import RequestBLS, bls_search
from census import Census
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Inches, Pt, RGBColor
from fredapi import Fred
from numpy import true_divide
from plotly.subplots import make_subplots
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from us import states
import us
from wikipedia.wikipedia import random

#Define file paths
dropbox_root                   =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
project_location               =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)','Research','Projects', 'Research Report Automation Project') 
main_output_location           =  os.path.join(project_location,'Output','Neighborhood Reports') #testing
# main_output_location           =  os.path.join(dropbox_root,'Research','Market Analysis','Neighborhood') #production
data_location                  =  os.path.join(project_location,'Data','Neighborhood Reports Data')
graphics_location              =  os.path.join(project_location,'Data','Graphics')
map_location                   =  os.path.join(project_location,'Data','Maps','Neighborhood Maps')

c    = Census('18335344cf4a0242ae9f7354489ef2f8860a9f61') #Census API Key
# https://maps.googleapis.com/maps/api/geocode/json?key=AIzaSyBMcoRFOW2rxAGxURCpA4gk10MROVVflLs&address=90%Jarvis%Place

#Directory Realted Functions
def CreateDirectory():

    global report_path,hood_folder_map,hood_folder
    
    state_folder_map         = os.path.join(map_location,state)
    hood_folder_map          = os.path.join(map_location,state,neighborhood)
    
    state_folder             = os.path.join(main_output_location,state)
    hood_folder              = os.path.join(main_output_location,state,neighborhood)

    for folder in [state_folder,hood_folder,state_folder_map,hood_folder_map]:
         if os.path.exists(folder):
            pass 
         else:
            os.mkdir(folder) 

    report_path = os.path.join(hood_folder,current_year + ' ' + state + ' - ' + neighborhood  + ' - hood' + '_draft.docx')

#Data Manipulation functions
def ConvertListElementsToFractionOfTotal(raw_list):
    #Convert list with raw totals into a list where each element is a fraction of the total
    total = sum(raw_list)

    converted_list = []
    for i in raw_list:
        converted_list.append(i/total * 100)
    
    return(converted_list)

#Data Gathering Related Functions

#Household Size
def GetPlaceHouseholdSizeData():
    total_number_households  =  c.sf1.state_place(fields=['H013001'],state_fips=state_fips,place=place_fips)[0]['H013001']
    average_household_size   =  c.sf1.state_place(fields=['H012001'],state_fips=state_fips,place=place_fips)[0]['H012001'] 


    neighborhood_1_people_households = c.sf1.state_place(fields=['H013002'],state_fips=state_fips,place=place_fips)[0]['H013002']
    neighborhood_2_people_households = c.sf1.state_place(fields=['H013003'],state_fips=state_fips,place=place_fips)[0]['H013003']
    neighborhood_3_people_households = c.sf1.state_place(fields=['H013004'],state_fips=state_fips,place=place_fips)[0]['H013004']
    neighborhood_4_people_households = c.sf1.state_place(fields=['H013005'],state_fips=state_fips,place=place_fips)[0]['H013005']
    neighborhood_5_people_households = c.sf1.state_place(fields=['H013006'],state_fips=state_fips,place=place_fips)[0]['H013006']
    neighborhood_6_people_households = c.sf1.state_place(fields=['H013007'],state_fips=state_fips,place=place_fips)[0]['H013007']
    neighborhood_7_people_households = c.sf1.state_place(fields=['H013008'],state_fips=state_fips,place=place_fips)[0]['H013008'] #7 or more

    neighborhood_household_size_distribution = [neighborhood_1_people_households,
                                               neighborhood_2_people_households,
                                               neighborhood_3_people_households,
                                               neighborhood_4_people_households,
                                               neighborhood_5_people_households,
                                               neighborhood_6_people_households,
                                               neighborhood_7_people_households]
            
    neighborhood_household_size_distribution = [(i/total_number_households * 100) for i in neighborhood_household_size_distribution] #convert from raw ammounts to fraction of total households
    return(neighborhood_household_size_distribution)
    
def GetCountyHouseholdSizeData():
    #Get County household size distribution
    county_total_number_households  =  c.sf1.state_county(fields=['H013001'],state_fips=state_fips,county_fips=county_fips)[0]['H013001']
    county_1_people_households = c.sf1.state_county(fields=['H013002'],state_fips=state_fips,county_fips=county_fips)[0]['H013002']
    county_2_people_households = c.sf1.state_county(fields=['H013003'],state_fips=state_fips,county_fips=county_fips)[0]['H013003']
    county_3_people_households = c.sf1.state_county(fields=['H013004'],state_fips=state_fips,county_fips=county_fips)[0]['H013004']
    county_4_people_households = c.sf1.state_county(fields=['H013005'],state_fips=state_fips,county_fips=county_fips)[0]['H013005']
    county_5_people_households = c.sf1.state_county(fields=['H013006'],state_fips=state_fips,county_fips=county_fips)[0]['H013006']
    county_6_people_households = c.sf1.state_county(fields=['H013007'],state_fips=state_fips,county_fips=county_fips)[0]['H013007']
    county_7_people_households = c.sf1.state_county(fields=['H013008'],state_fips=state_fips,county_fips=county_fips)[0]['H013008'] #7 or more

    county_household_size_distribution = [county_1_people_households,
                                               county_2_people_households,
                                               county_3_people_households,
                                               county_4_people_households,
                                               county_5_people_households,
                                               county_6_people_households,
                                               county_7_people_households]
    county_household_size_distribution = [(i/county_total_number_households * 100) for i in county_household_size_distribution] #convert from raw ammounts to fraction of total households
    return(county_household_size_distribution)

#Household Tenure
def GetPlaceHousingTenureData():
    #Occupied Housing Units by Tenure
    neighborhood_renter_households    = c.sf1.state_place(fields=['H004004'],state_fips=state_fips,place=place_fips)[0]['H004004']  
    neighborhood_owner_households     = c.sf1.state_place(fields=['H004003'],state_fips=state_fips,place=place_fips)[0]['H004003'] #Owned free and clear
    neighborhood_mortgage_households  = c.sf1.state_place(fields=['H004002'],state_fips=state_fips,place=place_fips)[0]['H004002'] #Owned with a mortgage or a loan
    neighborhood_tenure_total_households = neighborhood_renter_households + neighborhood_owner_households + neighborhood_mortgage_households
    
    neighborhood_tenure_distribution = [neighborhood_renter_households/neighborhood_tenure_total_households * 100,
                                       neighborhood_owner_households/neighborhood_tenure_total_households * 100,
                                       neighborhood_mortgage_households/neighborhood_tenure_total_households * 100]
    return(neighborhood_tenure_distribution)

def GetCountyHousingTenureData():
    county_renter_households          = c.sf1.state_county(fields=['H004004'],state_fips=state_fips,county_fips=county_fips)[0]['H004004']
    county_owner_households           = c.sf1.state_county(fields=['H004003'],state_fips=state_fips,county_fips=county_fips)[0]['H004003']
    county_mortgage_households        = c.sf1.state_county(fields=['H004002'],state_fips=state_fips,county_fips=county_fips)[0]['H004002']
    county_tenure_total_households    = county_renter_households + county_owner_households + county_mortgage_households
    
    county_tenure_distribution        = [ county_renter_households /county_tenure_total_households * 100, 
                                          county_owner_households /county_tenure_total_households * 100,
                                           county_mortgage_households/county_tenure_total_households * 100]
    return(county_tenure_distribution)

#Age Related Data Functions
def GetPlaceAgeData():
    print('Getting Place Age breakdown')
    #Return a list of ages across age groups for a given census place (7 digit FIPS, town/village/city) 

    #5 Year ACS age variables for men range:  B01001_003E - B01001_025E
    male_age_data = c.acs5.state_place(fields=["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,26)],state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the age values from the dictionary inside of it
    male_age_breakdown = []
    for field in ["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,26)]:
        male_age_breakdown.append(male_age_data[field])

    #5 Year ACS age variables for women range:  B01001_027E - B01001_049E
    female_age_data = c.acs5.state_place(fields=["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(27,50)],state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the age values from the dictionary inside of it
    female_age_breakdown = []
    for field in ["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(27,50)]:
        female_age_breakdown.append(female_age_data[field])
    
    
    total_pop = sum(male_age_breakdown) + sum(female_age_breakdown)

    total_age_breakdown = []
    for (men, women) in zip(male_age_breakdown, female_age_breakdown):
        total = (men + women)
        total_age_breakdown.append((total/total_pop) * 100)

    
    return(total_age_breakdown)

def GetCountyAgeData():
    print('Getting County Age data')
    #Return a list of ages across age groups for a given county (5 digit FIPS) 

    #5 Year ACS age variables for men range:  B01001_003E - B01001_025E
    male_age_data = c.acs5.state_county(fields=["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,26)],state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the age values from the dictionary inside of it
    male_age_breakdown = []
    for field in ["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,26)]:
        male_age_breakdown.append(male_age_data[field])

    #5 Year ACS age variables for women range:  B01001_027E - B01001_049E
    female_age_data = c.acs5.state_county(fields=["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(27,50)],state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the age values from the dictionary inside of it
    female_age_breakdown = []
    for field in ["B01001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(27,50)]:
        female_age_breakdown.append(female_age_data[field])
   
    total_pop = sum(male_age_breakdown) + sum(female_age_breakdown)

    total_age_breakdown = []
    for (men, women) in zip(male_age_breakdown, female_age_breakdown):
        total = (men + women)
        total_age_breakdown.append(total/total_pop * 100)


    return(total_age_breakdown)

#Housing related data functions
def GetPlaceHousingValues():
    return([])

def GetCountyHousingValues():
    return([])

#Number of Housing Units based on number of units in building
def GetPlaceNumberUnitsData():
    print('Getting Place housing units by number of units data')
    
    #5 Year ACS owner occupied number of units variables for men range:  B25032_003E - B25032_010E
    owner_occupied_fields_list = ["B25032_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,11)]
    owner_occupied_units_raw_data = c.acs5.state_place(fields = owner_occupied_fields_list,state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    owner_occupied_units_data = []
    for field in owner_occupied_fields_list:
        owner_occupied_units_data.append(owner_occupied_units_raw_data[field])

    #Now repeat for the renter occupied fields
    renter_occupied_fields_list = ["B25032_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(14,22)]
    
    #5 Year ACS owner occupied number of units variables for men range:   - B25032_014E  - B25032_021E 
    renter_occupied_units_raw_data = c.acs5.state_place(fields = renter_occupied_fields_list,state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    renter_occupied_units_data = []
    for field in renter_occupied_fields_list:
        renter_occupied_units_data.append(renter_occupied_units_raw_data[field])

    
    total_units = sum(owner_occupied_units_data) + sum(renter_occupied_units_data)

    total_unit_data = []
    for (oo, ro) in zip(owner_occupied_units_data, renter_occupied_units_data):
        total_unit_data.append(( (oo + ro )/total_units) * 100)

    
    return(total_unit_data)
    
def GetCountyNumberUnitsData():
    print('Getting County housing units by number of units data')
    
    #5 Year ACS owner occupied number of units variables for men range:  B25032_004E - B25032_010E
    owner_occupied_fields_list = ["B25032_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(3,11)]
    owner_occupied_units_raw_data = c.acs5.state_county(fields = owner_occupied_fields_list,state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    owner_occupied_units_data = []
    for field in owner_occupied_fields_list:
        owner_occupied_units_data.append(owner_occupied_units_raw_data[field])

    #Now repeat for the renter occupied fields
    renter_occupied_fields_list = ["B25032_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(14,22)]
    
    #5 Year ACS owner occupied number of units variables for men range: B25032_014E  - B25032_021E 
    renter_occupied_units_raw_data = c.acs5.state_county(fields = renter_occupied_fields_list,state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    renter_occupied_units_data = []
    for field in renter_occupied_fields_list:
        renter_occupied_units_data.append(renter_occupied_units_raw_data[field])

    
    total_units = sum(owner_occupied_units_data) + sum(renter_occupied_units_data)

    total_unit_data = []
    for (oo, ro) in zip(owner_occupied_units_data, renter_occupied_units_data):
        total_unit_data.append(( (oo + ro )/total_units) * 100)

    
    return(total_unit_data)

#Household Income data functions
def GetPlaceHouseholdIncomeValues():
    #5 Year ACS household income range:  B19001_002E -B19001_017E
    fields_list = ["B19001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,18)]
    household_income_data = c.acs5.state_place(fields=fields_list,state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    household_income_breakdown = []
    for field in fields_list:
        household_income_breakdown.append(household_income_data[field])


    total_pop = sum(household_income_breakdown) 

    total_income_breakdown = []
    for i in household_income_breakdown:
        total_income_breakdown.append((i/total_pop) * 100)

    assert len(total_income_breakdown) == 16
    return(total_income_breakdown)

def GetCountyHouseholdIncomeValues():
    #5 Year ACS household income range:  B19001_002E -B19001_017E
    household_income_data = c.acs5.state_county(fields=["B19001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,18)],state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    household_income_breakdown = []
    for field in["B19001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,18)]:
        household_income_breakdown.append(household_income_data[field])


    total_pop = sum(household_income_breakdown) 

    total_income_breakdown = []
    for i in household_income_breakdown:
        total_income_breakdown.append((i/total_pop) * 100)

    assert len(total_income_breakdown) == 16
    return(total_income_breakdown)

#Occupations Data
def GetPlaceTopOccupationsData():
    return([])
    # fields_list = ["B19001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,18)]
    # fields_list = []
    # data = c.acs5.state_place(fields=fields_list,state_fips=state_fips,place=place_fips)[0]
    
    # #Create an empty list and place the values from the dictionary inside of it
    # household_income_breakdown = []
    # for field in fields_list:
    #     household_income_breakdown.append(household_income_data[field])


    # total_pop = sum(household_income_breakdown) 

    # total_income_breakdown = []
    # for i in household_income_breakdown:
    #     total_income_breakdown.append((i/total_pop) * 100)

    # assert len(total_income_breakdown) == 16
    # return(total_income_breakdown)

def GetCountyTopOccupationsData():
    return([])
    # fields_list = ["B19001_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,18)]
    # fields_list = []
    # household_income_data = c.acs5.state_county(fields=fields_list,state_fips=state_fips,county_fips = county_fips)[0]
    
    # #Create an empty list and place the values from the dictionary inside of it
    # household_income_breakdown = []
    # for field in fields_list:
    #     household_income_breakdown.append(household_income_data[field])


    # total_pop = sum(household_income_breakdown) 

    # total_income_breakdown = []
    # for i in household_income_breakdown:
    #     total_income_breakdown.append((i/total_pop) * 100)

    # assert len(total_income_breakdown) == 16
    # return(total_income_breakdown)

#Year Housing Built Data
def GetPlaceHouseYearBuiltData():
    #5 Year ACS household year house built range:  B25034_002E -B25034_011E
    fields_list = ["B25034_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,12)]

    year_built_raw_data = c.acs5.state_place(fields=fields_list,state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    year_built_data = []
    for field in fields_list:
        year_built_data.append(year_built_raw_data[field])

    #Convert list with raw totals into a list where each element is a fraction of the total
    year_built_data = ConvertListElementsToFractionOfTotal(year_built_data)


    return(year_built_data)

def GetCountyHouseYearBuiltData():
    #5 Year ACS household year house built range:  B25034_002E -B25034_011E
    fields_list = ["B25034_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,12)]

    year_built_raw_data = c.acs5.state_county(fields=fields_list,state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    year_built_data = []
    for field in fields_list:
        year_built_data.append(year_built_raw_data[field])

    #Convert list with raw totals into a list where each element is a fraction of the total
    year_built_data = ConvertListElementsToFractionOfTotal(year_built_data)


    return(year_built_data)


#Travel Time to Work
def GetPlaceTravelTimeData():
    #5 Year ACS travel time range:   B08012_003E - B08012_013E
    fields_list = ["B08012_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,14)]

    travel_time_raw_data = c.acs5.state_place(fields=fields_list,state_fips=state_fips,place=place_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    travel_time_data = []
    for field in fields_list:
        travel_time_data.append(travel_time_raw_data[field])

    #Convert list with raw totals into a list where each element is a fraction of the total
    travel_time_data = ConvertListElementsToFractionOfTotal(travel_time_data)
    return(travel_time_data)

def GetCountyTravelTimeData():
    #5 Year ACS travel time range:   B08012_002E - B08012_013E
    fields_list = ["B08012_0" + ("0" *  (2 -len(str(i)))) + str(i) + "E" for i in range(2,14)]

    travel_time_raw_data = c.acs5.state_county(fields=fields_list,state_fips=state_fips,county_fips=county_fips)[0]
    
    #Create an empty list and place the values from the dictionary inside of it
    travel_time_data = []
    for field in fields_list:
        travel_time_data.append(travel_time_raw_data[field])

    #Convert list with raw totals into a list where each element is a fraction of the total
    travel_time_data = ConvertListElementsToFractionOfTotal(travel_time_data)
    return(travel_time_data)
   
#Travel Method to work
def GetPlaceTravelMethodData():
    neighborhood_method_to_work_total           = c.acs5.state_place(fields=['B08006_001E'],state_fips=state_fips,place=place_fips)[0]['B08006_001E']
    neighborhood_method_to_work_drove_alone     = c.acs5.state_place(fields=['B08006_003E'],state_fips=state_fips,place=place_fips)[0]['B08006_003E']
    neighborhood_method_to_work_carpooled       = c.acs5.state_place(fields=['B08006_004E'],state_fips=state_fips,place=place_fips)[0]['B08006_004E']
    neighborhood_method_to_work_walked          = c.acs5.state_place(fields=['B08006_015E'],state_fips=state_fips,place=place_fips)[0]['B08006_015E']
    neighborhood_method_to_work_publictrans     = c.acs5.state_place(fields=['B08006_008E'],state_fips=state_fips,place=place_fips)[0]['B08006_008E']
    neighborhood_method_to_work_workedhome      = c.acs5.state_place(fields=['B08006_017E'],state_fips=state_fips,place=place_fips)[0]['B08006_017E']
    neighborhood_method_to_work_other           = c.acs5.state_place(fields=['B08006_016E'],state_fips=state_fips,place=place_fips)[0]['B08006_016E']
    neighborhood_method_to_work_bike            = c.acs5.state_place(fields=['B08006_014E'],state_fips=state_fips,place=place_fips)[0]['B08006_014E']


    neighborhood_method_to_work_distribution   = [neighborhood_method_to_work_drove_alone,
                                                  neighborhood_method_to_work_carpooled,
                                                  neighborhood_method_to_work_publictrans,
                                                  neighborhood_method_to_work_walked,
                                                  neighborhood_method_to_work_workedhome,
                                                  neighborhood_method_to_work_bike,
                                                  neighborhood_method_to_work_other
                                                  ]

    neighborhood_method_to_work_distribution   = [i/neighborhood_method_to_work_total * 100 for i in neighborhood_method_to_work_distribution]  
    return(neighborhood_method_to_work_distribution) 

#Main data function
def GetData():
    #List of 5 Year American Community Survey Variables here: https://api.census.gov/data/2019/acs/acs5/variables.html
    #List of 2010 Census Variables here: https://api.census.gov/data/2010/dec/sf1/variables.html
    print('Getting Data')
    global overview_table_data
    global neighborhood_household_size_distribution,comparison_household_size_distribution
    global total_number_households, average_household_size
    global neighborhood_tenure_distribution, comparison_tenure_distribution
    global neighborhood_time_to_work_distribution, comparison_time_to_work_distribution
    global neighborhood_method_to_work_distribution
    global neighborhood_age_data,comparison_age_data
    global neighborhood_housing_value_data,comparison_housing_value_data
    global neighborhood_number_units_data,comparison_number_units_data
    global neighborhood_household_income_data, comparison_household_income_data
    global neighborhood_top_occupations_data,comparison_top_occupations_data
    global neighborhood_year_built_data, comparison_year_built_data   


    #Neighborhood households by size:
    neighborhood_household_size_distribution     = GetPlaceHouseholdSizeData()
    comparison_household_size_distribution       = GetCountyHouseholdSizeData()


    #Housing Tenure (owner occupied/renter)
    neighborhood_tenure_distribution = GetPlaceHousingTenureData()
    comparison_tenure_distribution   = GetCountyHousingTenureData()

    #Owner Occupied housing units by value
    neighborhood_housing_value_data = GetPlaceHousingValues()
    comparison_housing_value_data   = GetCountyHousingValues()


    #Housing Units by units in building
    neighborhood_number_units_data = GetPlaceNumberUnitsData()
    comparison_number_units_data   = GetCountyNumberUnitsData()


    #Housing Units by year structure built
    neighborhood_year_built_data = GetPlaceHouseYearBuiltData()
    comparison_year_built_data   = GetCountyHouseYearBuiltData()

    #Population by age data
    neighborhood_age_data = GetPlaceAgeData()
    comparison_age_data   = GetCountyAgeData()

    #Households by household income data
    neighborhood_household_income_data = GetPlaceHouseholdIncomeValues()
    comparison_household_income_data   = GetCountyHouseholdIncomeValues()

    #Top Employment Occupations
    neighborhood_top_occupations_data  = GetPlaceTopOccupationsData()
    comparison_top_occupations_data    = GetCountyTopOccupationsData()


    #Travel Time to Work
    neighborhood_time_to_work_distribution     =  GetPlaceTravelTimeData()
    comparison_time_to_work_distribution       =  GetCountyTravelTimeData()


    #Travel Mode to Work 
    neighborhood_method_to_work_distribution   = GetPlaceTravelMethodData()



    #Overview Table Data
    overview_table_data = [ ['','Area','2000 Census','2010 Census','Change','2021 Est.','Change','2026 Projected','Change'],
             ['Population',neighborhood,'','','','','','',''],
             ['',comparison_area,'','','','','','',''],
             ['Households',neighborhood,'','','','','','',''],
             ['',comparison_area,'','','','','','',''],
             ['Family Households',neighborhood,'','','','','','',''],
             ['',comparison_area,'','','','','','',''],
              ]



#Graph Related Functions
def SetGraphFormatVariables():
    global graph_width, graph_height, scale,tickfont_size,left_margin,right_margin,top_margin,bottom_margin,legend_position,paper_backgroundcolor,title_position

    #Set graph size and format variables
    marginInches = 1/18
    ppi = 96.85 
    width_inches = 6.5
    height_inches = 3.3

    graph_width  = (width_inches - marginInches)   * ppi
    graph_height = (height_inches  - marginInches) * ppi

    #Set scale for resolution 1 = no change, > 1 increases resolution. Very important for run time of main script. 
    scale = 3

    #Set tick font size (also controls legend font size)
    tickfont_size = 8 

    #Set Margin parameters/legend location
    left_margin   = 0
    right_margin  = 0
    top_margin    = 75
    bottom_margin = 10
    legend_position = 1.10

    #Paper color
    paper_backgroundcolor = 'white'

    #Title Position
    title_position = .95    

def CreateHouseholdSizeHistogram():
    print('Creating Household size graph')

    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(y=neighborhood_household_size_distribution,
           x=['1','2','3','4','5','6','7+'],
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )
    fig.add_trace(
    go.Bar(y=comparison_household_size_distribution,
           x=['1','2','3','4','5','6','7+'],
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Households by Household Size",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'household_size_graph.png'),engine='kaleido',scale=scale)

def CreateHouseholdTenureHistogram():
    print('Creating Household tenure graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(y=neighborhood_tenure_distribution,
           x=['Renter Occupied','Owner Occupied (Outright)','Owner Occupied (Mortgage)'],
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )
    fig.add_trace(
    go.Bar(y=comparison_tenure_distribution,
           x=['Renter Occupied','Owner Occupied (Outright)','Owner Occupied (Mortgage)'],
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Occupied Housing Units by Tenure",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'household_tenure_graph.png'),engine='kaleido',scale=scale)

def CreateHouseholdNumberUnitsInBuildingHistogram():
    print('Creating Household by number of units in structure graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    number_units_categories = ['1 Unit (Detached)','1 Unit (Attached)','2 Units','3-4 Units','5-9 Units','10-19 Units','20-49 Units','50 >= Units']

    #Add Bars with neighborhood distribution
    fig.add_trace(
    go.Bar(y=neighborhood_number_units_data,
           x=number_units_categories,
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )

    #Add Bars with comparison distribution
    fig.add_trace(
    go.Bar(y=comparison_number_units_data,
           x=number_units_categories,
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Housing Units by Units in Structure",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'household_units_in_structure_graph.png'),engine='kaleido',scale=scale)

def CreateHouseholdYearBuiltHistogram():
    print('Creating Household Year Built graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    year_built_categories = ['2014 >=','2010-2013','2000-2009','1990-1999','1980-1989','1970-1979','1960-1969','1950-1959','1940-1949','<= 1939']

    #Add Bars with neighborhood year built data
    fig.add_trace(
    go.Bar(y=neighborhood_year_built_data,
           x=year_built_categories,
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )

    #Add bars for comparison area
    fig.add_trace(
    go.Bar(y=comparison_year_built_data,
           x=year_built_categories,
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Housing Units by Year Structure Built",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'household_year_built_graph.png'),engine='kaleido',scale=scale)

def CreateHouseholdValueHistogram():
    print('Creating Household value graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    housing_value_categories = []

    #Add Bars with neighborhood house value distribution
    fig.add_trace(
    go.Bar(y=neighborhood_housing_value_data,
           x=housing_value_categories,
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )

    fig.add_trace(
    go.Bar(y=comparison_housing_value_data,
           x=housing_value_categories,
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Owner Occupied Housing Units by Value",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'household_value_graph.png'),engine='kaleido',scale=scale)

def CreatePopulationByAgeHistogram():
    print('Creating Population by Age Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    age_ranges = ['< 5','5-9','10-14','15-17','18-19','20','21','22-24','25-29','30-34','35-39','40-44','45-49','50-54','55-59','60-61','62-64','65-66','67-69','70-74','75-79','80-84','85+']
    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(y=neighborhood_age_data,
           x=age_ranges,
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )

    #Add bars with comparison area age distribution
    fig.add_trace(
    go.Bar(y=comparison_age_data,
           x=age_ranges,
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Population by Age",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )


    fig.update_xaxes(tickangle = 45)  
    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'population_by_age_graph.png'),engine='kaleido',scale=scale)

def CreatePopulationByIncomeHistogram():
    print('Creating Population by Income Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    income_categories = ['< $10,000',
                         '$10,000-14,999',
                         '$15,000-19,999',
                         '$20,000-24,999',
                         '$25,000-29,999',
                         '$30,000-34,999',
                         '$35,000-39,999',
                         '$40,000-44,999',
                         '$45,000-49,999',
                         '$50,000-59,999',
                         '$60,000-74,999',
                         '$75,000-99,999',
                         '$100,000-124,999',
                         '$125,000-149,999',
                         '$150,000-199,999',
                         '> $200,000']

    assert len(income_categories) == len(neighborhood_household_income_data) == len(comparison_household_income_data)
    
    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(y=neighborhood_household_income_data,
           x=income_categories,
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )

    #Add bars for comparison area        
    fig.add_trace(
    go.Bar(y=comparison_household_income_data,
           x=income_categories,
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )
    
    
    #Set Title
    fig.update_layout(
    title_text="Households by Household Income",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_xaxes(tickangle = 45)       
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'population_by_income_graph.png'),engine='kaleido',scale=scale)

def CreateTopOccupationsHistogram():
    print('Creating Top Occupations Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])
    
    occupations_categories = []

    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(y=neighborhood_top_occupations_data,
           x=occupations_categories,
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )

    
    
    #Set Title
    fig.update_layout(
    title_text="Top Employment Occupations",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'top_occupations_graph.png'),engine='kaleido',scale=scale)

def CreateTravelTimeHistogram():
    print('Creating Travel Time to work Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    #Add Bars with neighborhood household size distribution
    fig.add_trace(
    go.Bar(y=neighborhood_time_to_work_distribution,
           x=['< 5 Minutes','5-9 Minutes','10-14 Minutes','15-19 Minutes','20-24 Minutes','25-29 Minutes','30-34 Minutes','35-39 Minutes','40-44 Minutes','45-59 Minutes','60-89 Minutes','> 90 Minutes'],
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )
    fig.add_trace(
    go.Bar(y=comparison_time_to_work_distribution,
           x=['< 5 Minutes','5-9 Minutes','10-14 Minutes','15-19 Minutes','20-24 Minutes','25-29 Minutes','30-34 Minutes','35-39 Minutes','40-44 Minutes','45-59 Minutes','60-89 Minutes','> 90 Minutes'],
           name=comparison_area,
           marker_color="#B3C3FF")
            ,secondary_y=False
            )


    
    #Set Title
    fig.update_layout(
    title_text="Travel Time to Work",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'travel_time_graph.png'),engine='kaleido',scale=scale)

def CreateTravelModeHistogram():
    print('Creating Travel Mode to work Graph')
    fig = make_subplots(specs=[[{"secondary_y": False}]])

    
    fig.add_trace(
    go.Bar(y=neighborhood_method_to_work_distribution,
           x=['Drove Alone','Car Pooled','Public Transportation','Walked','Worked from Home','Biked','Other'],
           name=neighborhood,
           marker_color="#4160D3")
            ,secondary_y=False
            )
    
    #Set Title
    fig.update_layout(
    title_text="Travel Mode to Work",    

    title={
        'y':title_position,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    
    # #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = tickfont_size
                )

                      )
    
    fig.update_yaxes(title=None)

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor=paper_backgroundcolor,
    plot_bgcolor ="White"
                     )

    #Set size and margin
    fig.update_layout(
    margin=dict(l=left_margin, r=right_margin, t=top_margin, b= bottom_margin),
    height    = graph_height,
    width     = graph_width,
        
                    )



    #Add % to  axis ticks
    fig.update_yaxes(ticksuffix = '%', tickfont = dict(size=tickfont_size),tickformat='.1f',secondary_y=False)       
    fig.write_image(os.path.join(hood_folder,'travel_mode_graph.png'),engine='kaleido',scale=scale)

def CreateGraphs():
    print('Creating Graphs')
    CreateHouseholdSizeHistogram()
    CreateHouseholdTenureHistogram()
    CreateHouseholdValueHistogram()
    CreateHouseholdYearBuiltHistogram()
    CreateHouseholdNumberUnitsInBuildingHistogram()
    CreatePopulationByAgeHistogram()
    CreatePopulationByIncomeHistogram()
    CreateTopOccupationsHistogram()
    CreateTravelTimeHistogram()
    CreateTravelModeHistogram()

#Langauge Related Functions    
def CarLanguage():
    print('Writing Car Langauge')
    
    major_highways                = page.section('Major highways')
    major_Highways                = page.section('Major Highways')
    roadways                      = page.section('Roadways')
    highways                      = page.section('Highways')
    public_roadways               = page.section('Public roadways')
    major_roads                   = page.section('Major roads and highways')
    roads_and_highways            = page.section('Roads and highways')
    major_roads_and_Highways      = page.section('Major roads and Highways')
    car_language = ''
    for count,section in enumerate([major_highways,major_Highways,roadways,highways,public_roadways,major_roads,roads_and_highways,major_roads_and_Highways]):
        if (section != None) and (count == 0):
            car_language =  section 
        elif (section != None) and (count > 0):
            car_language = car_language + ' ' + "\n" + section 

    

    #If the wikipedia page is missiing all highway sections 
    if car_language == '':
        return(neighborhood + ' is not connected by any major highways or roads.')
    else:
        return(car_language)

def PlaneLanguage():
    print('Writing Plane Langauge')
    #Go though some common section names for airports
    airports              = page.section('Airports')
    air                   = page.section('Air')
    aviation              = page.section('Aviation')

    plane_language = ''
    for count,section in enumerate([airports,air,aviation]):
        if (section != None) and (count == 0):
            plane_language =  section 
        elif (section != None) and (count > 0):
            plane_language = plane_language + ' ' + "\n" + section 

    

    #If the wikipedia page is missiing all airport sections 
    if plane_language == '':
        return(neighborhood + ' is not served by any airport.')
    else:
        return(plane_language)
        
def BusLanguage():
    print('Writing Bus Langauge')
    bus                          =  page.section('Bus')
    intercity_bus                =  page.section('Intercity buses')
    public_Transportation        =  page.section('Public Transportation')
    
    #Add the text from the sections above to a single string variable
    bus_language = ''
    for count,section in enumerate([bus,intercity_bus,public_Transportation]):
        if (section != None) and (count == 0):
            bus_language =  section 
        elif (section != None) and (count > 0):
            bus_language = bus_language + ' ' + "\n" + section 

    
    #If the wikipedia page is missiing all airport sections return default phrase
    if bus_language == '':
        return(neighborhood + ' does not have public bus service.')
    else:
        return(bus_language)

def TrainLanguage():
    print('Writing Train Langauge')
    rail                         =  page.section('Rail')
    public_transportation        =  page.section('Public transportation')
    public_Transportation        =  page.section('Public Transportation')
    public_transport             =  page.section('Public transport')
    mass_transit                 =  page.section('Mass transit')
    rail_network                 =  page.section('Rail Network')

    #Add the text from the sections above to a single string variable
    train_language = ''
    for count,section in enumerate([rail,public_transportation,public_Transportation,public_transport,mass_transit,rail_network]):
        if (section != None) and (count == 0):
            train_language =  section 
        elif (section != None) and (count > 0):
            train_language = train_language + ' ' + "\n" + section 

    
    #If the wikipedia page is missiing all airport sections return default phrase
    if train_language == '':
        return(neighborhood + ' is not served by any commuter or light rail lines.')
    else:
        return(train_language)

def SummaryLangauge():
    return(wikipedia.summary((neighborhood + ',' + state)))

def OutlookLanguage():
    return('Neighborhood analysis can best be summarized by referring to neighborhood life cycles. ' +
          'Neighborhoods are perceived to go through four cycles, the first being growth, the second being stability, the third decline, and the fourth revitalization. ' +
          'It is our observation that the subject’s neighborhood is exhibiting several stages of the economic life, with an overall predominance of stability and both limited decline and limited revitalization in some sectors. ' +
          'The immediate area surrounding the subject, has had a historically low vacancy level and is located just to the south of the ------ submarket,' +
          """ which has multiple office and retail projects completed within the past two years and more development in the subject’s immediate vicinity either under construction or preparing to break ground."""+
          ' The proximity of the ________ and ________ will ensure the neighborhood will continue ' +
          'to attract growth in the long-term.')
    pass

def CreateLanguage():
    print('Creating Langauge')

    global bus_language,car_language,plane_language,train_language,transportation_language,summary_langauge,conclusion_langauge
    transportation_language         =  page.section('Transportation')
    bus_language   = BusLanguage()
    car_language   = CarLanguage()
    plane_language = PlaneLanguage()
    train_language = TrainLanguage()
    summary_langauge =  SummaryLangauge()
    conclusion_langauge = OutlookLanguage()
    pass


#Report document related functions
def SetPageMargins(document,margin_size):
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin   = Inches(margin_size)
        section.right_margin  = Inches(margin_size)

def SetDocumentStyle(document):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(9)

def AddTitle(document):
    main_title = document.add_heading('Neighborhood & Demographic Overview',level=0) 
    main_title.style = document.styles['Heading 1']
    main_title.paragraph_format.space_after  = Pt(6)
    main_title.paragraph_format.space_before = Pt(12)
    main_title_style = main_title.style
    main_title_style.font.name = "Avenir Next LT Pro Light"
    main_title_style.font.size = Pt(18)
    main_title_style.font.bold = False
    main_title_style.font.color.rgb = RGBColor.from_string('3F65AB')
    main_title_style.element.xml
    rFonts = main_title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

    title = document.add_heading(neighborhood + ' at a Glance',level=1)
    title.style = document.styles['Heading 2']
    title.paragraph_format.space_after  = Pt(6)
    title.paragraph_format.space_before = Pt(12)
    title_style = title.style
    title_style.font.name = "Avenir Next LT Pro Light"
    title_style.font.size = Pt(14)
    title_style.font.bold = False
    title_style.font.color.rgb = RGBColor.from_string('3F65AB')
    title_style.element.xml
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

def AddHeading(document,title,heading_level,heading_number,font_size): #Function we use to insert the headers other than the title header
            heading = document.add_heading(title,level=heading_level)
            heading.style = document.styles[heading_number]
            heading_style =  heading.style
            heading_style.font.name = "Avenir Next LT Pro Light"
            heading_style.font.size = Pt(font_size)
            heading_style.font.bold = False
            heading.paragraph_format.space_after  = Pt(6)
            heading.paragraph_format.space_before = Pt(12)

            #Color
            heading_style.font.color.rgb = RGBColor.from_string('3F65AB')            
            heading_style.element.xml
            rFonts = heading_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def Citation(document,text):
    citation_paragraph = document.add_paragraph()
    citation_paragraph.paragraph_format.space_after  = Pt(6)
    citation_paragraph.paragraph_format.space_before = Pt(6)
    run = citation_paragraph.add_run('Source: ' + text)
    font = run.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(8)
    font.italic = True
    font.color.rgb  = RGBColor.from_string('929292')
    citation_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text != 'Google Maps':
        blank_paragraph = document.add_paragraph('')
        blank_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def AddMap(document):
    print('Adding Map')
    #Add image of map
    if os.path.exists(os.path.join(hood_folder_map,'map.png')):
        map = document.add_picture(os.path.join(hood_folder_map,'map.png'),width=Inches(6.5))
    else:    
        try:
            #Search Google Maps for County
            options = webdriver.ChromeOptions()
            options.add_argument("--start-maximized")
            browser = webdriver.Chrome(executable_path=(os.path.join(os.environ['USERPROFILE'], 'Desktop','chromedriver.exe')),options=options)
            browser.get('https:google.com/maps')
            Place = browser.find_element_by_class_name("tactile-searchbox-input")
            Place.send_keys((neighborhood + ', ' + state))
            Submit = browser.find_element_by_xpath(
            "/html/body/jsl/div[3]/div[9]/div[3]/div[1]/div[1]/div[1]/div[2]/div[1]/button")
            Submit.click()
            time.sleep(5)
            zoomout = browser.find_element_by_xpath(
            """/html/body/jsl/div[3]/div[9]/div[22]/div[1]/div[2]/div[7]/div/button""")
            zoomout.click()
            time.sleep(7)

            if 'Leahy' in os.environ['USERPROFILE']: #differnet machines have different screen coordinates
                print('Using Mikes coordinates for screenshot')
                im2 = pyautogui.screenshot(region=(1089,276, 2405, 1754) ) #left, top, width, and height
            
            elif 'Dominic' in os.environ['USERPROFILE']:
                print('Using Doms coordinates for screenshot')
                im2 = pyautogui.screenshot(region=(3680,254,1968 ,1231) ) #left, top, width, and height
            
            else:
                im2 = pyautogui.screenshot(region=(1089,276, 2405, 1754) ) #left, top, width, and height

            time.sleep(.25)
            im2.save(os.path.join(hood_folder_map,'map.png'))
            im2.close()
            time.sleep(1)
            map = document.add_picture(os.path.join(hood_folder_map,'map.png'),width=Inches(6.5))
            browser.quit()
        except Exception as e:
            print(e)
            try:
                browser.quit()
            except:
                pass
    



       
def AddTable(document,data_for_table): #Function we use to insert our overview table into the report document
    #list of list where each list is a row for our table
     
    #make sure each list inside the list of lists has the same number of elements
    for row in data_for_table:
        for row2 in data_for_table:
            assert len(row) == len(row2)


    #create table object
    tab = document.add_table(rows=len(data_for_table), cols=len(data_for_table[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True
    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,data_for_table)): 

    
        row.height = Inches(0)

        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):
            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)

            # elif current_column == 1:
            #     cell.width = Inches(1.19)

            # elif current_column == 2:
            #     cell.width = Inches(0.8)



            #add border to top row
            if current_row == 1:
                    tcPr = cell._element.tcPr
                    tcBorders = OxmlElement("w:tcBorders")
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                    tcBorders.append(top)
                    tcPr.append(tcBorders)
            
            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(8)
                    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
                        
def IntroSection(document):
    AddTitle(document = document)
    # AddMap(document = document)
    Citation(document,'Google Maps')
    AddHeading(document = document, title = 'Summary',            heading_level = 2,heading_number='Heading 3',font_size=11)
    
    #Get summary section from wikipedia and add it 
    summary_paragraph           = document.add_paragraph(summary_langauge)
    summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # AddTable(document = document,data_for_table = overview_table_data )

def NeigborhoodSection(document):
    print('Writing Neighborhood Section')
    AddHeading(document = document, title = 'Neighborhood',            heading_level = 1,heading_number='Heading 2',font_size=14)
    AddHeading(document = document, title = 'Housing',                  heading_level = 2,heading_number='Heading 3',font_size=11)
    
    #Insert Household size graph
    if os.path.exists(os.path.join(hood_folder,'household_size_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_size_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Insert Household Tenure graph
    if os.path.exists(os.path.join(hood_folder,'household_tenure_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_tenure_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Insert Household value graph
    if os.path.exists(os.path.join(hood_folder,'household_value_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_value_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')

    #Insert household units by units in_structure graph
    if os.path.exists(os.path.join(hood_folder,'household_units_in_structure_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_units_in_structure_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
        
    #Insert household units by year built graph
    if os.path.exists(os.path.join(hood_folder,'household_year_built_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'household_year_built_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')

def DemographicsSection(document):
    print('Writing Neighborhood Section')
    AddHeading(document = document, title = 'Demographics',                                   heading_level = 1,heading_number='Heading 2',font_size=14)
    AddHeading(document = document, title = 'Population',                                     heading_level = 2,heading_number='Heading 3',font_size=11)
    
    #Insert population by age graph
    if os.path.exists(os.path.join(hood_folder,'population_by_age_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'population_by_age_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Insert populatin by income graph
    if os.path.exists(os.path.join(hood_folder,'population_by_income_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'population_by_income_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')







    #Employment and Transportation Subsection
    AddHeading(document = document, title = 'Employment and Transportation',                  heading_level = 2,heading_number='Heading 3',font_size=11)

    #Insert top occupations graph
    if os.path.exists(os.path.join(hood_folder,'top_occupations_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'top_occupations_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Insert Travel Time to Work graph
    if os.path.exists(os.path.join(hood_folder,'travel_time_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'travel_time_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Insert Transport Method to Work graph
    if os.path.exists(os.path.join(hood_folder,'travel_mode_graph.png')):
        fig = document.add_picture(os.path.join(hood_folder,'travel_mode_graph.png'),width=Inches(6.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        Citation(document,'U.S. Census Bureau')
    
    #Transportation Methods table
    table_paragraph = document.add_paragraph('Transportation Methods')
    table_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    transportation_paragraph = document.add_paragraph(transportation_language)

    #Insert the transit graphics(car, bus,plane, train)
    tab = document.add_table(rows=1, cols=2)
    for pic in ['car.png','train.png','bus.png','plane.png']:
        row_cells = tab.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        if pic == 'car.png':
            run.add_text(' ')
        run.add_picture(os.path.join(graphics_location,pic))
    


    transit_language = [car_language,train_language,bus_language,plane_language]
    # transit_language = ['car_language','train_language','bus_language','plane_language']

    #Loop through the rows in the table
    for current_row ,row in enumerate(tab.rows): 
        #loop through all cells in the current row
        for current_column,cell in enumerate(row.cells):
            if current_column == 1 and current_row > 0:
                cell.text = transit_language[current_row-1]

            if current_column == 0:
                cell.width = Inches(.2)
            else:
                cell.width = Inches(6)
 
def OutlookSection(document):
    print('Writing Outlook Section')
    AddHeading(document = document, title = 'Conclusion',            heading_level = 1,heading_number='Heading 2',font_size=14)
    conclusion_paragraph           = document.add_paragraph(conclusion_langauge)
    conclusion_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
def WriteReport():
    print('Writing Report')
    #Create Document
    document = Document()
    SetPageMargins(document   = document, margin_size=1)
    SetDocumentStyle(document = document)
    IntroSection(document = document)
    NeigborhoodSection(document     = document)
    DemographicsSection(document = document)
    OutlookSection(document = document)


    #Save report
    document.save(report_path)  

def CleanUpPNGs():
    print('Deleting PNG files')
    #Report writing done, delete figures
    files = os.listdir(hood_folder)
    for image in files:
        if image.endswith(".png"):
            os.remove(os.path.join(hood_folder, image))

def CreateDirectoryCSV():
    print('Creating CSV with file path information on all existing hood reports')
    dropbox_links                  = []
    dropbox_research_names         = []
    dropbox_analysis_types         = []
    dropbox_states                 = []
    dropbox_versions               = []
    dropbox_statuses               = []
    dropbox_document_names         = []


    for (dirpath, dirnames, filenames) in os.walk(main_output_location):
        if filenames == []:
            continue
        else:
            filenames = filenames[0]
        if filenames == 'Dropbox Neighborhoods.csv':
            continue


        dropbox_document_names.append(filenames)
        dropbox_analysis_types.append('Neighborhood')
        dropbox_link = dirpath.replace(dropbox_root,r'https://www.dropbox.com/home')
        dropbox_link = dropbox_link.replace("\\",r'/')    
        dropbox_links.append(dropbox_link)
        dropbox_versions.append(filenames[0:4])
        if '_draft' in filenames:
            file_status = 'Draft'
        else:
            file_status = 'Final'

        dropbox_statuses.append(file_status)

        research_name = filenames.split('-')[1].strip()
        state_name    = filenames[5:7]

        dropbox_research_names.append(research_name)
        dropbox_states.append(state_name)
        
        

    dropbox_df = pd.DataFrame({'Market Research Name':dropbox_research_names,
                           'Analysis Type': dropbox_analysis_types,
                           'State':         dropbox_states,
                           "Dropbox Links":dropbox_links,
                           'Version':dropbox_versions,
                           'Status':dropbox_statuses,
                           'Document Name': dropbox_document_names})
    dropbox_df = dropbox_df.sort_values(by=['State','Market Research Name'])

    dropbox_df.to_csv(os.path.join(main_output_location,'Dropbox Neighborhoods.csv'))

def Main():
    SetGraphFormatVariables()
    CreateDirectory()
    GetData()
    CreateGraphs()
    CreateLanguage()
    WriteReport()
    CleanUpPNGs()
    CreateDirectoryCSV()






#Decide if you want to export data in excel files in the county folder
data_export = False

# Get Input from User
# fips = input('Enter the 7 digit Census Place FIPS Code')
# fips = fips.replace('-','',1).strip()

fips = '1213275' #cocunut creek
# fips = '3643874' #lynbrook
state_fips = fips[0:2]
place_fips = fips[2:]
assert len(fips) == 7

# county_fips = input('Enter the 5 digit FIPS code for the comparison county')
# county_fips = county_fips.replace('-','',1).strip()
county_fips = '12011' #broward
# county_fips = '36059' #nassau
assert len(county_fips) == 5



#Get name of city
neighborhood = c.sf1.state_place(fields=['NAME'],state_fips=state_fips,place=place_fips)[0]['NAME']
state_full_name = neighborhood.split(',')[1].strip()
neighborhood = neighborhood.split(',')[0].strip().title()

#Name of State
state = us.states.lookup(state_full_name) #convert the full state name to the 2 letter abbreviation
state = state.abbr
assert len(state) == 2
comparison_area = c.sf1.state_county(fields=['NAME'],state_fips=county_fips[0:2],county_fips=county_fips[2:])[0]['NAME']
county_fips = county_fips [2:]
print('Preparing report for: ' + neighborhood)



todays_date = date.today()
current_year = str(todays_date.year)
page                          =  wikipedia.page((neighborhood + ',' + state))

Main()


