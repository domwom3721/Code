import os
from numpy import insert

import pandas as pd
from docx import Document
from docx.dml.color import ColorFormat
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Inches, Pt, RGBColor


def AddOverviewTable(document,number_rows,number_cols,row_data,col_width): #Function we use to insert our overview table into the report document

    #Make sure each row has the same number of items 
    for row in row_data:
        for row2 in row_data:
            assert len(row) == len(row2)


    #create table object
    tab = document.add_table(rows=len(row_data), cols=len(row_data[0]))
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    tab.allow_autofit = True
    #decide if we use standard style or custom

    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,row_data)): 
        assert (len(row_data_list) == number_cols) and (isinstance(row_data_list, list)) #make sure there is an item for each column in this row
        
        row.height = Inches(0.17)
        
        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):

            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)

            elif current_column == 1:
                cell.width = Inches(1.19)

            elif current_column == 2:
                cell.width = Inches(0.8)

            elif current_column == 3:
                cell.width = Inches(0.8)

            elif current_column == 4:
                cell.width = Inches(1.18)

            elif current_column == 5:
                cell.width = Inches(0.8)

            elif current_column == 6:
                cell.width = Inches(0.8)

            #add border to top row
            if current_row == 1:
                    tcPr = cell._element.tcPr

                    tcBorders = OxmlElement("w:tcBorders")
                
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                
                # left = OxmlElement('w:left')
                # left.set(qn('w:val'), 'nil')
                
                # bottom = OxmlElement('w:bottom')
                # bottom.set(qn('w:val'), 'nil')
                # bottom.set(qn('w:sz'), '4')
                # bottom.set(qn('w:space'), '0')
                # bottom.set(qn('w:color'), 'auto')

                # right = OxmlElement('w:right')
                # right.set(qn('w:val'), 'nil')

                    tcBorders.append(top)
                # tcBorders.append(left)
                # tcBorders.append(bottom)
                # tcBorders.append(right)
                    tcPr.append(tcBorders)

                     



                            
                
                

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(7)
                    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    #make first row bold
                    if current_row == 0: 
                        font = run.font
                        font.size= Pt(8)
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
    

def AddTable(document,row_data,col_width): #Function we use to insert our wide tables into report document 
    #Make sure each list in the list of lists have the same number of elements 
    for row in row_data:
        for row2 in row_data:
            assert len(row) == len(row2)
    number_rows = len(row_data)
    number_cols = len(row_data[0])


    assert number_rows == len(row_data) #make sure the data has a list of inputs for each row

    #create table object
    tab = document.add_table(rows=number_rows, cols=number_cols)
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER

    #loop through the rows in the table
    for current_row ,(row,row_data_list) in enumerate(zip(tab.rows,row_data)): 
        assert (len(row_data_list) == number_cols) and (isinstance(row_data_list, list)) #make sure there is an item for each column in this row

       
        row.height = Inches(0.17)
        
        #loop through all cells in the current row
        for current_column,(cell,cell_data) in enumerate(zip(row.cells,row_data_list)):
            cell.text = str(cell_data)

            if current_row == 0:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)
            else:
                cell.width = Inches(col_width)

            #add border to top row
            if current_row == 1:
                    tcPr = cell._element.tcPr

                    tcBorders = OxmlElement("w:tcBorders")
                
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')
                
                # left = OxmlElement('w:left')
                # left.set(qn('w:val'), 'nil')
                
                # bottom = OxmlElement('w:bottom')
                # bottom.set(qn('w:val'), 'nil')
                # bottom.set(qn('w:sz'), '4')
                # bottom.set(qn('w:space'), '0')
                # bottom.set(qn('w:color'), 'auto')

                # right = OxmlElement('w:right')
                # right.set(qn('w:val'), 'nil')

                    tcBorders.append(top)
                # tcBorders.append(left)
                # tcBorders.append(bottom)
                # tcBorders.append(right)
                    tcPr.append(tcBorders)


            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    font = run.font
                    if current_row == 0:
                        font.size= Pt(8)
                    else:
                        font.size = Pt(7)

                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'

def AddHeading(document,title,heading_level): #Function we use to insert the headers other than the title header
            heading = document.add_heading(title,level=heading_level)
            heading.style = document.styles['Heading 3']
            heading_style =  heading.style
            heading_style.font.name = "Avenir Next LT Pro"
            heading_style.font.size = Pt(11)
            heading_style.font.bold = False
            heading.paragraph_format.space_after  = Pt(6)
            heading.paragraph_format.space_before = Pt(12)

            #Color
            heading_style.font.color.rgb = RGBColor.from_string('3F65AB')            
            heading_style.element.xml
            rFonts = heading_style.element.rPr.rFonts
            rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def CreateRowDataForTable(data_frame,data_frame2,data_frame3,var1,var2,var3,modifier1,modifier2,modifier3,title): #Returns a list which will populate a row in the overview table
    if data_frame.equals(data_frame2):
        row_data_list = [title,
        data_frame[var1].iloc[-1],
        data_frame[var2].iloc[-1],
        data_frame[var3].iloc[-1],
        data_frame3[var1].iloc[-1],
        data_frame3[var2].iloc[-1],
        data_frame3[var3].iloc[-1]]


    else:
        row_data_list = [title,
        data_frame[var1].iloc[-1],
        data_frame[var2].iloc[-1],
        data_frame[var3].iloc[-1],
        data_frame2[var1].iloc[-1],
        data_frame2[var2].iloc[-1],
        data_frame2[var3].iloc[-1]]

    
    #Add modifiers to variables ($,%,bps)
    first_spot  = 1
    second_spot = 4
    #loop through the 3 modifiers passed into function, going to the places in the row where the modifiers are used ie: (1,4),(2,5),(3,6)
    for modifier in [modifier1,modifier2,modifier3]:
        if modifier == "$":

            if var1 == ('Total Sales Volume') or var1 == ('Asset Value/Unit') or var1 == ('Asset Value/Sqft') or var1 == ('Market Effective Rent/Unit'):
                row_data_list[first_spot] =  str("${:,.0f}".format(row_data_list[first_spot]))
                row_data_list[second_spot] = str("${:,.0f}".format(row_data_list[second_spot])) 
                
            else:
                row_data_list[first_spot] = modifier1 + '' + str(row_data_list[first_spot])
                row_data_list[second_spot] = modifier1 + '' + str(row_data_list[second_spot])

        elif modifier == '%':
            if var1 == ('Total Sales Volume') or  var1 == ('Sales Volume Transactions'): #only want whole number for percent change in these variables
                row_data_list[first_spot]  = str("{0:,.0f}".format(row_data_list[first_spot]))   + modifier
                row_data_list[second_spot] = str("{0:,.0f}".format(row_data_list[second_spot]))  + modifier
            else:
                row_data_list[first_spot]  = str(row_data_list[first_spot]) + modifier
                row_data_list[second_spot] = str(row_data_list[second_spot]) + modifier

        elif modifier == 'bps':
            row_data_list[first_spot]  = str("{:,.0f}".format(row_data_list[first_spot]))  + ' ' + modifier
            row_data_list[second_spot] = str("{:,.0f}".format(row_data_list[second_spot])) + ' ' + modifier

        else:
            if var1 == 'Sales Volume Transactions':
                row_data_list[first_spot] =  str("{:,.0f}".format(row_data_list[first_spot]))
                row_data_list[second_spot] = str("{:,.0f}".format(row_data_list[second_spot]))
            
            elif var1 == 'Absorption Units' or var1 == 'Net Absorption SF':
                row_data_list[first_spot] =  str("{:,.0f}".format(row_data_list[first_spot]))
                row_data_list[second_spot] = str("{:,.0f}".format(row_data_list[second_spot])) 

            else:
                row_data_list[first_spot]  = str(row_data_list[first_spot])  + ' ' + modifier
                row_data_list[second_spot] = str(row_data_list[second_spot]) + ' ' + modifier
            
        
        first_spot  +=1
        second_spot +=1
        
    for count,cell in enumerate(row_data_list):
        if cell == 'nan' or cell == 'inf' or cell == 'nan%' or cell == 'inf%' or cell =='nan bps' or cell =='inf bps':
            row_data_list[count] = 'NA'
    return(row_data_list)

def CreateRowDataForWideTable(data_frame,data_frame2,data_frame3,data_frame4,var1,modifier,sector): #Returns list of lists with data we use to fill rows in the wide table
    #This function takes a variable and returns a list of lists of that variables value over time in the market, submarket, and nation, 
    #and if we are doing a market the different quality slices. Each list in the list represents a row in a table for either rent or vacancy
    if data_frame.equals(data_frame2):
        market_or_submarket = 'market'
    else:
        market_or_submarket = 'submarket'


    #Sort the 4 dataframes
    data_frame  = data_frame.sort_values(by=['Period'],ascending = False) 
    data_frame2 = data_frame2.sort_values(by=['Period'],ascending = False) 
    data_frame3 = data_frame3.sort_values(by=['Period'],ascending = False) 
    data_frame4 = data_frame4.sort_values(by=['Slice','Period'],ascending = False) 

    if modifier == '$' and sector == 'Multifamily':
        data_frame[var1]  = data_frame[var1].map('${:,.0f}'.format)
        data_frame2[var1] = data_frame2[var1].map('${:,.0f}'.format)
        data_frame3[var1] = data_frame3[var1].map('${:,.0f}'.format)
        data_frame4[var1] = data_frame4[var1].map('${:,.0f}'.format)

    elif modifier == '$' and sector != 'Multifamily':
        data_frame[var1]  = data_frame[var1].map('${:,.2f}'.format)
        data_frame2[var1] = data_frame2[var1].map('${:,.2f}'.format)
        data_frame3[var1] = data_frame3[var1].map('${:,.2f}'.format)
        data_frame4[var1] = data_frame4[var1].map('${:,.2f}'.format)
   
    else:
        data_frame[var1]  = data_frame[var1].map('{:,.1f}%'.format)
        data_frame2[var1] = data_frame2[var1].map('{:,.1f}%'.format)
        data_frame3[var1] = data_frame3[var1].map('{:,.1f}%'.format)
        data_frame4[var1] = data_frame4[var1].map('{:,.1f}%'.format)

    

    #Keep all rows from the current year and the last quarter from all previous years
    data_frame_last2_quarters             = data_frame.head(2) #dataframe with 2 most recent quarters
    data_frame2_last2_quarters            = data_frame2.head(2) #dataframe with 2 most recent quarters
    data_frame3_last2_quarters            = data_frame3.head(2) #dataframe with 2 most recent quarters
    data_frame4_last2_quarters            = data_frame4.groupby(['Slice']).head(2) #dataframe with 2 most recent quarters

    #Keep all remaining rows in the dataframe
    data_frame                           = data_frame[2:]     
    data_frame                           = data_frame[data_frame['Period'].str.contains("Q4")]
    data_frame['Period']                 = data_frame['Period'].str.replace(r' Q4', '')

    data_frame2                           = data_frame2[2:]     
    data_frame2                           = data_frame2[data_frame2['Period'].str.contains("Q4")]
    data_frame2['Period']                 = data_frame2['Period'].str.replace(r' Q4', '')

    data_frame3                           = data_frame3[2:]     
    data_frame3                           = data_frame3[data_frame3['Period'].str.contains("Q4")]
    data_frame3['Period']                 = data_frame3['Period'].str.replace(r' Q4', '')

    tail_length                           =  int(len(data_frame4))/len(data_frame4['Slice'].unique()) - 2
    data_frame4                           = data_frame4.groupby(['Slice']).tail(tail_length)
    data_frame4                           = data_frame4[data_frame4['Period'].str.contains("Q4")]
    data_frame4['Period']                 = data_frame4['Period'].str.replace(r' Q4', '')
   

    #Append the 2 most recent quarters with the Q4 dataframe
    data_frame = data_frame_last2_quarters.append(data_frame)
    data_frame = data_frame.reset_index()

    data_frame2 = data_frame2_last2_quarters.append(data_frame2)
    data_frame2 = data_frame2.reset_index()

    data_frame3 = data_frame3_last2_quarters.append(data_frame3)
    data_frame3 = data_frame3.reset_index()

    data_frame4 = data_frame4_last2_quarters.append(data_frame4)
    data_frame4 = data_frame4.reset_index()

    #Cut down to the variables we are going to display in the table
    data_frame = data_frame[[var1,'Period']]
    data_frame2 = data_frame2[[var1,'Period']]
    data_frame3 = data_frame3[[var1,'Period']]
    data_frame4 = data_frame4[[var1,'Period','Slice']]

    #Sort again so that most recent quarter is last
    data_frame  = data_frame.sort_values(by=['Period'],ascending = True) 
    data_frame2 = data_frame2.sort_values(by=['Period'],ascending = True) 
    data_frame3 = data_frame3.sort_values(by=['Period'],ascending = True) 
    data_frame4 = data_frame4.sort_values(by=['Slice','Period'],ascending = True) 

    #Add empty rows to the top of the dataframes
    data = []
    data.insert(0, {'Period': '', var1: 'Submarket'})
    data_frame  = pd.concat([pd.DataFrame(data), data_frame], ignore_index=True)
    
    data = []
    data.insert(0, {'Period': '', var1: 'Market'})
    data_frame2  = pd.concat([pd.DataFrame(data), data_frame2], ignore_index=True)

    data = []
    data.insert(0, {'Period': '', var1: 'National'})
    data_frame3  = pd.concat([pd.DataFrame(data), data_frame3], ignore_index=True)

 
    
    #If we are doing a market
    if market_or_submarket == 'market':

        try:
            list_of_lists = [data_frame['Period'].tolist(),
                data_frame3[var1].tolist(),
                data_frame2[var1].tolist()]
            
            #Since it's a market, we are going to add a row for each slice we have in our slice data
            for slice in data_frame4['Slice'].unique():
                data_frame_temp = data_frame4.loc[(data_frame4['Slice'] == slice)]
                data_frame_temp = data_frame_temp.reset_index()
                slice_list_to_add = data_frame_temp[var1].tolist()
                slice_list_to_add.insert(0,slice)
                list_of_lists.append(slice_list_to_add)

            #make sure each list (row) in the list of lists (rows) have same number of items
            for list in list_of_lists:
                for list2 in list_of_lists:
                    assert len(list) == len(list2)
    
        #When we don't have full data for the slices
        except:
             list_of_lists = [data_frame['Period'].tolist(),
                data_frame3[var1].tolist(),
                data_frame2[var1].tolist()]


        return(list_of_lists)

    #If we are doing a submarket
    else:
        try:
            list_of_lists = [data_frame['Period'].tolist(),
                    data_frame3[var1].tolist(),
                    data_frame2[var1].tolist(),
                    data_frame[var1].tolist()]
                        #make sure each list (row) in the list of lists (rows) have same number of items
            for list in list_of_lists:
                for list2 in list_of_lists:
                    assert len(list) == len(list2)

            return(list_of_lists)
        

        except:
            #In some cases, the submarket does not have a full 10 years of data. To handle this, we make sure we display data for quarters  we
            # coverage and make the market and nation rows blank
            submarket_period_list = data_frame['Period'].tolist()
            submarket_list        =  data_frame[var1].tolist()
            
            #Create empty rows for market and nation
            market_list           =  ['Market']
            nation_list           =  ['National']
            for i in range(len(submarket_list) -1):
                market_list.append('') 
                nation_list.append('')

            #Make lists of the data and ensure they are all the same size
            list_of_lists = [submarket_period_list,nation_list,market_list,submarket_list]
            for list in list_of_lists:
                for list2 in list_of_lists:
                    assert len(list) == len(list2)

            return(list_of_lists)

        

        
















     # row_data = [ [],[],[] ] # list of lists we will fill with data for our table


        # if current_quarter == 'Q1': #market Q1
        #     #First row
        #     row_data[0].append('')

        #     for year in range(most_recent_year_less_10 ,most_recent_year):           
        #         row_data[0].append(str(year))
                
        #     row_data[0].append(most_recent_period)
        #     assert len(row_data[0]) == 12


            
        #     def fill_list(first_string,df,list_number): #Fills list with time series of variable values
        #         row_data[list_number].append(first_string)
        #         for year in range(most_recent_year_less_10 ,most_recent_year,):
        #             df_temp    = df.loc[(df['Year'] == year) & (df['Quarter']=='Q4')]
        #             try:
        #                 current_year_var_value = df_temp[var1].iloc[0] #Value of the variable in the fourth quarter of the year
        #                 if modifier == '%':
        #                     row_data[list_number].append(str(current_year_var_value)+modifier)
        #                 elif  modifier == '$':
        #                     if var1 == 'Market Effective Rent/Unit':
        #                         row_data[list_number].append( str("${:,.0f}".format(current_year_var_value)))
        #                     else:
        #                         row_data[list_number].append( str("${:,.2f}".format(current_year_var_value)))                            
        #             except:
        #                 row_data[list_number].append(str(''))

        #         #Now that we have filled in the list for the previous years we fill in with the most recent value of the variable    
                
        #         if modifier == '%':
        #             row_data[list_number].append( str(df[var1].iloc[-1]) + modifier     )
        #         elif modifier == '$':
        #             if var1 == 'Market Effective Rent/Unit':
        #                 row_data[list_number].append(   str("${:,.0f}".format(df[var1].iloc[-1]))     )
        #             else:                    
        #                 row_data[list_number].append(   str("${:,.2f}".format(df[var1].iloc[-1]))     )
        #         else:
        #             row_data[list_number].append('')

        #         assert len(row_data[list_number]) == 12 
                
                
                
        #     #Second Row (National)
        #     fill_list('National',data_frame3,1)    

            

        #     #Third row
        #     fill_list('Market',data_frame2,2)
         
            



    

        # else:#market not Q1
        #     #First row
        #     row_data[0].append('')

        #     for year in range(most_recent_year_less_10 ,most_recent_year):           
        #         row_data[0].append(str(year))

        #     row_data[0].append(second_most_recent_period)    
        #     row_data[0].append(most_recent_period)
        #     assert len(row_data[0]) == 13


            
        #     def fill_list(first_string,df,list_number): #Fills list with time series of variable values
        #         row_data[list_number].append(first_string)
        #         for year in range(most_recent_year_less_10 ,most_recent_year,):
        #             df_temp    = df.loc[(df['Year'] == year) & (df['Quarter']=='Q4')]
        #             try:
        #                 current_year_var_value = df_temp[var1].iloc[0] #Value of the variable in the fourth quarter of the year
        #                 if modifier == '%':
        #                     row_data[list_number].append(str(current_year_var_value)+modifier)
        #                 elif  modifier == '$':
        #                     if var1 == 'Market Effective Rent/Unit':
        #                         row_data[list_number].append( str("${:,.0f}".format(current_year_var_value)))
        #                     else:
        #                         row_data[list_number].append( str("${:,.2f}".format(current_year_var_value)))                            
        #             except:
        #                 row_data[list_number].append(str(''))

        #         #Now that we have filled in the list for the previous years we fill in with the most recent value of the variable    
                
        #         if modifier == '%':
        #              row_data[list_number].append( str(df[var1].iloc[-2]) + modifier     )
        #              row_data[list_number].append( str(df[var1].iloc[-1]) + modifier     )
        #         elif modifier == '$':
        #             if var1 == 'Market Effective Rent/Unit':
        #                 row_data[list_number].append(   str("${:,.0f}".format(df[var1].iloc[-2]))     )
        #                 row_data[list_number].append(   str("${:,.0f}".format(df[var1].iloc[-1]))     )
        #             else:
        #                 row_data[list_number].append(   str("${:,.2f}".format(df[var1].iloc[-2]))     )
        #                 row_data[list_number].append(   str("${:,.2f}".format(df[var1].iloc[-1]))     )
        #         else:
        #             row_data[list_number].append('')
        #             row_data[list_number].append('')

        #         assert len(row_data[list_number]) == 13 
                
                
                
        #     #Second Row (National)
        #     fill_list('National',data_frame3,1)    

            

        #     #Third row
        #     fill_list('Market',data_frame2,2)

        

        # return(row_data)






    #Submarkets
    # else:
        
        # row_data = [ [],[],[],[] ]
        
        # if current_quarter == 'Q1': #Submarket, Q1
        #     #First row
        #     row_data[0].append('')

        #     for year in range(most_recent_year_less_10 ,most_recent_year):
        #         row_data[0].append(str(year))
                
        #     row_data[0].append(most_recent_period)
        #     assert len(row_data[0]) == 12


            
        #     def fill_list(first_string,df,list_number):
        #         row_data[list_number].append(first_string)
        #         for year in range(most_recent_year_less_10 ,most_recent_year,):
        #             df_temp    = df.loc[(df['Year'] == year) & (df['Quarter']=='Q4')]
        #             try:
        #                 current_year_var_value = df_temp[var1].iloc[0]
        #                 if modifier == '%':
        #                     row_data[list_number].append(str(current_year_var_value)+modifier)
        #                 elif  modifier == '$':
        #                     if var1 == 'Market Effective Rent/Unit':
        #                         row_data[list_number].append( str("${:,.0f}".format(current_year_var_value))  )
        #                     else:
        #                         row_data[list_number].append( str("${:,.2f}".format(current_year_var_value))  )
                             
                        
        #             except:
        #                 row_data[list_number].append(str(''))
                    
        #         #row_data[list_number].append('')
        #         if modifier == '%':
        #             row_data[list_number].append( str(df[var1].iloc[-1]) + modifier     )
        #         elif modifier == '$':
        #             if var1 == 'Market Effective Rent/Unit':
        #                 row_data[list_number].append(   str("${:,.0f}".format(df[var1].iloc[-1]))     )
        #             else:
        #                 row_data[list_number].append(   str("${:,.2f}".format(df[var1].iloc[-1]))     )
        #         assert len(row_data[list_number]) == 12 
                
                
                
        #     #Second Row (National)
        #     fill_list('National',data_frame3,1)    

            

        #     #Third row
        #     fill_list('Market',data_frame2,2)

 

        #     #Fourth Row
        #     fill_list('Submarket',data_frame,3)

 
        
        # else: #Submarket, not Q1
        #     row_data[0].append('')

        #     for year in range(most_recent_year_less_10 ,most_recent_year):
        #         row_data[0].append(str(year))
        #     row_data[0].append(second_most_recent_period)    
        #     row_data[0].append(most_recent_period)
        #     assert len(row_data[0]) == 13


            
        #     def fill_list(first_string,df,list_number):
        #         row_data[list_number].append(first_string)
        #         for year in range(most_recent_year_less_10 ,most_recent_year,):
        #             df_temp    = df.loc[(df['Year'] == year) & (df['Quarter']=='Q4')]
        #             try:
        #                 current_year_var_value = df_temp[var1].iloc[0]
        #                 if modifier == '%':
        #                     row_data[list_number].append(str(current_year_var_value)+modifier)
        #                 elif  modifier == '$':
        #                     if var1 == 'Market Effective Rent/Unit':
        #                         row_data[list_number].append( str("${:,.0f}".format(current_year_var_value))  )
        #                     else:
        #                         row_data[list_number].append( str("${:,.2f}".format(current_year_var_value))  )
                             
                        
        #             except:
        #                 row_data[list_number].append(str(''))
                    
                
        #         if modifier == '%':
        #             row_data[list_number].append( str(df[var1].iloc[-2]) + modifier     )
        #             row_data[list_number].append( str(df[var1].iloc[-1]) + modifier     )
        #         elif modifier == '$':
        #              if var1 == 'Market Effective Rent/Unit':
        #                  row_data[list_number].append(   str("${:,.0f}".format(df[var1].iloc[-2]))     )
        #                  row_data[list_number].append(   str("${:,.0f}".format(df[var1].iloc[-1]))     )
        #              else:
        #                 row_data[list_number].append(   str("${:,.2f}".format(df[var1].iloc[-2]))     )
        #                 row_data[list_number].append(   str("${:,.2f}".format(df[var1].iloc[-1]))     )
        #         assert len(row_data[list_number]) == 13
                
                
                
        #     #Second Row (National)
        #     fill_list('National',data_frame3,1)    

            

        #     #Third row
        #     fill_list('Market',data_frame2,2)

 

        #     #Fourth Row
        #     fill_list('Submarket',data_frame,3)

        # return(row_data)


def AddMarketPerformanceTable(document,col_width,market_data_frame,sector): #Function we use to insert our wide tables into report document 

    #Convert market dataframe into a (mostly) annual dataset for a handful of variables
    
    #Start by declaring a list of varibles we want to display
    if sector == 'Multifamily':
        variables_of_interest = ['Period','Inventory Units','Absorption Units 12 Mo','Net Delivered Units 12 Mo','Vacancy Rate','Under Construction Units','Market Effective Rent/Unit']
    else:
        variables_of_interest = ['Period','Inventory SF','Net Absorption SF 12 Mo','Net Delivered SF 12 Mo','Vacancy Rate','Availability Rate','Under Construction SF','Market Rent/SF']

   
    
    #Now create an annual dataset where we keep the last period of each year besides the current year
    market_data_frame = market_data_frame.sort_values(by=['Year','Quarter'],ascending = False) 

    #Keep all rows from the current year and the last quarter from all previous years
    market_data_frame_last2_quarters     = market_data_frame.head(2) #dataframe with 2 most recent quarters
    market_data_frame                    = market_data_frame[2:]     #dataframe with remaining quarters
    market_data_frame                    = market_data_frame[market_data_frame['Period'].str.contains("Q4")]
    market_data_frame['Period']          = market_data_frame['Period'].str.replace(r' Q4', '')
    
    #Append the 2 most recent quarters with the Q4 dataframe
    market_data_frame = market_data_frame_last2_quarters.append(market_data_frame)
    market_data_frame = market_data_frame.reset_index()

    #Cut down to the variables we are going to display in the table
    market_data_frame = market_data_frame[variables_of_interest]

    #create table object
    number_rows = len(market_data_frame) + 1 #we add extra row for variable names at the top
    number_cols = len(market_data_frame.columns)

    tab = document.add_table(rows=number_rows, cols=number_cols)
    tab.alignment     = WD_TABLE_ALIGNMENT.CENTER
    for current_row,row in enumerate(tab.rows): 
        for current_column,cell in enumerate(row.cells):
            
            if current_row == 0:
                var_name = str(variables_of_interest[current_column])
                cell.text = var_name
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                
                
                

            else:
                current_variable = variables_of_interest[current_column]
                data = market_data_frame[current_variable].iloc[current_row-1] #look up the data value for the right period for the current variable
                if type(data) == str:
                    pass
                else:
                    if current_variable == 'Vacancy Rate' or current_variable == 'Availability Rate':
                        data = "{:,.1f}%".format(data) 

                    elif current_variable == 'Market Effective Rent/Unit':
                        data = "${:,.0f}".format(data)
                    
                    elif current_variable == 'Market Rent/SF':
                        data = "${:,.2f}".format(data)
                    
                    elif (current_variable == 'Inventory SF') or (current_variable == 'Inventory Units'):
                        data = "{:,.0f}".format(data)
                    else:
                        data = "{:,.0f}".format(data)

                cell.text = data
                

            

            #set column widths
            if current_column == 0:
                cell.width = Inches(1.25)
            else:
                cell.width = Inches(col_width)

            
            #add border to top row
            if current_row == 1:
                tcPr = cell._element.tcPr
                tcBorders = OxmlElement("w:tcBorders")
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
            
                tcBorders.append(top)
                # tcBorders.append(left)
                # tcBorders.append(bottom)
                # tcBorders.append(right)
                tcPr.append(tcBorders)

            #loop through the paragraphs in the cell and set font and style
            for paragraph in cell.paragraphs:
                if current_column > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    font = run.font
                    if current_row == 0:
                        font.size= Pt(8)
                    else:
                        font.size = Pt(7)

                    #make first row bold
                    if current_row == 0: 
                        font.bold = True
                        font.name = 'Avenir Next LT Pro Demi'
