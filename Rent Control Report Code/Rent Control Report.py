#By Mike Leahy, April 4 2022
    #Summary: Loops through an excel file with information on rent control laws in different cities and towns across the United States
    #         For each row (town), we create a word document and write sections for each column where there is text stored

import os
from numpy import NaN
import pandas as pd

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from datetime import date, datetime

#Define file pre-paths
dropbox_root                   =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
project_location               =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)','Research','Projects','Research Report Automation Project') 
main_output_location           =  os.path.join(dropbox_root,'Research', 'Resources', 'Appraisers', 'Rent Control') 
data_location                  =  os.path.join(dropbox_root,'Research', 'Resources', 'Appraisers', 'Rent Control')


def CreateDirectory(state, muncipality_name, document_name):
    global report_path
    state_folder             = os.path.join(main_output_location, state)
    town_folder              = os.path.join(main_output_location, state, muncipality_name)
    

    for folder in [state_folder, town_folder]:
        if os.path.exists(folder) == False:
            os.mkdir(folder) 

    report_path              = os.path.join(town_folder, document_name + '.docx')

def SetPageMargins(document, margin_size):
    sections = document.sections
    for section in sections:
        section.top_margin    = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin   = Inches(margin_size)
        section.right_margin  = Inches(margin_size)

def SetDocumentStyle(document):
    style     = document.styles['Normal']
    font      = style.font
    font.name = 'Avenir Next LT Pro (Body)'
    font.size = Pt(9)

def AddTitle(document):
    title                               = document.add_heading(document_title,level=1)
    title.style                         = document.styles['Heading 2']
    title.paragraph_format.space_after  = Pt(6)
    title.paragraph_format.space_before = Pt(12)
    title_style                         = title.style
    title_style.font.name               = "Avenir Next LT Pro Light"
    title_style.font.size               = Pt(14)
    title_style.font.bold               = False
    title_style.font.color.rgb          = RGBColor.from_string('3F65AB')
    title_style.element.xml
    rFonts                              = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro Light")

def AddHeading(document, title, heading_level): 
    #Function we use to insert the headers other than the title header
    heading                               = document.add_heading(title, level = heading_level)
    heading.style                         = document.styles['Heading 3']
    heading_style                         = heading.style
    heading_style.font.name               = "Avenir Next LT Pro"
    heading_style.font.size               = Pt(11)
    heading_style.font.bold               = False
    heading.paragraph_format.space_after  = Pt(6)
    heading.paragraph_format.space_before = Pt(12)

    #Color
    heading_style.font.color.rgb          = RGBColor.from_string('3F65AB')            
    heading_style.element.xml
    rFonts                                = heading_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Avenir Next LT Pro")

def AddDocumentParagraph(document, language_variable):
    assert type(language_variable) == list

    for paragraph in language_variable:
        
        #Skip blank paragraphs
        if paragraph == '':
            continue
        
        par                                               = document.add_paragraph(str(paragraph))
        par.alignment                                     = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_after                  = Pt(primary_space_after_paragraph)
        summary_format                                    = document.styles['Normal'].paragraph_format
        summary_format.line_spacing_rule                  = WD_LINE_SPACING.SINGLE
        style                                             = document.styles['Normal']
        font                                              = style.font
        font.name                                         = primary_font
        par.style                                         = document.styles['Normal']

def AddDocumentPicture(document, image_path, citation):
    if os.path.exists(image_path):
        fig                                         = document.add_picture(os.path.join(image_path),width=Inches(6.5))
        last_paragraph                              = document.paragraphs[-1] 
        last_paragraph.paragraph_format.space_after = Pt(0)
        last_paragraph.alignment                    = WD_ALIGN_PARAGRAPH.CENTER

def AddTableTitle(document, title):
    table_title_paragraph                               = document.add_paragraph(title)
    table_title_paragraph.alignment                     = WD_ALIGN_PARAGRAPH.CENTER
    table_title_paragraph.paragraph_format.space_after  = Pt(6)
    table_title_paragraph.paragraph_format.space_before = Pt(12)

    for run in table_title_paragraph.runs:
                    font      = run.font
                    font.name = 'Avenir Next LT Pro Medium'

def OutlookSection(document):
    print('Writing Outlook Section')
    AddHeading(          document = document, title = 'Conclusion',            heading_level = 2)

#Set formatting paramaters for reports
primary_font                  = 'Avenir Next LT Pro Light' 
primary_space_after_paragraph = 8
bowery_grey                   = "#D7DEEA"
bowery_dark_grey              = "#A6B0BF"
bowery_dark_blue              = "#4160D3"
bowery_light_blue             = "#B3C3FF"
bowery_black                  = "#404858"

todays_date                   = date.today()
current_year                  = todays_date.year



#Read in our excel file where we store text on rent control regulations
rent_control_df                    = pd.read_excel(os.path.join(data_location,'Rent Control Template.xlsx'))
rent_control_df                    = rent_control_df.fillna('')
rent_control_df['Date of Update:'] = rent_control_df['Date of Update:'].dt.strftime('%m/%d/%Y')

section_list = ['Overview',
                'Governance',
                'Applicable To',
                'Exemptions',
                'Permitted Annual Increases',
                'Rent Control',
                'Rent Stabilization',
                'Vacancy Decontrol',
                'Vacancy Bonus',
                'Capital Improvements',
                'Individual Apartment Improvements',
                'Preferential Rent',
                'Co-op & Condo Conversions',
                '421-A',
                'Warehousing',
                'Hardship Rental Increases',
                'Rent Roll Filing',
                'Increase in Services Charge',
                'Real Property Tax Credits & Rebates',
                'Parking Fees',
                'Utilities Increases',
                'CPI',
                'Rental of Vacant Units',
                'Suggested Paragraph for Appraisal Report',
                ]

#Loop through each row of the dataframe, create a report and directory for each
for i in range(len(rent_control_df)):
    state            = rent_control_df['State'].iloc[i]
    muncipality_name = rent_control_df['Municipality'].iloc[i]
    document_name    = rent_control_df['Document Title'].iloc[i]
    document_name    = str(current_year) + ' ' + state + ' - ' +  muncipality_name + ' - ' +  'Rent Control_draft'
    document_title   = rent_control_df['Document Title'].iloc[i]
    phone_number     = rent_control_df['Phone Number of Rent Leveling Board:'].iloc[i]
    website          = rent_control_df['Source/Website:'].iloc[i]
    date_of_update   = rent_control_df['Date of Update:'].iloc[i]
    date_of_update   = str(date_of_update).replace('nan','')

    print('Creating Report for: ', muncipality_name + ', ' + state)
    CreateDirectory(state = state, muncipality_name = muncipality_name, document_name = document_name)
    
    #Create Document
    document = Document()
    SetPageMargins(       document = document, margin_size = 1)
    SetDocumentStyle(     document = document)
    AddTitle(             document = document)

    #Loop through each column in the excel file and (if there is text stored) make a section for that column in the document and write the stored text to it




    for section in section_list:

        language = rent_control_df[section].iloc[i]
        if language != ('nan' and '' and  NaN) or section == 'Suggested Paragraph for Appraisal Report':
            AddHeading(document = document, title = section, heading_level = 2)
            AddDocumentParagraph(document = document, language_variable = [language])
            
            if section == 'Overview':
                AddDocumentParagraph(document = document, language_variable = [('Phone Number of Rent Leveling Board: ' + phone_number), 
                                                                               ('Website: '                             + website), 
                                                                               ( 'Date of Update: '                      + str(date_of_update)),
                                                                               ])

    
    

    #Save report
    document.save(report_path)  
    print('Report Complete')