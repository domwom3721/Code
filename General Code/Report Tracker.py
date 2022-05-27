#Date: 02/02/2022
#Author: Mike Leahy
#Summary: Uses our 3 report csv summary files and produces a summary of jobs produced by each member of the research team

import os
from sre_parse import State
import pandas as pd
from plotly.subplots import make_subplots
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt


#Define file paths
dropbox_root                        =  os.path.join(os.environ['USERPROFILE'], 'Dropbox (Bowery)') 
main_output_location                =  os.path.join(dropbox_root,'Research','Market Analysis') 
kpi_location                        =  os.path.join(dropbox_root,'Research','Projects','KPI') 

#Read in our 3 csv reports as dataframes
area_report_df                      = pd.read_csv(os.path.join(main_output_location,'Area','Dropbox Areas.csv'),        encoding='latin-1')
market_report_df                    = pd.read_csv(os.path.join(main_output_location,'Market','CoStar Markets.csv'),     encoding='latin-1')
hood_report_df                      = pd.read_csv(os.path.join(main_output_location,'Neighborhood','Dropbox Neighborhoods.csv'),encoding='latin-1')

#Restrict to Final Reports
area_report_df                       = area_report_df.loc[area_report_df['Status'] == 'Final']
market_report_df                     = market_report_df.loc[market_report_df['Status'] == 'Final']
hood_report_df                       = hood_report_df.loc[hood_report_df['Status'] == 'Final']


#Collapse down each dataframe to the total done by each team member
area_report_df['Total Reports']      = 1
market_report_df['Total Reports']    = 1
hood_report_df['Total Reports']      = 1

area_report_df                      = area_report_df.groupby(['Version']).agg({'Total Reports': 'sum'}).reset_index()
market_report_df                    = market_report_df.groupby(['Version']).agg({'Total Reports': 'sum'}).reset_index()
hood_report_df                      = hood_report_df.groupby(['Version']).agg({'Total Reports': 'sum'}).reset_index()

area_report_df['Type']              = 'Area'
market_report_df['Type']            = 'Market'
hood_report_df['Type']              = 'Hood'


#Append the 3 dataframes together
kpi_df                              = pd.concat([area_report_df,market_report_df,hood_report_df])
kpi_df                              = kpi_df.sort_values(by=['Type','Version'])

#Export as csv file
kpi_df.to_csv(os.path.join(kpi_location,'KPI.csv'),index=False)




#Create graphs for reports over time
marginInches = 1/18
ppi = 96.85 
width_inches = 6.5
height_inches = 3.3
graph_width  = (width_inches - marginInches)   * ppi
graph_height = (height_inches  - marginInches) * ppi
legend_position = 1.10

def HoodGraph():
    hood_df = kpi_df.loc[kpi_df['Type'] =='Hood']
    print(hood_df)
    fig = make_subplots(specs=[[{"secondary_y": True}]])


    fig.add_trace(
        go.Scatter(x=hood_df['Version'],
                y=hood_df['Total Reports'],
                name='Neighborhood Reports',
                line = dict(color="#4160D3", width = 1,dash = 'dash')
                )
        ,secondary_y=False)




    #Set Y-Axis format
    fig.update_yaxes( tickfont = dict(size=12),
                    linecolor='black'  
                    )



    #Set Title
    fig.update_layout(
    title_text="Total Neighborhood Reports by Year",    

    title={
        'y':0.9,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = 12
                )

                        )

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor='White',
    plot_bgcolor ="White"
                        )

    #Set size and margin
    fig.update_layout(
    height    = graph_height,
    width     = graph_width,
        
                    )

    fig.write_image(os.path.join(kpi_location,'hood reports.png'),engine='kaleido',scale=10)

def AreaGraph():
    area_df = kpi_df.loc[kpi_df['Type'] =='Area']
    print(area_df)
    fig = make_subplots(specs=[[{"secondary_y": True}]])


    fig.add_trace(
        go.Scatter(x=area_df['Version'],
                y=area_df['Total Reports'],
                name='Neighborhood Reports',
                line = dict(color="#4160D3", width = 1,dash = 'dash')
                )
        ,secondary_y=False)




    #Set Y-Axis format
    fig.update_yaxes( tickfont = dict(size=12),
                    linecolor='black'  
                    )



    #Set Title
    fig.update_layout(
    title_text="Total Area Reports by Quarter",    

    title={
        'y':0.9,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = 12
                )

                        )

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor='White',
    plot_bgcolor ="White"
                        )

    #Set size and margin
    fig.update_layout(
    height    = graph_height,
    width     = graph_width,
        
                    )

    fig.write_image(os.path.join(kpi_location,'area reports.png'),engine='kaleido',scale=10)

def MarketGraph():
    market_df = kpi_df.loc[kpi_df['Type'] =='Market']
    print(market_df)
    fig = make_subplots(specs=[[{"secondary_y": True}]])


    fig.add_trace(
        go.Scatter(x=market_df['Version'],
                y=market_df['Total Reports'],
                name='Market Reports',
                line = dict(color="#4160D3", width = 1,dash = 'dash')
                )
        ,secondary_y=False)




    #Set Y-Axis format
    fig.update_yaxes( tickfont = dict(size=12),
                    linecolor='black'  
                    )



    #Set Title
    fig.update_layout(
    title_text="Total Market Reports by Quarter",    

    title={
        'y':0.9,
        'x':0.5,
        'xanchor': 'center',
        'yanchor': 'top'},
                    
                    )
    #Set Legend Layout
    fig.update_layout(
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=legend_position ,
        xanchor="center",
        x=0.5,
        font_size = 12
                )

                        )

    #Set Font and Colors
    fig.update_layout(
    font_family="Avenir Next LT Pro",
    font_color='#262626',
    font_size = 10.5,
    paper_bgcolor='White',
    plot_bgcolor ="White"
                        )

    #Set size and margin
    fig.update_layout(
    height    = graph_height,
    width     = graph_width,
        
                    )

    fig.write_image(os.path.join(kpi_location,'market reports.png'),engine='kaleido',scale=10)

HoodGraph()
AreaGraph()
MarketGraph()

#Create word document with the png files inserted
document = Document()
fig = document.add_picture(os.path.join(kpi_location,'area reports.png'), width = Inches(4.5))
fig = document.add_picture(os.path.join(kpi_location,'hood reports.png'), width = Inches(4.5))
fig = document.add_picture(os.path.join(kpi_location,'market reports.png'), width = Inches(4.5))


document.save(os.path.join(kpi_location,'KPI.docx'))  
